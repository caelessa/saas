import mimetypes
import os, uuid, re, json, hashlib, unicodedata, base64, binascii, html, tempfile, subprocess, time
import requests
from datetime import datetime, date, timedelta, timezone
from pathlib import Path
from zoneinfo import ZoneInfo
from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory, send_file, abort, session
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy.orm import joinedload
from sqlalchemy.orm.attributes import set_committed_value
from sqlalchemy import inspect, text, or_
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from werkzeug.datastructures import FileStorage
from services.document_reader import extract_text, parse_cnh, parse_crlv
from services.storage_service import StorageService, StorageNotFoundError
from services.contract_service import gerar_numero_contrato, registrar_evento_contrato
from services.contract_state_service import ContractStateService, ContractStateError
from services.vehicle_state_service import VehicleStateService, VehicleStateError
from services.pdf_service import gerar_pdf_contrato
from services.signature_service import SignatureService, SignatureValidationError
from services.communication_service import CommunicationService, CommunicationError
from services.signature_provider_service import SignatureProviderService, SignatureProviderError
from services.alert_service import sync_operational_alerts, maintenance_indicator
from services.maintenance_notification_service import maintenance_message, reminder_datetime
from services.odometer_ocr_service import read_odometer
from io import BytesIO
from decimal import Decimal
from urllib.parse import quote
from docx import Document as DocxDocument
from functools import wraps
from itsdangerous import URLSafeTimedSerializer, BadSignature, SignatureExpired

BASE=Path(__file__).parent; UPLOAD=BASE/'uploads'; UPLOAD.mkdir(exist_ok=True)
APP_STARTED_AT=datetime.utcnow()
storage=StorageService(UPLOAD)
app=Flask(__name__)
app.config['SECRET_KEY']=os.getenv('SECRET_KEY','dev-change-me')
url=os.getenv('DATABASE_URL','sqlite:///'+str(BASE/'frota_facil.db'))
if url.startswith('postgres://'): url=url.replace('postgres://','postgresql://',1)
app.config['SQLALCHEMY_DATABASE_URI']=url
app.config['SQLALCHEMY_TRACK_MODIFICATIONS']=False
app.config['SQLALCHEMY_ENGINE_OPTIONS']={
 'pool_pre_ping': True,
 'pool_recycle': 240,
 'pool_timeout': 20,
}
app.config['MAX_CONTENT_LENGTH']=120*1024*1024
db=SQLAlchemy(app); login=LoginManager(app); login.login_view='entrar'

SAO_PAULO_TZ=ZoneInfo('America/Sao_Paulo')

# Meta Embedded Signup (credenciais do aplicativo Frota Fácil, nunca do tenant).
META_APP_ID=(os.getenv('META_APP_ID') or '').strip()
META_APP_SECRET=(os.getenv('META_APP_SECRET') or '').strip()
META_WHATSAPP_CONFIG_ID=(os.getenv('META_WHATSAPP_CONFIG_ID') or '').strip()
META_GRAPH_VERSION=(os.getenv('META_GRAPH_VERSION') or 'v23.0').strip()

def _tenant_timezone_name(tenant_id=None):
 try:
  if tenant_id is None:
   if current_user.is_authenticated:
    tenant_id=getattr(current_user,'tenant_id',None)
   if tenant_id is None:
    tenant_id=session.get('owner_tenant_id') or session.get('driver_portal_tenant_id')
  TenantModel=globals().get('Tenant')
  if tenant_id and TenantModel:
   tenant=TenantModel.query.get(int(tenant_id))
   nome=((getattr(tenant,'timezone_name',None) if tenant else None) or 'America/Sao_Paulo').strip()
   try:
    ZoneInfo(nome)
    return nome
   except Exception:
    pass
 except Exception:
  pass
 return 'America/Sao_Paulo'

def _tenant_zone(tenant_id=None):
 try:
  return ZoneInfo(_tenant_timezone_name(tenant_id))
 except Exception:
  return SAO_PAULO_TZ

def _as_tenant_time(value,tenant_id=None):
 if not value:
  return None
 if isinstance(value,date) and not isinstance(value,datetime):
  return value
 if value.tzinfo is None:
  value=value.replace(tzinfo=timezone.utc)
 return value.astimezone(_tenant_zone(tenant_id))

def _as_sao_paulo(value):
 return _as_tenant_time(value,None)

TENANT_TIMEZONE_OPTIONS=[
 ('America/Sao_Paulo','Brasília / São Paulo (UTC−03:00)'),
 ('America/Recife','Recife / Nordeste (UTC−03:00)'),
 ('America/Fortaleza','Fortaleza (UTC−03:00)'),
 ('America/Cuiaba','Cuiabá (UTC−04:00)'),
 ('America/Manaus','Manaus (UTC−04:00)'),
 ('America/Rio_Branco','Rio Branco / Acre (UTC−05:00)'),
]
TENANT_TIMEZONE_NAMES={name for name,_label in TENANT_TIMEZONE_OPTIONS}

def _message_db_time_as_utc_naive(value):
 # Compatibilidade: MessageQueue/MessageEvent usam horário local de São Paulo
 # sem tzinfo. Convertemos somente para apresentação antes do filtro
 # sp_datetime; nada é regravado no banco.
 if not value:
  return value
 if value.tzinfo is not None:
  return value.astimezone(timezone.utc).replace(tzinfo=None)
 return value.replace(tzinfo=SAO_PAULO_TZ).astimezone(timezone.utc).replace(tzinfo=None)

@app.template_filter('sp_datetime')
def sp_datetime(value,fmt='%d/%m/%Y %H:%M'):
 local=_as_tenant_time(value,None)
 return local.strftime(fmt) if local else '-'

@app.template_filter('brl')
def brl(value):
 try:
  n=Decimal(str(value or 0))
  return f'{n:,.2f}'.replace(',','X').replace('.',',').replace('X','.')
 except Exception:
  return '0,00'


class Tenant(db.Model):
 id=db.Column(db.Integer,primary_key=True); nome=db.Column(db.String(120),nullable=False); cnpj=db.Column(db.String(18)); ativo=db.Column(db.Boolean,default=True); conferir_km_motorista=db.Column(db.Boolean,default=False); cobrar_km_excedente=db.Column(db.Boolean,default=False)
 razao_social=db.Column(db.String(180)); nome_fantasia=db.Column(db.String(150)); inscricao_estadual=db.Column(db.String(30)); inscricao_municipal=db.Column(db.String(30)); telefone=db.Column(db.String(30)); email=db.Column(db.String(150)); responsavel_legal=db.Column(db.String(150)); logradouro=db.Column(db.String(180)); numero_endereco=db.Column(db.String(30)); complemento=db.Column(db.String(100)); bairro=db.Column(db.String(100)); cidade=db.Column(db.String(100)); uf=db.Column(db.String(2)); cep=db.Column(db.String(10)); logo_key=db.Column(db.String(255)); favicon_key=db.Column(db.String(255)); cor_primaria=db.Column(db.String(7)); cor_secundaria=db.Column(db.String(7)); timezone_name=db.Column(db.String(80),default='America/Sao_Paulo')
class User(UserMixin,db.Model):
 id=db.Column(db.Integer,primary_key=True); tenant_id=db.Column(db.Integer,db.ForeignKey('tenant.id'),nullable=False); nome=db.Column(db.String(100)); email=db.Column(db.String(120),unique=True,nullable=False); senha=db.Column(db.String(255)); perfil=db.Column(db.String(30),default='admin'); tenant=db.relationship('Tenant')
class Driver(db.Model):
 id=db.Column(db.Integer,primary_key=True); tenant_id=db.Column(db.Integer,index=True,nullable=False); nome=db.Column(db.String(150),nullable=False); cpf=db.Column(db.String(14)); rg=db.Column(db.String(30)); numero_cnh=db.Column(db.String(20)); categoria=db.Column(db.String(5)); data_nascimento=db.Column(db.String(10)); validade_cnh=db.Column(db.String(10)); telefone=db.Column(db.String(30)); telefone2=db.Column(db.String(30)); contato2_nome=db.Column(db.String(150)); contato2_parentesco=db.Column(db.String(40)); telefone3=db.Column(db.String(30)); contato3_nome=db.Column(db.String(150)); contato3_parentesco=db.Column(db.String(40)); email=db.Column(db.String(120)); endereco=db.Column(db.String(250)); logradouro=db.Column(db.String(160)); numero_endereco=db.Column(db.String(20)); complemento=db.Column(db.String(100)); bairro=db.Column(db.String(100)); cidade=db.Column(db.String(100)); uf=db.Column(db.String(2)); cep=db.Column(db.String(10)); status=db.Column(db.String(30),default='Ativo'); criado_em=db.Column(db.DateTime,default=datetime.utcnow)
class Investor(db.Model):
 id=db.Column(db.Integer,primary_key=True); tenant_id=db.Column(db.Integer,index=True,nullable=False); nome=db.Column(db.String(150),nullable=False); cpf_cnpj=db.Column(db.String(20)); telefone=db.Column(db.String(30)); telefone2=db.Column(db.String(30)); contato2_nome=db.Column(db.String(150)); contato2_parentesco=db.Column(db.String(40)); telefone3=db.Column(db.String(30)); contato3_nome=db.Column(db.String(150)); contato3_parentesco=db.Column(db.String(40)); email=db.Column(db.String(120)); regra_repasse=db.Column(db.String(30),default='Valor fixo'); observacoes=db.Column(db.Text)
class InvestorAccess(db.Model):
 __tablename__='investor_access'
 id=db.Column(db.Integer,primary_key=True)
 tenant_id=db.Column(db.Integer,nullable=False,index=True)
 investor_id=db.Column(db.Integer,db.ForeignKey('investor.id'),nullable=False,unique=True,index=True)
 email=db.Column(db.String(150),nullable=False,unique=True,index=True)
 senha=db.Column(db.String(255),nullable=False)
 ativo=db.Column(db.Boolean,default=True,index=True)
 ultimo_acesso_em=db.Column(db.DateTime)
 criado_em=db.Column(db.DateTime,default=datetime.utcnow)
 investor=db.relationship('Investor')

class DriverAccess(db.Model):
 __tablename__='driver_access'
 id=db.Column(db.Integer,primary_key=True)
 tenant_id=db.Column(db.Integer,nullable=False,index=True)
 driver_id=db.Column(db.Integer,db.ForeignKey('driver.id'),nullable=False,unique=True,index=True)
 email=db.Column(db.String(150),nullable=False,unique=True,index=True)
 senha=db.Column(db.String(255),nullable=False)
 ativo=db.Column(db.Boolean,default=True,index=True)
 ultimo_acesso_em=db.Column(db.DateTime)
 criado_em=db.Column(db.DateTime,default=datetime.utcnow)
 driver=db.relationship('Driver')

class InvestorVehicleRule(db.Model):
 __tablename__='investor_vehicle_rule'
 id=db.Column(db.Integer,primary_key=True)
 tenant_id=db.Column(db.Integer,index=True,nullable=False)
 investor_id=db.Column(db.Integer,db.ForeignKey('investor.id'),nullable=False,index=True)
 vehicle_id=db.Column(db.Integer,db.ForeignKey('vehicle.id'),nullable=False,index=True)
 percentual_proprietario=db.Column(db.Numeric(6,2),nullable=False)
 percentual_locadora=db.Column(db.Numeric(6,2),nullable=False)
 vigencia_inicio=db.Column(db.Date,nullable=False,default=date.today)
 vigencia_fim=db.Column(db.Date)
 observacoes=db.Column(db.Text)
 criado_em=db.Column(db.DateTime,default=datetime.utcnow)
 investor=db.relationship('Investor')
 vehicle=db.relationship('Vehicle')

class VehicleInvestment(db.Model):
 __tablename__='vehicle_investment'
 id=db.Column(db.Integer,primary_key=True)
 tenant_id=db.Column(db.Integer,nullable=False,index=True)
 investor_id=db.Column(db.Integer,db.ForeignKey('investor.id'),nullable=False,index=True)
 vehicle_id=db.Column(db.Integer,db.ForeignKey('vehicle.id'),nullable=False,unique=True,index=True)
 data_aquisicao=db.Column(db.Date)
 valor_aquisicao=db.Column(db.Numeric(14,2),default=0)
 capital_proprio=db.Column(db.Numeric(14,2),default=0)
 valor_financiado=db.Column(db.Numeric(14,2),default=0)
 saldo_devedor=db.Column(db.Numeric(14,2),default=0)
 valor_mercado=db.Column(db.Numeric(14,2),default=0)
 criado_em=db.Column(db.DateTime,default=datetime.utcnow)
 atualizado_em=db.Column(db.DateTime,default=datetime.utcnow,onupdate=datetime.utcnow)
 investor=db.relationship('Investor')
 vehicle=db.relationship('Vehicle')

class VehicleInvestmentHistory(db.Model):
 __tablename__='vehicle_investment_history'
 id=db.Column(db.Integer,primary_key=True)
 tenant_id=db.Column(db.Integer,nullable=False,index=True)
 investor_id=db.Column(db.Integer,db.ForeignKey('investor.id'),nullable=False,index=True)
 vehicle_id=db.Column(db.Integer,db.ForeignKey('vehicle.id'),nullable=False,index=True)
 data_aquisicao=db.Column(db.Date)
 valor_aquisicao=db.Column(db.Numeric(14,2),default=0)
 capital_proprio=db.Column(db.Numeric(14,2),default=0)
 valor_financiado=db.Column(db.Numeric(14,2),default=0)
 saldo_devedor=db.Column(db.Numeric(14,2),default=0)
 valor_mercado=db.Column(db.Numeric(14,2),default=0)
 registrado_em=db.Column(db.DateTime,default=datetime.utcnow,index=True)
 investor=db.relationship('Investor')
 vehicle=db.relationship('Vehicle')

class Vehicle(db.Model):
 id=db.Column(db.Integer,primary_key=True); tenant_id=db.Column(db.Integer,index=True,nullable=False); placa=db.Column(db.String(10),nullable=False); renavam=db.Column(db.String(20)); chassi=db.Column(db.String(30)); marca_modelo=db.Column(db.String(150)); ano_fabricacao=db.Column(db.String(4)); ano_modelo=db.Column(db.String(4)); cor=db.Column(db.String(30)); combustivel=db.Column(db.String(100)); motorizacao=db.Column(db.String(20)); km_atual=db.Column(db.Integer,default=0); status=db.Column(db.String(30),default='Disponível'); proprietario_legal=db.Column(db.String(150)); cpf_cnpj_proprietario=db.Column(db.String(20)); investor_id=db.Column(db.Integer,db.ForeignKey('investor.id')); valor_repasse=db.Column(db.Numeric(12,2),default=0); limite_km=db.Column(db.Integer); valor_km_excedente=db.Column(db.Numeric(10,2),default=0); rastreador_id=db.Column(db.String(80)); controlar_oleo=db.Column(db.Boolean,default=False); ultima_troca_oleo_km=db.Column(db.Integer); intervalo_oleo_km=db.Column(db.Integer,default=10000); alerta_oleo_km=db.Column(db.Integer,default=100); current_driver_id=db.Column(db.Integer,db.ForeignKey('driver.id')); current_contract_id=db.Column(db.Integer,db.ForeignKey('contract.id')); status_changed_at=db.Column(db.DateTime); status_reason=db.Column(db.String(255)); investor=db.relationship('Investor'); current_driver=db.relationship('Driver',foreign_keys=[current_driver_id]); current_contract=db.relationship('Contract',foreign_keys=[current_contract_id],post_update=True)
class Odometer(db.Model):
 id=db.Column(db.Integer,primary_key=True); tenant_id=db.Column(db.Integer,index=True,nullable=False); vehicle_id=db.Column(db.Integer,db.ForeignKey('vehicle.id')); km=db.Column(db.Integer,nullable=False); origem=db.Column(db.String(40)); data=db.Column(db.DateTime,default=datetime.utcnow); vehicle=db.relationship('Vehicle')
class MileageRequest(db.Model):
 id=db.Column(db.Integer,primary_key=True); tenant_id=db.Column(db.Integer,index=True,nullable=False); vehicle_id=db.Column(db.Integer,db.ForeignKey('vehicle.id'),nullable=False); driver_id=db.Column(db.Integer,db.ForeignKey('driver.id'),nullable=False); token=db.Column(db.String(64),unique=True,nullable=False,index=True); status=db.Column(db.String(30),default='Pendente'); expires_at=db.Column(db.DateTime); sent_at=db.Column(db.DateTime,default=datetime.utcnow); submitted_at=db.Column(db.DateTime); km=db.Column(db.Integer); previous_km=db.Column(db.Integer); photo=db.Column(db.String(255)); notes=db.Column(db.Text); vehicle=db.relationship('Vehicle'); driver=db.relationship('Driver')
class ContractTemplate(db.Model):
 id=db.Column(db.Integer,primary_key=True); tenant_id=db.Column(db.Integer,index=True,nullable=False); nome=db.Column(db.String(120)); descricao=db.Column(db.String(255)); versao=db.Column(db.Integer,default=1); padrao=db.Column(db.Boolean,default=False); tipo_veiculo=db.Column(db.String(30)); possui_limite_km=db.Column(db.Boolean,default=False); conteudo=db.Column(db.Text); nome_original=db.Column(db.String(255)); arquivo_original=db.Column(db.String(255)); hash_original=db.Column(db.String(64)); preparado_em=db.Column(db.DateTime); gestora_nome=db.Column(db.String(180)); gestora_fantasia=db.Column(db.String(120)); gestora_cnpj=db.Column(db.String(30)); gestora_endereco=db.Column(db.String(255)); parceira_nome=db.Column(db.String(180)); parceira_cnpj=db.Column(db.String(30)); parceira_endereco=db.Column(db.String(255)); ativo=db.Column(db.Boolean,default=True)
class Contract(db.Model):
 id=db.Column(db.Integer,primary_key=True)
 tenant_id=db.Column(db.Integer,index=True,nullable=False)
 numero_contrato=db.Column(db.String(30),unique=True,index=True)
 driver_id=db.Column(db.Integer,db.ForeignKey('driver.id'))
 vehicle_id=db.Column(db.Integer,db.ForeignKey('vehicle.id'))
 template_id=db.Column(db.Integer,db.ForeignKey('contract_template.id'))
 template_nome=db.Column(db.String(120))
 template_versao=db.Column(db.Integer,default=1)
 versao=db.Column(db.Integer,default=1)
 data_inicio=db.Column(db.String(10))
 hora_inicio=db.Column(db.String(5))
 data_fim=db.Column(db.String(10))
 periodicidade=db.Column(db.String(30))
 dia_vencimento=db.Column(db.String(30))
 valor_locacao=db.Column(db.Numeric(12,2))
 caucao=db.Column(db.Numeric(12,2))
 franquia=db.Column(db.Numeric(12,2))
 limite_km=db.Column(db.Integer)
 valor_km_excedente=db.Column(db.Numeric(10,2))
 multa_atraso_percentual=db.Column(db.Numeric(6,2))
 juros_mes_percentual=db.Column(db.Numeric(6,2))
 indice_correcao=db.Column(db.String(30))
 prazo_bloqueio_horas=db.Column(db.Integer)
 multa_diaria=db.Column(db.Numeric(12,2))
 taxa_adm_multas_percentual=db.Column(db.Numeric(6,2))
 nacionalidade=db.Column(db.String(60))
 estado_civil=db.Column(db.String(60))
 profissao=db.Column(db.String(100))
 cidade_assinatura=db.Column(db.String(100))
 status=db.Column(db.String(30),default='Gerado')
 texto_final=db.Column(db.Text)
 criado_por_id=db.Column(db.Integer,db.ForeignKey('user.id'))
 criado_em=db.Column(db.DateTime,default=datetime.utcnow)
 atualizado_em=db.Column(db.DateTime,default=datetime.utcnow,onupdate=datetime.utcnow)
 assinado_em=db.Column(db.DateTime)
 assinatura_id=db.Column(db.String(120))
 documento_id=db.Column(db.Integer,db.ForeignKey('document.id'))
 arquivo_pdf=db.Column(db.String(255))
 hash_documento=db.Column(db.String(64))
 arquivo_pdf_assinado=db.Column(db.String(255))
 hash_documento_assinado=db.Column(db.String(64))
 documento_assinado_id=db.Column(db.Integer)
 codigo_publico=db.Column(db.String(24),unique=True,index=True)
 enviado_whatsapp_em=db.Column(db.DateTime)
 visualizado_em=db.Column(db.DateTime)
 gerado_em=db.Column(db.DateTime)
 clicksign_envelope_id=db.Column(db.String(80),index=True)
 clicksign_document_id=db.Column(db.String(80))
 clicksign_signer_id=db.Column(db.String(80))
 clicksign_status=db.Column(db.String(30))
 clicksign_sent_at=db.Column(db.DateTime)
 driver=db.relationship('Driver')
 vehicle=db.relationship('Vehicle',foreign_keys=[vehicle_id])
 template=db.relationship('ContractTemplate')
 criado_por=db.relationship('User')

class Signature(db.Model):
 id=db.Column(db.Integer,primary_key=True)
 tenant_id=db.Column(db.Integer,index=True,nullable=False)
 contract_id=db.Column(db.Integer,db.ForeignKey('contract.id'),nullable=False,unique=True,index=True)
 driver_id=db.Column(db.Integer,db.ForeignKey('driver.id'),nullable=False,index=True)
 status=db.Column(db.String(30),default='Assinada')
 signatario_nome=db.Column(db.String(150),nullable=False)
 cpf_confirmado=db.Column(db.String(14))
 arquivo_assinatura=db.Column(db.String(255),nullable=False)
 hash_assinatura=db.Column(db.String(64),nullable=False)
 hash_documento=db.Column(db.String(64),nullable=False)
 ip=db.Column(db.String(64))
 user_agent=db.Column(db.Text)
 aceite_texto=db.Column(db.Text)
 assinado_em=db.Column(db.DateTime,default=datetime.utcnow,index=True)
 contract=db.relationship('Contract',backref=db.backref('signature',uselist=False,cascade='all, delete-orphan'))
 driver=db.relationship('Driver')

class ContractEvent(db.Model):
 id=db.Column(db.Integer,primary_key=True)
 tenant_id=db.Column(db.Integer,index=True,nullable=False)
 contract_id=db.Column(db.Integer,db.ForeignKey('contract.id'),nullable=False,index=True)
 user_id=db.Column(db.Integer,db.ForeignKey('user.id'))
 evento=db.Column(db.String(60),nullable=False)
 descricao=db.Column(db.Text)
 status_anterior=db.Column(db.String(30))
 status_novo=db.Column(db.String(30))
 criado_em=db.Column(db.DateTime,default=datetime.utcnow,index=True)
 user=db.relationship('User')
 contract=db.relationship('Contract',backref=db.backref('eventos',lazy='dynamic',cascade='all, delete-orphan'))

class VehicleEvent(db.Model):
 id=db.Column(db.Integer,primary_key=True)
 tenant_id=db.Column(db.Integer,index=True,nullable=False)
 vehicle_id=db.Column(db.Integer,db.ForeignKey('vehicle.id'),nullable=False,index=True)
 contract_id=db.Column(db.Integer,db.ForeignKey('contract.id'))
 driver_id=db.Column(db.Integer,db.ForeignKey('driver.id'))
 user_id=db.Column(db.Integer,db.ForeignKey('user.id'))
 evento=db.Column(db.String(60),nullable=False)
 descricao=db.Column(db.Text)
 status_anterior=db.Column(db.String(30))
 status_novo=db.Column(db.String(30))
 criado_em=db.Column(db.DateTime,default=datetime.utcnow,index=True)
 vehicle=db.relationship('Vehicle',foreign_keys=[vehicle_id])
 contract=db.relationship('Contract',foreign_keys=[contract_id])
 driver=db.relationship('Driver',foreign_keys=[driver_id])
 user=db.relationship('User',foreign_keys=[user_id])

class Document(db.Model):
 id=db.Column(db.Integer,primary_key=True); tenant_id=db.Column(db.Integer,index=True,nullable=False); tipo=db.Column(db.String(40)); entidade=db.Column(db.String(30)); entidade_id=db.Column(db.Integer); identificador=db.Column(db.String(180),index=True); numero_documento=db.Column(db.String(60),index=True); nome_original=db.Column(db.String(255)); arquivo=db.Column(db.String(255)); hash_sha256=db.Column(db.String(64)); status=db.Column(db.String(20),default='Ativo'); versao=db.Column(db.Integer,default=1); criado_em=db.Column(db.DateTime,default=datetime.utcnow)
class Maintenance(db.Model):
 id=db.Column(db.Integer,primary_key=True); tenant_id=db.Column(db.Integer,index=True,nullable=False); vehicle_id=db.Column(db.Integer,db.ForeignKey('vehicle.id')); tipo=db.Column(db.String(100)); data=db.Column(db.String(10)); km=db.Column(db.Integer); custo=db.Column(db.Numeric(12,2)); proxima_km=db.Column(db.Integer); proxima_data=db.Column(db.String(10)); proxima_hora=db.Column(db.String(5)); alerta_km_antes=db.Column(db.Integer,default=500); alerta_dias_antes=db.Column(db.Integer,default=7); observacoes=db.Column(db.Text); status=db.Column(db.String(20),default='Ativa',index=True); oficina=db.Column(db.String(160)); notificar_motorista=db.Column(db.Boolean,default=False); lembrete_um_dia=db.Column(db.Boolean,default=True); notificacao_agendamento_id=db.Column(db.Integer); notificacao_lembrete_id=db.Column(db.Integer); concluida_em=db.Column(db.DateTime); concluida_por_id=db.Column(db.Integer,db.ForeignKey('user.id')); vehicle=db.relationship('Vehicle'); concluida_por=db.relationship('User',foreign_keys=[concluida_por_id])
class Inspection(db.Model):
 id=db.Column(db.Integer,primary_key=True)
 tenant_id=db.Column(db.Integer,index=True,nullable=False)
 vehicle_id=db.Column(db.Integer,db.ForeignKey('vehicle.id'),nullable=False,index=True)
 driver_id=db.Column(db.Integer,db.ForeignKey('driver.id'))
 contract_id=db.Column(db.Integer,db.ForeignKey('contract.id'))
 token=db.Column(db.String(64),unique=True,nullable=False,index=True)
 status=db.Column(db.String(30),default='Pendente',index=True)
 tipo_vistoria=db.Column(db.String(20),default='guiada',index=True)
 requested_at=db.Column(db.DateTime,default=datetime.utcnow,index=True)
 expires_at=db.Column(db.DateTime,index=True)
 started_at=db.Column(db.DateTime)
 submitted_at=db.Column(db.DateTime)
 video_key=db.Column(db.String(255))
 video_mime=db.Column(db.String(80))
 duration_seconds=db.Column(db.Integer)
 painel_photo_key=db.Column(db.String(255))
 painel_photo_mime=db.Column(db.String(80))
 front_photo_key=db.Column(db.String(255))
 front_photo_mime=db.Column(db.String(80))
 right_photo_key=db.Column(db.String(255))
 right_photo_mime=db.Column(db.String(80))
 rear_photo_key=db.Column(db.String(255))
 rear_photo_mime=db.Column(db.String(80))
 left_photo_key=db.Column(db.String(255))
 left_photo_mime=db.Column(db.String(80))
 km_informada=db.Column(db.Integer)
 brightness_avg=db.Column(db.Numeric(8,2))
 brightness_status=db.Column(db.String(30))
 damage_analysis_status=db.Column(db.String(30),default='NAO_ANALISADA',index=True)
 damage_analysis_level=db.Column(db.String(30))
 damage_analysis_summary=db.Column(db.Text)
 damage_analysis_at=db.Column(db.DateTime)
 notes=db.Column(db.Text)
 vehicle=db.relationship('Vehicle',foreign_keys=[vehicle_id])
 driver=db.relationship('Driver',foreign_keys=[driver_id])
 contract=db.relationship('Contract',foreign_keys=[contract_id])

class InspectionAttempt(db.Model):
 id=db.Column(db.Integer,primary_key=True)
 inspection_id=db.Column(db.Integer,db.ForeignKey('inspection.id'),nullable=False,index=True)
 tenant_id=db.Column(db.Integer,index=True,nullable=False)
 video_key=db.Column(db.String(255),nullable=False)
 video_mime=db.Column(db.String(80))
 duration_seconds=db.Column(db.Integer)
 painel_photo_key=db.Column(db.String(255))
 painel_photo_mime=db.Column(db.String(80))
 km_informada=db.Column(db.Integer)
 brightness_avg=db.Column(db.Numeric(8,2))
 brightness_min=db.Column(db.Numeric(8,2))
 dark_ratio=db.Column(db.Numeric(8,4))
 submitted_at=db.Column(db.DateTime,default=datetime.utcnow,index=True)
 decision=db.Column(db.String(30),default='Pendente')
 decision_notes=db.Column(db.Text)
 decided_at=db.Column(db.DateTime)
 inspection=db.relationship('Inspection',backref=db.backref('attempts',lazy=True,order_by='InspectionAttempt.id.desc()'))

def _encerrar_vistorias_anteriores(item):
 """Encerra solicitações antigas já satisfeitas pela vistoria informada.

 Também alcança registros legados sem contract_id, desde que pertençam ao mesmo
 veículo e motorista. O filtro por id impede que um link antigo encerre uma
 solicitação nova criada intencionalmente depois dele.
 """
 if not item or not item.id or not item.tenant_id or not item.vehicle_id:
  return 0
 filtros=[
  Inspection.tenant_id==item.tenant_id,
  Inspection.vehicle_id==item.vehicle_id,
  Inspection.id<item.id,
  Inspection.status.in_(['Pendente','Regravar','Recebida']),
 ]
 if item.contract_id:
  vinculo=Inspection.contract_id==item.contract_id
  if item.driver_id:
   vinculo=or_(vinculo,db.and_(Inspection.contract_id.is_(None),Inspection.driver_id==item.driver_id))
  filtros.append(vinculo)
 elif item.driver_id:
  filtros.append(Inspection.driver_id==item.driver_id)
 else:
  return 0
 anteriores=Inspection.query.filter(*filtros).all()
 agora=datetime.utcnow()
 for anterior in anteriores:
  anterior.status='Encerrada'
  anterior.expires_at=agora
  anterior.notes=f'Encerrada automaticamente pela conclusão da vistoria #{item.id} do mesmo contrato.'
 return len(anteriores)

class Alert(db.Model):
 id=db.Column(db.Integer,primary_key=True); tenant_id=db.Column(db.Integer,index=True,nullable=False); titulo=db.Column(db.String(150)); mensagem=db.Column(db.Text); nivel=db.Column(db.String(20),default='info'); lido=db.Column(db.Boolean,default=False); criado_em=db.Column(db.DateTime,default=datetime.utcnow); source_key=db.Column(db.String(120),index=True); entidade=db.Column(db.String(40)); entidade_id=db.Column(db.Integer); action_url=db.Column(db.String(255)); atualizado_em=db.Column(db.DateTime,default=datetime.utcnow,onupdate=datetime.utcnow); resolvido_em=db.Column(db.DateTime)
class Integration(db.Model):
 id=db.Column(db.Integer,primary_key=True); tenant_id=db.Column(db.Integer,index=True,nullable=False); tipo=db.Column(db.String(40)); ativo=db.Column(db.Boolean,default=False); configuracao=db.Column(db.Text)

class MessageQueue(db.Model):
 id=db.Column(db.Integer,primary_key=True)
 tenant_id=db.Column(db.Integer,index=True,nullable=False)
 channel=db.Column(db.String(30),nullable=False,default='whatsapp')
 provider=db.Column(db.String(40),nullable=False,default='whatsapp_web')
 recipient=db.Column(db.String(40),nullable=False,index=True)
 recipient_name=db.Column(db.String(150))
 message_type=db.Column(db.String(50),default='texto')
 body=db.Column(db.Text,nullable=False)
 template_name=db.Column(db.String(120))
 template_parameters=db.Column(db.Text)
 related_entity=db.Column(db.String(40))
 related_entity_id=db.Column(db.Integer)
 status=db.Column(db.String(30),default='PENDENTE',index=True)
 external_id=db.Column(db.String(180),index=True)
 attempts=db.Column(db.Integer,default=0)
 error_message=db.Column(db.Text)
 scheduled_at=db.Column(db.DateTime,index=True)
 sent_at=db.Column(db.DateTime)
 created_at=db.Column(db.DateTime,default=datetime.utcnow,index=True)
 updated_at=db.Column(db.DateTime,default=datetime.utcnow,onupdate=datetime.utcnow)

class MessageEvent(db.Model):
 id=db.Column(db.Integer,primary_key=True)
 tenant_id=db.Column(db.Integer,index=True,nullable=False)
 message_id=db.Column(db.Integer,db.ForeignKey('message_queue.id'),nullable=False,index=True)
 event=db.Column(db.String(40),nullable=False)
 description=db.Column(db.Text)
 created_at=db.Column(db.DateTime,default=datetime.utcnow,index=True)
 message=db.relationship('MessageQueue',backref=db.backref('events',lazy='dynamic',cascade='all, delete-orphan'))

class BillingAudit(db.Model):
 id=db.Column(db.Integer,primary_key=True)
 tenant_id=db.Column(db.Integer,index=True,nullable=False)
 contract_id=db.Column(db.Integer,index=True,nullable=False)
 message_id=db.Column(db.Integer,db.ForeignKey('message_queue.id'),index=True)
 driver_name=db.Column(db.String(150))
 vehicle_label=db.Column(db.String(180))
 plate=db.Column(db.String(20))
 billing_date=db.Column(db.Date,index=True,nullable=False)
 base_amount=db.Column(db.Numeric(12,2),nullable=False,default=0)
 km_period=db.Column(db.Integer)
 km_limit=db.Column(db.Integer)
 km_excess=db.Column(db.Integer,default=0)
 excess_rate=db.Column(db.Numeric(12,2),default=0)
 excess_amount=db.Column(db.Numeric(12,2),default=0)
 total_amount=db.Column(db.Numeric(12,2),nullable=False,default=0)
 body=db.Column(db.Text,nullable=False)
 template_name=db.Column(db.String(120))
 provider=db.Column(db.String(40))
 status=db.Column(db.String(30),default='GERADA')
 external_id=db.Column(db.String(180))
 error_message=db.Column(db.Text)
 payment_status=db.Column(db.String(20),default='PENDENTE',index=True)
 paid_at=db.Column(db.DateTime,index=True)
 paid_by_id=db.Column(db.Integer,db.ForeignKey('user.id'))
 payment_method=db.Column(db.String(50))
 payment_notes=db.Column(db.Text)
 reminder_count=db.Column(db.Integer,default=0)
 last_reminder_at=db.Column(db.DateTime,index=True)
 closed_at=db.Column(db.DateTime)
 receipt_token=db.Column(db.String(64),unique=True,index=True)
 receipt_key=db.Column(db.String(255))
 receipt_name=db.Column(db.String(255))
 receipt_mime=db.Column(db.String(100))
 receipt_uploaded_at=db.Column(db.DateTime,index=True)
 created_at=db.Column(db.DateTime,default=datetime.utcnow,index=True)
 message=db.relationship('MessageQueue')


class FrotaAdmin(db.Model):
 __tablename__='frota_admin'
 id=db.Column(db.Integer,primary_key=True)
 nome=db.Column(db.String(120),nullable=False)
 email=db.Column(db.String(150),nullable=False,unique=True,index=True)
 senha=db.Column(db.String(255),nullable=False)
 ativo=db.Column(db.Boolean,default=True,index=True)
 perfil=db.Column(db.String(30),default='SUPORTE')
 ultimo_acesso_em=db.Column(db.DateTime)
 criado_em=db.Column(db.DateTime,default=datetime.utcnow,index=True)

class SupportTicket(db.Model):
 __tablename__='support_ticket'
 id=db.Column(db.Integer,primary_key=True)
 tenant_id=db.Column(db.Integer,index=True,nullable=False)
 user_id=db.Column(db.Integer,db.ForeignKey('user.id'),nullable=False,index=True)
 titulo=db.Column(db.String(180),nullable=False)
 categoria=db.Column(db.String(50),nullable=False,default='Duvida')
 prioridade=db.Column(db.String(20),nullable=False,default='Normal')
 descricao=db.Column(db.Text,nullable=False)
 pagina_origem=db.Column(db.String(255))
 status=db.Column(db.String(30),nullable=False,default='ABERTO',index=True)
 resposta=db.Column(db.Text)
 respondido_por_id=db.Column(db.Integer,db.ForeignKey('user.id'))
 respondido_por_admin_id=db.Column(db.Integer,db.ForeignKey('frota_admin.id'))
 respondido_por_nome=db.Column(db.String(120))
 criado_em=db.Column(db.DateTime,default=datetime.utcnow,index=True)
 atualizado_em=db.Column(db.DateTime,default=datetime.utcnow,onupdate=datetime.utcnow,index=True)
 resolvido_em=db.Column(db.DateTime)
 user=db.relationship('User',foreign_keys=[user_id])
 respondido_por=db.relationship('User',foreign_keys=[respondido_por_id])
 respondido_por_admin=db.relationship('FrotaAdmin',foreign_keys=[respondido_por_admin_id])
@login.user_loader
def load_user(uid):
 return User.query.options(joinedload(User.tenant)).filter_by(id=int(uid)).first()
def tid(): return current_user.tenant_id


def json_safe(value):
 if value is None or isinstance(value,(str,int,float,bool)):
  return value
 if isinstance(value,(datetime,date)):
  return value.isoformat()
 if isinstance(value,Decimal):
  return str(value)
 return str(value)

def model_rows(model, tenant_id):
 rows=[]
 for item in model.query.filter_by(tenant_id=tenant_id).all():
  rows.append({column.name:json_safe(getattr(item,column.name)) for column in item.__table__.columns})
 return rows

def tenant_backup_payload(tenant_id):
 tenant=Tenant.query.get(tenant_id)
 models=[Driver,Investor,InvestorAccess,InvestorVehicleRule,VehicleInvestment,VehicleInvestmentHistory,Vehicle,Odometer,MileageRequest,ContractTemplate,Contract,ContractEvent,Document,Maintenance,Inspection,Alert,Integration,MessageQueue,MessageEvent,BillingAudit,SupportTicket]
 return {
  'formato':'frota-facil-tenant-backup-v1',
  'gerado_em_utc':datetime.now(timezone.utc).isoformat(),
  'tenant':{column.name:json_safe(getattr(tenant,column.name)) for column in Tenant.__table__.columns},
  'usuarios':model_rows(User,tenant_id),
  'dados':{model.__tablename__:model_rows(model,tenant_id) for model in models},
 }



def corrigir_dados_crlv_ocr(texto, dados):
 """Aplica validações semânticas nos campos extraídos do CRLV.

 O parser original continua sendo a primeira fonte. Esta camada corrige falsos
 positivos comuns do layout do CRLV Digital, especialmente quando a linha de
 CHASSI é confundida com MARCA / MODELO / VERSÃO.
 """
 dados=dict(dados or {})
 linhas=[re.sub(r'\s+',' ',x).strip() for x in (texto or '').splitlines()]
 linhas=[x for x in linhas if x]

 # VIN/chassi: 17 caracteres, sem I/O/Q, conforme padrão usual do VIN.
 vin_re=re.compile(r'(?<![A-Z0-9])([A-HJ-NPR-Z0-9]{17})(?![A-Z0-9])',re.I)
 vins=[]
 for linha in linhas:
  for m in vin_re.finditer(linha.upper()):
   vin=m.group(1).upper()
   if any(c.isalpha() for c in vin) and any(c.isdigit() for c in vin):
    vins.append(vin)

 def parece_chassi(valor):
  v=(valor or '').upper().strip()
  return bool(vin_re.search(v)) or ('***' in v and any(ch.isdigit() for ch in v))

 # Se o parser trouxe chassi incompleto ou colocou a linha do chassi em marca/modelo,
 # usa o VIN detectado diretamente no texto.
 if vins:
  chassi_atual=(dados.get('chassi') or '').strip().upper()
  if not vin_re.fullmatch(chassi_atual):
   dados['chassi']=vins[0]

 marca_atual=(dados.get('marca_modelo') or '').strip()
 marca_invalida=(
  not marca_atual
  or parece_chassi(marca_atual)
  or marca_atual.count('*')>=2
  or bool(re.fullmatch(r'[\W_0-9]+',marca_atual))
 )

 if marca_invalida:
  # No CRLV Digital, a linha MARCA/MODELO/VERSÃO normalmente aparece
  # imediatamente antes da linha ESPÉCIE/TIPO (ex.: PASSAGEIRO AUTOMOVEL).
  especie_re=re.compile(
   r'^(PASSAGEIRO|CARGA|MISTO|ESPECIAL|TRACAO|COMPETICAO)\b.*\b'
   r'(AUTOMOVEL|CAMIONETA|CAMINHONETE|UTILITARIO|MOTOCICLETA|MOTONETA|'
   r'CAMINHAO|ONIBUS|MICROONIBUS|REBOQUE|SEMIRREBOQUE)\b',re.I
  )
  cabecalhos={
   'MARCA / MODELO / VERSÃO','MARCA/MODELO/VERSÃO','PLACA ANTERIOR / UF CHASSI',
   'PLACA ANTERIOR/UF CHASSI','COR PREDOMINANTE','ESPÉCIE / TIPO',
   'ESPECIE / TIPO','COMBUSTÍVEL','COMBUSTIVEL'
  }

  candidato=None
  for i,linha in enumerate(linhas):
   if especie_re.search(linha) and i>0:
    # olha até 3 linhas para trás porque alguns extratores inserem linhas vazias/ruído
    for j in range(i-1,max(-1,i-4),-1):
     c=linhas[j].strip()
     cu=c.upper()
     if cu in cabecalhos:
      continue
     if parece_chassi(c) or c.count('*')>=2:
      continue
     if not re.search(r'[A-ZÀ-Ü]',cu):
      continue
     if len(c)<3 or len(c)>90:
      continue
     # descarta campos que claramente são números/códigos soltos.
     if re.fullmatch(r'[\d\s./-]+',c):
      continue
     candidato=c
     break
   if candidato:
    break

  if candidato:
   dados['marca_modelo']=candidato

 return dados


def limpar_campo_ocr_veiculo(campo, valor):
 valor=(valor or '').strip()
 valor=re.sub(r'\s+',' ',valor)
 if campo=='combustivel':
  # Corrige concatenações comuns do OCR do CRLV, como:
  # ELETRICO/FONTE EXTERNAPARTICULAR
  valor=re.sub(r'(?i)(FONTE EXTERNA)(PARTICULAR|ALUGUEL|OFICIAL|APRENDIZAGEM)$',r'\1',valor)
  valor=re.sub(r'(?i)(ELETRICO)(PARTICULAR|ALUGUEL|OFICIAL|APRENDIZAGEM)$',r'\1',valor)
  return valor[:100]
 limites={
  'placa':10,'renavam':20,'chassi':40,'marca_modelo':120,
  'ano_fabricacao':10,'ano_modelo':10,'cor':40,'status':30,
  'proprietario_legal':180,'cpf_cnpj_proprietario':30,'rastreador_id':80,
 }
 return valor[:limites.get(campo,255)]

def slug_documento(value, fallback='SEM-NOME'):
 value=unicodedata.normalize('NFKD',value or '').encode('ascii','ignore').decode('ascii')
 value=re.sub(r'[^A-Za-z0-9]+','-',value.upper()).strip('-')
 return value[:60] or fallback

def extensao_segura(nome_original):
 ext=Path(nome_original or '').suffix.lower()
 return ext if ext in {'.pdf','.png','.jpg','.jpeg','.webp'} else '.bin'

def identificador_documento(tipo, entidade_id, referencia, ano=None):
 tipo=slug_documento(tipo,'DOC')
 referencia=slug_documento(referencia)
 ano=str(ano or datetime.now().year)
 return f'{tipo}-{int(entidade_id):06d}-{referencia}-{ano}'

def armazenar_documento_cadastro(tipo, entidade, entidade_id, referencia, numero_documento, nome_original, temp_key, mimetype, ano=None):
 """Move o arquivo temporário do OCR para a pasta definitiva e registra no banco."""
 if not temp_key:
  return None
 prefixo=f'{tid()}/temporarios/'
 if not temp_key.startswith(prefixo):
  raise ValueError('Chave temporária inválida.')
 conteudo=storage.download(temp_key)
 identificador=identificador_documento(tipo,entidade_id,referencia,ano)
 ext=extensao_segura(nome_original)
 chave=f'{tid()}/documentos/{entidade.lower()}s/{entidade_id}/{tipo.lower()}/{identificador}{ext}'
 storage.upload(BytesIO(conteudo),chave,mimetype or 'application/octet-stream')
 documento=Document(
  tenant_id=tid(),
  tipo=tipo,
  entidade=entidade,
  entidade_id=entidade_id,
  identificador=identificador,
  numero_documento=(numero_documento or '').strip() or None,
  nome_original=secure_filename(nome_original or f'{identificador}{ext}'),
  arquivo=chave,
  hash_sha256=hashlib.sha256(conteudo).hexdigest(),
  status='Ativo',
 )
 db.session.add(documento)
 storage.delete(temp_key)
 return documento

SAO_PAULO=ZoneInfo("America/Sao_Paulo")

def agora_sao_paulo_naive():
 return datetime.now(SAO_PAULO).replace(tzinfo=None)

def endereco_completo_motorista(driver):
 partes=[]
 if driver.logradouro:
  linha=driver.logradouro
  if driver.numero_endereco: linha+=f", nº {driver.numero_endereco}"
  partes.append(linha)
 if driver.complemento: partes.append(driver.complemento)
 local=", ".join([x for x in [driver.bairro,driver.cidade] if x])
 if driver.uf: local=(local+("/" if local else "")+driver.uf)
 if local: partes.append(local)
 if driver.cep: partes.append(f"CEP {driver.cep}")
 return ", ".join(partes) if partes else (driver.endereco or "A preencher")

def valor_extenso(valor):
 from num2words import num2words
 try: return num2words(Decimal(str(valor or 0)),lang="pt_BR",to="currency")
 except Exception: return "valor não informado"

def normalize_phone(value):
 """Normaliza telefones para o formato internacional usado pela Meta.

 Para números brasileiros:
 - aceita máscara, espaços e pontuação;
 - aceita 10/11 dígitos sem DDI e acrescenta 55;
 - aceita 12/13 dígitos já iniciados por 55;
 - retorna vazio para formatos claramente inválidos.
 """
 digits=re.sub(r'\D','',value or '')
 if not digits:
  return ''
 if digits.startswith('00'):
  digits=digits[2:]
 if len(digits) in (10,11):
  digits='55'+digits
 if len(digits) not in (12,13) or not digits.startswith('55'):
  return ''
 return digits

def normalizar_telefones_form(campos):
 """Normaliza campos de telefone recebidos de formulário.

 Retorna (valores_normalizados, campo_invalido). Campos vazios são aceitos.
 """
 valores={}
 for campo,valor in campos.items():
  bruto=(valor or '').strip()
  if not bruto:
   valores[campo]=''
   continue
  normalizado=normalize_phone(bruto)
  if not normalizado:
   return {},campo
  valores[campo]=normalizado
 return valores,None

def active_request(vehicle_id, driver_id):
 return MileageRequest.query.filter_by(tenant_id=tid(),vehicle_id=vehicle_id,driver_id=driver_id,status='Pendente').filter(MileageRequest.expires_at>datetime.utcnow()).order_by(MileageRequest.id.desc()).first()



def motorista_atual_veiculo(vehicle):
 # O vínculo gravado pela máquina de estados é a fonte primária.
 if vehicle.current_driver_id and vehicle.current_contract_id and vehicle.status in ('Reservado','Alugado'):
  return Driver.query.filter_by(id=vehicle.current_driver_id,tenant_id=vehicle.tenant_id).first()
 # Fallback para bases antigas que ainda não possuem current_driver_id preenchido.
 contrato=Contract.query.filter(Contract.tenant_id==vehicle.tenant_id,Contract.vehicle_id==vehicle.id,Contract.status.in_(['Gerado','Enviado','Visualizado','Assinado','Ativo'])).order_by(Contract.id.desc()).first()
 return contrato.driver if contrato else None

def criar_mensagem_whatsapp(*, tenant_id, driver, body, message_type, related_entity, related_entity_id, scheduled_at=None, template_name=None, template_parameters=None):
 telefone=normalize_phone(driver.telefone if driver else None)
 if not telefone:
  return None, None, 'Motorista sem telefone/WhatsApp válido.'
 integration=Integration.query.filter_by(tenant_id=tenant_id,tipo='whatsapp').first()
 cfg=CommunicationService.parse_config(integration)
 provider_cfg=(cfg.get('provider') or 'web').lower()
 fila=MessageQueue(
  tenant_id=tenant_id,channel='whatsapp',provider='whatsapp_business' if provider_cfg=='business' else 'whatsapp_web',
  recipient=telefone,recipient_name=driver.nome,message_type=message_type,body=body,template_name=template_name,template_parameters=json.dumps(template_parameters or [],ensure_ascii=False),
  related_entity=related_entity,related_entity_id=related_entity_id,status='AGENDADA' if scheduled_at else 'PENDENTE',
  scheduled_at=scheduled_at,created_at=agora_sao_paulo_naive(),updated_at=agora_sao_paulo_naive(),
 )
 db.session.add(fila); db.session.flush()
 if scheduled_at:
  return fila, None, None
 try:
  result=CommunicationService().send_whatsapp(phone=telefone,message=body,integration=integration,template_name=template_name,template_language=cfg.get('template_language') or 'pt_BR',template_parameters=template_parameters or [])
  fila.provider=result.provider; fila.status=result.status; fila.external_id=result.external_id; fila.attempts=(fila.attempts or 0)+1
  fila.sent_at=agora_sao_paulo_naive() if result.status=='ENVIADA' else None
  db.session.add(MessageEvent(tenant_id=tenant_id,message_id=fila.id,event=result.status,description='Mensagem processada pelo provedor configurado.',created_at=agora_sao_paulo_naive()))
  return fila, result.redirect_url, None
 except CommunicationError as exc:
  fila.status='FALHA'; fila.error_message=str(exc); fila.attempts=(fila.attempts or 0)+1
  db.session.add(MessageEvent(tenant_id=tenant_id,message_id=fila.id,event='FALHA',description=str(exc),created_at=agora_sao_paulo_naive()))
  return fila, None, str(exc)

def processar_mensagens_agendadas(tenant_id=None, limit=100):
 now=agora_sao_paulo_naive()
 q=MessageQueue.query.filter(MessageQueue.status=='AGENDADA',MessageQueue.scheduled_at.isnot(None),MessageQueue.scheduled_at<=now)
 if tenant_id is not None: q=q.filter(MessageQueue.tenant_id==tenant_id)
 items=q.order_by(MessageQueue.scheduled_at.asc()).limit(limit).all()
 processed=0
 for fila in items:
  integration=Integration.query.filter_by(tenant_id=fila.tenant_id,tipo='whatsapp').first()
  cfg=CommunicationService.parse_config(integration)
  # Freio de segurança: uma cobrança agendada nunca pode ser enviada se a
  # automação de cobranças deste tenant estiver desligada no momento do envio.
  # Isso também protege contra mensagens que ficaram na fila de um período em
  # que a automação ainda estava habilitada.
  eh_cobranca=(
   (fila.message_type or '').strip().lower() in ('lembrete_pagamento_semanal','cobranca','cobranca_semanal')
   or (fila.related_entity or '').strip().lower()=='cobranca'
  )
  billing_enabled=bool(cfg.get('billing_automation_enabled',cfg.get('automatic_billing_enabled',False)))
  if eh_cobranca and fila.related_entity_id:
   audit_pago=BillingAudit.query.filter_by(id=fila.related_entity_id,tenant_id=fila.tenant_id).first()
   if audit_pago and (audit_pago.payment_status or 'PENDENTE').upper()=='PAGO':
    fila.status='CANCELADA'; fila.updated_at=now
    db.session.add(MessageEvent(tenant_id=fila.tenant_id,message_id=fila.id,event='CANCELADA',description='Cobrança agendada cancelada porque o pagamento já foi baixado.',created_at=now))
    processed+=1; continue
  if eh_cobranca and not billing_enabled:
   fila.status='CANCELADA'; fila.updated_at=now
   db.session.add(MessageEvent(tenant_id=fila.tenant_id,message_id=fila.id,event='CANCELADA',description='Cobrança automática cancelada porque a automação de cobranças está desabilitada.',created_at=now))
   processed+=1; continue
  if (cfg.get('provider') or 'web').lower()!='business':
   fila.status='AGUARDANDO_MANUAL'; fila.updated_at=now
   db.session.add(MessageEvent(tenant_id=fila.tenant_id,message_id=fila.id,event='AGUARDANDO_MANUAL',description='WhatsApp Web não permite envio agendado automático.',created_at=now))
   processed+=1; continue
  try:
   try: params=json.loads(fila.template_parameters or '[]')
   except Exception: params=[]
   result=CommunicationService().send_whatsapp(phone=fila.recipient,message=fila.body,integration=integration,template_name=fila.template_name,template_language=cfg.get('template_language') or 'pt_BR',template_parameters=params)
   fila.provider=result.provider; fila.status=result.status; fila.external_id=result.external_id; fila.attempts=(fila.attempts or 0)+1; fila.sent_at=now; fila.updated_at=now
   db.session.add(MessageEvent(tenant_id=fila.tenant_id,message_id=fila.id,event=result.status,description='Mensagem agendada enviada automaticamente.',created_at=now))
  except CommunicationError as exc:
   fila.status='FALHA'; fila.error_message=str(exc); fila.attempts=(fila.attempts or 0)+1; fila.updated_at=now
   db.session.add(MessageEvent(tenant_id=fila.tenant_id,message_id=fila.id,event='FALHA',description=str(exc),created_at=now))
  processed+=1
 db.session.commit()
 return processed

def oil_status(v):
 if not v.controlar_oleo or v.ultima_troca_oleo_km is None or not v.intervalo_oleo_km:
  return {'state':'off','label':'Não configurado','remaining':None,'next_km':None}
 next_km=v.ultima_troca_oleo_km+v.intervalo_oleo_km
 remaining=next_km-(v.km_atual or 0)
 alert=v.alerta_oleo_km if v.alerta_oleo_km is not None else 100
 if remaining < 0:
  return {'state':'overdue','label':f'Vencida há {abs(remaining):,} km'.replace(',','.'),'remaining':remaining,'next_km':next_km}
 if remaining <= alert:
  return {'state':'warning','label':f'Faltam {remaining:,} km'.replace(',','.'),'remaining':remaining,'next_km':next_km}
 return {'state':'ok','label':f'Faltam {remaining:,} km'.replace(',','.'),'remaining':remaining,'next_km':next_km}

def recalcular_alertas(tenant_id):
 try:
  result=sync_operational_alerts(db.session,Alert,Maintenance,Vehicle,tenant_id,'/manutencoes','/veiculos')
  # Manutenções canceladas não podem voltar a gerar alertas operacionais.
  canceladas=[m.id for m in Maintenance.query.filter_by(tenant_id=tenant_id,status='Cancelada').all()]
  if canceladas:
   agora=agora_sao_paulo_naive()
   Alert.query.filter(
    Alert.tenant_id==tenant_id,Alert.entidade=='Manutenção',
    Alert.entidade_id.in_(canceladas),Alert.resolvido_em.is_(None)
   ).update({'resolvido_em':agora},synchronize_session=False)
  # Persiste os alertas recém-criados antes de tentar o envio imediato.
  db.session.commit()
  # Quando Alertas operacionais estiver ativado, tenta enviar imediatamente
  # os novos alertas ao motorista, sem aguardar o próximo ciclo do cron job.
  # O processador mantém a proteção contra duplicidade e ignora alertas cancelados/resolvidos.
  try:
   processar_alertas_automaticos(tenant_id)
  except Exception:
   db.session.rollback()
   app.logger.exception('Falha no envio imediato de alertas do tenant %s',tenant_id)
  return result
 except Exception:
  db.session.rollback()
  app.logger.exception('Falha ao recalcular alertas do tenant %s',tenant_id)
  return []

def migrate_schema():
 additions={
  'tenant':[('conferir_km_motorista','BOOLEAN DEFAULT FALSE'),('cobrar_km_excedente','BOOLEAN DEFAULT FALSE'),('razao_social','VARCHAR(180)'),('nome_fantasia','VARCHAR(150)'),('inscricao_estadual','VARCHAR(30)'),('inscricao_municipal','VARCHAR(30)'),('telefone','VARCHAR(30)'),('email','VARCHAR(150)'),('responsavel_legal','VARCHAR(150)'),('logradouro','VARCHAR(180)'),('numero_endereco','VARCHAR(30)'),('complemento','VARCHAR(100)'),('bairro','VARCHAR(100)'),('cidade','VARCHAR(100)'),('uf','VARCHAR(2)'),('cep','VARCHAR(10)'),('logo_key','VARCHAR(255)'),('favicon_key','VARCHAR(255)'),('cor_primaria','VARCHAR(7)'),('cor_secundaria','VARCHAR(7)'),('timezone_name',"VARCHAR(80) DEFAULT 'America/Sao_Paulo'")],
  'vehicle':[
   ('controlar_oleo','BOOLEAN DEFAULT FALSE'),('ultima_troca_oleo_km','INTEGER'),
   ('intervalo_oleo_km','INTEGER DEFAULT 10000'),('alerta_oleo_km','INTEGER DEFAULT 100'),
   ('current_driver_id','INTEGER'),('current_contract_id','INTEGER'),('status_changed_at','TIMESTAMP'),('status_reason','VARCHAR(255)'),('motorizacao','VARCHAR(20)'),
  ],
  'driver':[('telefone2','VARCHAR(30)'),('contato2_nome','VARCHAR(150)'),('contato2_parentesco','VARCHAR(40)'),('telefone3','VARCHAR(30)'),('contato3_nome','VARCHAR(150)'),('contato3_parentesco','VARCHAR(40)'),('logradouro','VARCHAR(160)'),('numero_endereco','VARCHAR(20)'),('complemento','VARCHAR(100)'),('bairro','VARCHAR(100)'),('cidade','VARCHAR(100)'),('uf','VARCHAR(2)'),('cep','VARCHAR(10)')],
  'contract_template':[
   ('descricao','VARCHAR(255)'),('versao','INTEGER DEFAULT 1'),('padrao','BOOLEAN DEFAULT FALSE'),
   ('nome_original','VARCHAR(255)'),('arquivo_original','VARCHAR(255)'),('hash_original','VARCHAR(64)'),('preparado_em','TIMESTAMP'),('gestora_nome','VARCHAR(180)'),('gestora_fantasia','VARCHAR(120)'),
   ('gestora_cnpj','VARCHAR(30)'),('gestora_endereco','VARCHAR(255)'),('parceira_nome','VARCHAR(180)'),
   ('parceira_cnpj','VARCHAR(30)'),('parceira_endereco','VARCHAR(255)'),
  ],
  'maintenance':[('alerta_km_antes','INTEGER DEFAULT 500'),('alerta_dias_antes','INTEGER DEFAULT 7'),('status',"VARCHAR(20) DEFAULT 'Ativa'"),('oficina','VARCHAR(160)'),('proxima_hora','VARCHAR(5)'),('notificar_motorista','BOOLEAN DEFAULT FALSE'),('lembrete_um_dia','BOOLEAN DEFAULT TRUE'),('notificacao_agendamento_id','INTEGER'),('notificacao_lembrete_id','INTEGER'),('concluida_em','TIMESTAMP'),('concluida_por_id','INTEGER')],
  'alert':[('source_key','VARCHAR(120)'),('entidade','VARCHAR(40)'),('entidade_id','INTEGER'),('action_url','VARCHAR(255)'),('atualizado_em','TIMESTAMP'),('resolvido_em','TIMESTAMP')],
  'inspection':[('tipo_vistoria',"VARCHAR(20) DEFAULT 'guiada'"),('painel_photo_key','VARCHAR(255)'),('painel_photo_mime','VARCHAR(80)'),('front_photo_key','VARCHAR(255)'),('front_photo_mime','VARCHAR(80)'),('right_photo_key','VARCHAR(255)'),('right_photo_mime','VARCHAR(80)'),('rear_photo_key','VARCHAR(255)'),('rear_photo_mime','VARCHAR(80)'),('left_photo_key','VARCHAR(255)'),('left_photo_mime','VARCHAR(80)'),('km_informada','INTEGER'),('damage_analysis_status',"VARCHAR(30) DEFAULT 'NAO_ANALISADA'"),('damage_analysis_level','VARCHAR(30)'),('damage_analysis_summary','TEXT'),('damage_analysis_at','TIMESTAMP')],
  'inspection_attempt':[('painel_photo_key','VARCHAR(255)'),('painel_photo_mime','VARCHAR(80)'),('km_informada','INTEGER')],
  'message_queue':[('template_parameters','TEXT')],
  'support_ticket':[('respondido_por_admin_id','INTEGER'),('respondido_por_nome','VARCHAR(120)')],
  'billing_audit':[
   ('payment_status',"VARCHAR(20) DEFAULT 'PENDENTE'"),('paid_at','TIMESTAMP'),('paid_by_id','INTEGER'),
   ('payment_method','VARCHAR(50)'),('payment_notes','TEXT'),('reminder_count','INTEGER DEFAULT 0'),
   ('last_reminder_at','TIMESTAMP'),('closed_at','TIMESTAMP'),('receipt_token','VARCHAR(64)'),('receipt_key','VARCHAR(255)'),('receipt_name','VARCHAR(255)'),('receipt_mime','VARCHAR(100)'),('receipt_uploaded_at','TIMESTAMP'),
  ],
  'contract':[
   ('template_nome','VARCHAR(120)'),('template_versao','INTEGER DEFAULT 1'),('hora_inicio','VARCHAR(5)'),
   ('periodicidade','VARCHAR(30)'),('dia_vencimento','VARCHAR(30)'),('multa_atraso_percentual','NUMERIC(6,2)'),
   ('juros_mes_percentual','NUMERIC(6,2)'),('indice_correcao','VARCHAR(30)'),('prazo_bloqueio_horas','INTEGER'),
   ('multa_diaria','NUMERIC(12,2)'),('taxa_adm_multas_percentual','NUMERIC(6,2)'),('nacionalidade','VARCHAR(60)'),
   ('estado_civil','VARCHAR(60)'),('profissao','VARCHAR(100)'),('cidade_assinatura','VARCHAR(100)'),
   ('numero_contrato','VARCHAR(30)'),('versao','INTEGER DEFAULT 1'),('criado_por_id','INTEGER'),
   ('criado_em','TIMESTAMP'),('atualizado_em','TIMESTAMP'),('assinado_em','TIMESTAMP'),
   ('assinatura_id','VARCHAR(120)'),('documento_id','INTEGER'),('arquivo_pdf','VARCHAR(255)'),('hash_documento','VARCHAR(64)'),('arquivo_pdf_assinado','VARCHAR(255)'),('hash_documento_assinado','VARCHAR(64)'),('documento_assinado_id','INTEGER'),('codigo_publico','VARCHAR(24)'),('enviado_whatsapp_em','TIMESTAMP'),('visualizado_em','TIMESTAMP'),('gerado_em','TIMESTAMP'),('clicksign_envelope_id','VARCHAR(80)'),('clicksign_document_id','VARCHAR(80)'),('clicksign_signer_id','VARCHAR(80)'),('clicksign_status','VARCHAR(30)'),('clicksign_sent_at','TIMESTAMP'),
  ],
 }
 inspector=inspect(db.engine)
 for table_name,fields in additions.items():
  if table_name not in inspector.get_table_names():
   continue
  columns={c['name'] for c in inspector.get_columns(table_name)}
  for name,definition in fields:
   if name not in columns:
    with db.engine.begin() as conn:
     conn.execute(text(f'ALTER TABLE {table_name} ADD COLUMN {name} {definition}'))

LOCADRIVERS_TEMPLATE_VERSION=2

def moeda_br(value):
 try:
  n=Decimal(str(value or 0))
 except Exception:
  n=Decimal('0')
 return f'{n:,.2f}'.replace(',','X').replace('.',',').replace('X','.')

def data_br(value):
 try:
  return datetime.strptime(value,'%Y-%m-%d').strftime('%d/%m/%Y')
 except Exception:
  return value or ''

def locadrivers_template_completo():
 return """CONTRATO PARTICULAR DE LOCAÇÃO DE VEÍCULO PARA TRANSPORTE PRIVADO POR APLICATIVO

IDENTIFICAÇÃO DAS PARTES

GESTORA DA LOCAÇÃO: {{gestora_nome}}, nome fantasia {{gestora_fantasia}}, inscrita no CNPJ sob nº {{gestora_cnpj}}, com endereço em {{gestora_endereco}}, doravante denominada GESTORA ou LOCADORA.

PROPRIETÁRIO DO VEÍCULO: {{proprietario_nome}}, CPF/CNPJ nº {{proprietario_documento}}, legítimo proprietário do veículo objeto deste contrato, doravante denominado PROPRIETÁRIO.

PARCEIRA OPERACIONAL: {{parceira_nome}}, inscrita no CNPJ sob nº {{parceira_cnpj}}, com endereço em {{parceira_endereco}}, doravante denominada PARCEIRA OPERACIONAL.

LOCATÁRIO: {{motorista_nome}}, {{nacionalidade}}, {{estado_civil}}, {{profissao}}, RG nº {{motorista_rg}}, CPF nº {{motorista_cpf}}, CNH nº {{motorista_cnh}}, residente e domiciliado em {{motorista_endereco}}, doravante denominado LOCATÁRIO.

As partes acima identificadas têm, entre si, justo e contratado o presente CONTRATO DE LOCAÇÃO DE VEÍCULO, mediante as cláusulas e condições seguintes.

CLÁUSULA 1ª — DO OBJETO

1.1. O presente contrato tem como objeto a locação do seguinte veículo:
Modelo: {{veiculo_modelo}}
Cor: {{veiculo_cor}}
Ano/Modelo: {{veiculo_ano_fabricacao}}/{{veiculo_ano_modelo}}
Placa: {{veiculo_placa}}
Renavam: {{veiculo_renavam}}
Quilometragem inicial: {{km_inicial}} km

CLÁUSULA 2ª — DO VALOR, PAGAMENTO, CAUÇÃO E INADIMPLÊNCIA

2.1. O LOCATÁRIO pagará o valor {{periodicidade_minuscula}} de R$ {{valor_locacao}} ({{valor_locacao_extenso}}), com vencimento toda {{dia_vencimento}}, via PIX ou boleto bancário.

2.2. No ato da assinatura, o LOCATÁRIO pagará caução no valor de R$ {{caucao}} ({{caucao_extenso}}), que poderá ser utilizada para abatimento de débitos contratuais.

2.3. A caução não cobre multas de trânsito, avarias ou danos ao veículo e será devolvida em até 60 (sessenta) dias úteis após o encerramento do contrato, desde que não existam pendências financeiras ou contratuais.

2.4. O atraso no pagamento acarretará multa de {{multa_atraso_percentual}}% sobre o valor em atraso, juros de {{juros_mes_percentual}}% ao mês sobre o saldo devedor e correção monetária pelo índice {{indice_correcao}}.

2.5. O atraso superior a {{prazo_bloqueio_horas}} horas poderá ocasionar bloqueio do veículo, rescisão contratual, retirada imediata do veículo e cobrança de multa diária de R$ {{multa_diaria}}, limitada ao valor FIPE do veículo.

2.6. O LOCATÁRIO autoriza eventual negativação junto aos órgãos de proteção ao crédito em caso de inadimplência, observadas as exigências legais aplicáveis.

CLÁUSULA 3ª — DO PRAZO

3.1. O presente contrato terá vigência de {{prazo_dias}} dias, com início em {{data_inicio_formatada}} às {{hora_inicio}}, e término previsto para {{data_fim_formatada}}, podendo ser prorrogado por acordo entre as partes até o limite máximo de 1 (um) ano.

CLÁUSULA 4ª — DA RESCISÃO

4.1. A rescisão antecipada pelo LOCATÁRIO antes do terceiro mês implicará multa equivalente à caução, no valor de R$ {{caucao}}.

4.2. Após o terceiro mês, o LOCATÁRIO poderá rescindir o contrato mediante aviso prévio de 7 (sete) dias.

4.3. A LOCADORA poderá rescindir o contrato a qualquer momento, devendo o veículo ser devolvido em até 1 (um) dia corrido após a comunicação.

CLÁUSULA 5ª — DAS CONDIÇÕES E DA VISTORIA DO VEÍCULO

5.1. O LOCATÁRIO declara receber o veículo em perfeitas condições de uso, conservação e funcionamento, conforme vistoria anexa, que integra este contrato.

5.2. O LOCATÁRIO obriga-se a enviar toda segunda-feira fotos externas do veículo, foto do painel com a quilometragem e foto da etiqueta de troca de óleo.

CLÁUSULA 6ª — DO USO DO VEÍCULO

6.1. O veículo deverá ser utilizado exclusivamente pelo LOCATÁRIO para transporte privado por aplicativo.

6.2. É expressamente proibido sublocar, ceder, emprestar ou permitir que terceiro conduza o veículo sem autorização; utilizar o veículo fora do Estado de São Paulo; retirar adesivos obrigatórios; ou utilizá-lo para fins ilícitos.

6.3. A quilometragem semanal fica limitada a {{limite_km}} km. A quilometragem excedente será cobrada no valor de R$ {{valor_km_excedente}} por quilômetro.

CLÁUSULA 7ª — DAS MULTAS DE TRÂNSITO

7.1. O LOCATÁRIO será integralmente responsável pelas multas e infrações de trânsito ocorridas durante a vigência da locação, inclusive pela identificação do condutor e pontuação na CNH.

7.2. As multas deverão ser reembolsadas à LOCADORA acrescidas de {{taxa_adm_multas_percentual}}% de taxa administrativa.

CLÁUSULA 8ª — DAS AVARIAS, SINISTROS E DANOS A TERCEIROS

8.1. O LOCATÁRIO responderá integralmente por danos e colisões, arranhões e mau uso, despesas de guincho e prejuízos causados a terceiros durante a vigência deste contrato.

8.2. Em caso de acidente, roubo ou furto, o LOCATÁRIO deverá comunicar imediatamente a LOCADORA e apresentar boletim de ocorrência em até 2 (dois) dias.

CLÁUSULA 9ª — DO SEGURO

9.1. O veículo possui seguro. Em caso de sinistro, o LOCATÁRIO será responsável pelo pagamento da franquia no valor de R$ {{franquia}} ({{franquia_extenso}}), sem prejuízo de outros valores não cobertos pela apólice quando decorrentes de sua responsabilidade.

CLÁUSULA 10ª — DA DEVOLUÇÃO

10.1. O veículo deverá ser devolvido nas mesmas condições em que foi entregue, ressalvado o desgaste natural decorrente do uso regular.

10.2. Havendo danos, o LOCATÁRIO deverá efetuar o pagamento dos reparos em até 3 (três) dias úteis após a apresentação dos orçamentos.

10.3. A não devolução poderá ensejar as medidas judiciais e criminais cabíveis.

CLÁUSULA 11ª — DAS DISPOSIÇÕES GERAIS

11.1. O PROPRIETÁRIO não possui responsabilidade pela gestão da locação, cobranças ou relação contratual operacional com o LOCATÁRIO.

11.2. Eventuais tolerâncias de qualquer das partes não constituem novação, renúncia ou alteração das condições deste contrato.

11.3. Este contrato não poderá ser cedido ou transferido sem autorização expressa e escrita da LOCADORA.

CLÁUSULA 12ª — DO FORO

12.1. Fica eleito o foro da Comarca de São Paulo/SP para dirimir quaisquer controvérsias oriundas deste contrato, com renúncia a qualquer outro, por mais privilegiado que seja.

E, por estarem de acordo, as partes assinam o presente instrumento.

{{cidade_assinatura}}, {{data_assinatura_formatada}}.


________________________________________
{{gestora_nome}}
GESTORA / LOCADORA


________________________________________
{{proprietario_nome}}
PROPRIETÁRIO


________________________________________
{{motorista_nome}}
LOCATÁRIO


________________________________________
{{parceira_nome}}
PARCEIRA OPERACIONAL


________________________________________        ________________________________________
TESTEMUNHA 1 — Nome/CPF                           TESTEMUNHA 2 — Nome/CPF
"""

def ensure_locadrivers_template(tenant_id, force_new_version=False):
 existing=ContractTemplate.query.filter_by(tenant_id=tenant_id,nome='Locadrivers Completo',versao=LOCADRIVERS_TEMPLATE_VERSION).first()
 if existing and not force_new_version:
  return existing
 ContractTemplate.query.filter_by(tenant_id=tenant_id,padrao=True).update({'padrao':False},synchronize_session=False)
 model=ContractTemplate(
  tenant_id=tenant_id,nome='Locadrivers Completo',descricao='Minuta integral Locadrivers com 12 cláusulas e preenchimento automático.',
  versao=LOCADRIVERS_TEMPLATE_VERSION,padrao=True,tipo_veiculo='Todos',possui_limite_km=True,
  conteudo=locadrivers_template_completo(),nome_original='minuta-locadrivers-v2',
  gestora_nome='LCADRIVER CORRETORA DE ALUGUEL DE VEÍCULOS LTDA',gestora_fantasia='LOCADRIVERS',
  gestora_cnpj='64.406.745/0001-38',gestora_endereco='Av. Deputado Emilio Carlos, nº 656, Vila Caldas, Carapicuíba/SP — CEP 06310-160',
  parceira_nome='L.C.DVS CORRETORA DE ALUGUEL DE VEÍCULOS LTDA',parceira_cnpj='48.758.670/0001-06',
  parceira_endereco='Alameda Araguaia, nº 2104, Alphaville Industrial, Barueri/SP — CEP 06455-000',ativo=True,
 )
 db.session.add(model)
 return model

def seed():
 db.create_all()
 migrate_schema()
 # Garante tabela de histórico após migrações e numera contratos antigos.
 db.create_all()
 for contrato in Contract.query.filter(Contract.numero_contrato.is_(None)).order_by(Contract.id).all():
  contrato.numero_contrato=gerar_numero_contrato(contrato.id,contrato.criado_em or datetime.utcnow())
  if contrato.status in (None,'Ativo'):
   contrato.status='Gerado'
  contrato.versao=contrato.versao or 1
 if not Tenant.query.first():
  t=Tenant(nome='Locadora Demonstração'); db.session.add(t); db.session.flush()
  u=User(tenant_id=t.id,nome='Administrador',email='admin@frotafacil.local',senha=generate_password_hash('admin123')); db.session.add(u)
  base='''CONTRATO DE LOCAÇÃO\nLOCATÁRIO: {{motorista_nome}}, CPF {{motorista_cpf}}, CNH {{motorista_cnh}}.\nVEÍCULO: {{veiculo_modelo}}, placa {{veiculo_placa}}, Renavam {{veiculo_renavam}}.\nVALOR: R$ {{valor_locacao}}. CAUÇÃO: R$ {{caucao}}.\nINÍCIO: {{data_inicio}}. TÉRMINO: {{data_fim}}.\nLIMITE DE KM: {{limite_km}}. EXCEDENTE: R$ {{valor_km_excedente}}/km.'''
  db.session.add_all([ContractTemplate(tenant_id=t.id,nome='Combustão com limite',tipo_veiculo='Combustão',possui_limite_km=True,conteudo=base),ContractTemplate(tenant_id=t.id,nome='Elétrico com limite',tipo_veiculo='Elétrico',possui_limite_km=True,conteudo=base+'\nO LOCATÁRIO se responsabiliza pela recarga e uso de equipamentos homologados.')]); db.session.flush()
 # Instala automaticamente a minuta completa v2 para todos os tenants.
 for tenant in Tenant.query.all():
  ensure_locadrivers_template(tenant.id)
 # Conta administrativa Frota Fácil independente das locadoras.
 # É criada somente quando as variáveis de ambiente são informadas e ainda não existe admin.
 admin_email=(os.getenv('FROTA_FACIL_ADMIN_EMAIL') or '').strip().lower()
 admin_password=os.getenv('FROTA_FACIL_ADMIN_PASSWORD') or ''
 admin_name=(os.getenv('FROTA_FACIL_ADMIN_NAME') or 'Administrador Frota Fácil').strip()
 if admin_email and admin_password and not FrotaAdmin.query.filter_by(email=admin_email).first():
  db.session.add(FrotaAdmin(nome=admin_name,email=admin_email,senha=generate_password_hash(admin_password),ativo=True,perfil='SUPORTE'))
 db.session.commit()


@app.route('/sobre-a-empresa')
def sobre_a_empresa():
    empresa = {
        'razao_social': (os.getenv('FROTA_FACIL_RAZAO_SOCIAL') or 'Gutos Car').strip(),
        'nome_fantasia': (os.getenv('FROTA_FACIL_NOME_FANTASIA') or 'Gutos Car').strip(),
        'cnpj': (os.getenv('FROTA_FACIL_CNPJ') or '').strip(),
        'endereco': (os.getenv('FROTA_FACIL_ENDERECO') or '').strip(),
        'telefone': (os.getenv('FROTA_FACIL_TELEFONE') or '').strip(),
        'email': (os.getenv('FROTA_FACIL_EMAIL') or '').strip(),
        'descricao': (os.getenv('FROTA_FACIL_DESCRICAO') or 'Empresa responsável pela operação piloto e uso da plataforma Frota Fácil para gestão de locação de veículos.').strip(),
    }
    return render_template('sobre_empresa.html', empresa=empresa)

@app.route('/criar-conta',methods=['GET','POST'])
def criar_conta():
 if current_user.is_authenticated: return redirect(url_for('dashboard'))
 if request.method=='POST':
  nome=request.form.get('nome','').strip(); empresa=request.form.get('empresa','').strip(); email=request.form.get('email','').strip().lower(); senha=request.form.get('senha','')
  if not nome or not empresa or not email or len(senha)<6:
   flash('Preencha todos os campos. A senha deve ter pelo menos 6 caracteres.','danger')
  elif User.query.filter_by(email=email).first():
   flash('Este e-mail já está cadastrado.','danger')
  else:
   t=Tenant(nome=empresa,ativo=True); db.session.add(t); db.session.flush()
   u=User(tenant_id=t.id,nome=nome,email=email,senha=generate_password_hash(senha),perfil='admin'); db.session.add(u)
   base='''CONTRATO DE LOCAÇÃO\nLOCATÁRIO: {{motorista_nome}}, CPF {{motorista_cpf}}, CNH {{motorista_cnh}}.\nVEÍCULO: {{veiculo_modelo}}, placa {{veiculo_placa}}, Renavam {{veiculo_renavam}}.\nVALOR: R$ {{valor_locacao}}. CAUÇÃO: R$ {{caucao}}.\nINÍCIO: {{data_inicio}}. TÉRMINO: {{data_fim}}.'''
   db.session.add(ContractTemplate(tenant_id=t.id,nome='Modelo básico',tipo_veiculo='Todos',possui_limite_km=False,conteudo=base,versao=1,padrao=False))
   ensure_locadrivers_template(t.id)
   db.session.commit(); login_user(u); flash('Conta criada. Sua base está limpa e pronta para os cadastros.','success'); return redirect(url_for('dashboard'))
 return render_template('criar_conta.html')

@app.route('/entrar',methods=['GET','POST'])
def entrar():
 if request.method=='POST':
  u=User.query.filter_by(email=request.form['email']).first()
  if u and check_password_hash(u.senha,request.form['senha']): login_user(u); return redirect(url_for('dashboard'))
  flash('E-mail ou senha inválidos.','danger')
 return render_template('login.html')
@app.route('/sair')
@login_required
def sair(): logout_user(); return redirect(url_for('entrar'))
@app.route('/')
@login_required
def dashboard():
 recalcular_alertas(tid())
 vehicles=Vehicle.query.filter_by(tenant_id=tid()).all()
 system_alerts=Alert.query.filter(Alert.tenant_id==tid(),Alert.resolvido_em.is_(None)).order_by(Alert.nivel.desc(),Alert.criado_em.desc()).limit(8).all()
 status_counts={status:0 for status in ['Disponível','Reservado','Alugado','Devolução','Manutenção','Inativo']}
 for vehicle in vehicles:
  status_counts[vehicle.status or 'Disponível']=status_counts.get(vehicle.status or 'Disponível',0)+1
 cards={
  'veiculos':len(vehicles),'motoristas':Driver.query.filter_by(tenant_id=tid()).count(),
  'contratos':Contract.query.filter(Contract.tenant_id==tid(),Contract.status.in_(['Rascunho','Gerado','Enviado','Visualizado','Assinado','Ativo'])).count(),
  'alertas':Alert.query.filter(Alert.tenant_id==tid(),Alert.resolvido_em.is_(None)).count(),
  'disponiveis':status_counts.get('Disponível',0),'reservados':status_counts.get('Reservado',0),
  'alugados':status_counts.get('Alugado',0),'manutencao':status_counts.get('Manutenção',0),
  'cobrancas_hoje':sum(1 for c in Contract.query.filter(Contract.tenant_id==tid(),Contract.status.in_(['Assinado','Ativo'])).all() if cobranca_vence_hoje(c)),
 }
 return render_template('dashboard.html',cards=cards,veiculos=sorted(vehicles,key=lambda v:v.id,reverse=True)[:6],alertas=system_alerts,oil_status=oil_status)

@app.route('/motoristas',methods=['GET','POST'])
@login_required
def motoristas():
 if request.method=='POST':
  campos=['nome','cpf','rg','numero_cnh','categoria','data_nascimento','validade_cnh','telefone','telefone2','contato2_nome','contato2_parentesco','telefone3','contato3_nome','contato3_parentesco','email','endereco','logradouro','numero_endereco','complemento','bairro','cidade','uf','cep','status']
  vals={k:request.form.get(k) for k in campos}
  telefones,telefone_invalido=normalizar_telefones_form({
   'telefone':vals.get('telefone'),
   'telefone2':vals.get('telefone2'),
   'telefone3':vals.get('telefone3'),
  })
  if telefone_invalido:
   flash(f'Telefone inválido no campo {telefone_invalido}. Informe DDD + número; o Frota Fácil acrescenta o código do Brasil automaticamente.','danger')
   return redirect(url_for('motoristas'))
  vals.update(telefones)
  d=Driver(tenant_id=tid(),**vals)
  db.session.add(d)
  try:
   db.session.flush()
   armazenar_documento_cadastro(
    tipo='CNH',
    entidade='Motorista',
    entidade_id=d.id,
    referencia=d.nome,
    numero_documento=d.numero_cnh,
    nome_original=request.form.get('_documento_nome'),
    temp_key=request.form.get('_documento_temp_key'),
    mimetype=request.form.get('_documento_mimetype'),
    ano=(d.validade_cnh or '')[-4:] if d.validade_cnh else None,
   )
   db.session.commit()
  except Exception:
   db.session.rollback()
   app.logger.exception('Falha ao cadastrar motorista e armazenar CNH')
   flash('Não foi possível concluir o cadastro e armazenar a CNH. Tente novamente.','danger')
   return redirect(url_for('motoristas'))
  # O Portal do Motorista pertence ao motorista, mas o convite automático só é
  # disparado após uma assinatura real de contrato. O envio manual continua disponível.
  flash('Motorista cadastrado e CNH armazenada automaticamente.','success')
  return redirect(url_for('motoristas'))
 q=(request.args.get('q') or '').strip()
 query=Driver.query.filter_by(tenant_id=tid())
 if q:
  like=f'%{q}%'
  query=query.filter(or_(Driver.nome.ilike(like),Driver.cpf.ilike(like),Driver.rg.ilike(like),Driver.numero_cnh.ilike(like),Driver.telefone.ilike(like)))
 return render_template('motoristas.html',items=query.order_by(Driver.nome).all(),q=q)

@app.route('/motoristas/<int:id>/editar',methods=['GET','POST'])
@login_required
def editar_motorista(id):
 d=Driver.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 if request.method=='POST':
  campos=['nome','cpf','rg','numero_cnh','categoria','data_nascimento','validade_cnh','telefone','telefone2','contato2_nome','contato2_parentesco','telefone3','contato3_nome','contato3_parentesco','email','endereco','logradouro','numero_endereco','complemento','bairro','cidade','uf','cep','status']
  vals={campo:request.form.get(campo) for campo in campos}
  telefones,telefone_invalido=normalizar_telefones_form({
   'telefone':vals.get('telefone'),
   'telefone2':vals.get('telefone2'),
   'telefone3':vals.get('telefone3'),
  })
  if telefone_invalido:
   flash(f'Telefone inválido no campo {telefone_invalido}. Informe DDD + número; o Frota Fácil acrescenta o código do Brasil automaticamente.','danger')
   return redirect(url_for('editar_motorista',id=d.id))
  vals.update(telefones)
  for campo,valor in vals.items():
   setattr(d,campo,valor)
  db.session.commit(); flash('Motorista atualizado com sucesso.','success'); return redirect(url_for('motoristas'))
 return render_template('editar_motorista.html',d=d)

@app.route('/motoristas/importar',methods=['POST'])
@login_required
def importar_motorista():
 f=request.files.get('arquivo')
 if not f:
  flash('Selecione um arquivo.','danger')
  return redirect(url_for('motoristas'))

 nome_original=secure_filename(f.filename or 'cnh')
 mimetype=f.mimetype or 'application/octet-stream'
 conteudo=f.read()
 temp_key=f'{tid()}/temporarios/{uuid.uuid4().hex}_{nome_original}'
 try:
  storage.upload(BytesIO(conteudo),temp_key,mimetype)
  # Libera a conexão durante o OCR.
  db.session.remove()
  arquivo_ocr=FileStorage(stream=BytesIO(conteudo),filename=nome_original,content_type=mimetype)
  texto=extract_text(arquivo_ocr, document_type='cnh')
  dados=parse_cnh(texto)
  return render_template('confirmar_motorista.html',dados=dados,documento_temp_key=temp_key,documento_nome=nome_original,documento_mimetype=mimetype)
 except Exception as exc:
  try: storage.delete(temp_key)
  except Exception: pass
  app.logger.exception('Falha ao processar CNH')
  flash(f'Não foi possível processar a CNH: {exc}','danger')
  return redirect(url_for('motoristas'))
 finally:
  db.session.remove()
@app.route('/motoristas/excluir/<int:id>',methods=['POST'])
@login_required
def excluir_motorista(id):
 x=Driver.query.filter_by(id=id,tenant_id=tid()).first_or_404(); db.session.delete(x); db.session.commit(); return redirect(url_for('motoristas'))

def regra_proprietario_veiculo(vehicle_id, referencia=None):
 referencia=referencia or date.today()
 q=InvestorVehicleRule.query.filter_by(tenant_id=tid(),vehicle_id=vehicle_id).filter(InvestorVehicleRule.vigencia_inicio<=referencia).filter(or_(InvestorVehicleRule.vigencia_fim.is_(None),InvestorVehicleRule.vigencia_fim>=referencia)).order_by(InvestorVehicleRule.vigencia_inicio.desc(),InvestorVehicleRule.id.desc())
 return q.first()

def resumo_financeiro_proprietario(investor_id):
 veiculos=Vehicle.query.filter_by(tenant_id=tid(),investor_id=investor_id).order_by(Vehicle.placa).all()
 ids={v.id for v in veiculos}
 contratos=Contract.query.filter(Contract.tenant_id==tid(),Contract.vehicle_id.in_(ids)).all() if ids else []
 contrato_por_id={c.id:c for c in contratos}
 auditorias=BillingAudit.query.filter(BillingAudit.tenant_id==tid(),BillingAudit.contract_id.in_(list(contrato_por_id))).order_by(BillingAudit.billing_date.desc()).all() if contratos else []
 receita_paga=Decimal('0'); receita_aberta=Decimal('0'); repasse_proprietario=Decimal('0'); receita_locadora=Decimal('0')
 por_veiculo={v.id:{'receita_paga':Decimal('0'),'receita_aberta':Decimal('0'),'repasse':Decimal('0'),'locadora':Decimal('0')} for v in veiculos}
 for a in auditorias:
  c=contrato_por_id.get(a.contract_id)
  if not c or c.vehicle_id not in ids: continue
  total=Decimal(str(a.total_amount or 0))
  regra=regra_proprietario_veiculo(c.vehicle_id,a.billing_date)
  pct=Decimal(str(regra.percentual_proprietario or 0)) if regra else Decimal('0')
  pct_l=Decimal(str(regra.percentual_locadora or 0)) if regra else Decimal('0')
  if (a.payment_status or '').upper()=='PAGO':
   receita_paga+=total; por_veiculo[c.vehicle_id]['receita_paga']+=total
   rep=total*pct/Decimal('100'); loc=total*pct_l/Decimal('100')
   repasse_proprietario+=rep; receita_locadora+=loc
   por_veiculo[c.vehicle_id]['repasse']+=rep; por_veiculo[c.vehicle_id]['locadora']+=loc
  else:
   receita_aberta+=total; por_veiculo[c.vehicle_id]['receita_aberta']+=total
 custos=Decimal('0')
 for m in Maintenance.query.filter(Maintenance.tenant_id==tid(),Maintenance.vehicle_id.in_(ids)).all() if ids else []:
  if m.custo:
   valor=Decimal(str(m.custo)); custos+=valor
   por_veiculo.setdefault(m.vehicle_id,{}).setdefault('custos',Decimal('0')); por_veiculo[m.vehicle_id]['custos']+=valor
 for v in veiculos: por_veiculo[v.id].setdefault('custos',Decimal('0'))
 return {'veiculos':veiculos,'receita_paga':receita_paga,'receita_aberta':receita_aberta,'repasse_proprietario':repasse_proprietario,'receita_locadora':receita_locadora,'custos':custos,'resultado_proprietario':repasse_proprietario-custos,'por_veiculo':por_veiculo}


def owner_portal_required(view):
 @wraps(view)
 def wrapped(*args,**kwargs):
  access_id=session.get('owner_access_id')
  investor_id=session.get('owner_investor_id')
  tenant_id=session.get('owner_tenant_id')
  if not access_id or not investor_id or not tenant_id:
   return redirect(url_for('portal_proprietario_entrar'))
  access=InvestorAccess.query.filter_by(id=access_id,investor_id=investor_id,tenant_id=tenant_id,ativo=True).first()
  if not access:
   session.pop('owner_access_id',None); session.pop('owner_investor_id',None); session.pop('owner_tenant_id',None)
   return redirect(url_for('portal_proprietario_entrar'))
  return view(*args,**kwargs)
 return wrapped

def regra_proprietario_veiculo_portal(tenant_id, investor_id, vehicle_id, data_ref):
 return InvestorVehicleRule.query.filter(
  InvestorVehicleRule.tenant_id==tenant_id,
  InvestorVehicleRule.investor_id==investor_id,
  InvestorVehicleRule.vehicle_id==vehicle_id,
  InvestorVehicleRule.vigencia_inicio<=data_ref,
  or_(InvestorVehicleRule.vigencia_fim.is_(None),InvestorVehicleRule.vigencia_fim>=data_ref)
 ).order_by(InvestorVehicleRule.vigencia_inicio.desc(),InvestorVehicleRule.id.desc()).first()

def resumo_portal_proprietario(tenant_id, investor_id, inicio, fim, vehicle_id=None, placa=None):
 """Resumo do portal com visão teórica (contratual) e realizado pago separados.

 O previsto considera o valor-base da locação e a regra de repasse vigente na
 data de cada competência. Excesso de KM não entra na previsão, pois só existe
 depois da leitura real. Valores pagos continuam disponíveis para comparação.
 """
 investor=Investor.query.filter_by(id=investor_id,tenant_id=tenant_id).first_or_404()
 vehicles_query=Vehicle.query.filter_by(tenant_id=tenant_id,investor_id=investor_id)
 if vehicle_id:
  vehicles_query=vehicles_query.filter(Vehicle.id==vehicle_id)
 if placa:
  vehicles_query=vehicles_query.filter(Vehicle.placa.ilike(f'%{placa}%'))
 vehicles=vehicles_query.order_by(Vehicle.placa).all()
 ids=[v.id for v in vehicles]
 rows={v.id:{'vehicle':v,'receita':Decimal('0'),'repasse':Decimal('0'),'pago':Decimal('0'),'repasse_pago':Decimal('0'),'custos':Decimal('0'),'resultado':Decimal('0'),'resultado_real':Decimal('0'),'aberto':Decimal('0'),'investment':None,'roi_periodo':None,'roi_aquisicao_periodo':None,'rentabilidade_mensal':None,'payback_meses':None,'capital_recuperado_pct':None,'desempenho':'Sem dados de investimento','ranking':None} for v in vehicles}
 contracts=Contract.query.filter(Contract.tenant_id==tenant_id,Contract.vehicle_id.in_(ids or [-1])).all()
 contract_map={c.id:c for c in contracts}
 monthly={}

 def month_bucket(d):
  key=d.strftime('%Y-%m')
  return monthly.setdefault(key,{'repasse':Decimal('0'),'pago':Decimal('0'),'custos':Decimal('0')})

 def parse_contract_date(value, fallback):
  try: return datetime.strptime(value or '','%Y-%m-%d').date()
  except Exception: return fallback

 def periodicidade_dias(c):
  txt=unicodedata.normalize('NFKD',str(c.periodicidade or 'Semanal')).encode('ascii','ignore').decode('ascii').lower()
  if 'diar' in txt: return 1
  if 'mens' in txt: return None
  return 7

 # 1) PREVISÃO TEÓRICA: nasce do contrato, não do status de pagamento.
 # Contratos cancelados/rascunhos não geram expectativa financeira.
 status_validos={'Gerado','Enviado','Visualizado','Assinado','Ativo','Encerrado','Finalizado'}
 for c in contracts:
  if c.vehicle_id not in rows or (c.status and c.status not in status_validos):
   continue
  c_ini=parse_contract_date(c.data_inicio,inicio)
  c_fim=parse_contract_date(c.data_fim,fim)
  periodo_ini=max(inicio,c_ini); periodo_fim=min(fim,c_fim)
  if periodo_ini>periodo_fim: continue
  base=Decimal(str(c.valor_locacao or 0))
  if base<=0: continue
  intervalo=periodicidade_dias(c)
  competencias=[]
  if intervalo is None:  # mensal: uma competência por mês enquanto o contrato estiver vigente
   cursor=date(periodo_ini.year,periodo_ini.month,1)
   limite=date(periodo_fim.year,periodo_fim.month,1)
   while cursor<=limite:
    ref=max(cursor,periodo_ini)
    competencias.append(ref)
    cursor=date(cursor.year+1,1,1) if cursor.month==12 else date(cursor.year,cursor.month+1,1)
  else:
   # ancora a recorrência na data inicial do contrato e avança até entrar no filtro
   ref=c_ini
   if ref<periodo_ini:
    saltos=max(0,(periodo_ini-ref).days//intervalo)
    ref=ref+timedelta(days=saltos*intervalo)
    while ref<periodo_ini: ref+=timedelta(days=intervalo)
   while ref<=periodo_fim:
    competencias.append(ref); ref+=timedelta(days=intervalo)
  for ref in competencias:
   rule=regra_proprietario_veiculo_portal(tenant_id,investor_id,c.vehicle_id,ref)
   pct=Decimal(str(rule.percentual_proprietario or 0)) if rule else Decimal('100')
   share=base*pct/Decimal('100')
   rows[c.vehicle_id]['receita']+=base
   rows[c.vehicle_id]['repasse']+=share
   month_bucket(ref)['repasse']+=share

 # 2) REALIZADO: pagamentos efetivamente baixados.
 # O filtro por veículo precisa ser aplicado pela relação estrutural da cobrança:
 # BillingAudit -> contract_id -> Contract.vehicle_id. A placa gravada na auditoria
 # é apenas uma fotografia histórica e não deve decidir a qual veículo o pagamento
 # pertence, pois registros antigos podem conter placa ausente/incorreta.
 # Como `contracts` acima já contém SOMENTE contratos dos veículos filtrados,
 # limitar a consulta aos IDs desses contratos garante que o KPI realizado respeite
 # exatamente o filtro de veículo/placa do portal.
 contract_ids=list(contract_map.keys())
 audits_raw=BillingAudit.query.filter(
  BillingAudit.tenant_id==tenant_id,
  BillingAudit.contract_id.in_(contract_ids or [-1]),
  BillingAudit.billing_date>=inicio,
  BillingAudit.billing_date<=fim
 ).order_by(BillingAudit.billing_date,BillingAudit.id).all()

 # Compatibilidade com reenvios antigos que podiam gerar auditorias duplicadas no
 # mesmo dia para o mesmo contrato. Mantém uma única ocorrência por competência,
 # priorizando um registro PAGO e, em empate, o mais recente.
 audits_por_cobranca={}
 for audit in audits_raw:
  chave=(audit.contract_id,audit.billing_date)
  atual=audits_por_cobranca.get(chave)
  audit_pago=(audit.payment_status or '').upper()=='PAGO'
  atual_pago=(atual.payment_status or '').upper()=='PAGO' if atual else False
  if atual is None or (audit_pago and not atual_pago) or (audit_pago==atual_pago and (audit.id or 0)>(atual.id or 0)):
   audits_por_cobranca[chave]=audit

 for audit in audits_por_cobranca.values():
  contract=contract_map.get(audit.contract_id)
  if not contract or contract.vehicle_id not in rows:
   continue
  vehicle=rows[contract.vehicle_id]['vehicle']
  total=Decimal(str(audit.total_amount or 0))
  row=rows[vehicle.id]
  pagamento_status=(audit.payment_status or '').upper()
  contrato_status=unicodedata.normalize('NFKD',str(contract.status or '')).encode('ascii','ignore').decode('ascii').strip().lower()

  # Contrato cancelado deixa de produzir saldo em aberto no Portal do Proprietário.
  # O histórico de valores realmente pagos é preservado, pois representa movimentação
  # financeira efetivamente ocorrida antes do cancelamento.
  contrato_cancelado=contrato_status in {'cancelado','cancelada','cancelled','canceled'}

  if pagamento_status=='PAGO':
   rule=regra_proprietario_veiculo_portal(tenant_id,investor_id,vehicle.id,audit.billing_date)
   pct=Decimal(str(rule.percentual_proprietario or 0)) if rule else Decimal('100')
   share=total*pct/Decimal('100')
   row['pago']+=total; row['repasse_pago']+=share; month_bucket(audit.billing_date)['pago']+=share
  elif not contrato_cancelado:
   row['aberto']+=total

 # 3) Custos reais registrados no período.
 for m in Maintenance.query.filter(Maintenance.tenant_id==tenant_id,Maintenance.vehicle_id.in_(ids or [-1])).all():
  try: d=datetime.strptime(m.data,'%Y-%m-%d').date()
  except Exception: continue
  if d<inicio or d>fim or m.vehicle_id not in rows: continue
  cost=Decimal(str(m.custo or 0))
  rows[m.vehicle_id]['custos']+=cost; month_bucket(d)['custos']+=cost

 investments={x.vehicle_id:x for x in VehicleInvestment.query.filter(
  VehicleInvestment.tenant_id==tenant_id,
  VehicleInvestment.investor_id==investor_id,
  VehicleInvestment.vehicle_id.in_(ids or [-1])
 ).all()}
 periodo_meses=max(Decimal('1'),Decimal(str(max(1,(fim-inicio).days+1)))/Decimal('30.4375'))

 total_receita=total_repasse=total_pago=total_repasse_pago=total_custos=total_aberto=Decimal('0')
 total_capital_proprio=Decimal('0')
 vehicle_rows=[]
 for v in vehicles:
  row=rows[v.id]
  row['resultado']=row['repasse']-row['custos']
  row['resultado_real']=row['repasse_pago']-row['custos']
  inv=investments.get(v.id); row['investment']=inv
  if inv:
   capital=Decimal(str(inv.capital_proprio or 0))
   aquisicao=Decimal(str(inv.valor_aquisicao or 0))
   if capital>0:
    total_capital_proprio+=capital
   if capital>0:
    row['roi_periodo']=(row['resultado_real']/capital)*Decimal('100')
    row['rentabilidade_mensal']=row['roi_periodo']/periodo_meses
    row['capital_recuperado_pct']=row['roi_periodo']
    if row['resultado_real']>0:
     media_mensal=row['resultado_real']/periodo_meses
     row['payback_meses']=capital/media_mensal if media_mensal>0 else None
   if aquisicao>0:
    row['roi_aquisicao_periodo']=(row['resultado_real']/aquisicao)*Decimal('100')
  total_receita+=row['receita']; total_repasse+=row['repasse']; total_pago+=row['pago']; total_repasse_pago+=row['repasse_pago']; total_custos+=row['custos']; total_aberto+=row['aberto']
  row['driver']=Driver.query.filter_by(id=v.current_driver_id,tenant_id=tenant_id).first() if v.current_driver_id else None
  row['contract']=Contract.query.filter(
   Contract.tenant_id==tenant_id,Contract.vehicle_id==v.id,
   Contract.status.in_(['Gerado','Enviado','Visualizado','Assinado','Ativo'])
  ).order_by(Contract.id.desc()).first()
  row['maintenance']=Maintenance.query.filter(
   Maintenance.tenant_id==tenant_id,Maintenance.vehicle_id==v.id,Maintenance.status.notin_(['Concluída','Cancelada'])
  ).order_by(Maintenance.id.desc()).first()
  row['alerts']=Alert.query.filter_by(tenant_id=tenant_id,entidade='vehicle',entidade_id=v.id).filter(Alert.resolvido_em.is_(None)).count()
  vehicle_rows.append(row)

 # Ranking de rentabilidade: prioriza ROI real quando o proprietário informou capital.
 ranked=[r for r in vehicle_rows if r.get('roi_periodo') is not None]
 ranked.sort(key=lambda r:r['roi_periodo'],reverse=True)
 for pos,row in enumerate(ranked,1):
  row['ranking']=pos
  if row['resultado_real']<=0:
   row['desempenho']='Baixo retorno'
  elif pos==1 and len(ranked)>1:
   row['desempenho']='Maior rentabilidade'
  else:
   row['desempenho']='Retorno positivo'
 for row in vehicle_rows:
  if row.get('roi_periodo') is None:
   row['desempenho']='Investimento pendente'

 months=[]
 cursor=date(inicio.year,inicio.month,1); limit=date(fim.year,fim.month,1)
 while cursor<=limit:
  key=cursor.strftime('%Y-%m'); item=monthly.get(key,{'repasse':Decimal('0'),'pago':Decimal('0'),'custos':Decimal('0')})
  result=item['repasse']-item['custos']
  months.append({'label':cursor.strftime('%m/%Y'),'repasse':float(item['repasse']),'pago':float(item['pago']),'custos':float(item['custos']),'resultado':float(result)})
  cursor=date(cursor.year+1,1,1) if cursor.month==12 else date(cursor.year,cursor.month+1,1)

 total_resultado_real=total_repasse_pago-total_custos
 total_roi_periodo=None
 total_rentabilidade_mensal=None
 total_payback_meses=None
 if total_capital_proprio>0:
  total_roi_periodo=(total_resultado_real/total_capital_proprio)*Decimal('100')
  total_rentabilidade_mensal=total_roi_periodo/periodo_meses
  if total_resultado_real>0:
   total_media_mensal=total_resultado_real/periodo_meses
   if total_media_mensal>0:
    total_payback_meses=total_capital_proprio/total_media_mensal

 return {
  'investor':investor,'vehicles':vehicle_rows,'months':months,
  'ranking':ranked,
  'totals':{
   'receita':total_receita,
   'repasse':total_repasse,
   'pago':total_pago,
   'repasse_pago':total_repasse_pago,
   'custos':total_custos,
   'resultado':total_repasse-total_custos,
   'resultado_real':total_resultado_real,
   'aberto':total_aberto,
   'capital_proprio':total_capital_proprio,
   'roi_periodo':total_roi_periodo,
   'rentabilidade_mensal':total_rentabilidade_mensal,
   'payback_meses':total_payback_meses,
  }
 }


def regra_repasse_locadora_portal(tenant_id, vehicle, data_ref):
 """Retorna os percentuais vigentes para a locadora e o proprietário.

 Veículos sem proprietário vinculado pertencem integralmente à locadora. Para
 veículos de proprietário, usa a regra comercial vigente na competência.
 """
 if not vehicle.investor_id:
  return Decimal('100'),Decimal('0')
 rule=InvestorVehicleRule.query.filter(
  InvestorVehicleRule.tenant_id==tenant_id,
  InvestorVehicleRule.investor_id==vehicle.investor_id,
  InvestorVehicleRule.vehicle_id==vehicle.id,
  InvestorVehicleRule.vigencia_inicio<=data_ref,
  or_(InvestorVehicleRule.vigencia_fim.is_(None),InvestorVehicleRule.vigencia_fim>=data_ref)
 ).order_by(InvestorVehicleRule.vigencia_inicio.desc(),InvestorVehicleRule.id.desc()).first()
 if not rule:
  # Conservador: se há proprietário mas a regra ainda não foi cadastrada,
  # não atribui receita líquida à locadora até a condição comercial ser definida.
  return Decimal('0'),Decimal('100')
 pct_prop=Decimal(str(rule.percentual_proprietario or 0))
 pct_loc=Decimal(str(rule.percentual_locadora if rule.percentual_locadora is not None else (Decimal('100')-pct_prop)))
 return pct_loc,pct_prop


def resumo_financeiro_locadora(tenant_id, inicio, fim):
 """Dashboard financeiro da locadora com visão teórica e realizada separadas."""
 vehicles=Vehicle.query.filter_by(tenant_id=tenant_id).order_by(Vehicle.placa).all()
 vehicle_map={v.id:v for v in vehicles}
 ids=list(vehicle_map.keys())
 rows={v.id:{'vehicle':v,'receita_teorica':Decimal('0'),'repasse_teorico':Decimal('0'),'locadora_teorica':Decimal('0'),'receita_paga':Decimal('0'),'repasse_pago':Decimal('0'),'locadora_real':Decimal('0'),'custos':Decimal('0'),'resultado_teorico':Decimal('0'),'resultado_real':Decimal('0')} for v in vehicles}
 monthly={}
 sem_regra=set()

 def bucket(d):
  key=d.strftime('%Y-%m')
  return monthly.setdefault(key,{'receita_teorica':Decimal('0'),'repasse_teorico':Decimal('0'),'locadora_teorica':Decimal('0'),'receita_paga':Decimal('0'),'repasse_pago':Decimal('0'),'locadora_real':Decimal('0'),'custos':Decimal('0')})

 def parse_date(value, fallback):
  try: return datetime.strptime(value or '','%Y-%m-%d').date()
  except Exception: return fallback

 def periodicidade_dias(c):
  txt=unicodedata.normalize('NFKD',str(c.periodicidade or 'Semanal')).encode('ascii','ignore').decode('ascii').lower()
  if 'diar' in txt: return 1
  if 'mens' in txt: return None
  if 'quinz' in txt: return 15
  return 7

 contracts=Contract.query.filter(Contract.tenant_id==tenant_id,Contract.vehicle_id.in_(ids or [-1])).all()
 contract_map={c.id:c for c in contracts}
 status_validos={'Gerado','Enviado','Visualizado','Assinado','Ativo','Encerrado','Finalizado'}

 # 1) TEÓRICO: competências previstas pelos contratos no intervalo selecionado.
 for c in contracts:
  v=vehicle_map.get(c.vehicle_id)
  if not v or (c.status and c.status not in status_validos): continue
  c_ini=parse_date(c.data_inicio,inicio)
  c_fim=parse_date(c.data_fim,fim)
  periodo_ini=max(inicio,c_ini); periodo_fim=min(fim,c_fim)
  if periodo_ini>periodo_fim: continue
  base=Decimal(str(c.valor_locacao or 0))
  if base<=0: continue
  intervalo=periodicidade_dias(c)
  competencias=[]
  if intervalo is None:
   cursor=date(periodo_ini.year,periodo_ini.month,1); limite=date(periodo_fim.year,periodo_fim.month,1)
   while cursor<=limite:
    competencias.append(max(cursor,periodo_ini))
    cursor=date(cursor.year+1,1,1) if cursor.month==12 else date(cursor.year,cursor.month+1,1)
  else:
   ref=c_ini
   if ref<periodo_ini:
    saltos=max(0,(periodo_ini-ref).days//intervalo); ref=ref+timedelta(days=saltos*intervalo)
    while ref<periodo_ini: ref+=timedelta(days=intervalo)
   while ref<=periodo_fim:
    competencias.append(ref); ref+=timedelta(days=intervalo)
  for ref in competencias:
   pct_loc,pct_prop=regra_repasse_locadora_portal(tenant_id,v,ref)
   if v.investor_id and pct_loc==0 and pct_prop==100:
    has_rule=InvestorVehicleRule.query.filter(InvestorVehicleRule.tenant_id==tenant_id,InvestorVehicleRule.investor_id==v.investor_id,InvestorVehicleRule.vehicle_id==v.id,InvestorVehicleRule.vigencia_inicio<=ref,or_(InvestorVehicleRule.vigencia_fim.is_(None),InvestorVehicleRule.vigencia_fim>=ref)).first()
    if not has_rule: sem_regra.add(v.id)
   rep=base*pct_prop/Decimal('100'); loc=base*pct_loc/Decimal('100')
   r=rows[v.id]; r['receita_teorica']+=base; r['repasse_teorico']+=rep; r['locadora_teorica']+=loc
   b=bucket(ref); b['receita_teorica']+=base; b['repasse_teorico']+=rep; b['locadora_teorica']+=loc

 # 2) REAL: cobranças efetivamente pagas. Excesso de KM entra aqui pelo total cobrado.
 audits=BillingAudit.query.filter(BillingAudit.tenant_id==tenant_id,BillingAudit.contract_id.in_(list(contract_map.keys()) or [-1]),BillingAudit.billing_date>=inicio,BillingAudit.billing_date<=fim).order_by(BillingAudit.billing_date).all()
 for a in audits:
  c=contract_map.get(a.contract_id); v=vehicle_map.get(c.vehicle_id) if c else None
  if not v or (a.payment_status or '').upper()!='PAGO': continue
  total=Decimal(str(a.total_amount or 0)); pct_loc,pct_prop=regra_repasse_locadora_portal(tenant_id,v,a.billing_date)
  rep=total*pct_prop/Decimal('100'); loc=total*pct_loc/Decimal('100')
  r=rows[v.id]; r['receita_paga']+=total; r['repasse_pago']+=rep; r['locadora_real']+=loc
  b=bucket(a.billing_date); b['receita_paga']+=total; b['repasse_pago']+=rep; b['locadora_real']+=loc

 # 3) Custos reais registrados no período.
 for m in Maintenance.query.filter(Maintenance.tenant_id==tenant_id,Maintenance.vehicle_id.in_(ids or [-1])).all():
  try: d=datetime.strptime(m.data or '','%Y-%m-%d').date()
  except Exception: continue
  if d<inicio or d>fim or m.vehicle_id not in rows: continue
  custo=Decimal(str(m.custo or 0)); rows[m.vehicle_id]['custos']+=custo; bucket(d)['custos']+=custo

 totals={k:Decimal('0') for k in ['receita_teorica','repasse_teorico','locadora_teorica','receita_paga','repasse_pago','locadora_real','custos','resultado_teorico','resultado_real']}
 vehicle_rows=[]
 for v in vehicles:
  r=rows[v.id]; r['resultado_teorico']=r['locadora_teorica']; r['resultado_real']=r['locadora_real']
  for k in totals: totals[k]+=r[k]
  r['driver']=Driver.query.filter_by(id=v.current_driver_id,tenant_id=tenant_id).first() if v.current_driver_id else None
  vehicle_rows.append(r)

 months=[]; cursor=date(inicio.year,inicio.month,1); limit=date(fim.year,fim.month,1)
 while cursor<=limit:
  item=monthly.get(cursor.strftime('%Y-%m'),{k:Decimal('0') for k in ['receita_teorica','repasse_teorico','locadora_teorica','receita_paga','repasse_pago','locadora_real','custos']})
  months.append({'label':cursor.strftime('%m/%Y'),'receita_teorica':float(item['receita_teorica']),'repasse_teorico':float(item['repasse_teorico']),'locadora_teorica':float(item['locadora_teorica']),'receita_paga':float(item['receita_paga']),'repasse_pago':float(item['repasse_pago']),'locadora_real':float(item['locadora_real']),'custos':float(item['custos']),'resultado_teorico':float(item['locadora_teorica']),'resultado_real':float(item['locadora_real'])})
  cursor=date(cursor.year+1,1,1) if cursor.month==12 else date(cursor.year,cursor.month+1,1)
 return {'vehicles':vehicle_rows,'months':months,'totals':totals,'sem_regra':len(sem_regra)}


@app.route('/financeiro-locadora')
@login_required
def financeiro_locadora():
 today=date.today()
 try: inicio=datetime.strptime(request.args.get('inicio') or f'{today.year}-01-01','%Y-%m-%d').date()
 except Exception: inicio=date(today.year,1,1)
 try: fim=datetime.strptime(request.args.get('fim') or today.isoformat(),'%Y-%m-%d').date()
 except Exception: fim=today
 if inicio>fim: inicio,fim=fim,inicio
 data=resumo_financeiro_locadora(tid(),inicio,fim)
 return render_template('financeiro_locadora.html',inicio=inicio.isoformat(),fim=fim.isoformat(),**data)


@app.route('/investidores',methods=['GET','POST'])
@login_required
def investidores():
 if request.method=='POST':
  telefone_bruto=request.form.get('telefone')
  telefone=normalize_phone(telefone_bruto) if (telefone_bruto or '').strip() else ''
  if (telefone_bruto or '').strip() and not telefone:
   flash('Telefone inválido. Informe DDD + número; o Frota Fácil acrescenta o código do Brasil automaticamente.','danger')
   return redirect(url_for('investidores'))
  x=Investor(tenant_id=tid(),nome=request.form['nome'],cpf_cnpj=request.form.get('cpf_cnpj'),telefone=telefone,email=request.form.get('email'),regra_repasse='Percentual por veículo',observacoes=request.form.get('observacoes'))
  db.session.add(x); db.session.commit()
  flash('Proprietário cadastrado. Agora você pode vincular os veículos e percentuais.','success')
  # Convite do portal disparado no momento do cadastro, quando habilitado.
  # O envio manual continua disponível em "Acesso ao portal" independentemente desta opção.
  integracao_whatsapp=Integration.query.filter_by(tenant_id=tid(),tipo='whatsapp').first()
  cfg_whatsapp=CommunicationService.parse_config(integracao_whatsapp)
  if cfg_whatsapp.get('owner_portal_auto_invite_enabled'):
   ok_convite,msg_convite=enviar_acesso_portal_proprietario_whatsapp(x)
   if ok_convite:
    flash('Convite para ativar o Portal do Proprietário enviado automaticamente pelo WhatsApp.','success')
   else:
    flash('Proprietário cadastrado, mas o convite automático do portal não foi enviado: '+msg_convite+' Você pode reenviar manualmente em Acesso ao portal.','warning')
  return redirect(url_for('proprietario_financeiro',id=x.id))
 q=(request.args.get('q') or '').strip(); query=Investor.query.filter_by(tenant_id=tid())
 if q:
  like=f'%{q}%'; query=query.filter(or_(Investor.nome.ilike(like),Investor.cpf_cnpj.ilike(like),Investor.telefone.ilike(like),Investor.email.ilike(like)))
 items=query.order_by(Investor.nome).all()
 resumos={x.id:resumo_financeiro_proprietario(x.id) for x in items}
 return render_template('investidores.html',items=items,resumos=resumos,q=q)

@app.route('/investidores/<int:id>/editar',methods=['GET','POST'])
@login_required
def editar_investidor(id):
 x=Investor.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 if request.method=='POST':
  telefone_bruto=request.form.get('telefone'); telefone=normalize_phone(telefone_bruto) if (telefone_bruto or '').strip() else ''
  if (telefone_bruto or '').strip() and not telefone:
   flash('Telefone inválido. Informe DDD + número.','danger'); return redirect(url_for('editar_investidor',id=x.id))
  for campo in ['nome','cpf_cnpj','email','observacoes']: setattr(x,campo,request.form.get(campo))
  x.telefone=telefone; x.regra_repasse='Percentual por veículo'; db.session.commit(); flash('Proprietário atualizado com sucesso.','success'); return redirect(url_for('proprietario_financeiro',id=x.id))
 return render_template('editar_investidor.html',x=x)

def aplicar_regra_proprietario_veiculo(*,proprietario,veiculo,pct,pct_l,inicio,motorizacao=None,observacoes=None):
 # Não permite tomar silenciosamente um veículo que já pertence a outro proprietário.
 if veiculo.investor_id and veiculo.investor_id!=proprietario.id:
  raise ValueError(f'O veículo {veiculo.placa} já está vinculado a outro proprietário.')
 anterior=InvestorVehicleRule.query.filter_by(tenant_id=tid(),investor_id=proprietario.id,vehicle_id=veiculo.id,vigencia_fim=None).order_by(InvestorVehicleRule.id.desc()).first()
 # Se a regra atual começou na mesma data, atualiza em vez de criar duplicidade.
 if anterior and anterior.vigencia_inicio==inicio:
  anterior.percentual_proprietario=pct
  anterior.percentual_locadora=pct_l
  anterior.observacoes=observacoes
 else:
  if anterior and anterior.vigencia_inicio<inicio:
   anterior.vigencia_fim=inicio-timedelta(days=1)
  db.session.add(InvestorVehicleRule(tenant_id=tid(),investor_id=proprietario.id,vehicle_id=veiculo.id,percentual_proprietario=pct,percentual_locadora=pct_l,vigencia_inicio=inicio,observacoes=observacoes))
 veiculo.investor_id=proprietario.id
 veiculo.proprietario_legal=proprietario.nome
 veiculo.cpf_cnpj_proprietario=proprietario.cpf_cnpj
 mot=(motorizacao or '').strip()
 if mot in ('Combustão','Elétrico','Híbrido'):
  veiculo.motorizacao=mot

@app.route('/investidores/<int:id>/financeiro',methods=['GET','POST'])
@login_required
def proprietario_financeiro(id):
 x=Investor.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 if request.method=='POST':
  try:
   pct=Decimal((request.form.get('percentual_proprietario') or '0').replace(',','.'))
   pct_l=Decimal((request.form.get('percentual_locadora') or str(Decimal('100')-pct)).replace(',','.'))
  except Exception:
   flash('Informe percentuais válidos.','danger'); return redirect(url_for('proprietario_financeiro',id=x.id))
  if pct<0 or pct_l<0 or pct>100 or pct_l>100 or (pct+pct_l)!=Decimal('100'):
   flash('Os percentuais do proprietário e da locadora devem somar exatamente 100%.','danger'); return redirect(url_for('proprietario_financeiro',id=x.id))
  inicio_txt=request.form.get('vigencia_inicio') or date.today().isoformat()
  try: inicio=datetime.strptime(inicio_txt,'%Y-%m-%d').date()
  except Exception: inicio=date.today()
  mot=(request.form.get('motorizacao') or '').strip()
  obs=(request.form.get('observacoes') or '').strip() or None
  acao=(request.form.get('acao') or 'individual').strip()
  try:
   if acao=='lote':
    ids=[]
    for raw in request.form.getlist('vehicle_ids'):
     try: ids.append(int(raw))
     except Exception: pass
    ids=list(dict.fromkeys(ids))
    if not ids:
     flash('Selecione pelo menos um veículo para aplicar a regra em lote.','warning'); return redirect(url_for('proprietario_financeiro',id=x.id))
    veiculos=Vehicle.query.filter(Vehicle.tenant_id==tid(),Vehicle.id.in_(ids)).order_by(Vehicle.placa).all()
    if len(veiculos)!=len(ids):
     raise ValueError('Um ou mais veículos selecionados não pertencem a esta locadora.')
    for v in veiculos:
     aplicar_regra_proprietario_veiculo(proprietario=x,veiculo=v,pct=pct,pct_l=pct_l,inicio=inicio,motorizacao=mot,observacoes=obs)
    db.session.commit(); flash(f'Regra aplicada com sucesso a {len(veiculos)} veículo(s).','success')
   else:
    vehicle_id=request.form.get('vehicle_id',type=int)
    v=Vehicle.query.filter_by(id=vehicle_id,tenant_id=tid()).first_or_404()
    aplicar_regra_proprietario_veiculo(proprietario=x,veiculo=v,pct=pct,pct_l=pct_l,inicio=inicio,motorizacao=mot,observacoes=obs)
    db.session.commit(); flash('Condição comercial do veículo salva.','success')
  except ValueError as exc:
   db.session.rollback(); flash(str(exc),'danger')
  except Exception:
   db.session.rollback(); app.logger.exception('Falha ao salvar condição comercial do proprietário %s',x.id); flash('Não foi possível salvar a condição comercial.','danger')
  return redirect(url_for('proprietario_financeiro',id=x.id))
 resumo=resumo_financeiro_proprietario(x.id)
 regras=InvestorVehicleRule.query.filter_by(tenant_id=tid(),investor_id=x.id).order_by(InvestorVehicleRule.vigencia_inicio.desc(),InvestorVehicleRule.id.desc()).all()
 disponiveis=Vehicle.query.filter(Vehicle.tenant_id==tid(),or_(Vehicle.investor_id==x.id,Vehicle.investor_id.is_(None))).order_by(Vehicle.placa).all()
 # previsão do contrato vigente por veículo
 previsoes={}
 for v in resumo['veiculos']:
  c=Contract.query.filter(Contract.tenant_id==tid(),Contract.vehicle_id==v.id,Contract.status.in_(['Assinado','Ativo'])).order_by(Contract.id.desc()).first()
  regra=regra_proprietario_veiculo(v.id,date.today())
  if c and regra:
   base=Decimal(str(c.valor_locacao or 0)); pp=Decimal(str(regra.percentual_proprietario or 0)); pl=Decimal(str(regra.percentual_locadora or 0))
   previsoes[v.id]={'contrato':c,'base':base,'proprietario':base*pp/Decimal('100'),'locadora':base*pl/Decimal('100'),'regra':regra}
 return render_template('proprietario_financeiro.html',x=x,resumo=resumo,regras=regras,veiculos_disponiveis=disponiveis,previsoes=previsoes)




def driver_portal_required(view):
 @wraps(view)
 def wrapped(*args,**kwargs):
  access_id=session.get('driver_access_id'); driver_id=session.get('driver_portal_driver_id'); tenant_id=session.get('driver_portal_tenant_id')
  if not access_id or not driver_id or not tenant_id:
   return redirect(url_for('portal_motorista_entrar'))
  access=DriverAccess.query.filter_by(id=access_id,driver_id=driver_id,tenant_id=tenant_id,ativo=True).first()
  if not access:
   session.pop('driver_access_id',None); session.pop('driver_portal_driver_id',None); session.pop('driver_portal_tenant_id',None)
   return redirect(url_for('portal_motorista_entrar'))
  return view(*args,**kwargs)
 return wrapped

def _driver_activation_serializer():
 return URLSafeTimedSerializer(app.config['SECRET_KEY'],salt='frota-facil-driver-activation-v1')

def gerar_link_ativacao_portal_motorista(driver,reset_nonce=None):
 if not driver: return None
 payload={'tenant_id':driver.tenant_id,'driver_id':driver.id}
 if reset_nonce: payload['reset_nonce']=reset_nonce
 token=_driver_activation_serializer().dumps(payload)
 return url_for('portal_motorista_ativar',token=token,_external=True)

def _dados_token_ativacao_motorista(token,max_age=604800):
 try:
  data=_driver_activation_serializer().loads(token,max_age=max_age)
 except SignatureExpired:
  return None,'Este link de ativação expirou. Solicite um novo link à locadora.',None
 except BadSignature:
  return None,'Este link de ativação é inválido.',None
 try:
  tenant_id=int(data.get('tenant_id')); driver_id=int(data.get('driver_id'))
 except (TypeError,ValueError,AttributeError):
  return None,'Este link de ativação é inválido.',None
 driver=Driver.query.filter_by(id=driver_id,tenant_id=tenant_id).first()
 if not driver: return None,'Motorista não encontrado para este link.',None
 return driver,None,data

def enviar_acesso_portal_motorista_whatsapp(driver):
 if not driver: return False,'Motorista não encontrado.'
 telefone=normalize_phone(driver.telefone)
 if not telefone: return False,'Motorista sem telefone/WhatsApp válido.'
 integration=Integration.query.filter_by(tenant_id=driver.tenant_id,tipo='whatsapp').first()
 cfg=CommunicationService.parse_config(integration)
 provider=(cfg.get('provider') or 'web').lower()
 if provider!='business' or not integration or not integration.ativo:
  return False,'O envio requer WhatsApp Business conectado.'
 template_name=(cfg.get('driver_portal_template_name') or '').strip() or 'acesso_portal_motorista'
 tenant=Tenant.query.get(driver.tenant_id)
 nome_locadora=(((tenant.nome_fantasia if tenant else '') or '').strip() or ((tenant.nome if tenant else '') or '').strip() or 'Locadora')
 access=DriverAccess.query.filter_by(tenant_id=driver.tenant_id,driver_id=driver.id).first()
 if access and (access.senha or '').startswith('RESET_PENDING:'):
  nonce=(access.senha or '').split(':',1)[1]
  link=gerar_link_ativacao_portal_motorista(driver,reset_nonce=nonce)
  instrucao='Abra o link para escolher novamente seu e-mail e criar uma nova senha.'
 elif access and access.ativo:
  link=url_for('portal_motorista_entrar',tenant=driver.tenant_id,_external=True)
  instrucao='Use seu e-mail e senha já cadastrados para entrar.'
 else:
  link=gerar_link_ativacao_portal_motorista(driver)
  instrucao='Abra o link para escolher seu e-mail e criar sua senha.'
 params=[driver.nome or '',nome_locadora,link]
 body=f'Olá, {driver.nome}. A {nome_locadora} disponibilizou seu Portal do Motorista. {instrucao} Acesse: {link}'
 fila=MessageQueue(tenant_id=driver.tenant_id,channel='whatsapp',provider='whatsapp_business',recipient=telefone,recipient_name=driver.nome,message_type='acesso_portal_motorista',body=body,template_name=template_name,template_parameters=json.dumps(params,ensure_ascii=False),related_entity='Motorista',related_entity_id=driver.id,status='PENDENTE',created_at=agora_sao_paulo_naive(),updated_at=agora_sao_paulo_naive())
 db.session.add(fila); db.session.flush()
 try:
  result=CommunicationService().send_whatsapp(phone=telefone,message=body,integration=integration,template_name=template_name,template_language=cfg.get('template_language') or 'pt_BR',template_parameters=params)
  fila.provider=result.provider; fila.status=result.status; fila.external_id=result.external_id; fila.attempts=(fila.attempts or 0)+1; fila.sent_at=agora_sao_paulo_naive() if result.status=='ENVIADA' else None
  db.session.add(MessageEvent(tenant_id=driver.tenant_id,message_id=fila.id,event=result.status,description='Acesso ao Portal do Motorista processado pelo WhatsApp.',created_at=agora_sao_paulo_naive()))
  db.session.commit()
  return (True,'Link do Portal do Motorista enviado pelo WhatsApp.') if result.status=='ENVIADA' else (False,f'O WhatsApp retornou status {result.status}.')
 except CommunicationError as exc:
  fila.status='FALHA'; fila.error_message=str(exc); fila.attempts=(fila.attempts or 0)+1
  db.session.add(MessageEvent(tenant_id=driver.tenant_id,message_id=fila.id,event='FALHA',description=str(exc),created_at=agora_sao_paulo_naive()))
  db.session.commit(); return False,str(exc)

@app.route('/motoristas/portal-acessos')
@login_required
def acessos_portal_motorista():
 drivers=Driver.query.filter_by(tenant_id=tid()).order_by(Driver.nome).all()
 access_map={a.driver_id:a for a in DriverAccess.query.filter_by(tenant_id=tid()).all()}
 return render_template('motorista_acessos.html',drivers=drivers,access_map=access_map)

@app.route('/motoristas/<int:id>/portal/whatsapp',methods=['POST'])
@login_required
def enviar_acesso_motorista_whatsapp(id):
 driver=Driver.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 ok,msg=enviar_acesso_portal_motorista_whatsapp(driver); flash(msg,'success' if ok else 'warning')
 return redirect(url_for('acessos_portal_motorista'))

@app.route('/motoristas/<int:id>/portal/redefinir',methods=['POST'])
@login_required
def redefinir_acesso_motorista(id):
 driver=Driver.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 access=DriverAccess.query.filter_by(tenant_id=tid(),driver_id=driver.id).first()
 if access:
  nonce=uuid.uuid4().hex+uuid.uuid4().hex; access.ativo=False; access.senha='RESET_PENDING:'+nonce; db.session.commit()
 ok,msg=enviar_acesso_portal_motorista_whatsapp(driver)
 flash(('Acesso anterior invalidado. '+msg) if access else msg,'success' if ok else 'warning')
 return redirect(url_for('acessos_portal_motorista'))

@app.route('/motoristas/<int:id>/portal/bloquear',methods=['POST'])
@login_required
def bloquear_acesso_motorista(id):
 driver=Driver.query.filter_by(id=id,tenant_id=tid()).first_or_404(); access=DriverAccess.query.filter_by(tenant_id=tid(),driver_id=driver.id).first()
 if access: access.ativo=False; db.session.commit(); flash('Acesso do motorista bloqueado.','success')
 return redirect(url_for('acessos_portal_motorista'))

@app.route('/portal-motorista/ativar/<token>',methods=['GET','POST'])
def portal_motorista_ativar(token):
 driver,erro,data=_dados_token_ativacao_motorista(token)
 if erro: return render_template('portal_motorista_ativar.html',driver=None,tenant=None,erro=erro,email_sugerido='',token=token),400
 tenant=Tenant.query.get(driver.tenant_id); access=DriverAccess.query.filter_by(tenant_id=driver.tenant_id,driver_id=driver.id).first()
 reset=bool(access and (access.senha or '').startswith('RESET_PENDING:')); token_nonce=(data or {}).get('reset_nonce')
 if access and reset:
  stored=(access.senha or '').split(':',1)[1]
  if not token_nonce or token_nonce!=stored: return render_template('portal_motorista_ativar.html',driver=driver,tenant=tenant,erro='Este link de redefinição não é mais válido.',email_sugerido=access.email or driver.email or '',token=token),400
 elif access:
  flash('Seu acesso ao Portal do Motorista já foi criado. Entre com seu e-mail e senha.','success'); return redirect(url_for('portal_motorista_entrar',tenant=driver.tenant_id))
 elif token_nonce:
  return render_template('portal_motorista_ativar.html',driver=driver,tenant=tenant,erro='Este link de redefinição não é mais válido.',email_sugerido=driver.email or '',token=token),400
 email_sugerido=((access.email if access else None) or driver.email or '').strip().lower()
 if request.method=='POST':
  email=(request.form.get('email') or '').strip().lower(); senha=request.form.get('senha') or ''; confirmar=request.form.get('confirmar_senha') or ''
  if not email or '@' not in email: flash('Informe um e-mail válido.','danger')
  elif len(senha)<6: flash('Crie uma senha com pelo menos 6 caracteres.','danger')
  elif senha!=confirmar: flash('A confirmação da senha não confere.','danger')
  else:
   duplicate=DriverAccess.query.filter(DriverAccess.email==email)
   if access: duplicate=duplicate.filter(DriverAccess.id!=access.id)
   if duplicate.first(): flash('Este e-mail já está vinculado a outro acesso de motorista.','danger')
   else:
    if access: access.email=email; access.senha=generate_password_hash(senha); access.ativo=True
    else: access=DriverAccess(tenant_id=driver.tenant_id,driver_id=driver.id,email=email,senha=generate_password_hash(senha),ativo=True); db.session.add(access)
    db.session.commit(); session['driver_access_id']=access.id; session['driver_portal_driver_id']=driver.id; session['driver_portal_tenant_id']=driver.tenant_id; access.ultimo_acesso_em=datetime.utcnow(); db.session.commit()
    flash('Acesso criado com sucesso. Bem-vindo ao Portal do Motorista!','success'); return redirect(url_for('portal_motorista'))
  email_sugerido=email
 return render_template('portal_motorista_ativar.html',driver=driver,tenant=tenant,erro=None,email_sugerido=email_sugerido,token=token)

@app.route('/portal-motorista/entrar',methods=['GET','POST'])
def portal_motorista_entrar():
 if session.get('driver_access_id'): return redirect(url_for('portal_motorista'))
 tenant_ref=request.args.get('tenant',type=int) or request.form.get('tenant',type=int); tenant_login=Tenant.query.get(tenant_ref) if tenant_ref else None
 if request.method=='POST':
  email=(request.form.get('email') or '').strip().lower(); access=DriverAccess.query.filter_by(email=email,ativo=True).first()
  if access and check_password_hash(access.senha,request.form.get('senha') or ''):
   driver=Driver.query.filter_by(id=access.driver_id,tenant_id=access.tenant_id).first()
   if driver:
    session['driver_access_id']=access.id; session['driver_portal_driver_id']=driver.id; session['driver_portal_tenant_id']=driver.tenant_id; access.ultimo_acesso_em=datetime.utcnow(); db.session.commit(); return redirect(url_for('portal_motorista'))
  flash('E-mail ou senha inválidos.','danger')
 return render_template('portal_motorista_login.html',tenant_login=tenant_login)

@app.route('/portal-motorista/sair')
def portal_motorista_sair():
 tenant_ref=session.get('driver_portal_tenant_id'); session.pop('driver_access_id',None); session.pop('driver_portal_driver_id',None); session.pop('driver_portal_tenant_id',None)
 return redirect(url_for('portal_motorista_entrar',tenant=tenant_ref)) if tenant_ref else redirect(url_for('portal_motorista_entrar'))

@app.route('/portal-motorista')
@driver_portal_required
def portal_motorista():
 tenant_id=int(session['driver_portal_tenant_id']); driver_id=int(session['driver_portal_driver_id']); driver=Driver.query.filter_by(id=driver_id,tenant_id=tenant_id).first_or_404(); tenant=Tenant.query.get(tenant_id)
 contracts=Contract.query.filter_by(tenant_id=tenant_id,driver_id=driver_id).order_by(Contract.id.desc()).all()
 contract_ids=[c.id for c in contracts]
 active_contracts=[c for c in contracts if (c.status or '').strip() in {'Assinado','Ativo'}]
 active_contract_ids={c.id for c in active_contracts}
 active_vehicle_ids={c.vehicle_id for c in active_contracts if c.vehicle_id}
 vehicle_ids={c.vehicle_id for c in contracts if c.vehicle_id}
 current_vehicle=Vehicle.query.filter_by(tenant_id=tenant_id,current_driver_id=driver_id).first()
 if not current_vehicle:
  active=next((c for c in active_contracts if c.vehicle_id),None); current_vehicle=active.vehicle if active else None
 audits_raw=BillingAudit.query.filter(BillingAudit.tenant_id==tenant_id,BillingAudit.contract_id.in_(contract_ids)).order_by(BillingAudit.billing_date.desc(),BillingAudit.id.desc()).all() if contract_ids else []
 contract_map={c.id:c for c in contracts}

 # Consolida registros antigos criados pelo bug de cobrança diária.
 # Para contratos semanais, cada semana é uma única competência financeira.
 # Para outras periodicidades, mantém uma cobrança por data.
 # Prioridade: PAGO > COMPROVANTE_RECEBIDO > demais; em empate, registro mais recente.
 def _portal_billing_key(audit):
  c=contract_map.get(audit.contract_id)
  periodicidade=unicodedata.normalize('NFKD',str(c.periodicidade or '')).encode('ascii','ignore').decode('ascii').lower() if c else ''
  d=audit.billing_date
  if d and 'seman' in periodicidade:
   iso=d.isocalendar()
   return (audit.contract_id,'semana',iso.year,iso.week)
  return (audit.contract_id,'data',d)

 def _portal_billing_rank(audit):
  status=(audit.payment_status or '').upper()
  if status=='PAGO': return 3
  if status=='COMPROVANTE_RECEBIDO': return 2
  return 1

 audits_por_competencia={}
 for audit in audits_raw:
  chave=_portal_billing_key(audit)
  atual=audits_por_competencia.get(chave)
  if atual is None or _portal_billing_rank(audit)>_portal_billing_rank(atual) or (_portal_billing_rank(audit)==_portal_billing_rank(atual) and (audit.id or 0)>(atual.id or 0)):
   audits_por_competencia[chave]=audit

 audits=sorted(audits_por_competencia.values(),key=lambda a:(a.billing_date or date.min,a.id or 0),reverse=True)

 # Situação atual: somente contratos ativos. Comprovante já enviado não é mais
 # uma ação pendente do motorista; aguarda apenas validação da locadora.
 current_audits=[a for a in audits if a.contract_id in active_contract_ids]
 pagos=sum(Decimal(str(a.total_amount or 0)) for a in current_audits if (a.payment_status or '').upper()=='PAGO')
 pendentes=sum(Decimal(str(a.total_amount or 0)) for a in current_audits if (a.payment_status or '').upper()!='PAGO')
 km_requests=MileageRequest.query.filter_by(tenant_id=tenant_id,driver_id=driver_id).order_by(MileageRequest.sent_at.desc()).limit(30).all()
 inspections=Inspection.query.filter_by(tenant_id=tenant_id,driver_id=driver_id).order_by(Inspection.requested_at.desc()).limit(30).all()
 maintenances=Maintenance.query.filter(Maintenance.tenant_id==tenant_id,Maintenance.vehicle_id.in_(list(vehicle_ids))).order_by(Maintenance.id.desc()).limit(30).all() if vehicle_ids else []

 # CRLV-e: somente documentos ativos dos veículos vinculados a contratos atuais
 # do motorista. Mantemos apenas o CRLV mais recente por veículo.
 crlv_by_vehicle={}
 if active_vehicle_ids:
  crlv_docs=Document.query.filter(
   Document.tenant_id==tenant_id,
   Document.entidade_id.in_(list(active_vehicle_ids)),
   Document.status=='Ativo',
   Document.tipo.ilike('%CRLV%'),
  ).order_by(Document.criado_em.desc(),Document.id.desc()).all()
  for doc in crlv_docs:
   entidade_norm=unicodedata.normalize('NFKD',str(doc.entidade or '')).encode('ascii','ignore').decode('ascii').lower()
   if entidade_norm!='veiculo':
    continue
   if doc.entidade_id not in crlv_by_vehicle:
    crlv_by_vehicle[doc.entidade_id]=doc

 audit_context={}
 for a in audits:
  c=contract_map.get(a.contract_id); v=c.vehicle if c else None
  if c:
   partes=[(v.placa if v else (a.plate or '')),(v.marca_modelo if v else (a.vehicle_label or 'Veículo')),(c.numero_contrato or f'Contrato #{c.id}')]
   if c.id not in active_contract_ids: partes.append('Histórico')
   a.vehicle_label=' · '.join(x for x in partes if x)
  audit_context[a.id]={
   'contrato':(c.numero_contrato if c else f'Contrato #{a.contract_id}'),
   'placa':(v.placa if v else (a.plate or '')),
   'veiculo':(v.marca_modelo if v else (a.vehicle_label or 'Veículo')),
   'atual':bool(c and c.id in active_contract_ids),
  }
 pendencias=[]
 for a in current_audits:
  status_pag=(a.payment_status or '').upper()
  if status_pag not in {'PAGO','COMPROVANTE_RECEBIDO'}:
   ctx=audit_context.get(a.id,{})
   ident=' · '.join(x for x in [ctx.get('placa'),ctx.get('veiculo'),ctx.get('contrato')] if x)
   pendencias.append({'tipo':'Pagamento','texto':f'{ident} — R$ {brl(a.total_amount)}','url':url_for('enviar_comprovante_pagamento',token=a.receipt_token) if a.receipt_token else None})
 for r in km_requests:
  if r.vehicle_id in active_vehicle_ids and (r.status or '').lower()=='pendente' and (not r.expires_at or r.expires_at>=datetime.utcnow()):
   c=next((x for x in active_contracts if x.vehicle_id==r.vehicle_id),None); numero=c.numero_contrato if c else ''
   pendencias.append({'tipo':'Quilometragem','texto':f'{r.vehicle.placa} · {r.vehicle.marca_modelo}'+(f' · {numero}' if numero else ''),'url':url_for('registrar_quilometragem_publica',token=r.token)})
 for i in inspections:
  if i.contract_id in active_contract_ids and (i.status or '').lower()=='pendente' and (not i.expires_at or i.expires_at>=datetime.utcnow()):
   c=contract_map.get(i.contract_id); numero=c.numero_contrato if c else ''
   pendencias.append({'tipo':'Vistoria','texto':f'{i.vehicle.placa} · {i.vehicle.marca_modelo}'+(f' · {numero}' if numero else ''),'url':url_for('vistoria_publica',token=i.token)})
 return render_template('portal_motorista.html',tenant=tenant,driver=driver,current_vehicle=current_vehicle,contracts=contracts,active_contract_ids=active_contract_ids,audits=audits,audit_context=audit_context,pagos=pagos,pendentes=pendentes,km_requests=km_requests,inspections=inspections,maintenances=maintenances,pendencias=pendencias,crlv_by_vehicle=crlv_by_vehicle)

def _portal_motorista_crlv_autorizado(document_id):
 tenant_id=int(session['driver_portal_tenant_id'])
 driver_id=int(session['driver_portal_driver_id'])
 doc=Document.query.filter_by(id=document_id,tenant_id=tenant_id,status='Ativo').first_or_404()
 entidade_norm=unicodedata.normalize('NFKD',str(doc.entidade or '')).encode('ascii','ignore').decode('ascii').lower()
 tipo_norm=unicodedata.normalize('NFKD',str(doc.tipo or '')).encode('ascii','ignore').decode('ascii').upper()
 if entidade_norm!='veiculo' or 'CRLV' not in tipo_norm or not doc.entidade_id:
  abort(404)
 autorizado=Contract.query.filter(
  Contract.tenant_id==tenant_id,
  Contract.driver_id==driver_id,
  Contract.vehicle_id==doc.entidade_id,
  Contract.status.in_(['Assinado','Ativo']),
 ).first()
 if not autorizado:
  abort(404)
 return doc

@app.route('/portal-motorista/documentos/<int:document_id>/crlv')
@driver_portal_required
def portal_motorista_visualizar_crlv(document_id):
 doc=_portal_motorista_crlv_autorizado(document_id)
 try:
  conteudo=storage.download(doc.arquivo)
 except StorageNotFoundError:
  abort(404)
 except Exception:
  app.logger.exception('Falha ao visualizar CRLV %s no Portal do Motorista',doc.id)
  abort(503)
 mimetype=mimetypes.guess_type(doc.nome_original or '')[0] or 'application/octet-stream'
 return send_file(BytesIO(conteudo),as_attachment=False,download_name=doc.nome_original or 'CRLV.pdf',mimetype=mimetype)

@app.route('/portal-motorista/documentos/<int:document_id>/crlv/baixar')
@driver_portal_required
def portal_motorista_baixar_crlv(document_id):
 doc=_portal_motorista_crlv_autorizado(document_id)
 try:
  conteudo=storage.download(doc.arquivo)
 except StorageNotFoundError:
  abort(404)
 except Exception:
  app.logger.exception('Falha ao baixar CRLV %s no Portal do Motorista',doc.id)
  abort(503)
 mimetype=mimetypes.guess_type(doc.nome_original or '')[0] or 'application/octet-stream'
 return send_file(BytesIO(conteudo),as_attachment=True,download_name=doc.nome_original or 'CRLV.pdf',mimetype=mimetype)

def _owner_activation_serializer():
 return URLSafeTimedSerializer(app.config['SECRET_KEY'],salt='frota-facil-owner-activation-v1')

def gerar_link_ativacao_portal_proprietario(investor,reset_nonce=None):
 """Gera link assinado para o proprietário escolher e-mail e senha do portal.

 Quando ``reset_nonce`` é informado, o link fica vinculado à redefinição atual,
 impedindo que um convite antigo seja reutilizado após uma nova redefinição.
 """
 if not investor:
  return None
 payload={'tenant_id':investor.tenant_id,'investor_id':investor.id}
 if reset_nonce:
  payload['reset_nonce']=reset_nonce
 token=_owner_activation_serializer().dumps(payload)
 return url_for('portal_proprietario_ativar',token=token,_external=True)

def _dados_token_ativacao_proprietario(token,max_age=604800):
 try:
  data=_owner_activation_serializer().loads(token,max_age=max_age)
 except SignatureExpired:
  return None,'Este link de ativação expirou. Solicite um novo link à locadora.',None
 except BadSignature:
  return None,'Este link de ativação é inválido.',None
 try:
  tenant_id=int(data.get('tenant_id'))
  investor_id=int(data.get('investor_id'))
 except (TypeError,ValueError,AttributeError):
  return None,'Este link de ativação é inválido.',None
 investor=Investor.query.filter_by(id=investor_id,tenant_id=tenant_id).first()
 if not investor:
  return None,'Proprietário não encontrado para este link.',None
 return investor,None,data

def enviar_acesso_portal_proprietario_whatsapp(investor):
 """Envia ao proprietário o link do Portal do Proprietário via template Meta.

 O nome do template pode ser configurado por ``owner_portal_template_name`` na
 integração WhatsApp. Enquanto essa opção ainda não estiver exposta na tela de
 Integrações, usa o template aprovado ``acesso_portal_proprietario`` como padrão.
 Parâmetros do template: 1 proprietário, 2 locadora, 3 link (ativação ou login).
 """
 if not investor:
  return False,'Proprietário não encontrado.'
 telefone=normalize_phone(investor.telefone)
 if not telefone:
  return False,'Proprietário sem telefone/WhatsApp válido.'
 integration=Integration.query.filter_by(tenant_id=investor.tenant_id,tipo='whatsapp').first()
 cfg=CommunicationService.parse_config(integration)
 provider=(cfg.get('provider') or 'web').lower()
 if provider!='business' or not integration or not integration.ativo:
  return False,'O envio automático requer WhatsApp Business conectado.'
 template_name=(cfg.get('owner_portal_template_name') or '').strip() or 'acesso_portal_proprietario'
 tenant=Tenant.query.get(investor.tenant_id)
 nome_locadora=(
  (((tenant.nome_fantasia if tenant else '') or '').strip())
  or (((tenant.nome if tenant else '') or '').strip())
  or 'Locadora'
 )
 access=InvestorAccess.query.filter_by(tenant_id=investor.tenant_id,investor_id=investor.id).first()
 reset_nonce=None
 if access and (access.senha or '').startswith('RESET_PENDING:'):
  reset_nonce=(access.senha or '').split(':',1)[1]
  link=gerar_link_ativacao_portal_proprietario(investor,reset_nonce=reset_nonce)
  instrucao='Abra o link para escolher novamente seu e-mail de acesso e criar uma nova senha.'
 elif access and access.ativo:
  link=url_for('portal_proprietario_entrar',tenant=investor.tenant_id,_external=True)
  instrucao='Use seu e-mail e senha já cadastrados para entrar.'
 else:
  link=gerar_link_ativacao_portal_proprietario(investor)
  instrucao='Abra o link para escolher seu e-mail de acesso e criar sua senha.'
 params=[investor.nome or '',nome_locadora,link]
 body=(f'Olá, {investor.nome}. A {nome_locadora} disponibilizou seu Portal do Proprietário. '
       f'{instrucao} Acesse: {link}')
 fila=MessageQueue(
  tenant_id=investor.tenant_id,channel='whatsapp',provider='whatsapp_business',
  recipient=telefone,recipient_name=investor.nome,message_type='acesso_portal_proprietario',
  body=body,template_name=template_name,template_parameters=json.dumps(params,ensure_ascii=False),
  related_entity='Proprietario',related_entity_id=investor.id,status='PENDENTE',
  created_at=agora_sao_paulo_naive(),updated_at=agora_sao_paulo_naive(),
 )
 db.session.add(fila); db.session.flush()
 try:
  result=CommunicationService().send_whatsapp(
   phone=telefone,message=body,integration=integration,template_name=template_name,
   template_language=cfg.get('template_language') or 'pt_BR',template_parameters=params,
  )
  fila.provider=result.provider; fila.status=result.status; fila.external_id=result.external_id
  fila.attempts=(fila.attempts or 0)+1
  fila.sent_at=agora_sao_paulo_naive() if result.status=='ENVIADA' else None
  db.session.add(MessageEvent(
   tenant_id=investor.tenant_id,message_id=fila.id,event=result.status,
   description='Acesso ao Portal do Proprietário processado pelo WhatsApp.',
   created_at=agora_sao_paulo_naive(),
  ))
  db.session.commit()
  if result.status=='ENVIADA':
   return True,'Link do Portal do Proprietário enviado pelo WhatsApp.'
  return False,f'O WhatsApp retornou status {result.status}.'
 except CommunicationError as exc:
  fila.status='FALHA'; fila.error_message=str(exc); fila.attempts=(fila.attempts or 0)+1
  db.session.add(MessageEvent(
   tenant_id=investor.tenant_id,message_id=fila.id,event='FALHA',description=str(exc),
   created_at=agora_sao_paulo_naive(),
  ))
  db.session.commit()
  return False,str(exc)

@app.route('/investidores/<int:id>/acesso/whatsapp',methods=['POST'])
@login_required
def enviar_acesso_proprietario_whatsapp(id):
 investor=Investor.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 access=InvestorAccess.query.filter_by(tenant_id=tid(),investor_id=investor.id,ativo=True).first()
 ok,msg=enviar_acesso_portal_proprietario_whatsapp(investor)
 flash(msg,'success' if ok else 'warning')
 return redirect(url_for('acesso_proprietario',id=id))

@app.route('/investidores/<int:id>/acesso/redefinir',methods=['POST'])
@login_required
def redefinir_acesso_proprietario(id):
 investor=Investor.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 access=InvestorAccess.query.filter_by(tenant_id=tid(),investor_id=investor.id).first()
 if not access:
  ok,msg=enviar_acesso_portal_proprietario_whatsapp(investor)
  flash(msg,'success' if ok else 'warning')
  return redirect(url_for('acesso_proprietario',id=id))
 # Mantém o mesmo registro para preservar o vínculo, mas invalida imediatamente
 # a senha atual e exige uma nova ativação. O nonce também revoga convites de
 # redefinição anteriores sem precisar de nova coluna no banco.
 reset_nonce=uuid.uuid4().hex+uuid.uuid4().hex
 access.ativo=False
 access.senha='RESET_PENDING:'+reset_nonce
 db.session.commit()
 ok,msg=enviar_acesso_portal_proprietario_whatsapp(investor)
 if ok:
  flash('Acesso anterior invalidado. Novo link de ativação enviado pelo WhatsApp.','success')
 else:
  flash('Acesso anterior invalidado, mas o WhatsApp não foi enviado: '+msg,'warning')
 return redirect(url_for('acesso_proprietario',id=id))

@app.route('/investidores/<int:id>/acesso',methods=['GET','POST'])
@login_required
def acesso_proprietario(id):
 investor=Investor.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 access=InvestorAccess.query.filter_by(tenant_id=tid(),investor_id=investor.id).first()
 senha_temporaria=None
 acesso_novo=False
 if request.method=='POST':
  action=(request.form.get('acao') or 'salvar').strip()
  if action=='bloquear' and access:
   access.ativo=False; db.session.commit(); flash('Acesso do proprietário bloqueado.','success')
   return redirect(url_for('acesso_proprietario',id=id))
  if action=='ativar' and access:
   access.ativo=True; db.session.commit(); flash('Acesso do proprietário ativado.','success')
   return redirect(url_for('acesso_proprietario',id=id))
  email=(request.form.get('email') or investor.email or '').strip().lower()
  if not email:
   flash('Informe um e-mail para o acesso do proprietário.','danger')
   return redirect(url_for('acesso_proprietario',id=id))
  duplicate=InvestorAccess.query.filter(InvestorAccess.email==email)
  if access: duplicate=duplicate.filter(InvestorAccess.id!=access.id)
  if duplicate.first():
   flash('Este e-mail já está vinculado a outro acesso de proprietário.','danger')
   return redirect(url_for('acesso_proprietario',id=id))
  nova_senha=(request.form.get('senha') or '').strip()
  if not access:
   if not nova_senha:
    nova_senha=uuid.uuid4().hex[:10]
    senha_temporaria=nova_senha
   access=InvestorAccess(tenant_id=tid(),investor_id=investor.id,email=email,senha=generate_password_hash(nova_senha),ativo=True)
   db.session.add(access)
   acesso_novo=True
  else:
   access.email=email; access.ativo=True
   if nova_senha:
    access.senha=generate_password_hash(nova_senha)
    senha_temporaria=nova_senha
  investor.email=investor.email or email
  db.session.commit()
  envio_portal_msg=None
  envio_portal_ok=False
  if acesso_novo:
   envio_portal_ok,envio_portal_msg=enviar_acesso_portal_proprietario_whatsapp(investor)
  if senha_temporaria:
   return render_template('proprietario_acesso.html',investor=investor,access=access,senha_temporaria=senha_temporaria,envio_portal_msg=envio_portal_msg,envio_portal_ok=envio_portal_ok)
  flash('Acesso do proprietário atualizado.','success')
  return redirect(url_for('acesso_proprietario',id=id))
 return render_template('proprietario_acesso.html',investor=investor,access=access,senha_temporaria=None,envio_portal_msg=None,envio_portal_ok=False)


@app.route('/portal-proprietario/ativar/<token>',methods=['GET','POST'])
def portal_proprietario_ativar(token):
 investor,erro,token_data=_dados_token_ativacao_proprietario(token)
 if erro:
  return render_template('portal_proprietario_ativar.html',investor=None,tenant=None,erro=erro,email_sugerido='',token=token),400
 tenant=Tenant.query.get(investor.tenant_id)
 access=InvestorAccess.query.filter_by(tenant_id=investor.tenant_id,investor_id=investor.id).first()
 reset_em_andamento=bool(access and (access.senha or '').startswith('RESET_PENDING:'))
 token_reset_nonce=(token_data or {}).get('reset_nonce')
 if access and reset_em_andamento:
  stored_nonce=(access.senha or '').split(':',1)[1]
  if not token_reset_nonce or token_reset_nonce!=stored_nonce:
   erro='Este link de redefinição não é mais válido. Solicite um novo link à locadora.'
   return render_template('portal_proprietario_ativar.html',investor=investor,tenant=tenant,erro=erro,email_sugerido=access.email or investor.email or '',token=token),400
 elif access:
  flash('Seu acesso ao Portal do Proprietário já foi criado. Entre com seu e-mail e senha.','success')
  return redirect(url_for('portal_proprietario_entrar',tenant=investor.tenant_id))
 elif token_reset_nonce:
  erro='Este link de redefinição não é mais válido. Solicite um novo link à locadora.'
  return render_template('portal_proprietario_ativar.html',investor=investor,tenant=tenant,erro=erro,email_sugerido=investor.email or '',token=token),400
 email_sugerido=((access.email if access else None) or investor.email or '').strip().lower()
 if request.method=='POST':
  email=(request.form.get('email') or '').strip().lower()
  senha=request.form.get('senha') or ''
  confirmar=request.form.get('confirmar_senha') or ''
  if not email or '@' not in email or '.' not in email.rsplit('@',1)[-1]:
   flash('Informe um e-mail válido para acessar o portal.','danger')
   return render_template('portal_proprietario_ativar.html',investor=investor,tenant=tenant,erro=None,email_sugerido=email,token=token)
  if len(senha)<6:
   flash('Crie uma senha com pelo menos 6 caracteres.','danger')
   return render_template('portal_proprietario_ativar.html',investor=investor,tenant=tenant,erro=None,email_sugerido=email,token=token)
  if senha!=confirmar:
   flash('A confirmação da senha não confere.','danger')
   return render_template('portal_proprietario_ativar.html',investor=investor,tenant=tenant,erro=None,email_sugerido=email,token=token)
  duplicate=InvestorAccess.query.filter(InvestorAccess.email==email)
  if access:
   duplicate=duplicate.filter(InvestorAccess.id!=access.id)
  if duplicate.first():
   flash('Este e-mail já está vinculado a outro acesso de proprietário. Escolha outro e-mail.','danger')
   return render_template('portal_proprietario_ativar.html',investor=investor,tenant=tenant,erro=None,email_sugerido=email,token=token)
  if access and reset_em_andamento:
   access.email=email
   access.senha=generate_password_hash(senha)
   access.ativo=True
  else:
   access=InvestorAccess(
    tenant_id=investor.tenant_id,investor_id=investor.id,email=email,
    senha=generate_password_hash(senha),ativo=True,
   )
   db.session.add(access)
  db.session.commit()
  session['owner_access_id']=access.id
  session['owner_investor_id']=access.investor_id
  session['owner_tenant_id']=access.tenant_id
  access.ultimo_acesso_em=datetime.utcnow(); db.session.commit()
  flash('Acesso criado com sucesso. Bem-vindo ao Portal do Proprietário!','success')
  return redirect(url_for('portal_proprietario'))
 return render_template('portal_proprietario_ativar.html',investor=investor,tenant=tenant,erro=None,email_sugerido=email_sugerido,token=token)

@app.route('/portal-proprietario/entrar',methods=['GET','POST'])
def portal_proprietario_entrar():
 if session.get('owner_access_id'):
  return redirect(url_for('portal_proprietario'))
 tenant_ref=request.args.get('tenant',type=int) or request.form.get('tenant',type=int)
 tenant_login=Tenant.query.get(tenant_ref) if tenant_ref else None
 if request.method=='POST':
  email=(request.form.get('email') or '').strip().lower()
  access=InvestorAccess.query.filter_by(email=email,ativo=True).first()
  if access and check_password_hash(access.senha,request.form.get('senha') or ''):
   investor=Investor.query.filter_by(id=access.investor_id,tenant_id=access.tenant_id).first()
   if investor:
    session['owner_access_id']=access.id
    session['owner_investor_id']=access.investor_id
    session['owner_tenant_id']=access.tenant_id
    access.ultimo_acesso_em=datetime.utcnow(); db.session.commit()
    return redirect(url_for('portal_proprietario'))
  flash('E-mail ou senha inválidos.','danger')
 return render_template('portal_proprietario_login.html',tenant_login=tenant_login)

@app.route('/portal-proprietario/sair')
def portal_proprietario_sair():
 tenant_ref=session.get('owner_tenant_id')
 session.pop('owner_access_id',None); session.pop('owner_investor_id',None); session.pop('owner_tenant_id',None)
 return redirect(url_for('portal_proprietario_entrar',tenant=tenant_ref)) if tenant_ref else redirect(url_for('portal_proprietario_entrar'))

@app.route('/portal-proprietario')
@owner_portal_required
def portal_proprietario():
 tenant_id=int(session['owner_tenant_id']); investor_id=int(session['owner_investor_id'])
 today=date.today()
 try: inicio=datetime.strptime(request.args.get('inicio') or f'{today.year}-01-01','%Y-%m-%d').date()
 except Exception: inicio=date(today.year,1,1)
 try: fim=datetime.strptime(request.args.get('fim') or today.isoformat(),'%Y-%m-%d').date()
 except Exception: fim=today
 if inicio>fim: inicio,fim=fim,inicio
 vehicle_id=request.args.get('vehicle_id',type=int)
 placa=(request.args.get('placa') or '').strip().upper()[:10]
 available_vehicles=Vehicle.query.filter_by(tenant_id=tenant_id,investor_id=investor_id).order_by(Vehicle.placa).all()
 if vehicle_id and not any(v.id==vehicle_id for v in available_vehicles):
  vehicle_id=None
 data=resumo_portal_proprietario(tenant_id,investor_id,inicio,fim,vehicle_id=vehicle_id,placa=placa)
 tenant=Tenant.query.get(tenant_id)
 html=render_template('portal_proprietario.html',tenant=tenant,inicio=inicio.isoformat(),fim=fim.isoformat(),vehicle_id=vehicle_id,placa=placa,available_vehicles=available_vehicles,**data)

 # Compatibilidade com o template já publicado: substitui apenas o rodapé da
 # tabela de rentabilidade, permitindo distribuir a correção somente no app.py.
 totals=data.get('totals') or {}
 def _fmt_total(valor):
  try:
   return brl(valor)
  except Exception:
   return '0,00'
 def _fmt_pct(valor):
  if valor is None: return '—'
  try: return f"{Decimal(str(valor)):.2f}%".replace('.',',')
  except Exception: return '—'
 def _fmt_payback(valor):
  if valor is None: return '—'
  try: return f"{Decimal(str(valor)):.1f} meses".replace('.',',')
  except Exception: return '—'

 footer_totalizadores=(
  '<tfoot><tr>'
  '<td colspan="3"><strong>Totais do período</strong></td>'
  f'<td class="money"><strong>R$ {_fmt_total(totals.get("receita"))}</strong></td>'
  f'<td class="money"><strong>R$ {_fmt_total(totals.get("repasse_pago"))}</strong></td>'
  f'<td class="money"><strong>R$ {_fmt_total(totals.get("custos"))}</strong></td>'
  f'<td class="money"><strong>R$ {_fmt_total(totals.get("resultado_real"))}</strong></td>'
  f'<td class="money"><strong>R$ {_fmt_total(totals.get("capital_proprio"))}</strong></td>'
  f'<td class="money"><strong>{_fmt_pct(totals.get("roi_periodo"))}</strong></td>'
  f'<td class="money"><strong>{_fmt_pct(totals.get("rentabilidade_mensal"))}</strong></td>'
  f'<td class="money"><strong>{_fmt_payback(totals.get("payback_meses"))}</strong></td>'
  '<td>—</td><td></td>'
  '</tr></tfoot>'
 )
 html=re.sub(r'<tfoot>.*?</tfoot>',footer_totalizadores,html,count=1,flags=re.S)
 return html


def _decimal_portal(valor):
 txt=(valor or '').strip().replace('R$','').replace(' ','')
 if not txt: return Decimal('0')
 if ',' in txt:
  txt=txt.replace('.','').replace(',','.')
 try: return Decimal(txt)
 except Exception: raise ValueError('Valor inválido')

@app.route('/portal-proprietario/veiculos/<int:vehicle_id>/investimento',methods=['GET','POST'])
@owner_portal_required
def portal_proprietario_investimento(vehicle_id):
 tenant_id=int(session['owner_tenant_id']); investor_id=int(session['owner_investor_id'])
 vehicle=Vehicle.query.filter_by(id=vehicle_id,tenant_id=tenant_id,investor_id=investor_id).first_or_404()
 investment=VehicleInvestment.query.filter_by(tenant_id=tenant_id,investor_id=investor_id,vehicle_id=vehicle.id).first()
 if request.method=='POST':
  try:
   data_txt=(request.form.get('data_aquisicao') or '').strip()
   data_aquisicao=datetime.strptime(data_txt,'%Y-%m-%d').date() if data_txt else None
   vals={
    'valor_aquisicao':_decimal_portal(request.form.get('valor_aquisicao')),
    'capital_proprio':_decimal_portal(request.form.get('capital_proprio')),
    'valor_financiado':_decimal_portal(request.form.get('valor_financiado')),
    'saldo_devedor':_decimal_portal(request.form.get('saldo_devedor')),
    'valor_mercado':_decimal_portal(request.form.get('valor_mercado')),
   }
   if any(v<0 for v in vals.values()): raise ValueError('Os valores não podem ser negativos.')
  except Exception as exc:
   flash(str(exc) if str(exc) else 'Confira os valores informados.','danger')
   return redirect(url_for('portal_proprietario_investimento',vehicle_id=vehicle.id))
  if not investment:
   investment=VehicleInvestment(tenant_id=tenant_id,investor_id=investor_id,vehicle_id=vehicle.id)
   db.session.add(investment)
  investment.data_aquisicao=data_aquisicao
  for k,v in vals.items(): setattr(investment,k,v)
  investment.atualizado_em=datetime.utcnow()
  db.session.flush()
  db.session.add(VehicleInvestmentHistory(
   tenant_id=tenant_id,investor_id=investor_id,vehicle_id=vehicle.id,
   data_aquisicao=data_aquisicao,valor_aquisicao=vals['valor_aquisicao'],capital_proprio=vals['capital_proprio'],
   valor_financiado=vals['valor_financiado'],saldo_devedor=vals['saldo_devedor'],valor_mercado=vals['valor_mercado']
  ))
  db.session.commit()
  flash('Dados de investimento atualizados. O Frota Fácil já usará esses valores nas análises.','success')
  return redirect(url_for('portal_proprietario',vehicle_id=vehicle.id))
 history=VehicleInvestmentHistory.query.filter_by(tenant_id=tenant_id,investor_id=investor_id,vehicle_id=vehicle.id).order_by(VehicleInvestmentHistory.registrado_em.desc()).limit(20).all()
 tenant=Tenant.query.get(tenant_id)
 return render_template('portal_proprietario_investimento.html',tenant=tenant,vehicle=vehicle,investment=investment,history=history)

@app.route('/portal-proprietario/veiculos/<int:vehicle_id>/relatorio')
@owner_portal_required
def portal_proprietario_relatorio_veiculo(vehicle_id):
 tenant_id=int(session['owner_tenant_id']); investor_id=int(session['owner_investor_id'])
 vehicle=Vehicle.query.filter_by(id=vehicle_id,tenant_id=tenant_id,investor_id=investor_id).first_or_404()
 today=date.today()
 try: inicio=datetime.strptime(request.args.get('inicio') or f'{today.year}-01-01','%Y-%m-%d').date()
 except Exception: inicio=date(today.year,1,1)
 try: fim=datetime.strptime(request.args.get('fim') or today.isoformat(),'%Y-%m-%d').date()
 except Exception: fim=today
 if inicio>fim: inicio,fim=fim,inicio
 # Calcula o relatório do veículo dentro do conjunto completo do proprietário,
 # para preservar a posição real dele no ranking do período.
 data=resumo_portal_proprietario(tenant_id,investor_id,inicio,fim)
 row=next((r for r in data['vehicles'] if r['vehicle'].id==vehicle.id),None)
 tenant=Tenant.query.get(tenant_id); investor=Investor.query.filter_by(id=investor_id,tenant_id=tenant_id).first_or_404()
 return render_template('portal_proprietario_relatorio.html',tenant=tenant,investor=investor,vehicle=vehicle,row=row,inicio=inicio,fim=fim,gerado_em=agora_sao_paulo_naive())


def _veiculo_portal_proprietario(vehicle_id):
 tenant_id=int(session['owner_tenant_id'])
 investor_id=int(session['owner_investor_id'])
 return Vehicle.query.options(
  joinedload(Vehicle.current_driver),joinedload(Vehicle.current_contract)
 ).filter_by(id=vehicle_id,tenant_id=tenant_id,investor_id=investor_id).first_or_404()

def _ajustar_eventos_vistoria_para_status_atual(eventos, vistorias):
 """Atualiza apenas a exibição de eventos antigos de vistoria para não manter 'Aguardando aprovação' após conclusão.
 Também cobre registros já existentes, sem exigir alteração manual no banco.
 """
 for evento in eventos:
  if (evento.evento or '').strip() not in ('Vistoria em vídeo recebida','Vistoria fotográfica recebida'):
   continue
  descricao=evento.descricao or ''
  if 'Aguardando aprovação' not in descricao:
   continue
  candidatos=[]
  for vistoria in vistorias:
   if vistoria.vehicle_id!=evento.vehicle_id:
    continue
   if evento.contract_id and vistoria.contract_id and evento.contract_id!=vistoria.contract_id:
    continue
   if evento.driver_id and vistoria.driver_id and evento.driver_id!=vistoria.driver_id:
    continue
   referencia=vistoria.submitted_at or vistoria.requested_at
   if referencia and evento.criado_em:
    distancia=abs((referencia-evento.criado_em).total_seconds())
    if distancia>86400:
     continue
   else:
    distancia=0
   candidatos.append((distancia,vistoria.id,vistoria))
  if not candidatos:
   continue
  vistoria=min(candidatos,key=lambda x:(x[0],-x[1]))[2]
  status=(vistoria.status or '').strip().lower()
  if status in ('concluída','concluida','recebida'):
   evento.descricao=descricao.replace('Aguardando aprovação.','Vistoria concluída com o envio.').replace('Aguardando aprovação','Vistoria concluída com o envio')
  elif status=='aprovada':
   evento.descricao=descricao.replace('Aguardando aprovação.','Posteriormente aprovada pela locadora.').replace('Aguardando aprovação','Posteriormente aprovada pela locadora')
  elif status in ('regravar','rejeitada','rejeitado'):
   evento.descricao=descricao.replace('Aguardando aprovação.','Nova gravação solicitada pela locadora.').replace('Aguardando aprovação','Nova gravação solicitada pela locadora')
 return eventos

def _dados_historico_portal_proprietario(vehicle):
 """Histórico visível ao proprietário, sempre restrito ao tenant e ao veículo dele."""
 tenant_id=int(session['owner_tenant_id'])
 eventos=VehicleEvent.query.options(
  joinedload(VehicleEvent.contract),joinedload(VehicleEvent.driver)
 ).filter_by(tenant_id=tenant_id,vehicle_id=vehicle.id).order_by(VehicleEvent.criado_em.desc()).all()
 contratos=Contract.query.options(joinedload(Contract.driver)).filter_by(
  tenant_id=tenant_id,vehicle_id=vehicle.id
 ).order_by(Contract.id.desc()).all()
 odometros=Odometer.query.filter_by(
  tenant_id=tenant_id,vehicle_id=vehicle.id
 ).order_by(Odometer.data.desc(),Odometer.id.desc()).all()
 manutencoes=Maintenance.query.filter_by(
  tenant_id=tenant_id,vehicle_id=vehicle.id
 ).order_by(Maintenance.id.desc()).all()
 vistorias=Inspection.query.options(joinedload(Inspection.driver),joinedload(Inspection.contract)).filter_by(
  tenant_id=tenant_id,vehicle_id=vehicle.id
 ).order_by(Inspection.requested_at.desc(),Inspection.id.desc()).all()
 _ajustar_eventos_vistoria_para_status_atual(eventos,vistorias)
 documentos=Document.query.filter_by(
  tenant_id=tenant_id,entidade='Veículo',entidade_id=vehicle.id,status='Ativo'
 ).order_by(Document.criado_em.desc(),Document.id.desc()).all()

 # Linha do tempo unificada, contendo apenas informações adequadas ao proprietário.
 timeline=[]
 for x in eventos:
  timeline.append({
   'data':x.criado_em,'tipo':'Ocorrência','titulo':x.evento or 'Evento do veículo',
   'descricao':x.descricao or '', 'km':None
  })
 for x in odometros:
  timeline.append({
   'data':x.data,'tipo':'Quilometragem','titulo':f'{x.km:,} km'.replace(',','.'),
   'descricao':x.origem or 'Leitura registrada','km':x.km
  })
 maintenance_ids=[x.id for x in manutencoes]
 documentos_manutencao={}
 if maintenance_ids:
  docs_manut=Document.query.filter(Document.tenant_id==tenant_id,Document.entidade=='Manutenção',Document.entidade_id.in_(maintenance_ids),Document.status=='Ativo').order_by(Document.criado_em.desc(),Document.id.desc()).all()
  for doc in docs_manut: documentos_manutencao.setdefault(doc.entidade_id,[]).append(doc)
 for x in manutencoes:
  try: dt=datetime.strptime(x.data,'%Y-%m-%d') if x.data else None
  except Exception: dt=None
  timeline.append({
   'data':dt,'tipo':'Manutenção','titulo':x.tipo or 'Manutenção',
   'descricao':(' · '.join([p for p in [
    f'KM: {x.km:,}'.replace(',','.') if x.km is not None else None,
    f'Oficina: {x.oficina}' if x.oficina else None,
    f'Custo: R$ {brl(x.custo)}' if x.custo is not None else None,
    x.observacoes or None
   ] if p])), 'km':x.km,'maintenance_id':x.id
  })
 for x in contratos:
  try: dt=datetime.strptime(x.data_inicio,'%Y-%m-%d') if x.data_inicio else x.criado_em
  except Exception: dt=x.criado_em
  timeline.append({
   'data':dt,'tipo':'Contrato','titulo':x.numero_contrato or f'Contrato #{x.id}',
   'descricao':f"Motorista: {x.driver.nome if x.driver else '-'} · Status: {x.status or '-'}",
   'km':None
  })
 for x in vistorias:
  timeline.append({
   'data':x.submitted_at or x.requested_at,'tipo':'Vistoria','titulo':'Vistoria do veículo',
   'descricao':f"Status: {x.status or '-'}" + (f" · Motorista: {x.driver.nome}" if x.driver else ''),
   'km':None
  })
 for x in documentos:
  timeline.append({
   'data':x.criado_em,'tipo':'Documento','titulo':x.tipo or 'Documento',
   'descricao':x.nome_original or x.identificador or '', 'km':None
  })
 timeline.sort(key=lambda item:item['data'] or datetime.min,reverse=True)
 return {
  'eventos':eventos,'contratos':contratos,'odometros':odometros,'manutencoes':manutencoes,
  'vistorias':vistorias,'documentos':documentos,'timeline':timeline,
  'documentos_manutencao':documentos_manutencao
 }

@app.route('/portal-proprietario/veiculos/<int:vehicle_id>/historico')
@owner_portal_required
def portal_proprietario_historico_veiculo(vehicle_id):
 vehicle=_veiculo_portal_proprietario(vehicle_id)
 tenant=Tenant.query.get(int(session['owner_tenant_id']))
 investor=Investor.query.filter_by(
  id=int(session['owner_investor_id']),tenant_id=int(session['owner_tenant_id'])
 ).first_or_404()
 data=_dados_historico_portal_proprietario(vehicle)
 return render_template(
  'portal_proprietario_historico.html',
  tenant=tenant,investor=investor,vehicle=vehicle,**data
 )

@app.route('/portal-proprietario/manutencoes/<int:maintenance_id>/documentos/<int:document_id>')
@owner_portal_required
def portal_proprietario_documento_manutencao(maintenance_id,document_id):
 tenant_id=int(session['owner_tenant_id']); investor_id=int(session['owner_investor_id'])
 m=Maintenance.query.filter_by(id=maintenance_id,tenant_id=tenant_id).first_or_404()
 Vehicle.query.filter_by(id=m.vehicle_id,tenant_id=tenant_id,investor_id=investor_id).first_or_404()
 doc=Document.query.filter_by(id=document_id,tenant_id=tenant_id,entidade='Manutenção',entidade_id=m.id,status='Ativo').first_or_404()
 try: conteudo=storage.download(doc.arquivo)
 except StorageNotFoundError: abort(404)
 except Exception:
  app.logger.exception('Falha ao abrir comprovante de manutenção no portal %s',doc.id); abort(503)
 return send_file(BytesIO(conteudo),as_attachment=False,download_name=doc.nome_original,mimetype=_mimetype_documento(doc.nome_original))


@app.route('/portal-proprietario/veiculos/<int:vehicle_id>/historico/imprimir')
@owner_portal_required
def portal_proprietario_imprimir_historico_veiculo(vehicle_id):
 vehicle=_veiculo_portal_proprietario(vehicle_id)
 tenant=Tenant.query.get(int(session['owner_tenant_id']))
 investor=Investor.query.filter_by(
  id=int(session['owner_investor_id']),tenant_id=int(session['owner_tenant_id'])
 ).first_or_404()
 data=_dados_historico_portal_proprietario(vehicle)
 return render_template(
  'portal_proprietario_historico_impressao.html',
  tenant=tenant,investor=investor,vehicle=vehicle,
  gerado_em=agora_sao_paulo_naive(),**data
 )


@app.route('/health')
def health():
 try:
  db.session.execute(text('SELECT 1'))
  return {'status':'ok','application':'ok','database':'ok','timestamp':datetime.now(timezone.utc).isoformat()},200
 except Exception:
  db.session.rollback()
  app.logger.exception('Healthcheck: banco indisponível')
  return {'status':'degraded','application':'ok','database':'error','timestamp':datetime.now(timezone.utc).isoformat()},503


def normalizar_identificador_veiculo(valor):
 """Normaliza placa/RENAVAM/chassi para comparação de duplicidade."""
 return re.sub(r'[^A-Za-z0-9]','',str(valor or '')).upper().strip()

def localizar_veiculo_duplicado(tenant_id, *, placa=None, renavam=None, chassi=None, excluir_id=None):
 """Bloqueia duplicidade dentro do tenant por placa, RENAVAM ou chassi.

 A comparação é normalizada em Python para também encontrar registros antigos
 gravados com pontos, hífens, espaços ou diferenças de maiúsculas/minúsculas.
 """
 identificadores={
  'placa':normalizar_identificador_veiculo(placa),
  'renavam':normalizar_identificador_veiculo(renavam),
  'chassi':normalizar_identificador_veiculo(chassi),
 }
 query=Vehicle.query.filter_by(tenant_id=tenant_id)
 if excluir_id is not None:
  query=query.filter(Vehicle.id!=excluir_id)
 for existente in query.all():
  atuais={
   'placa':normalizar_identificador_veiculo(existente.placa),
   'renavam':normalizar_identificador_veiculo(existente.renavam),
   'chassi':normalizar_identificador_veiculo(existente.chassi),
  }
  for campo,valor in identificadores.items():
   if valor and atuais.get(campo)==valor:
    return campo,existente
 return None,None

def mensagem_duplicidade_veiculo(campo, existente):
 rotulos={'placa':'placa','renavam':'RENAVAM','chassi':'chassi'}
 identificacao=(existente.placa or existente.marca_modelo or f'ID {existente.id}')
 return f'Este veículo já está cadastrado nesta locadora: {rotulos.get(campo,campo)} já pertence ao veículo {identificacao}.'

@app.route('/veiculos',methods=['GET','POST'])
@login_required
def veiculos():
 if request.method=='POST':
  campos_veiculo=['placa','renavam','chassi','marca_modelo','ano_fabricacao','ano_modelo','cor','combustivel','motorizacao','status','proprietario_legal','cpf_cnpj_proprietario','rastreador_id']
  vals={k:limpar_campo_ocr_veiculo(k,request.form.get(k)) for k in campos_veiculo}
  # Canonicaliza os identificadores antes de validar e gravar.
  vals['placa']=normalizar_identificador_veiculo(vals.get('placa'))
  vals['renavam']=normalizar_identificador_veiculo(vals.get('renavam')) or None
  vals['chassi']=normalizar_identificador_veiculo(vals.get('chassi')) or None
  campo_dup,veiculo_dup=localizar_veiculo_duplicado(tid(),placa=vals.get('placa'),renavam=vals.get('renavam'),chassi=vals.get('chassi'))
  if veiculo_dup:
   flash(mensagem_duplicidade_veiculo(campo_dup,veiculo_dup),'danger')
   return redirect(url_for('veiculos'))
  v=Vehicle(tenant_id=tid(),**vals,km_atual=int(request.form.get('km_atual') or 0),investor_id=request.form.get('investor_id') or None,valor_repasse=request.form.get('valor_repasse') or 0,limite_km=request.form.get('limite_km') or None,valor_km_excedente=request.form.get('valor_km_excedente') or 0,controlar_oleo=bool(request.form.get('controlar_oleo')),ultima_troca_oleo_km=request.form.get('ultima_troca_oleo_km') or None,intervalo_oleo_km=request.form.get('intervalo_oleo_km') or 10000,alerta_oleo_km=request.form.get('alerta_oleo_km') or 100)
  # RC 1.0.15: o cadastro antes chamado Investidor passa a ser a fonte do Proprietário do veículo.
  if v.investor_id:
   proprietario=Investor.query.filter_by(id=v.investor_id,tenant_id=tid()).first()
   if proprietario:
    v.proprietario_legal=proprietario.nome
    v.cpf_cnpj_proprietario=proprietario.cpf_cnpj
  db.session.add(v)
  try:
   db.session.flush()
   db.session.add(Odometer(tenant_id=tid(),vehicle_id=v.id,km=v.km_atual,origem='Cadastro'))
   armazenar_documento_cadastro(
    tipo='CRLV',
    entidade='Veículo',
    entidade_id=v.id,
    referencia=v.placa,
    numero_documento=v.renavam,
    nome_original=request.form.get('_documento_nome'),
    temp_key=request.form.get('_documento_temp_key'),
    mimetype=request.form.get('_documento_mimetype'),
    ano=v.ano_modelo,
   )
   db.session.commit()
  except Exception:
   db.session.rollback()
   app.logger.exception('Falha ao cadastrar veículo e armazenar CRLV')
   flash('Não foi possível concluir o cadastro e armazenar o CRLV. Tente novamente.','danger')
   return redirect(url_for('veiculos'))
  flash('Veículo cadastrado e CRLV armazenado automaticamente.','success')
  return redirect(url_for('veiculos'))
 q=(request.args.get('q') or '').strip()
 proprietario_filtro=(request.args.get('proprietario') or '').strip()
 query=Vehicle.query.filter_by(tenant_id=tid())
 if q:
  like=f'%{q}%'
  query=query.filter(or_(Vehicle.placa.ilike(like),Vehicle.renavam.ilike(like),Vehicle.chassi.ilike(like),Vehicle.marca_modelo.ilike(like),Vehicle.proprietario_legal.ilike(like)))
 if proprietario_filtro=='sem':
  query=query.filter(Vehicle.investor_id.is_(None))
 elif proprietario_filtro:
  try:
   proprietario_id=int(proprietario_filtro)
  except (TypeError,ValueError):
   proprietario_id=None
  if proprietario_id and Investor.query.filter_by(id=proprietario_id,tenant_id=tid()).first():
   query=query.filter(Vehicle.investor_id==proprietario_id)
  else:
   proprietario_filtro=''
 investidores=Investor.query.filter_by(tenant_id=tid()).order_by(Investor.nome).all()
 return render_template('veiculos.html',items=query.order_by(Vehicle.placa).all(),investidores=investidores,motoristas=Driver.query.filter_by(tenant_id=tid(),status='Ativo').order_by(Driver.nome).all(),oil_status=oil_status,q=q,proprietario_filtro=proprietario_filtro)
@app.route('/veiculos/importar',methods=['POST'])
@login_required
def importar_veiculo():
 f=request.files.get('arquivo')
 if not f or not f.filename:
  flash('Selecione um arquivo.','danger')
  return redirect(url_for('veiculos'))
 nome_original=secure_filename(f.filename or 'crlv')
 mimetype=f.mimetype or 'application/octet-stream'
 conteudo=f.read()
 temp_key=f'{tid()}/temporarios/{uuid.uuid4().hex}_{nome_original}'
 try:
  storage.upload(BytesIO(conteudo),temp_key,mimetype)
  arquivo_ocr=FileStorage(stream=BytesIO(conteudo),filename=nome_original,content_type=mimetype)
  texto_ocr=extract_text(arquivo_ocr)
  dados=parse_crlv(texto_ocr)
  dados=corrigir_dados_crlv_ocr(texto_ocr,dados)
  return render_template('confirmar_veiculo.html',dados=dados,investidores=Investor.query.filter_by(tenant_id=tid()).all(),documento_temp_key=temp_key,documento_nome=nome_original,documento_mimetype=mimetype)
 except Exception as exc:
  try: storage.delete(temp_key)
  except Exception: pass
  app.logger.exception('Falha ao processar CRLV')
  flash(f'Não foi possível processar o CRLV: {exc}','danger')
  return redirect(url_for('veiculos'))

@app.route('/veiculos/<int:id>/editar',methods=['GET','POST'])
@login_required
def editar_veiculo(id):
 v=Vehicle.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 if request.method=='POST':
  novos={campo:request.form.get(campo) for campo in ['placa','renavam','chassi','marca_modelo','ano_fabricacao','ano_modelo','cor','combustivel','motorizacao','status','proprietario_legal','cpf_cnpj_proprietario','rastreador_id']}
  novos['placa']=normalizar_identificador_veiculo(novos.get('placa'))
  novos['renavam']=normalizar_identificador_veiculo(novos.get('renavam')) or None
  novos['chassi']=normalizar_identificador_veiculo(novos.get('chassi')) or None
  campo_dup,veiculo_dup=localizar_veiculo_duplicado(tid(),placa=novos.get('placa'),renavam=novos.get('renavam'),chassi=novos.get('chassi'),excluir_id=v.id)
  if veiculo_dup:
   flash(mensagem_duplicidade_veiculo(campo_dup,veiculo_dup),'danger')
   return redirect(url_for('editar_veiculo',id=v.id))
  for campo,valor in novos.items():
   setattr(v,campo,valor)
  v.investor_id=request.form.get('investor_id') or None
  if v.investor_id:
   proprietario=Investor.query.filter_by(id=v.investor_id,tenant_id=tid()).first()
   if proprietario:
    v.proprietario_legal=proprietario.nome
    v.cpf_cnpj_proprietario=proprietario.cpf_cnpj
  v.valor_repasse=request.form.get('valor_repasse') or 0
  v.limite_km=request.form.get('limite_km') or None
  v.valor_km_excedente=request.form.get('valor_km_excedente') or 0
  v.controlar_oleo=bool(request.form.get('controlar_oleo'))
  v.ultima_troca_oleo_km=request.form.get('ultima_troca_oleo_km') or None
  v.intervalo_oleo_km=request.form.get('intervalo_oleo_km') or 10000
  v.alerta_oleo_km=request.form.get('alerta_oleo_km') or 100
  db.session.commit()
  flash('Veículo atualizado com sucesso.','success')
  return redirect(url_for('veiculos'))
 return render_template('editar_veiculo.html',v=v,investidores=Investor.query.filter_by(tenant_id=tid()).order_by(Investor.nome).all())

@app.route('/veiculos/<int:id>/excluir',methods=['POST'])
@login_required
def excluir_veiculo(id):
 v=Vehicle.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 if v.status in ('Reservado','Alugado','Devolução') or Contract.query.filter(Contract.tenant_id==tid(),Contract.vehicle_id==v.id,Contract.status.in_(['Rascunho','Gerado','Enviado','Visualizado','Assinado','Ativo'])).first():
  flash('Não é possível excluir um veículo reservado, alugado ou com contrato vigente.','danger')
  return redirect(url_for('veiculos'))
 contratos=Contract.query.filter_by(tenant_id=tid(),vehicle_id=v.id).count()
 if contratos:
  flash('Este veículo possui histórico contratual e não pode ser excluído definitivamente. Mesmo que os contratos estejam cancelados ou encerrados, altere o status para Vendido ou Inativo para preservar a rastreabilidade.','danger')
  return redirect(url_for('veiculos'))

 # Regras de repasse são configurações dependentes do veículo.
 # Se o veículo puder ser excluído (sem histórico contratual), removemos essas regras
 # antes do DELETE do veículo para respeitar a chave estrangeira.
 InvestorVehicleRule.query.filter_by(tenant_id=tid(),vehicle_id=v.id).delete(synchronize_session=False)

 # Se já houver histórico operacional que deve ser preservado, não removemos o veículo.
 if Inspection.query.filter_by(tenant_id=tid(),vehicle_id=v.id).first() or VehicleEvent.query.filter_by(tenant_id=tid(),vehicle_id=v.id).first():
  db.session.rollback()
  flash('Este veículo possui histórico operacional e não pode ser excluído definitivamente. Altere o status para Vendido ou Inativo para preservar a rastreabilidade.','danger')
  return redirect(url_for('veiculos'))

 # Remove as fotos de quilometragem antes dos registros.
 # Fotos novas ficam no Cloudflare R2; nomes antigos sem "/" continuam
 # compatíveis com o armazenamento local usado anteriormente.
 solicitacoes=MileageRequest.query.filter_by(tenant_id=tid(),vehicle_id=v.id).all()
 for solicitacao in solicitacoes:
  if not solicitacao.photo:
   continue
  try:
   if '/' in solicitacao.photo:
    storage.delete(solicitacao.photo)
   else:
    foto=UPLOAD/str(solicitacao.tenant_id)/'odometros'/solicitacao.photo
    if foto.exists():
     foto.unlink()
  except Exception:
   app.logger.warning('Não foi possível remover a foto %s',solicitacao.photo)

 # Remove dados operacionais dependentes do veículo.
 MileageRequest.query.filter_by(tenant_id=tid(),vehicle_id=v.id).delete(synchronize_session=False)
 Odometer.query.filter_by(tenant_id=tid(),vehicle_id=v.id).delete(synchronize_session=False)
 Maintenance.query.filter_by(tenant_id=tid(),vehicle_id=v.id).delete(synchronize_session=False)
 documentos=Document.query.filter_by(tenant_id=tid(),entidade='veiculo',entidade_id=v.id).all()
 for documento in documentos:
  if documento.arquivo:
   try:
    storage.delete(documento.arquivo)
   except Exception:
    app.logger.warning('Não foi possível remover o documento %s',documento.arquivo)
  db.session.delete(documento)
 db.session.delete(v)
 db.session.commit()
 flash('Veículo excluído com sucesso.','success')
 return redirect(url_for('veiculos'))

@app.route('/veiculos/<int:id>/km',methods=['POST'])
@login_required
def atualizar_km(id):
 v=Vehicle.query.filter_by(id=id,tenant_id=tid()).first_or_404(); km=int(request.form['km']); v.km_atual=km; db.session.add(Odometer(tenant_id=tid(),vehicle_id=v.id,km=km,origem=request.form.get('origem','Manual'))); db.session.commit(); recalcular_alertas(tid()); flash('Quilometragem atualizada e alertas recalculados.','success'); return redirect(url_for('veiculos'))

@app.route('/veiculos/<int:id>/oleo',methods=['POST'])
@login_required
def configurar_oleo(id):
 v=Vehicle.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 v.controlar_oleo=bool(request.form.get('controlar_oleo'))
 v.ultima_troca_oleo_km=int(request.form['ultima_troca_oleo_km']) if request.form.get('ultima_troca_oleo_km') else None
 v.intervalo_oleo_km=int(request.form.get('intervalo_oleo_km') or 10000)
 v.alerta_oleo_km=int(request.form.get('alerta_oleo_km') or 100)
 db.session.commit(); flash('Plano de troca de óleo atualizado.','success'); return redirect(url_for('veiculos'))

@app.route('/veiculos/<int:id>/solicitar-km',methods=['POST'])
@login_required
def solicitar_km(id):
 v=Vehicle.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 d=motorista_atual_veiculo(v)
 # Seleção manual permanece como fallback para veículos sem contrato vigente.
 if not d and request.form.get('driver_id'):
  d=Driver.query.filter_by(id=request.form.get('driver_id'),tenant_id=tid()).first()
 if not d:
  flash('Não encontrei motorista vinculado a contrato vigente deste veículo. Selecione um motorista manualmente.','warning')
  return redirect(url_for('veiculos'))
 telefone=normalize_phone(d.telefone)
 if not telefone:
  flash('Cadastre um telefone/WhatsApp válido para o motorista.','danger'); return redirect(url_for('veiculos'))
 req=active_request(v.id,d.id)
 if not req:
  req=MileageRequest(tenant_id=tid(),vehicle_id=v.id,driver_id=d.id,token=uuid.uuid4().hex+uuid.uuid4().hex,expires_at=datetime.utcnow()+timedelta(days=7),previous_km=v.km_atual)
  db.session.add(req); db.session.commit()
 link=url_for('registrar_quilometragem_publica',token=req.token,_external=True)
 mensagem=f'Olá, {d.nome}! Precisamos da quilometragem atual do veículo {v.placa}. Abra o link, tire uma foto do painel e informe o km: {link}'
 integration=Integration.query.filter_by(tenant_id=tid(),tipo='whatsapp').first()
 cfg=CommunicationService.parse_config(integration)
 template_name=(cfg.get('mileage_template_name') or '').strip() or None
 template_params=[d.nome,v.marca_modelo or 'Veículo',v.placa,link]
 provider_cfg=(cfg.get('provider') or 'web').lower()
 fila=MessageQueue(
  tenant_id=tid(),channel='whatsapp',provider='whatsapp_business' if provider_cfg=='business' else 'whatsapp_web',recipient=telefone,
  recipient_name=d.nome,message_type='solicitacao_km',body=mensagem,template_name=template_name,template_parameters=json.dumps(template_params,ensure_ascii=False),
  related_entity='Veiculo',related_entity_id=v.id,status='PENDENTE',
  created_at=agora_sao_paulo_naive(),updated_at=agora_sao_paulo_naive(),
 )
 db.session.add(fila); db.session.flush()
 try:
  result=CommunicationService().send_whatsapp(phone=telefone,message=mensagem,integration=integration,template_name=template_name,template_language=cfg.get('template_language') or 'pt_BR',template_parameters=template_params)
  fila.provider=result.provider; fila.status=result.status; fila.external_id=result.external_id
  fila.attempts=(fila.attempts or 0)+1; fila.sent_at=agora_sao_paulo_naive() if result.status=='ENVIADA' else None
  db.session.add(MessageEvent(tenant_id=tid(),message_id=fila.id,event=result.status,description='Solicitação de quilometragem processada pelo provedor configurado.',created_at=agora_sao_paulo_naive()))
  db.session.commit()
  if result.redirect_url:
   return redirect(result.redirect_url)
  flash('Solicitação de quilometragem enviada automaticamente pela WhatsApp Business Platform.','success')
  return redirect(url_for('veiculos'))
 except CommunicationError as exc:
  fila.status='FALHA'; fila.error_message=str(exc); fila.attempts=(fila.attempts or 0)+1
  db.session.add(MessageEvent(tenant_id=tid(),message_id=fila.id,event='FALHA',description=str(exc),created_at=agora_sao_paulo_naive()))
  db.session.commit(); flash(str(exc),'danger')
  return redirect(url_for('veiculos'))

@app.route('/km/<token>',methods=['GET','POST'])
def registrar_quilometragem_publica(token):
 req=MileageRequest.query.options(joinedload(MileageRequest.vehicle),joinedload(MileageRequest.driver)).filter_by(token=token).first_or_404()
 if req.status in ('Concluído','Aguardando conferência'): return render_template('quilometragem_sucesso.html',req=req,ja_enviado=True)
 if req.expires_at and req.expires_at<datetime.utcnow():
  return render_template('quilometragem_publica.html',req=req,expirado=True),410
 if request.method=='POST':
  try: km=int(request.form.get('km',''))
  except ValueError:
   flash('Informe uma quilometragem válida.','danger'); return render_template('quilometragem_publica.html',req=req,expirado=False)
  if km < (req.vehicle.km_atual or 0):
   flash(f'A quilometragem não pode ser menor que a última leitura ({req.vehicle.km_atual:,} km).','danger'); return render_template('quilometragem_publica.html',req=req,expirado=False)
  foto=request.files.get('foto')
  if not foto or not foto.filename:
   flash('A foto do painel é obrigatória.','danger'); return render_template('quilometragem_publica.html',req=req,expirado=False)
  ext=Path(secure_filename(foto.filename)).suffix.lower()
  if ext not in ('.jpg','.jpeg','.png','.webp'):
   flash('Envie uma foto JPG, PNG ou WEBP.','danger'); return render_template('quilometragem_publica.html',req=req,expirado=False)
  chave=f'{req.tenant_id}/odometros/{uuid.uuid4().hex}{ext}'
  try:
   storage.upload(foto.stream,chave,foto.mimetype)
   req.km=km
   req.photo=chave
   req.notes=request.form.get('observacoes')
   req.submitted_at=datetime.utcnow()
   if req.vehicle and req.vehicle.id:
    tenant=Tenant.query.get(req.tenant_id)
   else:
    tenant=None
   if tenant and tenant.conferir_km_motorista:
    req.status='Aguardando conferência'
   else:
    req.status='Concluído'
    req.vehicle.km_atual=km
    db.session.add(Odometer(tenant_id=req.tenant_id,vehicle_id=req.vehicle_id,km=km,origem='Motorista via link'))
   db.session.commit()
  except Exception:
   db.session.rollback()
   app.logger.exception('Falha ao armazenar foto do painel')
   try:
    storage.delete(chave)
   except Exception:
    pass
   flash('Não foi possível armazenar a foto do painel. Tente novamente.','danger')
   return render_template('quilometragem_publica.html',req=req,expirado=False)
  recalcular_alertas(req.tenant_id)
  return redirect(url_for('registrar_quilometragem_publica',token=token))
 return render_template('quilometragem_publica.html',req=req,expirado=False)

def _frota_admin_atual():
 admin_id=session.get('frota_admin_id')
 if not admin_id:
  return None
 try:
  return FrotaAdmin.query.filter_by(id=int(admin_id),ativo=True).first()
 except Exception:
  return None

def frota_admin_required(view):
 @wraps(view)
 def wrapped(*args,**kwargs):
  admin=_frota_admin_atual()
  if not admin:
   session.pop('frota_admin_id',None)
   return redirect(url_for('frota_admin_entrar',next=request.path))
  return view(*args,**kwargs)
 return wrapped

@app.route('/admin-frota/entrar',methods=['GET','POST'])
def frota_admin_entrar():
 if _frota_admin_atual():
  return redirect(url_for('frota_admin_suporte'))
 if request.method=='POST':
  email=(request.form.get('email') or '').strip().lower()
  senha=request.form.get('senha') or ''
  admin=FrotaAdmin.query.filter_by(email=email,ativo=True).first()
  if admin and check_password_hash(admin.senha,senha):
   session['frota_admin_id']=admin.id
   admin.ultimo_acesso_em=datetime.utcnow(); db.session.commit()
   return redirect(url_for('frota_admin_suporte'))
  flash('E-mail ou senha administrativos inválidos.','danger')
 return render_template('admin_frota_login.html')

@app.route('/admin-frota/sair')
def frota_admin_sair():
 session.pop('frota_admin_id',None)
 flash('Sessão administrativa encerrada.','success')
 return redirect(url_for('frota_admin_entrar'))

@app.route('/ajuda')
@login_required
def ajuda():
 abertos=SupportTicket.query.filter_by(tenant_id=tid()).filter(SupportTicket.status.in_(['ABERTO','EM_ANALISE'])).count()
 return render_template('ajuda.html',abertos=abertos,support_admin=False)

_HELP_STOPWORDS={
 'a','ao','aos','as','com','como','da','das','de','do','dos','e','em','eu','faz','fazer','me','na','nas','no','nos',
 'o','os','ou','para','por','que','se','sem','um','uma','voce'
}

def _help_normalize(value):
 value=unicodedata.normalize('NFKD',value or '')
 return ''.join(c for c in value if not unicodedata.combining(c)).lower()

def _help_terms(value):
 return {t for t in re.findall(r'[a-z0-9]+',_help_normalize(value)) if len(t)>2 and t not in _HELP_STOPWORDS}

def _help_read_static(filename):
 path=Path(app.static_folder)/filename
 try:
  return path.read_text(encoding='utf-8').strip()
 except (OSError,UnicodeError) as exc:
  print('Assistente de ajuda - arquivo indisponivel:',filename,repr(exc))
  return ''

def _help_rank_blocks(text_value,question,limit=6):
 if not text_value: return []
 terms=_help_terms(question)
 blocks=[b.strip() for b in re.split(r'\n\s*\n+',text_value) if b.strip()]
 scored=[]
 for idx,block in enumerate(blocks):
  overlap=len(terms & _help_terms(block))
  if overlap: scored.append((overlap,idx,block))
 scored.sort(key=lambda item:(-item[0],item[1]))
 return [block for _,_,block in scored[:limit]]

def _help_context(question):
 knowledge=_help_read_static('Base_Conhecimento_IA_Frota_Facil.txt')
 manual=_help_read_static('Manual_Frota_Facil_Usuario.txt')
 if not knowledge:
  return ''
 rules=knowledge.split('[ACESSO E NAVEGAÇÃO]',1)[0].strip()
 knowledge_hits=_help_rank_blocks(knowledge,question,8)
 manual_hits=_help_rank_blocks(manual,question,4)
 parts=[rules]
 if knowledge_hits: parts.append('TRECHOS PRIORITÁRIOS DA BASE:\n'+'\n\n'.join(knowledge_hits))
 if manual_hits: parts.append('TRECHOS COMPLEMENTARES DO MANUAL:\n'+'\n\n'.join(manual_hits))
 parts.append('INSTRUÇÃO FINAL: responda em português do Brasil, em texto simples, sem Markdown. Seja curto e prático. Se os trechos não sustentarem a resposta, encaminhe para Ajuda e Suporte > Suporte.')
 return '\n\n'.join(parts)

def _help_clean_answer(value):
 value=(value or '').replace('\\r\\n','\n').replace('\\n','\n')
 value=re.sub(r'\*\*(.*?)\*\*',r'\1',value,flags=re.S)
 value=re.sub(r'__(.*?)__',r'\1',value,flags=re.S)
 value=re.sub(r'(?m)^\s*#{1,6}\s*','',value)
 value=re.sub(r'(?m)^\s*[-*]\s+','• ',value)
 value=re.sub(r'\n{3,}','\n\n',value)
 return value.strip()

@app.route('/ajuda/assistente',methods=['GET','POST'])
@login_required
def ajuda_assistente():
 pergunta=''; resposta=''; erro=''
 if request.method=='POST':
  pergunta=(request.form.get('pergunta') or '').strip()[:1200]
  if not pergunta: erro='Digite uma pergunta sobre o Frota Fácil.'
  else:
   api_key=(os.getenv('OPENAI_API_KEY') or '').strip()
   if not api_key: erro='O assistente ainda não está configurado. Abra uma solicitação de suporte.'
   else:
    try:
     base=_help_context(pergunta)
     if not base: raise RuntimeError('BASE_AJUDA_INDISPONIVEL')
     payload={'model':(os.getenv('FROTA_FACIL_HELP_MODEL') or 'gpt-5.6-luna').strip(),'input':[{'role':'system','content':[{'type':'input_text','text':base}]},{'role':'user','content':[{'type':'input_text','text':pergunta}]}],'max_output_tokens':600}
     resp=requests.post('https://api.openai.com/v1/responses',headers={'Authorization':f'Bearer {api_key}','Content-Type':'application/json'},json=payload,timeout=45)
     if resp.status_code>=400:
      try:
       api_error=(resp.json().get('error') or {}).get('code') or (resp.json().get('error') or {}).get('type') or 'sem_codigo'
      except Exception: api_error='resposta_invalida'
      raise RuntimeError('HTTP %s %s' % (resp.status_code,api_error))
     partes=[]
     for out in resp.json().get('output',[]):
      for c in out.get('content',[]):
       if c.get('type')=='output_text' and c.get('text'): partes.append(c.get('text').strip())
     resposta=_help_clean_answer('\n'.join(partes))
     if not resposta: erro='Não consegui gerar uma resposta agora. Você pode abrir uma solicitação de suporte.'
    except Exception as exc:
     print('Assististente de ajuda:',repr(exc)); erro='O assistente está temporariamente indisponível. Você pode abrir uma solicitação de suporte.'
 return render_template('ajuda_assistente.html',pergunta=pergunta,resposta=resposta,erro=erro)

@app.route('/ajuda/manual')
@login_required
def manual_usuario():
 return render_template('manual_usuario.html')

@app.route('/ajuda/manual/download/pdf')
@login_required
def manual_usuario_download_pdf():
 return send_from_directory(app.static_folder,'Manual_Frota_Facil_Usuario.pdf',as_attachment=True,download_name='Manual_Frota_Facil_Usuario.pdf')

@app.route('/ajuda/manual/download/docx')
@login_required
def manual_usuario_download_docx():
 return send_from_directory(app.static_folder,'Manual_Frota_Facil_Usuario.docx',as_attachment=True,download_name='Manual_Frota_Facil_Usuario.docx')

@app.route('/ajuda/suporte',methods=['GET','POST'])
@login_required
def suporte():
 if request.method=='POST':
  titulo=(request.form.get('titulo') or '').strip()
  descricao=(request.form.get('descricao') or '').strip()
  categoria=(request.form.get('categoria') or 'Duvida').strip()
  prioridade=(request.form.get('prioridade') or 'Normal').strip()
  pagina=(request.form.get('pagina_origem') or '').strip()[:255] or None
  categorias={'Duvida','Problema','Sugestao','Financeiro','Integracao','Outro'}
  prioridades={'Baixa','Normal','Alta','Urgente'}
  if categoria not in categorias: categoria='Outro'
  if prioridade not in prioridades: prioridade='Normal'
  if not titulo or not descricao:
   flash('Informe o assunto e descreva a solicitação.','danger')
  else:
   ticket=SupportTicket(tenant_id=tid(),user_id=current_user.id,titulo=titulo[:180],categoria=categoria,prioridade=prioridade,descricao=descricao,pagina_origem=pagina,status='ABERTO')
   db.session.add(ticket); db.session.commit()
   flash(f'Solicitação #{ticket.id} aberta com sucesso.','success')
   return redirect(url_for('suporte_ticket',ticket_id=ticket.id))
 tickets=SupportTicket.query.filter_by(tenant_id=tid()).order_by(SupportTicket.criado_em.desc()).limit(100).all()
 return render_template('suporte.html',tickets=tickets,origem=(request.args.get('origem') or '').strip()[:255],support_admin=False)

@app.route('/ajuda/suporte/<int:ticket_id>')
@login_required
def suporte_ticket(ticket_id):
 ticket=SupportTicket.query.filter_by(id=ticket_id,tenant_id=tid()).first_or_404()
 return render_template('suporte_ticket.html',ticket=ticket,support_admin=False)

@app.route('/ajuda/suporte/<int:ticket_id>/encerrar',methods=['POST'])
@login_required
def suporte_ticket_encerrar(ticket_id):
 ticket=SupportTicket.query.filter_by(id=ticket_id,tenant_id=tid()).first_or_404()
 if ticket.status!='RESOLVIDO':
  ticket.status='RESOLVIDO'; ticket.resolvido_em=datetime.utcnow(); ticket.atualizado_em=datetime.utcnow(); db.session.commit()
  flash('Solicitação marcada como resolvida.','success')
 return redirect(url_for('suporte_ticket',ticket_id=ticket.id))

# Compatibilidade: o endereço antigo não expõe mais a central dentro de contas de locadoras.
@app.route('/ajuda/suporte-central')
@login_required
def suporte_central_legado():
 flash('A Central de Suporte agora utiliza uma conta administrativa Frota Fácil separada.','info')
 return redirect(url_for('ajuda'))

def _admin_system_health():
 """Resumo operacional seguro, sem expor segredos nem chamar APIs externas."""
 cards=[]
 agora_utc=datetime.utcnow()
 agora_fila=agora_sao_paulo_naive()

 versao=(os.getenv('FROTA_FACIL_VERSION') or os.getenv('RENDER_GIT_COMMIT') or 'não informada').strip()
 if len(versao)>12 and versao!='não informada': versao=versao[:12]
 cards.append({
  'nome':'Aplicação','icone':'●','status':'ok','rotulo':'Online',
  'resumo':f'Versão {versao}',
  'detalhe':'Iniciada em '+_as_sao_paulo(APP_STARTED_AT).strftime('%d/%m/%Y %H:%M'),
 })

 try:
  inicio=time.perf_counter()
  db.session.execute(text('SELECT 1'))
  latencia=max(0,int((time.perf_counter()-inicio)*1000))
  db_status='ok' if latencia<500 else 'warning'
  cards.append({'nome':'Banco de dados','icone':'◆','status':db_status,'rotulo':'Conectado','resumo':f'Resposta em {latencia} ms','detalhe':'Consulta de integridade concluída.'})
 except Exception as exc:
  db.session.rollback()
  app.logger.exception('Falha no diagnóstico administrativo do banco')
  cards.append({'nome':'Banco de dados','icone':'◆','status':'danger','rotulo':'Falha','resumo':'Sem resposta','detalhe':'Não foi possível concluir a consulta de integridade.'})

 try:
  integracoes=Integration.query.filter_by(tipo='whatsapp').all()
  business=0
  for integracao in integracoes:
   try:
    cfg=CommunicationService.parse_config(integracao)
    if integracao.ativo and (cfg.get('provider') or 'web').lower()=='business': business+=1
   except Exception:
    continue
  falhas_24h=MessageQueue.query.filter(MessageQueue.channel=='whatsapp',MessageQueue.status=='FALHA',MessageQueue.updated_at>=agora_fila-timedelta(hours=24)).count()
  ultimo=MessageQueue.query.filter(MessageQueue.channel=='whatsapp',MessageQueue.sent_at.isnot(None)).order_by(MessageQueue.sent_at.desc()).first()
  if falhas_24h: wa_status='danger'; wa_rotulo='Atenção'
  elif business: wa_status='ok'; wa_rotulo='Operacional'
  else: wa_status='warning'; wa_rotulo='Não conectado'
  ultimo_txt=('Último envio: '+_as_sao_paulo(_message_db_time_as_utc_naive(ultimo.sent_at)).strftime('%d/%m %H:%M')) if ultimo and ultimo.sent_at else 'Nenhum envio registrado.'
  cards.append({'nome':'WhatsApp/Meta','icone':'✆','status':wa_status,'rotulo':wa_rotulo,'resumo':f'{business} locadora(s) Business · {falhas_24h} falha(s) em 24h','detalhe':ultimo_txt})
 except Exception:
  db.session.rollback(); app.logger.exception('Falha no diagnóstico administrativo do WhatsApp')
  cards.append({'nome':'WhatsApp/Meta','icone':'✆','status':'danger','rotulo':'Falha','resumo':'Diagnóstico indisponível','detalhe':'Não foi possível consultar os registros de mensagens.'})

 try:
  chave_openai=bool((os.getenv('OPENAI_API_KEY') or '').strip())
  ultima_analise=Inspection.query.filter(Inspection.damage_analysis_at.isnot(None)).order_by(Inspection.damage_analysis_at.desc()).first()
  if not chave_openai: ai_status='danger'; ai_rotulo='Não configurada'; ai_resumo='Chave ausente'
  elif ultima_analise and ultima_analise.damage_analysis_status=='FALHA': ai_status='warning'; ai_rotulo='Revisar'; ai_resumo='A análise mais recente falhou'
  else: ai_status='ok'; ai_rotulo='Configurada'; ai_resumo='Chave disponível'
  ai_detalhe=('Última análise: '+_as_sao_paulo(ultima_analise.damage_analysis_at).strftime('%d/%m/%Y %H:%M')) if ultima_analise else 'Nenhuma análise registrada ainda.'
  cards.append({'nome':'OpenAI','icone':'✦','status':ai_status,'rotulo':ai_rotulo,'resumo':ai_resumo,'detalhe':ai_detalhe})
 except Exception:
  db.session.rollback(); app.logger.exception('Falha no diagnóstico administrativo da OpenAI')
  cards.append({'nome':'OpenAI','icone':'✦','status':'warning','rotulo':'Indisponível','resumo':'Diagnóstico não concluído','detalhe':'Nenhuma chave ou conteúdo sensível foi exibido.'})

 try:
  ativas=0
  for integracao in Integration.query.filter_by(tipo='whatsapp',ativo=True).all():
   try:
    cfg=CommunicationService.parse_config(integracao)
    if cfg.get('automation_enabled') and any((cfg.get('km_automation_enabled'),cfg.get('billing_automation_enabled'),cfg.get('inspection_automation_enabled'))): ativas+=1
   except Exception:
    continue
  ultimo_auto=MessageEvent.query.filter(MessageEvent.description.ilike('%automat%')).order_by(MessageEvent.created_at.desc()).first()
  auto_status='ok' if ativas and ultimo_auto else ('warning' if ativas else 'neutral')
  auto_rotulo='Em atividade' if ativas and ultimo_auto else ('Aguardando execução' if ativas else 'Desativadas')
  auto_detalhe=('Última atividade: '+_as_sao_paulo(_message_db_time_as_utc_naive(ultimo_auto.created_at)).strftime('%d/%m/%Y %H:%M')) if ultimo_auto else 'Nenhuma execução automática registrada.'
  cards.append({'nome':'Automações','icone':'↻','status':auto_status,'rotulo':auto_rotulo,'resumo':f'{ativas} locadora(s) com automação ativa','detalhe':auto_detalhe})
 except Exception:
  db.session.rollback(); app.logger.exception('Falha no diagnóstico administrativo das automações')
  cards.append({'nome':'Automações','icone':'↻','status':'warning','rotulo':'Indisponível','resumo':'Diagnóstico não concluído','detalhe':'Consulte os registros do serviço de automação.'})

 try:
  pendentes=MessageQueue.query.filter(MessageQueue.status.in_(['PENDENTE','AGENDADA'])).count()
  falhas=MessageQueue.query.filter(MessageQueue.status=='FALHA',MessageQueue.updated_at>=agora_fila-timedelta(hours=24)).count()
  fila_status='danger' if falhas else ('warning' if pendentes>100 else 'ok')
  fila_rotulo='Com falhas' if falhas else ('Acumulada' if pendentes>100 else 'Normal')
  cards.append({'nome':'Fila de mensagens','icone':'≡','status':fila_status,'rotulo':fila_rotulo,'resumo':f'{pendentes} aguardando · {falhas} falha(s) em 24h','detalhe':'Inclui mensagens pendentes e agendadas de todas as locadoras.'})
 except Exception:
  db.session.rollback(); app.logger.exception('Falha no diagnóstico administrativo da fila')
  cards.append({'nome':'Fila de mensagens','icone':'≡','status':'danger','rotulo':'Falha','resumo':'Diagnóstico indisponível','detalhe':'Não foi possível consultar a fila.'})

 ordem={'danger':3,'warning':2,'neutral':1,'ok':0}
 pior=max(cards,key=lambda card:ordem.get(card['status'],0))['status']
 resumo_geral={'danger':'Há itens que exigem atenção','warning':'Sistema operacional com alertas','neutral':'Sistema operacional','ok':'Todos os indicadores normais'}[pior]
 return cards,resumo_geral,pior

@app.route('/admin-frota/suporte')
@frota_admin_required
def frota_admin_suporte():
 admin=_frota_admin_atual()
 status=(request.args.get('status') or '').strip().upper()
 tenant_id=request.args.get('locadora',type=int)
 prioridade=(request.args.get('prioridade') or '').strip()
 busca=(request.args.get('q') or '').strip()
 q=SupportTicket.query
 if status in ('ABERTO','EM_ANALISE','RESOLVIDO'):
  q=q.filter_by(status=status)
 if tenant_id:
  q=q.filter_by(tenant_id=tenant_id)
 if prioridade in ('Baixa','Normal','Alta','Urgente'):
  q=q.filter_by(prioridade=prioridade)
 if busca:
  like=f'%{busca}%'
  q=q.filter(or_(SupportTicket.titulo.ilike(like),SupportTicket.descricao.ilike(like)))
 tickets=q.order_by(SupportTicket.criado_em.desc()).limit(500).all()
 tenant_ids={x.tenant_id for x in tickets}
 tenants={t.id:t for t in Tenant.query.filter(Tenant.id.in_(tenant_ids)).all()} if tenant_ids else {}
 todas_locadoras=Tenant.query.order_by(Tenant.nome.asc()).all()
 counts={
  'ABERTO':SupportTicket.query.filter_by(status='ABERTO').count(),
  'EM_ANALISE':SupportTicket.query.filter_by(status='EM_ANALISE').count(),
  'RESOLVIDO':SupportTicket.query.filter_by(status='RESOLVIDO').count(),
 }
 health_cards,health_summary,health_status=_admin_system_health()
 return render_template('admin_frota_suporte.html',admin=admin,tickets=tickets,tenants=tenants,todas_locadoras=todas_locadoras,status_filtro=status,prioridade_filtro=prioridade,locadora_filtro=tenant_id,busca=busca,counts=counts,health_cards=health_cards,health_summary=health_summary,health_status=health_status,health_checked_at=datetime.utcnow())

@app.route('/admin-frota/suporte/<int:ticket_id>')
@frota_admin_required
def frota_admin_ticket(ticket_id):
 admin=_frota_admin_atual()
 ticket=SupportTicket.query.get_or_404(ticket_id)
 tenant=Tenant.query.get(ticket.tenant_id)
 return render_template('admin_frota_ticket.html',admin=admin,ticket=ticket,tenant=tenant)

@app.route('/admin-frota/suporte/<int:ticket_id>/responder',methods=['POST'])
@frota_admin_required
def frota_admin_responder(ticket_id):
 admin=_frota_admin_atual()
 ticket=SupportTicket.query.get_or_404(ticket_id)
 resposta=(request.form.get('resposta') or '').strip()
 status=(request.form.get('status') or 'EM_ANALISE').strip().upper()
 if status not in ('ABERTO','EM_ANALISE','RESOLVIDO'):
  status='EM_ANALISE'
 ticket.status=status
 if resposta:
  ticket.resposta=resposta
  ticket.respondido_por_admin_id=admin.id
  ticket.respondido_por_nome=admin.nome
  ticket.respondido_por_id=None
 if status=='RESOLVIDO':
  ticket.resolvido_em=datetime.utcnow()
 elif ticket.resolvido_em:
  ticket.resolvido_em=None
 ticket.atualizado_em=datetime.utcnow(); db.session.commit()
 flash(f'Solicitação #{ticket.id} atualizada.','success')
 return redirect(url_for('frota_admin_ticket',ticket_id=ticket.id))

@app.route('/configuracoes')
@login_required
def configuracoes():
 return render_template('configuracoes.html')

@app.route('/configuracoes/locadora',methods=['GET','POST'])
@login_required
def configuracoes_locadora():
 tenant=Tenant.query.get_or_404(tid())
 if request.method=='POST':
  timezone_name=(request.form.get('timezone_name') or tenant.timezone_name or 'America/Sao_Paulo').strip()
  if timezone_name not in TENANT_TIMEZONE_NAMES:
   flash('Selecione um fuso horário válido.','danger'); return redirect(url_for('configuracoes_locadora'))
  try:
   ZoneInfo(timezone_name)
  except Exception:
   flash('O fuso horário selecionado não está disponível no servidor.','danger'); return redirect(url_for('configuracoes_locadora'))
  tenant.timezone_name=timezone_name
  campos=['razao_social','nome_fantasia','cnpj','inscricao_estadual','inscricao_municipal','telefone','email','responsavel_legal','logradouro','numero_endereco','complemento','bairro','cidade','uf','cep','cor_primaria','cor_secundaria']
  for campo in campos:
   valor=(request.form.get(campo) or '').strip() or None
   if campo in ('cor_primaria','cor_secundaria') and valor and not re.fullmatch(r'#[0-9A-Fa-f]{6}',valor):
    flash('As cores devem estar no formato #RRGGBB.','danger'); return redirect(url_for('configuracoes_locadora'))
   setattr(tenant,campo,valor)
  if tenant.nome_fantasia: tenant.nome=tenant.nome_fantasia
  for field,form_name,prefix in [('logo_key','logo','logo'),('favicon_key','favicon','favicon')]:
   f=request.files.get(form_name)
   if f and f.filename:
    ext=Path(secure_filename(f.filename)).suffix.lower()
    allowed={'.png':'image/png','.jpg':'image/jpeg','.jpeg':'image/jpeg','.webp':'image/webp','.ico':'image/x-icon'}
    if ext not in allowed:
     flash(f'Formato inválido para {form_name}. Use PNG, JPG, WEBP ou ICO.','danger'); return redirect(url_for('configuracoes_locadora'))
    data=f.read(); key=f'{tid()}/configuracoes/identidade/{prefix}_{uuid.uuid4().hex}{ext}'
    storage.upload(BytesIO(data),key,allowed[ext]); old=getattr(tenant,field)
    setattr(tenant,field,key)
    if old:
     try: storage.delete(old)
     except Exception: pass
  db.session.commit(); flash('Dados e identidade visual da locadora atualizados.','success'); return redirect(url_for('configuracoes_locadora'))
 html_config=render_template('configuracoes_locadora.html',tenant=tenant)
 if 'name="timezone_name"' not in html_config:
  options=[]
  atual=(tenant.timezone_name or 'America/Sao_Paulo').strip()
  for value,label in TENANT_TIMEZONE_OPTIONS:
   selected=' selected' if value==atual else ''
   options.append(f'<option value="{value}"{selected}>{label}</option>')
  bloco=(
   '<div class="mb-3">'
   '<label class="form-label fw-semibold" for="timezone_name">Fuso horário da locadora</label>'
   '<select class="form-select" id="timezone_name" name="timezone_name">'+''.join(options)+'</select>'
   '<div class="form-text">Usado nos horários exibidos e nas automações deste tenant.</div>'
   '</div>'
  )
  if '</form>' in html_config:
   pos=html_config.rfind('</form>')
   html_config=html_config[:pos]+bloco+html_config[pos:]
 return html_config

@app.route('/configuracoes/locadora/arquivo/<tipo>')
@login_required
def arquivo_identidade_locadora(tipo):
 tenant=Tenant.query.get_or_404(tid()); key=tenant.logo_key if tipo=='logo' else tenant.favicon_key if tipo=='favicon' else None
 if not key: abort(404)
 data=storage.download(key); ext=Path(key).suffix.lower(); mime={'.png':'image/png','.jpg':'image/jpeg','.jpeg':'image/jpeg','.webp':'image/webp','.ico':'image/x-icon'}.get(ext,'application/octet-stream')
 return send_file(BytesIO(data),mimetype=mime,download_name=Path(key).name,as_attachment=False)


@app.route('/identidade/<int:tenant_id>/<tipo>')
def identidade_publica_tenant(tenant_id,tipo):
 tenant=Tenant.query.get_or_404(tenant_id)
 key=tenant.logo_key if tipo=='logo' else tenant.favicon_key if tipo=='favicon' else None
 if not key:
  abort(404)
 try:
  data=storage.download(key)
 except StorageNotFoundError:
  abort(404)
 except Exception:
  abort(503)
 ext=Path(key).suffix.lower()
 mime={'.png':'image/png','.jpg':'image/jpeg','.jpeg':'image/jpeg','.webp':'image/webp','.ico':'image/x-icon'}.get(ext,'application/octet-stream')
 return send_file(BytesIO(data),mimetype=mime,download_name=Path(key).name,as_attachment=False)

@app.context_processor
def identidade_visual_contexto():
 """Disponibiliza a identidade visual correta sem misturar tenants."""
 tenant_visual=None
 if current_user.is_authenticated:
  tenant_visual=current_user.tenant
 elif session.get('owner_tenant_id'):
  try:
   tenant_visual=Tenant.query.get(int(session.get('owner_tenant_id')))
  except Exception:
   tenant_visual=None
 return {'tenant_visual':tenant_visual}


@app.route('/configuracoes/automacoes',methods=['GET','POST'])
@login_required
def configuracoes_automacoes():
 item=Integration.query.filter_by(tenant_id=tid(),tipo='whatsapp').first()
 cfg=CommunicationService.parse_config(item)
 if request.method=='POST':
  if not item:
   item=Integration(tenant_id=tid(),tipo='whatsapp',ativo=False,configuracao='{}'); db.session.add(item)
  cfg=CommunicationService.parse_config(item)

  def _hora(nome,padrao):
   try: return max(0,min(23,int(request.form.get(nome) or padrao)))
   except Exception: return padrao
  def _intervalo(nome,padrao):
   try: return max(1,min(12,int(request.form.get(nome) or padrao)))
   except Exception: return padrao
  def _dias(nome,fallback):
   valores=[]
   for raw in request.form.getlist(nome):
    try:
     n=int(raw)
     if 0<=n<=6 and n not in valores: valores.append(n)
    except Exception: pass
   return sorted(valores) if valores else list(fallback)

  try: old_weekday=max(0,min(6,int(cfg.get('automation_weekday',0))))
  except Exception: old_weekday=0
  old_start=_hora('automation_start_hour',int(cfg.get('automation_start_hour',7) or 7))
  old_end=_hora('automation_end_hour',int(cfg.get('automation_end_hour',20) or 20))
  try: old_interval=max(1,min(12,int(cfg.get('reminder_interval_hours',1) or 1)))
  except Exception: old_interval=1

  km_enabled=request.form.get('km_automation_enabled')=='1'
  billing_enabled=request.form.get('billing_automation_enabled')=='1'
  inspection_enabled=request.form.get('inspection_automation_enabled')=='1'
  inspection_type=(request.form.get('inspection_automation_type') or cfg.get('inspection_automation_type') or 'fotos').strip().lower()
  if inspection_type not in ('fotos','simples','guiada'): inspection_type='fotos'
  damage_detection_enabled=request.form.get('inspection_damage_detection_enabled')=='1'
  km_days=_dias('km_automation_weekdays',cfg.get('km_automation_weekdays') or [old_weekday])
  billing_days=_dias('billing_automation_weekdays',cfg.get('billing_automation_weekdays') or [old_weekday])
  inspection_days=_dias('inspection_automation_weekdays',cfg.get('inspection_automation_weekdays') or [old_weekday])
  km_start=_hora('km_start_hour',int(cfg.get('km_start_hour',old_start) or old_start))
  km_end=_hora('km_end_hour',int(cfg.get('km_end_hour',old_end) or old_end))
  billing_start=_hora('billing_start_hour',int(cfg.get('billing_start_hour',old_start) or old_start))
  billing_end=_hora('billing_end_hour',int(cfg.get('billing_end_hour',old_end) or old_end))
  inspection_start=_hora('inspection_start_hour',int(cfg.get('inspection_start_hour',10) or 10))
  inspection_end=_hora('inspection_end_hour',int(cfg.get('inspection_end_hour',20) or 20))
  km_interval=_intervalo('km_reminder_interval_hours',int(cfg.get('km_reminder_interval_hours',old_interval) or old_interval))
  billing_interval=_intervalo('billing_reminder_interval_hours',int(cfg.get('billing_reminder_interval_hours',old_interval) or old_interval))
  try: inspection_interval=max(1,min(168,int(request.form.get('inspection_reminder_interval_hours') or cfg.get('inspection_reminder_interval_hours',3) or 3)))
  except Exception: inspection_interval=3
  try: inspection_expiry=max(24,min(168,int(request.form.get('inspection_expiry_hours') or cfg.get('inspection_expiry_hours',168) or 168)))
  except Exception: inspection_expiry=168

  if km_enabled and not request.form.getlist('km_automation_weekdays'):
   flash('Selecione pelo menos um dia para a solicitação de KM/foto.','danger'); return redirect(url_for('configuracoes_automacoes'))
  if billing_enabled and not request.form.getlist('billing_automation_weekdays'):
   flash('Selecione pelo menos um dia para as cobranças.','danger'); return redirect(url_for('configuracoes_automacoes'))
  if inspection_enabled and not request.form.getlist('inspection_automation_weekdays'):
   flash('Selecione pelo menos um dia para a vistoria automática.','danger'); return redirect(url_for('configuracoes_automacoes'))
  if km_end < km_start:
   flash('Na automação de KM, o horário final não pode ser anterior ao inicial.','danger'); return redirect(url_for('configuracoes_automacoes'))
  if billing_end < billing_start:
   flash('Na automação de cobranças, o horário final não pode ser anterior ao inicial.','danger'); return redirect(url_for('configuracoes_automacoes'))
  if inspection_end < inspection_start:
   flash('Na automação de vistoria, o horário final não pode ser anterior ao inicial.','danger'); return redirect(url_for('configuracoes_automacoes'))

  # Se a automação de cobranças acabou de ser desligada, elimina da fila
  # qualquer cobrança ainda não enviada. Assim o botão funciona como um freio
  # imediato, sem afetar o histórico financeiro nem as baixas manuais.
  if not billing_enabled:
   agora_cancelamento=agora_sao_paulo_naive()
   filas_cobranca=MessageQueue.query.filter(
    MessageQueue.tenant_id==tid(),
    MessageQueue.status.in_(['AGENDADA','PENDENTE','AGUARDANDO_MANUAL']),
   ).filter(
    db.or_(
     MessageQueue.related_entity=='Cobranca',
     MessageQueue.message_type.in_(['lembrete_pagamento_semanal','cobranca','cobranca_semanal'])
    )
   ).all()
   for fila_cobranca in filas_cobranca:
    fila_cobranca.status='CANCELADA'; fila_cobranca.updated_at=agora_cancelamento
    db.session.add(MessageEvent(tenant_id=tid(),message_id=fila_cobranca.id,event='CANCELADA',description='Cobrança cancelada ao desabilitar a automação de cobranças.',created_at=agora_cancelamento))

  cfg.update({
   'automation_enabled':request.form.get('automation_enabled')=='1',
   'km_automation_enabled':km_enabled,
   'km_automation_weekdays':km_days,
   'km_start_hour':km_start,'km_end_hour':km_end,'km_reminder_interval_hours':km_interval,
   'billing_automation_enabled':billing_enabled,
   'billing_automation_weekdays':billing_days,
   'billing_start_hour':billing_start,'billing_end_hour':billing_end,'billing_reminder_interval_hours':billing_interval,
   'inspection_automation_enabled':inspection_enabled,
   'inspection_automation_type':inspection_type,
   'inspection_damage_detection_enabled':damage_detection_enabled,
   'inspection_automation_weekdays':inspection_days,
   'inspection_start_hour':inspection_start,'inspection_end_hour':inspection_end,
   'inspection_reminder_interval_hours':inspection_interval,'inspection_expiry_hours':inspection_expiry,
   'owner_portal_auto_invite_enabled':request.form.get('owner_portal_auto_invite_enabled')=='1',
    'driver_portal_auto_invite_enabled':request.form.get('driver_portal_auto_invite_enabled')=='1',
   # Mantém as chaves antigas durante a transição para não quebrar outras rotinas/RCs.
   'automatic_km_enabled':km_enabled,'automatic_billing_enabled':billing_enabled,
   'automatic_alerts_enabled':request.form.get('automatic_alerts_enabled')=='1',
  })
  item.configuracao=json.dumps(cfg,ensure_ascii=False); db.session.commit(); flash('Automações atualizadas.','success'); return redirect(url_for('configuracoes_automacoes'))
 return render_template('configuracoes_automacoes.html',cfg=cfg,provider=(cfg.get('provider') or 'web'))

@app.route('/configuracoes/quilometragem',methods=['GET','POST'])
@login_required
def configuracoes_quilometragem():
 tenant=Tenant.query.get_or_404(tid())
 if request.method=='POST':
  tenant.conferir_km_motorista=request.form.get('conferir_km_motorista')=='1'
  db.session.commit()
  flash('Preferência de conferência de quilometragem atualizada.','success')
  return redirect(url_for('configuracoes_quilometragem'))
 return render_template('configuracoes_quilometragem.html',tenant=tenant)

@app.route('/quilometragens/conferencia')
@login_required
def conferencia_quilometragens():
 items=MileageRequest.query.options(joinedload(MileageRequest.vehicle),joinedload(MileageRequest.driver)).filter_by(tenant_id=tid(),status='Aguardando conferência').order_by(MileageRequest.submitted_at.asc()).all()
 return render_template('conferencia_quilometragens.html',items=items)

@app.route('/quilometragens/<int:id>/conferir',methods=['POST'])
@login_required
def conferir_quilometragem(id):
 req=MileageRequest.query.options(joinedload(MileageRequest.vehicle)).filter_by(id=id,tenant_id=tid(),status='Aguardando conferência').first_or_404()
 acao=request.form.get('acao')
 if acao=='rejeitar':
  req.status='Rejeitado'
  req.notes=((req.notes or '')+'\nRejeitado pelo administrador: '+(request.form.get('motivo') or 'Solicitada nova foto.')).strip()
  db.session.commit()
  flash('Leitura rejeitada. Gere uma nova solicitação de KM para o motorista.','success')
  return redirect(url_for('conferencia_quilometragens'))
 try:
  km=int(request.form.get('km') or req.km or 0)
 except ValueError:
  flash('Informe uma quilometragem válida.','danger'); return redirect(url_for('conferencia_quilometragens'))
 if km < (req.previous_km or 0):
  flash(f'A KM confirmada não pode ser menor que a leitura anterior ({req.previous_km:,} km).','danger'); return redirect(url_for('conferencia_quilometragens'))
 req.km=km; req.status='Concluído'; req.vehicle.km_atual=km
 origem='Administrador confirmou KM do motorista' if km==request.form.get('km_original',type=int) else 'Administrador corrigiu KM do motorista'
 db.session.add(Odometer(tenant_id=tid(),vehicle_id=req.vehicle_id,km=km,origem=origem))
 db.session.commit(); recalcular_alertas(tid())
 flash('Quilometragem conferida e veículo atualizado.','success')
 return redirect(url_for('conferencia_quilometragens'))

@app.route('/ferramentas/ocr-painel', methods=['GET','POST'])
@login_required
def ocr_painel_teste():
 resultado=None
 veiculos=Vehicle.query.filter_by(tenant_id=tid()).order_by(Vehicle.placa).all()
 veiculo_id=request.form.get('vehicle_id',type=int) if request.method=='POST' else None
 veiculo=None
 if veiculo_id:
  veiculo=Vehicle.query.filter_by(id=veiculo_id,tenant_id=tid()).first()
 if request.method=='POST':
  foto=request.files.get('foto')
  if not foto or not foto.filename:
   flash('Selecione uma foto do painel.','danger')
  else:
   ext=Path(secure_filename(foto.filename)).suffix.lower()
   if ext not in ('.jpg','.jpeg','.png','.webp'):
    flash('Envie uma foto JPG, PNG ou WEBP.','danger')
   else:
    try:
     data=foto.read()
     if not data:
      raise ValueError('arquivo vazio')
     resultado=read_odometer(data, previous_km=(veiculo.km_atual if veiculo else None))
    except Exception:
     app.logger.exception('Falha no teste OCR do painel')
     flash('Não foi possível processar a imagem. Tente outra foto.','danger')
 return render_template('ocr_painel_teste.html',veiculos=veiculos,veiculo=veiculo,resultado=resultado)

@app.route('/quilometragens')
@login_required
def quilometragens():
 items=MileageRequest.query.options(joinedload(MileageRequest.vehicle),joinedload(MileageRequest.driver)).filter_by(tenant_id=tid()).order_by(MileageRequest.id.desc()).all()
 return render_template('quilometragens.html',items=items)

@app.route('/quilometragens/<int:id>/foto')
@login_required
def foto_quilometragem(id):
 req=MileageRequest.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 if not req.photo:
  abort(404)

 # Compatibilidade com fotos antigas salvas localmente antes do R2.
 if '/' not in req.photo:
  return send_from_directory(UPLOAD/str(tid())/'odometros',req.photo)

 try:
  conteudo=storage.download(req.photo)
 except StorageNotFoundError:
  abort(404)
 except Exception:
  app.logger.exception('Falha ao baixar foto de quilometragem %s',req.id)
  abort(503)

 extensao=Path(req.photo).suffix.lower()
 mimetypes={
  '.jpg':'image/jpeg',
  '.jpeg':'image/jpeg',
  '.png':'image/png',
  '.webp':'image/webp',
 }
 return send_file(
  BytesIO(conteudo),
  mimetype=mimetypes.get(extensao,'application/octet-stream'),
  download_name=f'painel_{req.vehicle_id}_{req.id}{extensao}',
  as_attachment=False,
 )

def gerar_codigo_publico_contrato():
 for _ in range(10):
  codigo='FF-'+uuid.uuid4().hex[:4].upper()+'-'+uuid.uuid4().hex[:4].upper()+'-'+uuid.uuid4().hex[:4].upper()
  if not Contract.query.filter_by(codigo_publico=codigo).first():
   return codigo
 raise RuntimeError('Não foi possível gerar código público único.')

def garantir_codigo_publico_contrato(contrato):
 """Garante código público para contratos antigos criados antes da Sprint 0.9.2C."""
 if contrato.codigo_publico:
  return contrato.codigo_publico
 contrato.codigo_publico=gerar_codigo_publico_contrato()
 db.session.flush()
 return contrato.codigo_publico


def telefone_whatsapp(valor):
 return normalize_phone(valor)

def _dados_historico_veiculo(v):
 eventos=VehicleEvent.query.options(joinedload(VehicleEvent.user),joinedload(VehicleEvent.contract),joinedload(VehicleEvent.driver)).filter_by(tenant_id=tid(),vehicle_id=v.id).order_by(VehicleEvent.criado_em.desc()).all()
 contratos=Contract.query.options(joinedload(Contract.driver)).filter_by(tenant_id=tid(),vehicle_id=v.id).order_by(Contract.id.desc()).all()
 odometros=Odometer.query.filter_by(tenant_id=tid(),vehicle_id=v.id).order_by(Odometer.data.desc(),Odometer.id.desc()).all()
 manutencoes=Maintenance.query.filter_by(tenant_id=tid(),vehicle_id=v.id).order_by(Maintenance.id.desc()).all()
 vistorias=Inspection.query.options(joinedload(Inspection.driver),joinedload(Inspection.contract)).filter_by(tenant_id=tid(),vehicle_id=v.id).order_by(Inspection.requested_at.desc(),Inspection.id.desc()).all()
 _ajustar_eventos_vistoria_para_status_atual(eventos,vistorias)
 documentos=Document.query.filter_by(tenant_id=tid(),entidade='Veículo',entidade_id=v.id,status='Ativo').order_by(Document.criado_em.desc(),Document.id.desc()).all()
 return {
  'eventos':eventos,'contratos':contratos,'odometros':odometros,
  'manutencoes':manutencoes,'vistorias':vistorias,'documentos':documentos,
 }

@app.route('/veiculos/<int:id>/historico')
@login_required
def historico_veiculo(id):
 v=Vehicle.query.options(joinedload(Vehicle.current_driver),joinedload(Vehicle.current_contract)).filter_by(id=id,tenant_id=tid()).first_or_404()
 dados=_dados_historico_veiculo(v)
 return render_template('veiculo_historico.html',v=v,**dados)

@app.route('/veiculos/<int:id>/historico/imprimir')
@login_required
def imprimir_historico_veiculo(id):
 v=Vehicle.query.options(joinedload(Vehicle.current_driver),joinedload(Vehicle.current_contract)).filter_by(id=id,tenant_id=tid()).first_or_404()
 dados=_dados_historico_veiculo(v)
 return render_template('veiculo_historico_impressao.html',v=v,gerado_em=agora_sao_paulo_naive(),**dados)

def garantir_token_comprovante(audit):
 if audit.receipt_token: return audit.receipt_token
 for _ in range(10):
  token=uuid.uuid4().hex+uuid.uuid4().hex
  if not BillingAudit.query.filter_by(receipt_token=token).first():
   audit.receipt_token=token; db.session.flush(); return token
 raise RuntimeError('Não foi possível gerar link único para comprovante.')

def url_comprovante_cobranca(audit):
 return url_for('enviar_comprovante_pagamento',token=garantir_token_comprovante(audit),_external=True)

@app.route('/pagamento/<token>',methods=['GET','POST'])
def enviar_comprovante_pagamento(token):
 audit=BillingAudit.query.filter_by(receipt_token=token).first_or_404()
 contrato=Contract.query.options(joinedload(Contract.driver),joinedload(Contract.vehicle)).filter_by(id=audit.contract_id,tenant_id=audit.tenant_id).first()
 tenant=Tenant.query.get(audit.tenant_id)
 if request.method=='POST':
  if (audit.payment_status or 'PENDENTE')=='PAGO': return render_template('enviar_comprovante.html',audit=audit,contrato=contrato,tenant=tenant,concluido=True)
  arquivo=request.files.get('comprovante')
  if not arquivo or not arquivo.filename:
   flash('Selecione o comprovante.','danger'); return render_template('enviar_comprovante.html',audit=audit,contrato=contrato,tenant=tenant)
  nome=secure_filename(arquivo.filename); ext=Path(nome).suffix.lower()
  permitidos={'.pdf':'application/pdf','.jpg':'image/jpeg','.jpeg':'image/jpeg','.png':'image/png','.webp':'image/webp'}
  if ext not in permitidos:
   flash('Envie PDF, JPG, PNG ou WEBP.','danger'); return render_template('enviar_comprovante.html',audit=audit,contrato=contrato,tenant=tenant)
  data=arquivo.read()
  if not data or len(data)>15*1024*1024:
   flash('O comprovante deve ter até 15 MB.','danger'); return render_template('enviar_comprovante.html',audit=audit,contrato=contrato,tenant=tenant)
  chave=f'{audit.tenant_id}/documentos/comprovantes/{audit.id}/{uuid.uuid4().hex}_{nome}'
  try:
   storage.upload(BytesIO(data),chave,permitidos[ext])
   if audit.receipt_key:
    try: storage.delete(audit.receipt_key)
    except Exception: pass
   audit.receipt_key=chave; audit.receipt_name=nome; audit.receipt_mime=permitidos[ext]; audit.receipt_uploaded_at=agora_sao_paulo_naive()
   if (audit.payment_status or 'PENDENTE')!='PAGO': audit.payment_status='COMPROVANTE_RECEBIDO'
   db.session.commit()
  except Exception:
   db.session.rollback(); app.logger.exception('Falha ao armazenar comprovante da cobrança %s',audit.id)
   flash('Não foi possível armazenar o comprovante. Tente novamente.','danger'); return render_template('enviar_comprovante.html',audit=audit,contrato=contrato,tenant=tenant)
  return render_template('enviar_comprovante.html',audit=audit,contrato=contrato,tenant=tenant,concluido=True)
 return render_template('enviar_comprovante.html',audit=audit,contrato=contrato,tenant=tenant,concluido=False)

@app.route('/cobrancas/auditoria/<int:id>/comprovante')
@login_required
def visualizar_comprovante_pagamento(id):
 audit=BillingAudit.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 if not audit.receipt_key: abort(404)
 try: data=storage.download(audit.receipt_key)
 except StorageNotFoundError: abort(404)
 except Exception: abort(503)
 return send_file(BytesIO(data),mimetype=audit.receipt_mime or 'application/octet-stream',download_name=audit.receipt_name or f'comprovante-{audit.id}',as_attachment=False)

@app.route('/cobrancas',methods=['GET','POST'])
@login_required
def cobrancas():
 tenant=Tenant.query.get_or_404(tid())
 if request.method=='POST':
  tenant.cobrar_km_excedente=bool(request.form.get('cobrar_km_excedente'))
  db.session.commit()
  flash('Configuração de cobrança de KM excedente atualizada.','success')
  return redirect(url_for('cobrancas'))
 contratos_ativos=Contract.query.options(joinedload(Contract.driver),joinedload(Contract.vehicle)).filter(
  Contract.tenant_id==tid(),Contract.status.in_(['Assinado','Ativo'])
 ).order_by(Contract.id.desc()).all()
 items=[]
 for c in contratos_ativos:
  # Esta RC é focada no fluxo semanal. Contratos de outra periodicidade continuam fora da cobrança automática semanal.
  periodicidade=unicodedata.normalize('NFKD',str(c.periodicidade or '')).encode('ascii','ignore').decode('ascii').lower()
  if periodicidade and 'seman' not in periodicidade:
   continue
  info=calcular_cobranca_semanal(c)
  items.append({'contract':c,'info':info,'vence_hoje':cobranca_vence_hoje(c)})
 hoje=[x for x in items if x['vence_hoje']]
 auditoria=BillingAudit.query.filter_by(tenant_id=tid()).order_by(BillingAudit.created_at.desc()).limit(200).all()
 alterou=False
 for audit in auditoria:
  if not audit.receipt_token: garantir_token_comprovante(audit); alterou=True
 if alterou: db.session.commit()
 rendered_html=render_template('cobrancas.html',items=items,hoje=hoje,auditoria=auditoria,tenant=tenant)
 # Baixa manual da competência semanal atual.
 # Busca diretamente no banco (sem o limite visual de 200 registros da auditoria)
 # para que nenhum contrato vigente desapareça da baixa manual por possuir histórico antigo.
 contratos_vigentes_ids={c.id for c in contratos_ativos}
 agora_local=datetime.now(_tenant_zone(tid()))
 inicio_semana=agora_local.date()-timedelta(days=agora_local.weekday())
 fim_semana=inicio_semana+timedelta(days=6)
 pendentes_baixa=BillingAudit.query.filter(
  BillingAudit.tenant_id==tid(),
  BillingAudit.contract_id.in_(contratos_vigentes_ids),
  BillingAudit.billing_date>=inicio_semana,
  BillingAudit.billing_date<=fim_semana,
  BillingAudit.payment_status!='PAGO'
 ).order_by(BillingAudit.billing_date.desc(),BillingAudit.id.desc()).all() if contratos_vigentes_ids else []
 if pendentes_baixa:
  linhas=[]
  for a in pendentes_baixa:
   total='R$ {:,.2f}'.format(float(a.total_amount or 0)).replace(',', 'X').replace('.', ',').replace('X', '.')
   data_ref=a.billing_date.strftime('%d/%m/%Y') if a.billing_date else '-'
   form_action=url_for('marcar_cobranca_paga',id=a.id)
   linha=(
    '<tr>'
    '<td>'+str(a.id)+'</td>'
    '<td>'+html.escape(a.driver_name or '-')+'</td>'
    '<td>'+html.escape(a.vehicle_label or '-')+'</td>'
    '<td>'+html.escape(a.plate or '-')+'</td>'
    '<td>'+data_ref+'</td>'
    '<td>'+total+'</td>'
    '<td><form method="post" action="'+form_action+'" style="margin:0">'
    '<input type="hidden" name="payment_method" value="Baixa manual">'
    '<input type="hidden" name="payment_notes" value="Pagamento confirmado manualmente pela locadora.">'
    '<button type="submit" class="btn btn-success btn-sm" onclick="return confirm(\'Confirmar pagamento desta cobrança?\')">✓ Marcar como pago</button>'
    '</form></td>'
    '</tr>'
   )
   linhas.append(linha)
  painel=(
   '<div class="container-fluid mt-4"><div class="card shadow-sm"><div class="card-body">'
   '<h5 class="card-title mb-1">Baixa manual de cobranças existentes</h5>'
   '<p class="text-muted small">Marca como paga uma cobrança que já existe na auditoria, sem enviar mensagem e sem exigir comprovante.</p>'
   '<div class="table-responsive"><table class="table table-sm align-middle">'
   '<thead><tr><th>ID</th><th>Motorista</th><th>Veículo</th><th>Placa</th><th>Data</th><th>Total</th><th>Ação</th></tr></thead>'
   '<tbody>'+''.join(linhas)+'</tbody></table></div></div></div></div>'
  )
  if '</main>' in rendered_html:
   rendered_html=rendered_html.replace('</main>',painel+'</main>',1)
  elif '</body>' in rendered_html:
   rendered_html=rendered_html.replace('</body>',painel+'</body>',1)
  else:
   rendered_html=rendered_html+painel

 # Contratos vigentes sem BillingAudit da semana atual. Permite criar a competência
 # diretamente como PAGA, sem gerar mensagem, link de comprovante ou fila WhatsApp.
 # Cobranças antigas não impedem que a competência atual apareça aqui.
 contratos_sem_auditoria=[]
 for item in items:
  c=item['contract']
  auditoria_semana=BillingAudit.query.filter(
   BillingAudit.tenant_id==tid(),
   BillingAudit.contract_id==c.id,
   BillingAudit.billing_date>=inicio_semana,
   BillingAudit.billing_date<=fim_semana
  ).first()
  if auditoria_semana:
   continue
  contratos_sem_auditoria.append(item)

 if contratos_sem_auditoria:
  linhas_sem=[]
  for item in contratos_sem_auditoria:
   c=item['contract']
   info=item['info']
   total='R$ {:,.2f}'.format(float(info.get('total') or 0)).replace(',', 'X').replace('.', ',').replace('X', '.')
   motorista=c.driver.nome if c.driver else '-'
   veiculo=c.vehicle.marca_modelo if c.vehicle else '-'
   placa=c.vehicle.placa if c.vehicle else '-'
   form_action=url_for('criar_baixa_manual_cobranca',id=c.id)
   linha=(
    '<tr>'
    '<td>'+str(c.id)+'</td>'
    '<td>'+html.escape(motorista or '-')+'</td>'
    '<td>'+html.escape(veiculo or '-')+'</td>'
    '<td>'+html.escape(placa or '-')+'</td>'
    '<td>'+total+'</td>'
    '<td><form method="post" action="'+form_action+'" style="margin:0">'
    '<button type="submit" class="btn btn-primary btn-sm" onclick="return confirm(\'Criar esta cobrança e marcar como paga, sem enviar mensagem?\')">✓ Criar e marcar como pago</button>'
    '</form></td>'
    '</tr>'
   )
   linhas_sem.append(linha)
  periodo_semana=inicio_semana.strftime('%d/%m/%Y')+' a '+fim_semana.strftime('%d/%m/%Y')
  painel_sem=(
   '<div class="container-fluid mt-4"><div class="card shadow-sm"><div class="card-body">'
   '<h5 class="card-title mb-1">Contratos sem cobrança nesta semana</h5>'
   '<p class="text-muted small">Período '+periodo_semana+'. Cria a cobrança e já registra o pagamento, sem WhatsApp e sem comprovante.</p>'
   '<div class="table-responsive"><table class="table table-sm align-middle">'
   '<thead><tr><th>Contrato</th><th>Motorista</th><th>Veículo</th><th>Placa</th><th>Total</th><th>Ação</th></tr></thead>'
   '<tbody>'+''.join(linhas_sem)+'</tbody></table></div></div></div></div>'
  )
  if '</main>' in rendered_html:
   rendered_html=rendered_html.replace('</main>',painel_sem+'</main>',1)
  elif '</body>' in rendered_html:
   rendered_html=rendered_html.replace('</body>',painel_sem+'</body>',1)
  else:
   rendered_html=rendered_html+painel_sem
 # Coerência visual: se a competência atual já foi paga, o botão de lembrete
 # da tabela principal deixa de ser acionável. A proteção real continua no backend.
 pagos_semana=BillingAudit.query.filter(
  BillingAudit.tenant_id==tid(),
  BillingAudit.contract_id.in_(contratos_vigentes_ids),
  BillingAudit.billing_date>=inicio_semana,
  BillingAudit.billing_date<=fim_semana,
  BillingAudit.payment_status=='PAGO'
 ).all() if contratos_vigentes_ids else []
 contratos_pagos=sorted({a.contract_id for a in pagos_semana if a.contract_id})
 if contratos_pagos:
  ids_json=json.dumps(contratos_pagos)
  script_pago=(
   '<script>(function(){var pagos=new Set('+ids_json+');'
   'document.querySelectorAll("form[action]").forEach(function(f){'
   'var m=(f.getAttribute("action")||"").match(/\\/cobrancas\\/(\\d+)\\/whatsapp(?:$|\\?)/);'
   'if(!m||!pagos.has(parseInt(m[1],10)))return;'
   'var b=f.querySelector("button,input[type=submit]");'
   'if(b){var span=document.createElement("span");span.className="badge bg-success";span.textContent="Pago nesta semana";b.replaceWith(span);}'
   '});})();</script>'
  )
  if '</body>' in rendered_html:
   rendered_html=rendered_html.replace('</body>',script_pago+'</body>',1)
  else:
   rendered_html+=script_pago
 return rendered_html

@app.route('/cobrancas/<int:id>/whatsapp',methods=['POST'])
@login_required
def cobranca_whatsapp(id):
 c=Contract.query.options(joinedload(Contract.driver),joinedload(Contract.vehicle)).filter_by(id=id,tenant_id=tid()).first_or_404()
 if c.status not in ('Assinado','Ativo'):
  flash('Somente contratos vigentes podem gerar lembrete de cobrança.','warning'); return redirect(url_for('cobrancas'))
 fila,audit,redirect_url,err=gerar_e_enviar_cobranca(c,automatico=False)
 db.session.commit()
 if err: flash('Não foi possível enviar a cobrança: '+err,'danger')
 elif redirect_url: return redirect(redirect_url)
 else: flash('Lembrete de cobrança enviado pelo provedor configurado.','success')
 return redirect(url_for('cobrancas'))


def enviar_contrato_whatsapp_automatico(c, user_id=None):
 """Envia o contrato automaticamente após a geração sem comprometer a criação do contrato."""
 if not c or not c.driver or not c.vehicle or not c.arquivo_pdf:
  return False,'Contrato sem dados suficientes para envio automático.'

 ja_enviado=MessageQueue.query.filter(
  MessageQueue.tenant_id==c.tenant_id,
  MessageQueue.related_entity=='Contrato',
  MessageQueue.related_entity_id==c.id,
  MessageQueue.message_type=='contrato',
  MessageQueue.status.in_(['ENVIADA','PENDENTE','AGENDADA'])
 ).order_by(MessageQueue.id.desc()).first()
 if ja_enviado:
  return True,'Contrato já possui envio de WhatsApp registrado.'

 telefone=telefone_whatsapp(c.driver.telefone)
 if not telefone:
  return False,'Contrato gerado, mas o motorista não possui telefone válido para envio automático.'

 integration=Integration.query.filter_by(tenant_id=c.tenant_id,tipo='whatsapp').first()
 cfg=CommunicationService.parse_config(integration)
 provider_cfg=(cfg.get('provider') or 'web').lower()
 if provider_cfg!='business':
  return False,'Contrato gerado. O envio automático requer WhatsApp Business conectado.'

 template_name=(cfg.get('contract_template_name') or '').strip() or None
 if not template_name:
  return False,'Contrato gerado, mas o template de contrato do WhatsApp não está configurado.'

 codigo_publico=garantir_codigo_publico_contrato(c)
 link=url_for('contrato_publico',codigo=codigo_publico,_external=True)
 mensagem=(f'Olá, {c.driver.nome}! Segue o contrato {c.numero_contrato} referente ao veículo '
           f'{c.vehicle.marca_modelo} - placa {c.vehicle.placa}. Clique no link para visualizar o documento oficial: {link}')
 template_parameters=[
  c.driver.nome or '',
  c.numero_contrato or '',
  c.vehicle.marca_modelo or '',
  c.vehicle.placa or '',
  link,
 ]

 fila=MessageQueue(
  tenant_id=c.tenant_id,channel='whatsapp',provider='whatsapp_business',recipient=telefone,
  recipient_name=c.driver.nome,message_type='contrato',body=mensagem,template_name=template_name,
  template_parameters=json.dumps(template_parameters,ensure_ascii=False),
  related_entity='Contrato',related_entity_id=c.id,status='PENDENTE',
  created_at=agora_sao_paulo_naive(),updated_at=agora_sao_paulo_naive(),
 )
 db.session.add(fila)
 db.session.flush()

 try:
  result=CommunicationService().send_whatsapp(
   phone=telefone,message=mensagem,integration=integration,
   template_name=template_name,
   template_language=cfg.get('template_language') or 'pt_BR',
   template_parameters=template_parameters,
  )
  fila.provider=result.provider
  fila.status=result.status
  fila.external_id=result.external_id
  fila.attempts=(fila.attempts or 0)+1
  fila.sent_at=agora_sao_paulo_naive() if result.status=='ENVIADA' else None
  fila.updated_at=agora_sao_paulo_naive()
  db.session.add(MessageEvent(
   tenant_id=c.tenant_id,message_id=fila.id,event=result.status,
   description='Contrato enviado automaticamente após a geração.',
   created_at=agora_sao_paulo_naive()
  ))

  if result.status=='ENVIADA':
   c.enviado_whatsapp_em=agora_sao_paulo_naive()
   if c.status in ('Gerado','Rascunho'):
    try:
     ContractStateService(db.session,ContractEvent,VehicleEvent).transition(
      contract=c,new_status='Enviado',user_id=user_id,now=agora_sao_paulo_naive()
     )
    except (ContractStateError,VehicleStateError):
     app.logger.exception('Contrato %s enviado, mas falhou a transição para Enviado.',c.id)
   registrar_evento_contrato(
    db.session,ContractEvent,tenant_id=c.tenant_id,contract_id=c.id,user_id=user_id,
    evento='WHATSAPP_ENVIADO_AUTOMATICO',
    descricao=f'Contrato {c.numero_contrato} enviado automaticamente para {c.driver.nome}.',
    status_novo=c.status
   )
  db.session.commit()
  if result.status=='ENVIADA':
   return True,'Contrato gerado e enviado automaticamente pelo WhatsApp.'
  return False,f'Contrato gerado, mas o WhatsApp retornou status {result.status}.'
 except CommunicationError as exc:
  fila.status='FALHA'
  fila.error_message=str(exc)
  fila.attempts=(fila.attempts or 0)+1
  fila.updated_at=agora_sao_paulo_naive()
  db.session.add(MessageEvent(
   tenant_id=c.tenant_id,message_id=fila.id,event='FALHA',
   description=str(exc),created_at=agora_sao_paulo_naive()
  ))
  db.session.commit()
  return False,f'Contrato gerado, mas o envio automático falhou: {exc}'
 except Exception:
  db.session.rollback()
  app.logger.exception('Falha inesperada no envio automático do contrato %s',c.id)
  return False,'Contrato gerado, mas ocorreu uma falha inesperada no envio automático.'


@app.route('/contratos',methods=['GET','POST'])
@login_required
def contratos():
 if request.method=='POST':
  d=Driver.query.filter_by(id=request.form['driver_id'],tenant_id=tid()).first_or_404()
  v=Vehicle.query.filter_by(id=request.form['vehicle_id'],tenant_id=tid()).first_or_404()
  t=ContractTemplate.query.filter_by(id=request.form['template_id'],tenant_id=tid(),ativo=True).first_or_404()
  data_inicio=request.form.get('data_inicio','')
  data_fim=request.form.get('data_fim','')
  try:
   prazo_dias=(datetime.strptime(data_fim,'%Y-%m-%d')-datetime.strptime(data_inicio,'%Y-%m-%d')).days
  except Exception:
   prazo_dias=90
  proprietario_nome=v.proprietario_legal or (v.investor.nome if v.investor else 'A preencher')
  proprietario_documento=v.cpf_cnpj_proprietario or (v.investor.cpf_cnpj if v.investor else 'A preencher')
  valor_locacao=request.form.get('valor_locacao') or 0
  caucao=request.form.get('caucao') or 0
  franquia=request.form.get('franquia') or 0
  valor_km=request.form.get('valor_km_excedente') or 0
  periodicidade=request.form.get('periodicidade','Semanal')
  repl={
   'gestora_nome':t.gestora_nome or '', 'gestora_fantasia':t.gestora_fantasia or '', 'gestora_cnpj':t.gestora_cnpj or '', 'gestora_endereco':t.gestora_endereco or '',
   'parceira_nome':t.parceira_nome or '', 'parceira_cnpj':t.parceira_cnpj or '', 'parceira_endereco':t.parceira_endereco or '',
   'proprietario_nome':proprietario_nome,'proprietario_documento':proprietario_documento,
   'motorista_nome':d.nome,'motorista_cpf':d.cpf or 'A preencher','motorista_rg':d.rg or 'A preencher','motorista_cnh':d.numero_cnh or 'A preencher','motorista_endereco':endereco_completo_motorista(d),
   'nacionalidade':request.form.get('nacionalidade') or 'brasileiro','estado_civil':request.form.get('estado_civil') or 'solteiro','profissao':request.form.get('profissao') or 'motorista',
   'veiculo_modelo':v.marca_modelo or 'A preencher','veiculo_cor':v.cor or 'A preencher','veiculo_ano_fabricacao':v.ano_fabricacao or 'A preencher','veiculo_ano_modelo':v.ano_modelo or 'A preencher',
   'veiculo_placa':v.placa or 'A preencher','veiculo_renavam':v.renavam or 'A preencher','km_inicial':v.km_atual or 0,
   'periodicidade':periodicidade,'periodicidade_minuscula':periodicidade.lower(),'dia_vencimento':request.form.get('dia_vencimento','segunda-feira'),
   'valor_locacao':moeda_br(valor_locacao),'valor_locacao_extenso':valor_extenso(valor_locacao),'caucao':moeda_br(caucao),'caucao_extenso':valor_extenso(caucao),
   'franquia':moeda_br(franquia),'franquia_extenso':valor_extenso(franquia),'limite_km':request.form.get('limite_km') or '','valor_km_excedente':moeda_br(valor_km),
   'multa_atraso_percentual':request.form.get('multa_atraso_percentual') or '6','juros_mes_percentual':request.form.get('juros_mes_percentual') or '1','indice_correcao':request.form.get('indice_correcao') or 'IGPM',
   'prazo_bloqueio_horas':request.form.get('prazo_bloqueio_horas') or '48','multa_diaria':moeda_br(request.form.get('multa_diaria') or 500),'taxa_adm_multas_percentual':request.form.get('taxa_adm_multas_percentual') or '20',
   'data_inicio_formatada':data_br(data_inicio),'hora_inicio':request.form.get('hora_inicio') or '09:00','data_fim_formatada':data_br(data_fim),'prazo_dias':prazo_dias,
   'cidade_assinatura':request.form.get('cidade_assinatura') or 'Carapicuíba/SP','data_assinatura_formatada':data_br(data_inicio),
  }
  # Regra operacional: um veículo só pode ter um contrato vigente por vez.
  conflito=Contract.query.filter(Contract.tenant_id==tid(),Contract.vehicle_id==v.id,Contract.status.in_(['Rascunho','Gerado','Enviado','Visualizado','Assinado','Ativo'])).first()
  if conflito or v.status not in ('Disponível',None,''):
   flash(f'O veículo {v.placa} não está disponível para um novo contrato.','danger')
   return redirect(url_for('contratos'))
  # O número definitivo depende do ID; primeiro preservamos um marcador no texto.
  repl['numero_contrato']='{{numero_contrato}}'
  texto_final=t.conteudo or ''
  for key,value in repl.items():
   texto_final=texto_final.replace('{{'+key+'}}',str(value))
  c=Contract(
   tenant_id=tid(),driver_id=d.id,vehicle_id=v.id,template_id=t.id,template_nome=t.nome,template_versao=t.versao or 1,
   data_inicio=data_inicio,hora_inicio=repl['hora_inicio'],data_fim=data_fim,periodicidade=periodicidade,dia_vencimento=repl['dia_vencimento'],
   valor_locacao=valor_locacao,caucao=caucao,franquia=franquia,limite_km=request.form.get('limite_km') or None,valor_km_excedente=valor_km,
   multa_atraso_percentual=request.form.get('multa_atraso_percentual') or 6,juros_mes_percentual=request.form.get('juros_mes_percentual') or 1,
   indice_correcao=repl['indice_correcao'],prazo_bloqueio_horas=request.form.get('prazo_bloqueio_horas') or 48,multa_diaria=request.form.get('multa_diaria') or 500,
   taxa_adm_multas_percentual=request.form.get('taxa_adm_multas_percentual') or 20,nacionalidade=repl['nacionalidade'],estado_civil=repl['estado_civil'],
   profissao=repl['profissao'],cidade_assinatura=repl['cidade_assinatura'],texto_final=texto_final,
   status='Gerado',versao=1,criado_por_id=current_user.id,criado_em=agora_sao_paulo_naive(),codigo_publico=gerar_codigo_publico_contrato(),
  )
  db.session.add(c)
  db.session.flush()
  c.numero_contrato=gerar_numero_contrato(c.id,c.criado_em)
  c.texto_final=(c.texto_final or '').replace('{{numero_contrato}}',c.numero_contrato)
  if not c.limite_km:
   c.texto_final=c.texto_final.replace('6.3. A quilometragem semanal fica limitada a  km. A quilometragem excedente será cobrada no valor de R$ '+moeda_br(c.valor_km_excedente)+' por quilômetro.','6.3. Não há limite semanal de quilometragem neste contrato.')
  codigo_publico=garantir_codigo_publico_contrato(c)
  pdf_bytes=gerar_pdf_contrato(c.numero_contrato,c.texto_final,codigo_publico=codigo_publico,url_validacao=url_for('validar_contrato_publico',codigo=codigo_publico,_external=True))
  chave_pdf=f'{tid()}/documentos/contratos/{c.numero_contrato}.pdf'
  storage.upload(BytesIO(pdf_bytes),chave_pdf,'application/pdf')
  hash_pdf=hashlib.sha256(pdf_bytes).hexdigest()
  doc=Document(tenant_id=tid(),tipo='Contrato',entidade='Contrato',entidade_id=c.id,identificador=c.numero_contrato,numero_documento=c.numero_contrato,nome_original=f'{c.numero_contrato}.pdf',arquivo=chave_pdf,hash_sha256=hash_pdf,status='Ativo',versao=c.versao or 1,criado_em=agora_sao_paulo_naive())
  db.session.add(doc); db.session.flush()
  c.documento_id=doc.id; c.arquivo_pdf=chave_pdf; c.hash_documento=hash_pdf; c.gerado_em=agora_sao_paulo_naive()
  registrar_evento_contrato(
   db.session,ContractEvent,tenant_id=tid(),contract_id=c.id,user_id=current_user.id,
   evento='CONTRATO_GERADO',descricao=f'Contrato {c.numero_contrato} gerado com o modelo {c.template_nome} v{c.template_versao}.',
   status_novo='Gerado'
  )
  try:
   VehicleStateService(db.session,VehicleEvent).reserve(
    vehicle=v,contract=c,driver=d,user_id=current_user.id,
    reason=f'Reservado pelo contrato {c.numero_contrato}.',now=agora_sao_paulo_naive()
   )
   db.session.commit()
  except (VehicleStateError,ContractStateError) as exc:
   db.session.rollback()
   flash(str(exc),'danger')
   return redirect(url_for('contratos'))
  enviado_auto,mensagem_auto=enviar_contrato_whatsapp_automatico(c,user_id=current_user.id)
  if enviado_auto:
   flash(f'Contrato {c.numero_contrato} gerado; veículo {v.placa} reservado e contrato enviado automaticamente pelo WhatsApp.','success')
  else:
   flash(f'Contrato {c.numero_contrato} gerado; veículo {v.placa} reservado.','success')
   flash(mensagem_auto,'warning')
  return redirect(url_for('contrato_detalhe',id=c.id))
 hoje=date.today(); fim=hoje+timedelta(days=90)
 q=(request.args.get('q') or '').strip()
 consulta=Contract.query.options(joinedload(Contract.driver),joinedload(Contract.vehicle)).filter(Contract.tenant_id==tid())
 if q:
  termo=f'%{q}%'
  consulta=consulta.join(Driver,Contract.driver_id==Driver.id).join(Vehicle,Contract.vehicle_id==Vehicle.id).filter(or_(
   Contract.numero_contrato.ilike(termo),Contract.status.ilike(termo),Driver.nome.ilike(termo),Driver.cpf.ilike(termo),
   Driver.numero_cnh.ilike(termo),Vehicle.placa.ilike(termo),Vehicle.marca_modelo.ilike(termo),Vehicle.renavam.ilike(termo)
  ))
 return render_template('contratos.html',items=consulta.order_by(Contract.id.desc()).all(),motoristas=Driver.query.filter_by(tenant_id=tid()).all(),veiculos=Vehicle.query.filter_by(tenant_id=tid(),status='Disponível').order_by(Vehicle.placa).all(),modelos=ContractTemplate.query.filter_by(tenant_id=tid(),ativo=True).order_by(ContractTemplate.padrao.desc(),ContractTemplate.id.desc()).all(),hoje=hoje.isoformat(),fim_padrao=fim.isoformat(),q=q)

@app.route('/contratos/<int:id>')
@login_required
def contrato_detalhe(id):
 c=Contract.query.options(joinedload(Contract.driver),joinedload(Contract.vehicle),joinedload(Contract.criado_por)).filter_by(id=id,tenant_id=tid()).first_or_404()
 eventos=ContractEvent.query.options(joinedload(ContractEvent.user)).filter_by(tenant_id=tid(),contract_id=c.id).order_by(ContractEvent.criado_em.desc()).all()
 documento=Document.query.filter_by(id=c.documento_id,tenant_id=tid()).first() if c.documento_id else None
 signature_item=_integration('signature')
 signature_cfg=_integration_config(signature_item)
 signature_ready,signature_message=SignatureProviderService.readiness(SignatureProviderService.from_integration(signature_item))
 return render_template('contrato_detalhe.html',c=c,eventos=eventos,documento=documento,signature=c.signature,signature_cfg=signature_cfg,signature_ready=signature_ready,signature_message=signature_message)


@app.route('/integracoes/clicksign/testar',methods=['POST'])
@login_required
def testar_clicksign():
 item=_integration('signature')
 try:
  client=SignatureProviderService.clicksign_client(item)
  ok,message=client.test_connection()
  flash(message,'success' if ok else 'danger')
 except SignatureProviderError as exc:
  flash(str(exc),'danger')
 return redirect(url_for('integracoes'))

@app.route('/contratos/<int:id>/clicksign/enviar',methods=['POST'])
@login_required
def contrato_clicksign_enviar(id):
 c=Contract.query.options(joinedload(Contract.driver)).filter_by(id=id,tenant_id=tid()).first_or_404()
 if not c.arquivo_pdf:
  flash('Gere o PDF do contrato antes de enviá-lo à Clicksign.','danger')
  return redirect(url_for('contrato_detalhe',id=id))
 if c.clicksign_envelope_id:
  flash('Este contrato já possui um envelope Clicksign. Use Atualizar status.','info')
  return redirect(url_for('contrato_detalhe',id=id))
 signer_email=(request.form.get('signer_email') or c.driver.email or '').strip()
 try:
  pdf_bytes=storage.download(c.arquivo_pdf)
  integration=_integration('signature')
  client=SignatureProviderService.clicksign_client(integration)
  result=client.create_signature_flow(
   envelope_name=f'Frota Fácil - {c.numero_contrato}',
   filename=f'{c.numero_contrato}.pdf',pdf_bytes=pdf_bytes,
   signer_name=c.driver.nome,signer_email=signer_email,signer_cpf=c.driver.cpf,
  )
  c.clicksign_envelope_id=result.envelope_id
  c.clicksign_document_id=result.document_id
  c.clicksign_signer_id=result.signer_id
  c.clicksign_status=result.status
  c.clicksign_sent_at=agora_sao_paulo_naive()
  if not c.driver.email and signer_email:
   c.driver.email=signer_email
  registrar_evento_contrato(db.session,ContractEvent,tenant_id=tid(),contract_id=c.id,user_id=current_user.id,
   evento='CLICKSIGN_ENVIADO',descricao=f'Contrato enviado para assinatura pela Clicksign. Envelope {result.envelope_id}.',status_novo=c.status)
  db.session.commit()
  flash(f'Contrato enviado para a Clicksign ({result.status}). Verifique o e-mail do signatário.','success')
 except StorageNotFoundError:
  flash('O PDF oficial não foi encontrado no armazenamento.','danger')
 except SignatureProviderError as exc:
  db.session.rollback(); flash(str(exc),'danger')
 except Exception:
  db.session.rollback(); app.logger.exception('Falha ao enviar contrato à Clicksign'); flash('Falha inesperada ao enviar o contrato à Clicksign.','danger')
 return redirect(url_for('contrato_detalhe',id=id))

@app.route('/contratos/<int:id>/clicksign/status',methods=['POST'])
@login_required
def contrato_clicksign_status(id):
 c=Contract.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 if not c.clicksign_envelope_id:
  flash('Este contrato ainda não possui envelope Clicksign.','danger')
  return redirect(url_for('contrato_detalhe',id=id))
 try:
  client=SignatureProviderService.clicksign_client(_integration('signature'))
  body=client.envelope_details(c.clicksign_envelope_id)
  attrs=((body.get('data') or {}).get('attributes') or {}) if isinstance(body,dict) else {}
  status=str(attrs.get('status') or c.clicksign_status or 'desconhecido')
  c.clicksign_status=status
  registrar_evento_contrato(db.session,ContractEvent,tenant_id=tid(),contract_id=c.id,user_id=current_user.id,
   evento='CLICKSIGN_STATUS',descricao=f'Status Clicksign consultado: {status}.',status_novo=c.status)
  # Quando o envelope fecha, consideramos a assinatura externa concluída.
  if status=='closed' and c.status not in ('Assinado','Ativo','Encerrado'):
   now=agora_sao_paulo_naive()
   states=ContractStateService(db.session,ContractEvent,VehicleEvent)
   states.transition(contract=c,new_status='Assinado',user_id=current_user.id,now=now)
   states.transition(contract=c,new_status='Ativo',user_id=current_user.id,now=now)
  db.session.commit(); flash(f'Status Clicksign: {status}.','success')
 except (SignatureProviderError,ContractStateError,VehicleStateError) as exc:
  db.session.rollback(); flash(str(exc),'danger')
 return redirect(url_for('contrato_detalhe',id=id))

@app.route('/contratos/<int:id>/status',methods=['POST'])
@login_required
def contrato_status(id):
 c=Contract.query.options(joinedload(Contract.vehicle),joinedload(Contract.driver)).filter_by(id=id,tenant_id=tid()).first_or_404()
 novo=(request.form.get('status') or '').strip()
 vehicle_destination=(request.form.get('vehicle_destination') or 'Disponível').strip()
 try:
  service=ContractStateService(db.session,ContractEvent,VehicleEvent)
  service.transition(
   contract=c,new_status=novo,user_id=current_user.id,now=agora_sao_paulo_naive(),
   vehicle_destination=vehicle_destination,
  )
  db.session.commit()
 except (ContractStateError,VehicleStateError) as exc:
  db.session.rollback()
  flash(str(exc),'danger')
  return redirect(url_for('contrato_detalhe',id=id))
 flash(f'Status do contrato {c.numero_contrato} atualizado para {novo}.','success')
 return redirect(url_for('contrato_detalhe',id=id))


@app.route('/contratos/<int:id>/gerar-pdf',methods=['POST'])
@login_required
def gerar_pdf_contrato_existente(id):
 c=Contract.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 if c.arquivo_pdf and c.documento_id:
  flash('O PDF deste contrato já está armazenado.','info')
  return redirect(url_for('contrato_detalhe',id=id))
 chave_pdf=f'{tid()}/documentos/contratos/{c.numero_contrato}.pdf'
 try:
  codigo_publico=garantir_codigo_publico_contrato(c)
  pdf_bytes=gerar_pdf_contrato(c.numero_contrato,c.texto_final,codigo_publico=codigo_publico,url_validacao=url_for('validar_contrato_publico',codigo=codigo_publico,_external=True))
  storage.upload(BytesIO(pdf_bytes),chave_pdf,'application/pdf')
  hash_pdf=hashlib.sha256(pdf_bytes).hexdigest()
  doc=Document(tenant_id=tid(),tipo='Contrato',entidade='Contrato',entidade_id=c.id,identificador=c.numero_contrato,numero_documento=c.numero_contrato,nome_original=f'{c.numero_contrato}.pdf',arquivo=chave_pdf,hash_sha256=hash_pdf,status='Ativo',versao=c.versao or 1,criado_em=agora_sao_paulo_naive())
  db.session.add(doc); db.session.flush()
  c.documento_id=doc.id; c.arquivo_pdf=chave_pdf; c.hash_documento=hash_pdf; c.gerado_em=agora_sao_paulo_naive()
  registrar_evento_contrato(db.session,ContractEvent,tenant_id=tid(),contract_id=c.id,user_id=current_user.id,evento='PDF_ARMAZENADO',descricao=f'PDF do contrato {c.numero_contrato} armazenado no Cloudflare R2.',status_novo=c.status)
  db.session.commit()
  flash('PDF gerado e enviado para a Central de Documentos.','success')
 except Exception:
  db.session.rollback()
  try: storage.delete(chave_pdf)
  except Exception: pass
  app.logger.exception('Falha ao gerar PDF do contrato %s',c.id)
  flash('Não foi possível gerar e armazenar o PDF.','danger')
 return redirect(url_for('contrato_detalhe',id=id))

@app.route('/contratos/<int:id>/pdf')
@login_required
def contrato_pdf(id):
 c=Contract.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 if not c.arquivo_pdf: abort(404)
 usar_assinado=bool(c.arquivo_pdf_assinado) and request.args.get('original')!='1'
 chave=c.arquivo_pdf_assinado if usar_assinado else c.arquivo_pdf
 nome=f'{c.numero_contrato}_ASSINADO.pdf' if usar_assinado else f'{c.numero_contrato}.pdf'
 try: conteudo=storage.download(chave)
 except StorageNotFoundError: abort(404)
 return send_file(BytesIO(conteudo),as_attachment=True,download_name=nome,mimetype='application/pdf')

@app.route('/contratos/<int:id>/texto')
@login_required
def contrato_texto(id):
 c=Contract.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 return send_file(BytesIO((c.texto_final or '').encode('utf-8')),as_attachment=True,download_name=f'{c.numero_contrato or ("contrato_"+str(c.id))}.txt',mimetype='text/plain; charset=utf-8')

@app.route('/contratos/<int:id>/whatsapp')
@login_required
def contrato_whatsapp(id):
 c=Contract.query.options(joinedload(Contract.driver),joinedload(Contract.vehicle)).filter_by(id=id,tenant_id=tid()).first_or_404()
 if not c.arquivo_pdf:
  flash('Gere e armazene o PDF antes de enviá-lo.','warning')
  return redirect(url_for('contrato_detalhe',id=id))
 telefone=telefone_whatsapp(c.driver.telefone)
 if not telefone:
  flash('O motorista não possui telefone válido cadastrado.','danger')
  return redirect(url_for('contrato_detalhe',id=id))
 codigo_publico=garantir_codigo_publico_contrato(c)
 link=url_for('contrato_publico',codigo=codigo_publico,_external=True)
 mensagem=(f'Olá, {c.driver.nome}! Segue o contrato {c.numero_contrato} referente ao veículo '
           f'{c.vehicle.marca_modelo} - placa {c.vehicle.placa}. Clique no link para visualizar o documento oficial: {link}')
 c.enviado_whatsapp_em=agora_sao_paulo_naive()
 try:
  if c.status in ('Gerado','Rascunho'):
   ContractStateService(db.session,ContractEvent,VehicleEvent).transition(contract=c,new_status='Enviado',user_id=current_user.id,now=agora_sao_paulo_naive())
  registrar_evento_contrato(db.session,ContractEvent,tenant_id=tid(),contract_id=c.id,user_id=current_user.id,
   evento='WHATSAPP_PREPARADO',descricao=f'Mensagem do contrato {c.numero_contrato} preparada para o WhatsApp de {c.driver.nome}.',status_novo=c.status)
  db.session.commit()
 except (ContractStateError,VehicleStateError) as exc:
  db.session.rollback(); flash(str(exc),'danger'); return redirect(url_for('contrato_detalhe',id=id))
 integration=Integration.query.filter_by(tenant_id=tid(),tipo='whatsapp').first()
 cfg=CommunicationService.parse_config(integration)
 provider_cfg=(cfg.get('provider') or 'web').lower()
 template_name=(cfg.get('contract_template_name') or '').strip() or None
 template_parameters=[
  c.driver.nome or '',
  c.numero_contrato or '',
  c.vehicle.marca_modelo or '',
  c.vehicle.placa or '',
  link,
 ]
 fila=MessageQueue(
  tenant_id=tid(),channel='whatsapp',provider='whatsapp_business' if provider_cfg=='business' else 'whatsapp_web',recipient=telefone,
  recipient_name=c.driver.nome,message_type='contrato',body=mensagem,template_name=template_name,
  template_parameters=json.dumps(template_parameters,ensure_ascii=False),
  related_entity='Contrato',related_entity_id=c.id,status='PENDENTE',
  created_at=agora_sao_paulo_naive(),updated_at=agora_sao_paulo_naive(),
 )
 db.session.add(fila); db.session.flush()
 try:
  result=CommunicationService().send_whatsapp(
   phone=telefone,
   message=mensagem,
   integration=integration,
   template_name=template_name,
   template_language=cfg.get('template_language') or 'pt_BR',
   template_parameters=template_parameters,
  )
  fila.provider=result.provider; fila.status=result.status; fila.external_id=result.external_id
  fila.attempts=(fila.attempts or 0)+1; fila.sent_at=agora_sao_paulo_naive() if result.status=='ENVIADA' else None
  db.session.add(MessageEvent(tenant_id=tid(),message_id=fila.id,event=result.status,description='Mensagem de contrato processada pelo provedor configurado.',created_at=agora_sao_paulo_naive()))
  db.session.commit()
  if result.redirect_url:
   return redirect(result.redirect_url)
  flash('Contrato enviado automaticamente pela WhatsApp Business Platform.','success')
  return redirect(url_for('contrato_detalhe',id=id))
 except CommunicationError as exc:
  fila.status='FALHA'; fila.error_message=str(exc); fila.attempts=(fila.attempts or 0)+1
  db.session.add(MessageEvent(tenant_id=tid(),message_id=fila.id,event='FALHA',description=str(exc),created_at=agora_sao_paulo_naive()))
  db.session.commit(); flash(str(exc),'danger')
  return redirect(url_for('contrato_detalhe',id=id))

@app.route('/contrato-publico/<codigo>')
def contrato_publico(codigo):
 c=Contract.query.options(joinedload(Contract.driver),joinedload(Contract.vehicle)).filter_by(codigo_publico=codigo).first_or_404()
 if not c.arquivo_pdf:
  abort(404)
 if not c.visualizado_em:
  c.visualizado_em=agora_sao_paulo_naive()
  try:
   if c.status in ('Gerado','Enviado'):
    ContractStateService(db.session,ContractEvent,VehicleEvent).transition(contract=c,new_status='Visualizado',user_id=None,now=agora_sao_paulo_naive())
   registrar_evento_contrato(db.session,ContractEvent,tenant_id=c.tenant_id,contract_id=c.id,user_id=None,
    evento='CONTRATO_PUBLICO_ABERTO',descricao=f'Página pública do contrato {c.numero_contrato} visualizada.',status_novo=c.status)
   db.session.commit()
  except (ContractStateError,VehicleStateError):
   db.session.rollback()
 documento=Document.query.filter_by(id=c.documento_id,tenant_id=c.tenant_id).first() if c.documento_id else None
 page_html=render_template('contrato_publico.html',c=c,documento=documento,signature=c.signature)
 # UX: orienta o motorista e oferece atalho direto para a área de assinatura.
 if not c.signature and c.status not in ('Assinado','Ativo','Encerrado'):
  guia_assinatura='''
  <div id="ff-signature-guide" style="position:sticky;top:0;z-index:9999;background:#fff7d6;border:1px solid #e7c84b;border-radius:12px;padding:14px 16px;margin:12px auto;max-width:980px;box-shadow:0 4px 14px rgba(0,0,0,.12);font-family:Arial,sans-serif;color:#2b2b2b">
    <div style="font-size:18px;font-weight:700;margin-bottom:5px">✍️ Para assinar seu contrato</div>
    <div style="font-size:15px;line-height:1.45">Leia o contrato e <strong>role a página até o final</strong>. O campo para fazer sua assinatura está no final do documento.</div>
    <button type="button" id="ff-go-signature" style="margin-top:10px;border:0;border-radius:9px;padding:11px 16px;font-weight:700;cursor:pointer;background:#1f6feb;color:white;font-size:15px">⬇️ Ir direto para assinatura</button>
  </div>
  <script>
  (function(){
    function assinaturaTarget(){
      var input=document.querySelector('input[name="assinatura_data"]');
      if(input && input.closest('form')) return input.closest('form');
      var form=document.querySelector('form[action*="/assinar"]');
      if(form) return form;
      var canvas=document.querySelector('canvas');
      return canvas || null;
    }
    function irParaAssinatura(){
      var alvo=assinaturaTarget();
      if(alvo){
        alvo.scrollIntoView({behavior:'smooth',block:'start'});
        setTimeout(function(){
          var campo=alvo.querySelector ? alvo.querySelector('input:not([type="hidden"]), canvas, button') : null;
          if(campo && campo.focus) campo.focus({preventScroll:true});
        },500);
      }else{
        window.scrollTo({top:document.documentElement.scrollHeight,behavior:'smooth'});
      }
    }
    document.addEventListener('DOMContentLoaded',function(){
      var btn=document.getElementById('ff-go-signature');
      if(btn) btn.addEventListener('click',irParaAssinatura);
      var form=assinaturaTarget();
      if(form){
        var dica=document.createElement('div');
        dica.style.cssText='font-weight:700;margin:0 0 10px;color:#1f2937;font-size:15px';
        dica.textContent='Assine no quadro abaixo usando o dedo ou o mouse.';
        form.insertBefore(dica,form.firstChild);
      }
    });
  })();
  </script>
  '''
  body_match=re.search(r'<body\b[^>]*>',page_html,flags=re.I)
  if body_match:
   pos=body_match.end()
   page_html=page_html[:pos]+guia_assinatura+page_html[pos:]
  else:
   page_html=guia_assinatura+page_html
 return page_html

@app.route('/contrato-publico/<codigo>/assinar',methods=['POST'])
def assinar_contrato_publico(codigo):
 c=Contract.query.options(joinedload(Contract.driver),joinedload(Contract.vehicle)).filter_by(codigo_publico=codigo).first_or_404()
 if not c.arquivo_pdf or not c.hash_documento:
  flash('O documento oficial ainda não está disponível para assinatura.','danger')
  return redirect(url_for('contrato_publico',codigo=codigo))
 if c.signature or c.status in ('Assinado','Ativo','Encerrado'):
  flash('Este contrato já foi assinado.','success')
  return redirect(url_for('contrato_publico',codigo=codigo))
 assinatura_data=request.form.get('assinatura_data','')
 nome=(request.form.get('signatario_nome') or '').strip()
 cpf=(request.form.get('cpf_confirmado') or '').strip()
 aceite=request.form.get('aceite')=='1'
 try:
  service=SignatureService(storage=storage)
  png_bytes=service.validate_and_decode(assinatura_data)
  service.validate_identity(driver=c.driver,nome=nome,cpf=cpf,aceite=aceite)
  now=agora_sao_paulo_naive()
  key=f'{c.tenant_id}/documentos/assinaturas/{c.numero_contrato}_{uuid.uuid4().hex[:10]}.png'
  storage.upload(BytesIO(png_bytes),key,'image/png')
  assinatura=Signature(
   tenant_id=c.tenant_id,contract_id=c.id,driver_id=c.driver_id,status='Assinada',
   signatario_nome=nome,cpf_confirmado=cpf,arquivo_assinatura=key,
   hash_assinatura=hashlib.sha256(png_bytes).hexdigest(),hash_documento=c.hash_documento,
   ip=service.client_ip(request),user_agent=(request.headers.get('User-Agent') or '')[:1000],
   aceite_texto='Declaro que li integralmente e concordo com o contrato apresentado.',assinado_em=now,
  )
  db.session.add(assinatura)
  db.session.flush()
  c.assinatura_id=str(assinatura.id)

  # Gera uma segunda versão do documento contendo a assinatura visual.
  # O PDF original permanece preservado para auditoria.
  assinado_em_br=now.strftime('%d/%m/%Y %H:%M')
  pdf_assinado=gerar_pdf_contrato(
   c.numero_contrato,c.texto_final,codigo_publico=c.codigo_publico,
   url_validacao=url_for('validar_contrato_publico',codigo=c.codigo_publico,_external=True),
   assinatura_png=png_bytes,
   assinatura_info={'nome':nome,'cpf':cpf,'assinado_em':assinado_em_br,'codigo':c.codigo_publico,'hash_assinatura':hashlib.sha256(png_bytes).hexdigest()},
  )
  chave_pdf_assinado=f'{c.tenant_id}/documentos/contratos/{c.numero_contrato}_ASSINADO.pdf'
  storage.upload(BytesIO(pdf_assinado),chave_pdf_assinado,'application/pdf')
  hash_pdf_assinado=hashlib.sha256(pdf_assinado).hexdigest()
  doc_assinado=Document(
   tenant_id=c.tenant_id,tipo='Contrato Assinado',entidade='Contrato',entidade_id=c.id,
   identificador=f'{c.numero_contrato} - Assinado',numero_documento=c.numero_contrato,
   nome_original=f'{c.numero_contrato}_ASSINADO.pdf',arquivo=chave_pdf_assinado,hash_sha256=hash_pdf_assinado,
   status='Ativo',versao=(c.versao or 1),criado_em=now,
  )
  db.session.add(doc_assinado); db.session.flush()
  c.arquivo_pdf_assinado=chave_pdf_assinado
  c.hash_documento_assinado=hash_pdf_assinado
  c.documento_assinado_id=doc_assinado.id

  states=ContractStateService(db.session,ContractEvent,VehicleEvent)
  states.transition(contract=c,new_status='Assinado',user_id=None,now=now)
  states.transition(contract=c,new_status='Ativo',user_id=None,now=now)
  registrar_evento_contrato(db.session,ContractEvent,tenant_id=c.tenant_id,contract_id=c.id,user_id=None,
   evento='ASSINATURA_ELETRONICA',descricao=f'Contrato assinado eletronicamente por {nome}. IP registrado: {assinatura.ip or "não informado"}.',
   status_anterior='Visualizado',status_novo='Ativo').criado_em=now
  db.session.commit()
  # Convite automático do Portal do Motorista: somente após assinatura eletrônica real.
  # Se já houver DriverAccess ativo, o helper reutiliza o acesso e envia o link de login.
  try:
   integration=Integration.query.filter_by(tenant_id=c.tenant_id,tipo='whatsapp').first()
   cfg_portal=CommunicationService.parse_config(integration)
   if c.driver and bool(cfg_portal.get('driver_portal_auto_invite_enabled',False)):
    ok_portal,msg_portal=enviar_acesso_portal_motorista_whatsapp(c.driver)
    if not ok_portal:
     app.logger.warning('Contrato %s assinado, mas convite do Portal do Motorista falhou: %s',c.id,msg_portal)
  except Exception:
   app.logger.exception('Contrato %s assinado, mas houve falha no convite automático do Portal do Motorista',c.id)
 except (SignatureValidationError,ContractStateError,VehicleStateError) as exc:
  db.session.rollback()
  try:
   if 'key' in locals(): storage.delete(key)
   if 'chave_pdf_assinado' in locals(): storage.delete(chave_pdf_assinado)
  except Exception: pass
  flash(str(exc),'danger')
  return redirect(url_for('contrato_publico',codigo=codigo))
 except Exception:
  db.session.rollback()
  try:
   if 'key' in locals(): storage.delete(key)
   if 'chave_pdf_assinado' in locals(): storage.delete(chave_pdf_assinado)
  except Exception: pass
  app.logger.exception('Falha ao assinar contrato público')
  flash('Não foi possível concluir a assinatura. Tente novamente.','danger')
  return redirect(url_for('contrato_publico',codigo=codigo))
 return redirect(url_for('contrato_publico',codigo=codigo,assinado='1'))

@app.route('/contrato-publico/<codigo>/pdf')
def contrato_publico_pdf(codigo):
 c=Contract.query.filter_by(codigo_publico=codigo).first_or_404()
 if not c.arquivo_pdf:
  abort(404)
 chave=c.arquivo_pdf_assinado or c.arquivo_pdf
 nome=f'{c.numero_contrato}_ASSINADO.pdf' if c.arquivo_pdf_assinado else f'{c.numero_contrato}.pdf'
 try:
  conteudo=storage.download(chave)
 except StorageNotFoundError:
  abort(404)
 download=request.args.get('download')=='1'
 disposition='attachment' if download else 'inline'
 response=send_file(BytesIO(conteudo),as_attachment=download,download_name=nome,mimetype='application/pdf',max_age=0)
 response.headers['Content-Disposition']=f'{disposition}; filename="{nome}"'
 response.headers['X-Content-Type-Options']='nosniff'
 response.headers['Cache-Control']='private, no-store, max-age=0'
 return response

@app.route('/validar/contrato/<codigo>')
def validar_contrato_publico(codigo):
 # Compatibilidade: links antigos enviados pelo WhatsApp apontavam diretamente
 # para esta rota. Sem o parâmetro ?verificar=1, eles agora abrem a página
 # pública completa do contrato em vez de mostrar apenas a autenticação.
 c=Contract.query.filter_by(codigo_publico=codigo).first_or_404()
 if request.args.get('verificar') != '1':
  return redirect(url_for('contrato_publico',codigo=c.codigo_publico))
 c=Contract.query.options(joinedload(Contract.driver),joinedload(Contract.vehicle)).filter_by(id=c.id).first_or_404()
 documento=Document.query.filter_by(id=c.documento_id,tenant_id=c.tenant_id).first() if c.documento_id else None
 return render_template('validar_contrato.html',c=c,documento=documento)

@app.route('/contrato/<codigo>')
def contrato_publico_alias(codigo):
 # Alias curto e estável para compartilhamento externo.
 return redirect(url_for('contrato_publico',codigo=codigo))

CONTRACT_MARKERS=[
 ('motorista_nome','Nome do locatário/motorista'),('motorista_cpf','CPF do motorista'),('motorista_rg','RG do motorista'),
 ('motorista_cnh','CNH do motorista'),('motorista_endereco','Endereço do motorista'),('nacionalidade','Nacionalidade do motorista'),('estado_civil','Estado civil do motorista'),('profissao','Profissão do motorista'),('proprietario_nome','Nome do proprietário'),
 ('proprietario_documento','CPF/CNPJ do proprietário'),('veiculo_modelo','Modelo do veículo'),('veiculo_cor','Cor do veículo'),
 ('veiculo_ano_fabricacao','Ano de fabricação'),('veiculo_ano_modelo','Ano modelo'),('veiculo_placa','Placa'),('veiculo_renavam','Renavam'),
 ('km_inicial','KM inicial'),('periodicidade','Periodicidade'),('periodicidade_minuscula','Periodicidade em minúsculas'),
 ('valor_locacao','Valor da locação'),('valor_locacao_extenso','Valor da locação por extenso'),
 ('caucao','Caução'),('caucao_extenso','Caução por extenso'),('franquia','Franquia'),('franquia_extenso','Franquia por extenso'),
 ('limite_km','Limite de KM'),('valor_km_excedente','Valor por KM excedente'),('data_inicio_formatada','Data inicial'),
 ('data_fim_formatada','Data final'),('hora_inicio','Hora inicial'),('dia_vencimento','Dia de vencimento'),('cidade_assinatura','Cidade de assinatura'),
 ('data_assinatura_formatada','Data da assinatura'),('prazo_dias','Prazo em dias'),
 ('gestora_nome','Razão social da gestora'),('gestora_fantasia','Nome fantasia da gestora'),('gestora_cnpj','CNPJ da gestora'),
 ('gestora_endereco','Endereço da gestora'),('parceira_nome','Razão social da parceira'),('parceira_cnpj','CNPJ da parceira'),
 ('parceira_endereco','Endereço da parceira'),
]

def extrair_texto_contrato_bytes(data,nome_original):
 ext=Path(nome_original or '').suffix.lower()
 if ext=='.txt': return data.decode('utf-8',errors='replace').strip()
 if ext=='.docx':
  doc=DocxDocument(BytesIO(data)); blocos=[]
  for par in doc.paragraphs:
   if par.text.strip(): blocos.append(par.text.strip())
  for tabela in doc.tables:
   for linha in tabela.rows:
    celulas=[c.text.strip() for c in linha.cells]
    if any(celulas): blocos.append(' | '.join(celulas))
  return '\n'.join(blocos).strip()
 if ext=='.pdf':
  f=FileStorage(stream=BytesIO(data),filename=nome_original,content_type='application/pdf')
  return (extract_text(f,document_type='contract') or '').strip()
 raise ValueError('Formato não suportado. Envie DOCX, PDF ou TXT.')

def _sub_contexto(texto,pattern,marker,flags=re.I):
 marcador='{{'+marker+'}}'; encontrados=0
 def repl(m):
  nonlocal encontrados
  alvo=m.group(2) if m.lastindex and m.lastindex>=2 else ''
  # Não conta nem substitui novamente um campo que já virou marcador.
  if '{{' in alvo and '}}' in alvo:
   return m.group(0)
  encontrados+=1
  prefix=m.group(1) if m.lastindex and m.lastindex>=1 else ''
  suffix=m.group(3) if m.lastindex and m.lastindex>=3 else ''
  return prefix+marcador+suffix
 novo=re.sub(pattern,repl,texto,flags=flags)
 return novo,encontrados

def preparar_contrato_com_marcadores(texto):
 """Prepara contratos preenchidos por blocos semânticos e depois aplica fallbacks globais.

 A ordem importa: primeiro PARTES, depois VEÍCULO, FINANCEIRO e por último regras globais.
 Isso reduz falsos positivos como CPF do locatário confundido com documento do proprietário.
 """
 t=html.unescape(texto or '').replace('\r\n','\n').replace('\r','\n').replace('\xa0',' ')
 # Artefatos frequentes de extração de PDF/Word.
 for a,b in [(' Ɵ','ti'),('Ɵ','ti'),('ﬁ','fi'),('ﬂ','fl'),('–','–')]:
  t=t.replace(a,b)
 t=re.sub(r'[ \t]+\n','\n',t)
 t=re.sub(r'\n{3,}','\n\n',t)
 detectados=[]

 def registrar(marker,n):
  if n: detectados.append({'marker':marker,'quantidade':n})

 def aplicar(pattern,marker,flags=re.I):
  nonlocal t
  t,n=_sub_contexto(t,pattern,marker,flags=flags)
  registrar(marker,n); return n

 def aplicar_bloco(inicio,fim,regras,flags=re.I):
  """Aplica regras somente dentro de um bloco identificado por cabeçalho/contexto."""
  nonlocal t
  m=re.search(inicio,t,flags)
  if not m: return 0
  pos_ini=m.start()
  resto=t[m.end():]
  mf=re.search(fim,resto,flags) if fim else None
  pos_fim=m.end()+(mf.start() if mf else min(len(resto),1600))
  bloco=t[pos_ini:pos_fim]
  total=0
  for pattern,marker in regras:
   bloco,n=_sub_contexto(bloco,pattern,marker,flags=flags)
   registrar(marker,n); total+=n
  t=t[:pos_ini]+bloco+t[pos_fim:]
  return total

 # 1) BLOCO GESTORA/LOCADORA
 aplicar_bloco(
  r'(?:GESTORA\s+DA\s+LOCA[CÇ][AÃ]O|GESTORA|LOCADORA)\s*:',
  r'\n\s*(?:PROPRIET[ÁA]RIO|PARCEIRA\s+OPERACIONAL|LOCAT[ÁA]RIO|LOCATARIO)\b',
  [
   (r'((?:GESTORA\s+DA\s+LOCA[CÇ][AÃ]O|GESTORA|LOCADORA)\s*:\s*)([^\n,;]+)(?=\s*[,;\n])','gestora_nome'),
   (r'((?:nome\s+fantasia)\s+)([^,;\n]{2,120})(?=\s*[,;])','gestora_fantasia'),
   (r'((?:CNPJ)(?:/MF)?(?:\s*(?:sob\s+o\s+)?(?:n[ºo°]|número|numero|:))?\s*)([0-9.\-/]{14,22})(\b)','gestora_cnpj'),
   (r'((?:com\s+endere[cç]o(?:\s+comercial)?\s+(?:em|localizado\s+em)?\s*))([\s\S]{8,320}?)(,?\s*doravante|\.(?=\s*(?:\n|$)))','gestora_endereco'),
  ])

 # 2) BLOCO PROPRIETÁRIO
 aplicar_bloco(
  r'PROPRIET[ÁA]RIO(?:\s+DO\s+VE[ÍI]CULO)?\s*:',
  r'\n\s*(?:PARCEIRA\s+OPERACIONAL|LOCAT[ÁA]RIO|LOCATARIO|DO\s+OBJETO|CL[ÁA]USULA)\b',
  [
   (r'((?:PROPRIET[ÁA]RIO(?:\s+DO\s+VE[ÍI]CULO)?\s*:\s*))([^\n,;]+)(?=\s*[,;\n])','proprietario_nome'),
   (r'((?:CPF/CNPJ|CNPJ|CPF)(?:/MF)?(?:\s*(?:sob\s+o\s+)?(?:n[ºo°]|número|numero|:))?\s*)([0-9.\-/]{11,22})(\b)','proprietario_documento'),
  ])

 # 3) BLOCO PARCEIRA OPERACIONAL
 aplicar_bloco(
  r'PARCEIRA\s+OPERACIONAL[^:\n]*:',
  r'\n\s*(?:LOCAT[ÁA]RIO|LOCATARIO|DO\s+OBJETO|CL[ÁA]USULA)\b',
  [
   (r'((?:PARCEIRA\s+OPERACIONAL[^:\n]*:\s*))([^\n,;]+)(?=\s*[,;\n])','parceira_nome'),
   (r'((?:CNPJ)(?:/MF)?(?:\s*(?:sob\s+o\s+)?(?:n[ºo°]|número|numero|:))?\s*)([0-9.\-/]{14,22})(\b)','parceira_cnpj'),
   (r'((?:com\s+endere[cç]o(?:\s+comercial)?\s+(?:em|localizado\s+em)?\s*))([\s\S]{8,320}?)(,?\s*doravante|\.(?=\s*(?:\n|$)))','parceira_endereco'),
  ])

 # 4) BLOCO LOCATÁRIO/MOTORISTA. Captura inclusive endereço quebrado em linhas.
 aplicar_bloco(
  r'(?:LOCAT[ÁA]RIO|LOCATARIO|MOTORISTA)\s*:',
  r'\n\s*(?:DO\s+OBJETO|OBJETO\s+DO\s+CONTRATO|CL[ÁA]USULA\s+1|VE[ÍI]CULO\s*:|DADOS\s+DO\s+VE[ÍI]CULO)\b',
  [
   (r'((?:LOCAT[ÁA]RIO|LOCATARIO|MOTORISTA)\s*:\s*)([^\n,;]+)(?=\s*[,;\n])','motorista_nome'),
   (r'((?:LOCAT[ÁA]RIO|LOCATARIO|MOTORISTA)\s*:\s*\{\{motorista_nome\}\}\s*,?\s*)(brasileir[oa]|estrangeir[oa]|portugu[eê]s|portuguesa)(?=\s*[,;])','nacionalidade'),
   (r'((?:\{\{nacionalidade\}\}|brasileir[oa]|estrangeir[oa]|portugu[eê]s|portuguesa)\s*,?\s*)(solteir[oa]|casad[oa]|divorciad[oa]|vi[uú]v[oa]|separad[oa])(?=\s*[,;])','estado_civil'),
   (r'((?:\{\{estado_civil\}\}|solteir[oa]|casad[oa]|divorciad[oa]|vi[uú]v[oa]|separad[oa])\s*,?\s*)([^,;\n]{3,80})(?=\s*,\s*(?:RG|portador))','profissao'),
   (r'((?:RG)(?:\s*(?:sob\s+o\s+)?(?:n[ºo°]|número|numero|:))?\s*)([0-9.\-Xx]{5,20})(\b)','motorista_rg'),
   (r'((?:CPF(?:/MF)?)(?:\s*(?:sob\s+o\s+)?(?:n[ºo°]|número|numero|:))?\s*)([0-9.\-/]{11,20})(\b)','motorista_cpf'),
   (r'((?:CNH)(?:\s*(?:sob\s+o\s+)?(?:n[ºo°]|número|numero|:))?\s*)([0-9.\-]{8,20})(\b)','motorista_cnh'),
   (r'((?:residente\s+e\s+domiciliad[oa]\s+(?:em|à|a)\s+))([\s\S]{8,520}?)(,?\s*doravante|\.(?=\s*(?:\n|$)))','motorista_endereco'),
   (r'((?:ENDERE[CÇ]O(?:\s+DO\s+(?:LOCAT[ÁA]RIO|MOTORISTA))?\s*:\s*))([\s\S]{8,420}?)(?=\n(?:CPF|RG|CNH|VE[ÍI]CULO|PLACA|CL[ÁA]USULA)\b|\.(?:\s*\n|$))','motorista_endereco'),
  ])

 # 5) BLOCO VEÍCULO / OBJETO
 aplicar_bloco(
  r'(?:DO\s+OBJETO|OBJETO\s+DO\s+CONTRATO|VE[ÍI]CULO\s*:|autom[oó]vel\s+descrito\s+a\s+seguir)',
  r'\n\s*(?:DO\s+VALOR|DA\s+COBRAN[CÇ]A|CL[ÁA]USULA\s+2|CL[ÁA]USULA\s+3|PAGAMENTO)\b',
  [
   (r'((?:MARCA\s*/?\s*MODELO|MODELO|VE[ÍI]CULO)\s*:\s*)([^\n|;,]{2,140})(?=\s*(?:[|;,\n]|$))','veiculo_modelo'),
   (r'((?:objeto\s+(?:do|deste)\s+contrato[^:]{0,220}:\s*))([\s\S]{2,180}?)(?=\s+COR\s*:)','veiculo_modelo'),
   (r'((?:autom[oó]vel\s+descrito\s+a\s+seguir\s*:\s*))([\s\S]{2,180}?)(?=\s+COR\s*:)','veiculo_modelo'),
   (r'((?:COR)\s*:\s*)([^\n|;,]{2,40}?)(?=\s*(?:ANO|[|;,\n]|$))','veiculo_cor'),
   (r'((?:ANO(?:\s+DE\s+FABRICA[CÇ][AÃ]O)?|FABRICA[CÇ][AÃ]O)\s*:\s*)([0-9]{4})(\b)','veiculo_ano_fabricacao'),
   (r'((?:ANO\s+MODELO|MODELO\s+ANO)\s*:\s*)([0-9]{4})(\b)','veiculo_ano_modelo'),
   (r'((?:PLACA)\s*:\s*)([A-Z]{3}[0-9][A-Z0-9][0-9]{2})(\b)','veiculo_placa'),
   (r'((?:RENAVAM)\s*:\s*)([0-9.\-]{7,20})(\b)','veiculo_renavam'),
   (r'((?:QUILOMETRAGEM\s+INICIAL|KM\s+INICIAL)\s*:\s*)([0-9.]{1,12})(\s*km\b)?','km_inicial'),
  ])

 # 6) BLOCO FINANCEIRO
 aplicar_bloco(
  r'(?:DO\s+VALOR\s+DO\s+ALUGUEL|DA\s+COBRAN[CÇ]A|COBRAN[CÇ]A\s+E\s+GEST[AÃ]O\s+FINANCEIRA|pagará|pagara|CAU[CÇ][AÃ]O)',
  r'\n\s*(?:CL[ÁA]USULA\s+[4-9]|DAS\s+OBRIGA[CÇ][ÕO]ES|DA\s+MANUTEN[CÇ][AÃ]O|DO\s+SEGURO)\b',
  [
   (r'((?:VALOR\s+(?:SEMANAL|MENSAL|DA\s+LOCA[CÇ][AÃ]O|DO\s+ALUGUEL)|ALUGUEL)\s*(?::|DE)?\s*R?\$?\s*)([0-9.]+(?:,[0-9]{2})?)(\b)','valor_locacao'),
   (r'((?:quantia|valor)\s+(?:semanal|mensal|di[áa]ri[oa])\s+(?:de|no\s+valor\s+de)\s+R?\$?\s*)([0-9.]+(?:,[0-9]{2})?)(\b)','valor_locacao'),
   (r'((?:pagará|pagara)[^\n]{0,130}?R?\$\s*)([0-9.]+(?:,[0-9]{2})?)(\b)','valor_locacao'),
   (r'((?:CAU[CÇ][AÃ]O)[^\n]{0,150}?R?\$\s*)([0-9.]+(?:,[0-9]{2})?)(\b)','caucao'),
   (r'((?:FRANQUIA)[^\n]{0,150}?R?\$\s*)([0-9.]+(?:,[0-9]{2})?)(\b)','franquia'),
   (r'((?:LIMITE(?:\s+SEMANAL)?(?:\s+DE)?\s+KM|QUILOMETRAGEM\s+SEMANAL)[^0-9\n]{0,50})([0-9.]{1,10})(\s*km\b)','limite_km'),
   (r'((?:KM\s+EXCEDENTE|QUIL[ÔO]METRO\s+EXCEDENTE)[^\nR$]{0,80}R?\$?\s*)([0-9.]+(?:,[0-9]{2})?)(\b)','valor_km_excedente'),
   (r'((?:quantia|valor)\s+)(semanal|mensal|di[áa]ri[oa])(\s+(?:de|no\s+valor))','periodicidade_minuscula'),
   (r'((?:periodicidade\s*:\s*))(semanal|mensal|di[áa]ri[oa])(\b)','periodicidade'),
   (r'((?:sempre\s+(?:às|as|à|a)\s+|vencimento\s+(?:toda|todo|às|as|à|a)?\s*))((?:segundas?|ter[cç]as?|quartas?|quintas?|sextas?|s[áa]bados?|domingos?)(?:-feiras?)?)(\b)','dia_vencimento'),
  ])

 # Valores por extenso logo após marcadores monetários.
 for base_marker,extenso_marker in [('valor_locacao','valor_locacao_extenso'),('caucao','caucao_extenso'),('franquia','franquia_extenso')]:
  mb='{{'+base_marker+'}}'; me='{{'+extenso_marker+'}}'
  pattern=r'('+re.escape(mb)+r'\s*\()([\s\S]{3,180}?)(\))'
  t,n=re.subn(pattern,lambda m:m.group(1)+me+m.group(3),t,flags=re.I)
  registrar(extenso_marker,n)

 # 7) FALLBACKS GLOBAIS apenas para campos bem rotulados que ainda possam estar fora dos blocos.
 for pattern,marker in [
  (r'((?:PLACA)\s*:\s*)([A-Z]{3}[0-9][A-Z0-9][0-9]{2})(\b)','veiculo_placa'),
  (r'((?:RENAVAM)\s*:\s*)([0-9.\-]{7,20})(\b)','veiculo_renavam'),
  (r'((?:CNH)\s*(?:n[ºo°]|número|numero|:)?\s*)([0-9.\-]{8,20})(\b)','motorista_cnh'),
 ]:
  marcador='{{'+marker+'}}'
  if marcador not in t: aplicar(pattern,marker)

 # Consolida e mantém a ordem em que cada marcador foi encontrado.
 consolidados={}; ordem=[]
 for item in detectados:
  k=item['marker']
  if k not in consolidados: ordem.append(k); consolidados[k]=0
  consolidados[k]+=item['quantidade']
 return t,[{'marker':k,'quantidade':consolidados[k]} for k in ordem]

def salvar_original_modelo(data,nome_original):
 nome=secure_filename(nome_original or 'contrato')
 chave=f'{tid()}/modelos_contrato/originais/{uuid.uuid4().hex}_{nome}'
 storage.upload(BytesIO(data),chave,'application/octet-stream')
 return chave,hashlib.sha256(data).hexdigest()

@app.route('/modelos-contrato')
@login_required
def modelos_contrato():
 return render_template('modelos_contrato.html',items=ContractTemplate.query.filter_by(tenant_id=tid()).order_by(ContractTemplate.nome,ContractTemplate.versao.desc()).all())

@app.route('/modelos-contrato/locadrivers',methods=['POST'])
@login_required
def instalar_locadrivers():
 model=ensure_locadrivers_template(tid())
 db.session.commit()
 flash(f'Modelo Locadrivers Completo v{model.versao} instalado e definido como padrão.','success')
 return redirect(url_for('modelos_contrato'))

@app.route('/modelos-contrato/novo',methods=['GET','POST'])
@login_required
def novo_modelo_contrato():
 if request.method=='POST':
  return salvar_modelo_contrato(None)
 return render_template('modelo_contrato_form.html',modelo=None,conteudo_preparado='',marcadores=CONTRACT_MARKERS,detectados=[])

@app.route('/modelos-contrato/<int:id>/editar',methods=['GET','POST'])
@login_required
def editar_modelo_contrato(id):
 modelo=ContractTemplate.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 if request.method=='POST':
  return salvar_modelo_contrato(modelo)
 return render_template('modelo_contrato_form.html',modelo=modelo,conteudo_preparado=modelo.conteudo or '',marcadores=CONTRACT_MARKERS,detectados=[])

def salvar_modelo_contrato(modelo):
 acao=(request.form.get('acao') or 'salvar').strip()
 conteudo=request.form.get('conteudo','').strip()
 arquivo=request.files.get('arquivo')
 nome_original=(request.form.get('_nome_original') or '').strip() or None
 arquivo_original=(request.form.get('_arquivo_original') or '').strip() or None
 hash_original=(request.form.get('_hash_original') or '').strip() or None
 detectados=[]
 if arquivo and arquivo.filename:
  nome_original=secure_filename(arquivo.filename); data=arquivo.read()
  try:
   extraido=extrair_texto_contrato_bytes(data,nome_original)
   if not extraido: raise ValueError('O arquivo não contém texto legível.')
   arquivo_original,hash_original=salvar_original_modelo(data,nome_original)
   conteudo,detectados=preparar_contrato_com_marcadores(extraido)
  except Exception as exc:
   app.logger.exception('Falha ao preparar modelo de contrato')
   flash(f'Não foi possível preparar o contrato: {exc}','warning')
   return render_template('modelo_contrato_form.html',modelo=modelo,conteudo_preparado=conteudo,marcadores=CONTRACT_MARKERS,detectados=[])
 if acao=='preparar':
  if not conteudo: flash('Selecione um DOCX ou PDF preenchido para preparar.','warning')
  else: flash(f'Contrato analisado. {sum(x["quantidade"] for x in detectados)} campo(s) convertido(s) em marcadores. Revise antes de salvar.','success')
  return render_template('modelo_contrato_form.html',modelo=modelo,conteudo_preparado=conteudo,marcadores=CONTRACT_MARKERS,detectados=detectados,nome_original=nome_original,arquivo_original=arquivo_original,hash_original=hash_original)
 if not conteudo:
  flash('Informe o conteúdo do contrato ou use “Analisar e preparar contrato”.','danger')
  return render_template('modelo_contrato_form.html',modelo=modelo,conteudo_preparado=conteudo,marcadores=CONTRACT_MARKERS,detectados=detectados,nome_original=nome_original,arquivo_original=arquivo_original,hash_original=hash_original)
 versao=(modelo.versao or 1)+1 if modelo else 1
 novo=ContractTemplate(tenant_id=tid(),nome=request.form['nome'].strip(),descricao=request.form.get('descricao'),versao=versao,padrao=bool(request.form.get('padrao')),tipo_veiculo=request.form.get('tipo_veiculo','Todos'),possui_limite_km=bool(request.form.get('possui_limite_km')),conteudo=conteudo,nome_original=nome_original or (modelo.nome_original if modelo else None),arquivo_original=arquivo_original or (modelo.arquivo_original if modelo else None),hash_original=hash_original or (modelo.hash_original if modelo else None),preparado_em=agora_sao_paulo_naive() if arquivo_original else (modelo.preparado_em if modelo else None),gestora_nome=request.form.get('gestora_nome'),gestora_fantasia=request.form.get('gestora_fantasia'),gestora_cnpj=request.form.get('gestora_cnpj'),gestora_endereco=request.form.get('gestora_endereco'),parceira_nome=request.form.get('parceira_nome'),parceira_cnpj=request.form.get('parceira_cnpj'),parceira_endereco=request.form.get('parceira_endereco'),ativo=True)
 if novo.padrao: ContractTemplate.query.filter_by(tenant_id=tid(),padrao=True).update({'padrao':False},synchronize_session=False)
 db.session.add(novo); db.session.commit(); flash('Nova versão do modelo salva.','success')
 return redirect(url_for('modelos_contrato'))

@app.route('/modelos-contrato/<int:id>/alternar',methods=['POST'])
@login_required
def alternar_modelo_contrato(id):
 modelo=ContractTemplate.query.filter_by(id=id,tenant_id=tid()).first_or_404(); modelo.ativo=not modelo.ativo
 if not modelo.ativo: modelo.padrao=False
 db.session.commit(); return redirect(url_for('modelos_contrato'))

@app.route('/modelos-contrato/<int:id>/padrao',methods=['POST'])
@login_required
def definir_modelo_padrao(id):
 modelo=ContractTemplate.query.filter_by(id=id,tenant_id=tid(),ativo=True).first_or_404()
 ContractTemplate.query.filter_by(tenant_id=tid(),padrao=True).update({'padrao':False},synchronize_session=False); modelo.padrao=True; db.session.commit()
 return redirect(url_for('modelos_contrato'))

@app.route('/documentos',methods=['GET','POST'])
@login_required
def documentos():
 if request.method=='POST':
  f=request.files.get('arquivo')
  if not f or not f.filename:
   flash('Selecione um arquivo.','danger')
   return redirect(url_for('documentos'))
  nome_original=secure_filename(f.filename)
  if not nome_original:
   flash('Nome de arquivo inválido.','danger')
   return redirect(url_for('documentos'))
  chave=f'{tid()}/documentos/{uuid.uuid4().hex}_{nome_original}'
  try:
   conteudo=f.read()
   storage.upload(BytesIO(conteudo),chave,f.mimetype)
   entidade_id=request.form.get('entidade_id') or None
   identificador=identificador_documento(request.form['tipo'],entidade_id or 0,nome_original)
   db.session.add(Document(tenant_id=tid(),tipo=request.form['tipo'],entidade=request.form['entidade'],entidade_id=entidade_id,identificador=identificador,nome_original=nome_original,arquivo=chave,hash_sha256=hashlib.sha256(conteudo).hexdigest(),status='Ativo'))
   db.session.commit()
  except Exception:
   db.session.rollback()
   app.logger.exception('Falha ao armazenar documento')
   try: storage.delete(chave)
   except Exception: pass
   flash('Não foi possível armazenar o documento. Verifique a configuração do Cloudflare R2.','danger')
   return redirect(url_for('documentos'))
  flash('Documento armazenado de forma persistente.','success')
  return redirect(url_for('documentos'))
 q=(request.args.get('q') or '').strip()
 consulta=Document.query.filter_by(tenant_id=tid())
 contratos_relacionados={c.id:c for c in Contract.query.options(joinedload(Contract.driver),joinedload(Contract.vehicle)).filter_by(tenant_id=tid()).all()}
 if q:
  termo=f'%{q}%'
  ids_contratos=[c.id for c in contratos_relacionados.values() if any(q.lower() in (str(v or '').lower()) for v in [c.numero_contrato,c.driver.nome if c.driver else '',c.driver.cpf if c.driver else '',c.vehicle.placa if c.vehicle else '',c.vehicle.marca_modelo if c.vehicle else ''])]
  filtros=[Document.identificador.ilike(termo),Document.numero_documento.ilike(termo),Document.nome_original.ilike(termo),Document.tipo.ilike(termo),Document.entidade.ilike(termo)]
  if ids_contratos: filtros.append(db.and_(Document.entidade=='Contrato',Document.entidade_id.in_(ids_contratos)))
  consulta=consulta.filter(or_(*filtros))
 return render_template('documentos.html',items=consulta.order_by(Document.id.desc()).all(),motoristas=Driver.query.filter_by(tenant_id=tid()).all(),veiculos=Vehicle.query.filter_by(tenant_id=tid()).all(),storage_backend=storage.backend_name,q=q,contratos_relacionados=contratos_relacionados)

@app.route('/documentos/<int:id>/baixar')
@login_required
def baixar_documento(id):
 d=Document.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 try:
  conteudo=storage.download(d.arquivo)
 except StorageNotFoundError:
  abort(404)
 except Exception:
  app.logger.exception('Falha ao baixar documento %s',d.id)
  abort(503)
 return send_file(BytesIO(conteudo),as_attachment=True,download_name=d.nome_original,mimetype='application/octet-stream')

@app.route('/documentos/<int:id>/excluir',methods=['POST'])
@login_required
def excluir_documento(id):
 d=Document.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 try:
  storage.delete(d.arquivo)
  db.session.delete(d)
  db.session.commit()
  flash('Documento excluído.','success')
 except Exception:
  db.session.rollback()
  app.logger.exception('Falha ao excluir documento %s',d.id)
  flash('Não foi possível excluir o documento.','danger')
 return redirect(url_for('documentos'))



def _documentos_manutencao_query(tenant_id, maintenance_id):
 return Document.query.filter_by(tenant_id=tenant_id,entidade='Manutenção',entidade_id=maintenance_id,status='Ativo').order_by(Document.criado_em.desc(),Document.id.desc())

def _mimetype_documento(nome):
 tipo,_=mimetypes.guess_type(nome or '')
 return tipo or 'application/octet-stream'

@app.route('/manutencoes/comprovantes',methods=['GET','POST'])
@login_required
def comprovantes_manutencao():
 if request.method=='POST':
  maintenance_id=request.form.get('maintenance_id',type=int)
  m=Maintenance.query.filter_by(id=maintenance_id,tenant_id=tid()).first_or_404()
  tipo=(request.form.get('tipo') or 'Comprovante de manutenção').strip()[:40]
  arquivos=[f for f in request.files.getlist('arquivos') if f and f.filename]
  if not arquivos:
   flash('Selecione pelo menos um comprovante.','danger')
   return redirect(url_for('comprovantes_manutencao',maintenance_id=m.id))
  permitidas={'.pdf','.jpg','.jpeg','.png','.webp'}
  salvos=[]; enviados=[]
  try:
   for f in arquivos:
    nome_original=secure_filename(f.filename)
    ext=Path(nome_original).suffix.lower()
    if not nome_original or ext not in permitidas:
     raise ValueError('Envie somente arquivos PDF, JPG, JPEG, PNG ou WEBP.')
    conteudo=f.read()
    if not conteudo:
     raise ValueError(f'O arquivo {nome_original} está vazio.')
    if len(conteudo)>15*1024*1024:
     raise ValueError(f'O arquivo {nome_original} excede o limite de 15 MB.')
    chave=f'{tid()}/manutencoes/{m.id}/documentos/{uuid.uuid4().hex}_{nome_original}'
    storage.upload(BytesIO(conteudo),chave,f.mimetype or _mimetype_documento(nome_original))
    enviados.append(chave)
    doc=Document(tenant_id=tid(),tipo=tipo,entidade='Manutenção',entidade_id=m.id,identificador=identificador_documento(tipo,m.id,nome_original),nome_original=nome_original,arquivo=chave,hash_sha256=hashlib.sha256(conteudo).hexdigest(),status='Ativo')
    db.session.add(doc); salvos.append(doc)
   db.session.commit()
   flash(f'{len(salvos)} comprovante(s) anexado(s) à manutenção.','success')
  except ValueError as exc:
   db.session.rollback()
   for chave in enviados:
    try: storage.delete(chave)
    except Exception: pass
   flash(str(exc),'danger')
  except Exception:
   db.session.rollback()
   for chave in enviados:
    try: storage.delete(chave)
    except Exception: pass
   app.logger.exception('Falha ao anexar comprovantes à manutenção %s',m.id)
   flash('Não foi possível armazenar os comprovantes.','danger')
  return redirect(url_for('comprovantes_manutencao',maintenance_id=m.id))
 items=Maintenance.query.options(joinedload(Maintenance.vehicle)).filter_by(tenant_id=tid()).order_by(Maintenance.id.desc()).all()
 ids=[m.id for m in items]
 docs=Document.query.filter(Document.tenant_id==tid(),Document.entidade=='Manutenção',Document.entidade_id.in_(ids or [-1]),Document.status=='Ativo').order_by(Document.criado_em.desc(),Document.id.desc()).all()
 docs_por_manutencao={}
 for doc in docs: docs_por_manutencao.setdefault(doc.entidade_id,[]).append(doc)
 return render_template('comprovantes_manutencao.html',items=items,docs_por_manutencao=docs_por_manutencao,selecionada=request.args.get('maintenance_id',type=int))

@app.route('/manutencoes/<int:maintenance_id>/documentos/<int:document_id>')
@login_required
def visualizar_documento_manutencao(maintenance_id,document_id):
 Maintenance.query.filter_by(id=maintenance_id,tenant_id=tid()).first_or_404()
 doc=Document.query.filter_by(id=document_id,tenant_id=tid(),entidade='Manutenção',entidade_id=maintenance_id,status='Ativo').first_or_404()
 try: conteudo=storage.download(doc.arquivo)
 except StorageNotFoundError: abort(404)
 except Exception:
  app.logger.exception('Falha ao abrir comprovante de manutenção %s',doc.id); abort(503)
 return send_file(BytesIO(conteudo),as_attachment=False,download_name=doc.nome_original,mimetype=_mimetype_documento(doc.nome_original))

@app.route('/manutencoes/<int:maintenance_id>/documentos/<int:document_id>/excluir',methods=['POST'])
@login_required
def excluir_documento_manutencao(maintenance_id,document_id):
 Maintenance.query.filter_by(id=maintenance_id,tenant_id=tid()).first_or_404()
 doc=Document.query.filter_by(id=document_id,tenant_id=tid(),entidade='Manutenção',entidade_id=maintenance_id,status='Ativo').first_or_404()
 try:
  storage.delete(doc.arquivo); db.session.delete(doc); db.session.commit(); flash('Comprovante excluído.','success')
 except Exception:
  db.session.rollback(); app.logger.exception('Falha ao excluir comprovante de manutenção %s',doc.id); flash('Não foi possível excluir o comprovante.','danger')
 return redirect(url_for('comprovantes_manutencao',maintenance_id=maintenance_id))

@app.route('/manutencoes',methods=['GET','POST'])
@login_required
def manutencoes():
 if request.method=='POST':
  custo_txt=(request.form.get('custo') or '').strip()
  try:
   custo=Decimal(custo_txt.replace('.','').replace(',','.')) if custo_txt else Decimal('0')
  except Exception:
   flash('Custo inválido. Use, por exemplo, 350,00.','danger'); return redirect(url_for('manutencoes'))
  situacao=(request.form.get('situacao') or 'Agendada').strip()
  ja_realizada=(situacao == 'Realizada')
  m=Maintenance(
   tenant_id=tid(),vehicle_id=request.form['vehicle_id'],tipo=request.form['tipo'],
   data=request.form.get('data'),km=request.form.get('km') or None,custo=custo,
   proxima_km=request.form.get('proxima_km') or None,proxima_data=request.form.get('proxima_data'),proxima_hora=(request.form.get('proxima_hora') or '').strip() or None,
   alerta_km_antes=request.form.get('alerta_km_antes') or 500,
   alerta_dias_antes=request.form.get('alerta_dias_antes') or 7,
   observacoes=request.form.get('observacoes'),oficina=(request.form.get('oficina') or '').strip() or None,
   notificar_motorista=(False if ja_realizada else bool(request.form.get('notificar_motorista'))),lembrete_um_dia=(False if ja_realizada else bool(request.form.get('lembrete_um_dia'))),
   status=('Concluída' if ja_realizada else 'Ativa'),
   concluida_em=(datetime.utcnow() if ja_realizada else None),
   concluida_por_id=(current_user.id if ja_realizada else None),
  )
  db.session.add(m); db.session.flush()
  redirect_whatsapp=None
  notification_warning=None
  if ja_realizada and (m.proxima_km or m.proxima_data):
   prox=Maintenance(tenant_id=tid(),vehicle_id=m.vehicle_id,tipo=m.tipo,status='Ativa',proxima_km=m.proxima_km,proxima_data=m.proxima_data,proxima_hora=m.proxima_hora,alerta_km_antes=m.alerta_km_antes or 500,alerta_dias_antes=m.alerta_dias_antes or 7,observacoes='Próximo ciclo criado a partir de manutenção histórica.',oficina=m.oficina,notificar_motorista=False,lembrete_um_dia=False)
   db.session.add(prox)
   m.proxima_km=None; m.proxima_data=None; m.proxima_hora=None
  if ja_realizada:
   v=Vehicle.query.filter_by(id=m.vehicle_id,tenant_id=tid()).first()
   descricao=f"{m.tipo or 'Manutenção'} cadastrada como já realizada"
   if m.data: descricao+=f" em {data_br(m.data)}"
   if m.km is not None: descricao+=f" com {int(m.km):,} km".replace(',','.')
   if m.custo: descricao+=f". Custo R$ {m.custo:.2f}".replace('.',',')
   if m.oficina: descricao+=f". Oficina: {m.oficina}"
   if v:
    db.session.add(VehicleEvent(tenant_id=tid(),vehicle_id=v.id,user_id=current_user.id,evento='Manutenção histórica',descricao=descricao,status_anterior=v.status,status_novo=v.status))
  if (not ja_realizada) and m.notificar_motorista:
   v=Vehicle.query.filter_by(id=m.vehicle_id,tenant_id=tid()).first()
   d=motorista_atual_veiculo(v) if v else None
   if not d:
    notification_warning='Manutenção salva, mas não há motorista vinculado a contrato vigente para receber o WhatsApp.'
   else:
    integration=Integration.query.filter_by(tenant_id=tid(),tipo='whatsapp').first()
    cfg=CommunicationService.parse_config(integration)
    template_name=(cfg.get('maintenance_template_name') or '').strip() or None
    params=[d.nome,v.marca_modelo or 'Veículo',v.placa,m.tipo or 'Manutenção',data_br(m.proxima_data) if m.proxima_data else 'a definir',m.proxima_hora or 'a definir']
    body=maintenance_message(driver_name=d.nome,vehicle=v,maintenance=m,reminder=False)
    fila,redirect_whatsapp,err=criar_mensagem_whatsapp(tenant_id=tid(),driver=d,body=body,message_type='manutencao_agendada',related_entity='Manutencao',related_entity_id=m.id,template_name=template_name,template_parameters=params)
    if fila: m.notificacao_agendamento_id=fila.id
    if err: notification_warning='Manutenção salva, mas o envio imediato falhou: '+err
    if m.lembrete_um_dia and m.proxima_data:
     reminder_at=reminder_datetime(m)
     if reminder_at and reminder_at>agora_sao_paulo_naive():
      reminder_template=(cfg.get('maintenance_reminder_template_name') or template_name or '').strip() or None
      reminder_body=maintenance_message(driver_name=d.nome,vehicle=v,maintenance=m,reminder=True)
      rfila,_,_=criar_mensagem_whatsapp(tenant_id=tid(),driver=d,body=reminder_body,message_type='lembrete_manutencao',related_entity='Manutencao',related_entity_id=m.id,scheduled_at=reminder_at,template_name=reminder_template,template_parameters=params)
      if rfila: m.notificacao_lembrete_id=rfila.id
  db.session.commit()
  recalcular_alertas(tid())
  flash(('Manutenção já realizada registrada no histórico do veículo.' if ja_realizada else 'Manutenção registrada e monitoramento de alertas ativado.'),'success')
  if notification_warning: flash(notification_warning,'warning')
  if redirect_whatsapp:
   return redirect(redirect_whatsapp)
  return redirect(url_for('manutencoes'))
 # Processa automaticamente a fila vencida quando o gestor acessa o módulo; em produção o mesmo serviço pode ser chamado por cron.
 try: processar_mensagens_agendadas(tid(),limit=50)
 except Exception: app.logger.exception('Falha ao processar lembretes agendados')
 recalcular_alertas(tid())
 items=Maintenance.query.filter_by(tenant_id=tid()).order_by(Maintenance.id.desc()).all()
 alerts=Alert.query.filter(Alert.tenant_id==tid(),Alert.resolvido_em.is_(None),Alert.entidade=='Manutenção').order_by(Alert.criado_em.desc()).all()
 hoje_sp=datetime.now(ZoneInfo('America/Sao_Paulo')).date().isoformat()
 return render_template('manutencoes.html',items=items,veiculos=Vehicle.query.filter_by(tenant_id=tid()).all(),alertas=alerts,maintenance_indicator=maintenance_indicator,hoje_sp=hoje_sp)

@app.route('/manutencoes/<int:id>/concluir',methods=['POST'])
@login_required
def concluir_manutencao(id):
 m=Maintenance.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 if (m.status or 'Ativa') == 'Concluída':
  flash('Esta manutenção já foi concluída.','info')
  return redirect(url_for('manutencoes') + f'#manutencao-{m.id}')

 v=Vehicle.query.filter_by(id=m.vehicle_id,tenant_id=tid()).first_or_404()
 data_realizada=(request.form.get('data_realizada') or datetime.now(ZoneInfo('America/Sao_Paulo')).date().isoformat()).strip()
 km_realizada=request.form.get('km_realizada') or v.km_atual or None
 custo_txt=(request.form.get('custo_realizado') or '').strip()
 try:
  custo_realizado=Decimal(custo_txt.replace('.','').replace(',','.')) if custo_txt else (m.custo or Decimal('0'))
 except Exception:
  flash('Custo inválido. Use, por exemplo, 350,00.','danger')
  return redirect(url_for('manutencoes') + f'#manutencao-{m.id}')

 oficina=(request.form.get('oficina') or '').strip()
 obs_conclusao=(request.form.get('observacoes_conclusao') or '').strip()
 m.data=data_realizada
 m.km=int(km_realizada) if km_realizada not in (None,'') else None
 m.custo=custo_realizado
 m.oficina=oficina or m.oficina
 if obs_conclusao:
  m.observacoes=((m.observacoes + '\n') if m.observacoes else '') + 'Conclusão: ' + obs_conclusao
 m.status='Concluída'
 m.concluida_em=datetime.utcnow()
 m.concluida_por_id=current_user.id

 # A previsão antiga pertence à manutenção concluída e não deve continuar gerando alerta.
 m.proxima_km=None
 m.proxima_data=None

 descricao=f"{m.tipo or 'Manutenção'} concluída em {data_realizada}"
 if m.km is not None:
  descricao+=f" com {m.km:,} km".replace(',','.')
 if custo_realizado:
  descricao+=f". Custo R$ {custo_realizado:.2f}".replace('.',',')
 if oficina:
  descricao+=f". Oficina: {oficina}"
 if obs_conclusao:
  descricao+=f". {obs_conclusao}"
 db.session.add(VehicleEvent(tenant_id=tid(),vehicle_id=v.id,user_id=current_user.id,evento='Manutenção concluída',descricao=descricao,status_anterior=v.status,status_novo=v.status))

 # Se o usuário informar uma nova previsão, cria um novo ciclo sem apagar o histórico.
 nova_km=request.form.get('nova_proxima_km') or None
 nova_data=(request.form.get('nova_proxima_data') or '').strip() or None
 if nova_km or nova_data:
  prox=Maintenance(
   tenant_id=tid(),vehicle_id=v.id,tipo=m.tipo,status='Ativa',
   proxima_km=int(nova_km) if nova_km else None,proxima_data=nova_data,proxima_hora=(request.form.get('nova_proxima_hora') or m.proxima_hora or '').strip() or None,
   alerta_km_antes=request.form.get('novo_alerta_km_antes') or m.alerta_km_antes or 500,
   alerta_dias_antes=request.form.get('novo_alerta_dias_antes') or m.alerta_dias_antes or 7,
   observacoes='Gerada automaticamente após conclusão da manutenção anterior.',
   notificar_motorista=m.notificar_motorista,lembrete_um_dia=m.lembrete_um_dia,
  )
  db.session.add(prox)

 db.session.commit()
 recalcular_alertas(tid())
 flash('Manutenção concluída e registrada no histórico do veículo.','success')
 return redirect(url_for('manutencoes') + f'#manutencao-{m.id}')

@app.route('/manutencoes/gerenciar')
@login_required
def gerenciar_manutencoes():
 items=Maintenance.query.options(joinedload(Maintenance.vehicle)).filter_by(tenant_id=tid()).order_by(Maintenance.id.desc()).all()
 ids=[m.id for m in items]
 docs=Document.query.filter(Document.tenant_id==tid(),Document.entidade=='Manutenção',Document.entidade_id.in_(ids or [-1]),Document.status=='Ativo').all()
 docs_count={}
 for d in docs: docs_count[d.entidade_id]=docs_count.get(d.entidade_id,0)+1
 return render_template('gerenciar_manutencoes.html',items=items,docs_count=docs_count)

@app.route('/manutencoes/<int:id>/cancelar',methods=['POST'])
@login_required
def cancelar_manutencao(id):
 m=Maintenance.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 status=(m.status or 'Ativa').strip()
 if status=='Concluída':
  flash('Uma manutenção concluída não pode ser cancelada.','warning'); return redirect(url_for('gerenciar_manutencoes'))
 if status=='Cancelada':
  flash('Esta manutenção já está cancelada.','info'); return redirect(url_for('gerenciar_manutencoes'))
 motivo=(request.form.get('motivo') or '').strip()[:500]
 agora=agora_sao_paulo_naive()
 m.status='Cancelada'; m.notificar_motorista=False; m.lembrete_um_dia=False
 if motivo:
  m.observacoes=((m.observacoes+'\n') if m.observacoes else '')+'Cancelamento: '+motivo
 # Cancela lembretes ainda não enviados. Mensagens já entregues permanecem no histórico.
 MessageQueue.query.filter(
  MessageQueue.tenant_id==tid(),MessageQueue.related_entity=='Manutencao',MessageQueue.related_entity_id==m.id,
  MessageQueue.status.in_(['AGENDADA','PENDENTE'])
 ).update({'status':'CANCELADA','updated_at':agora},synchronize_session=False)
 Alert.query.filter(
  Alert.tenant_id==tid(),Alert.entidade=='Manutenção',Alert.entidade_id==m.id,Alert.resolvido_em.is_(None)
 ).update({'resolvido_em':agora},synchronize_session=False)
 if m.vehicle_id:
  v=Vehicle.query.filter_by(id=m.vehicle_id,tenant_id=tid()).first()
  if v:
   descricao=f"{m.tipo or 'Manutenção'} cancelada"+(f". Motivo: {motivo}" if motivo else '')
   db.session.add(VehicleEvent(tenant_id=tid(),vehicle_id=v.id,user_id=current_user.id,evento='Manutenção cancelada',descricao=descricao,status_anterior=v.status,status_novo=v.status))
 db.session.commit(); flash('Manutenção cancelada. Os lembretes pendentes foram interrompidos.','success')
 return redirect(url_for('gerenciar_manutencoes'))

@app.route('/manutencoes/<int:id>/excluir',methods=['POST'])
@login_required
def excluir_manutencao(id):
 m=Maintenance.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 if (m.status or '').strip()=='Concluída':
  flash('Manutenção concluída faz parte do histórico e não pode ser excluída.','warning'); return redirect(url_for('gerenciar_manutencoes'))
 if _documentos_manutencao_query(tid(),m.id).first():
  flash('Esta manutenção possui comprovantes. Exclua os comprovantes antes de excluir o cadastro.','warning'); return redirect(url_for('gerenciar_manutencoes'))
 agora=agora_sao_paulo_naive()
 MessageQueue.query.filter(
  MessageQueue.tenant_id==tid(),MessageQueue.related_entity=='Manutencao',MessageQueue.related_entity_id==m.id,
  MessageQueue.status.in_(['AGENDADA','PENDENTE'])
 ).update({'status':'CANCELADA','updated_at':agora},synchronize_session=False)
 Alert.query.filter(
  Alert.tenant_id==tid(),Alert.entidade=='Manutenção',Alert.entidade_id==m.id,Alert.resolvido_em.is_(None)
 ).update({'resolvido_em':agora},synchronize_session=False)
 db.session.delete(m); db.session.commit(); flash('Cadastro de manutenção excluído.','success')
 return redirect(url_for('gerenciar_manutencoes'))

def enviar_vistoria_whatsapp_automatico(item):
 """Envia automaticamente o link da vistoria usando o template aprovado da Meta.

 Parâmetros do template:
 1 motorista, 2 locadora, 3 veículo, 4 placa, 5 link da vistoria.
 A criação/regravação da vistoria nunca é desfeita se o WhatsApp falhar.
 """
 if not item or not item.driver or not item.vehicle:
  return False,'Vistoria criada, mas faltam dados do motorista ou veículo para o envio.'
 telefone=normalize_phone(item.driver.telefone)
 if not telefone:
  return False,'Vistoria criada, mas o motorista não possui telefone/WhatsApp válido.'
 integration=Integration.query.filter_by(tenant_id=item.tenant_id,tipo='whatsapp').first()
 cfg=CommunicationService.parse_config(integration)
 if (cfg.get('provider') or 'web').lower()!='business':
  return False,'Vistoria criada. O envio automático requer WhatsApp Business conectado.'
 template_name=(cfg.get('inspection_template_name') or '').strip() or None
 if not template_name:
  return False,'Vistoria criada, mas o template de vistoria do WhatsApp não está configurado.'
 link=url_for('vistoria_publica',token=item.token,_external=True)
 tenant=Tenant.query.filter_by(id=item.tenant_id).first()
 nome_locadora=(
  (((tenant.nome_fantasia if tenant else '') or '').strip())
  or (((tenant.nome if tenant else '') or '').strip())
  or 'Locadora'
 )
 template_parameters=[
  item.driver.nome or '',
  nome_locadora,
  item.vehicle.marca_modelo or '',
  item.vehicle.placa or '',
  link,
 ]
 mensagem=(f'Olá, {item.driver.nome}. A locadora {nome_locadora} solicita uma vistoria do veículo '
           f'{item.vehicle.marca_modelo}, placa {item.vehicle.placa}. '
           f'Para realizar a vistoria, acesse este link: {link} e siga as instruções exibidas na tela.')
 fila=MessageQueue(
  tenant_id=item.tenant_id,channel='whatsapp',provider='whatsapp_business',recipient=telefone,
  recipient_name=item.driver.nome,message_type='vistoria',body=mensagem,template_name=template_name,
  template_parameters=json.dumps(template_parameters,ensure_ascii=False),
  related_entity='Vistoria',related_entity_id=item.id,status='PENDENTE',
  created_at=agora_sao_paulo_naive(),updated_at=agora_sao_paulo_naive(),
 )
 db.session.add(fila); db.session.flush()
 try:
  result=CommunicationService().send_whatsapp(
   phone=telefone,message=mensagem,integration=integration,
   template_name=template_name,
   template_language=cfg.get('template_language') or 'pt_BR',
   template_parameters=template_parameters,
  )
  fila.provider=result.provider; fila.status=result.status; fila.external_id=result.external_id
  fila.attempts=(fila.attempts or 0)+1
  fila.sent_at=agora_sao_paulo_naive() if result.status=='ENVIADA' else None
  fila.updated_at=agora_sao_paulo_naive()
  db.session.add(MessageEvent(
   tenant_id=item.tenant_id,message_id=fila.id,event=result.status,
   description='Link da vistoria enviado automaticamente pelo WhatsApp.',
   created_at=agora_sao_paulo_naive()
  ))
  db.session.commit()
  if result.status=='ENVIADA':
   return True,'Vistoria criada e link enviado automaticamente pelo WhatsApp.'
  return False,f'Vistoria criada, mas o WhatsApp retornou status {result.status}.'
 except CommunicationError as exc:
  fila.status='FALHA'; fila.error_message=str(exc); fila.attempts=(fila.attempts or 0)+1; fila.updated_at=agora_sao_paulo_naive()
  db.session.add(MessageEvent(tenant_id=item.tenant_id,message_id=fila.id,event='FALHA',description=str(exc),created_at=agora_sao_paulo_naive()))
  db.session.commit()
  return False,f'Vistoria criada, mas o envio automático falhou: {exc}'
 except Exception:
  db.session.rollback()
  app.logger.exception('Falha inesperada no envio automático da vistoria %s',item.id)
  return False,'Vistoria criada, mas ocorreu uma falha inesperada no envio automático.'



def _extrair_frames_vistoria(video_bytes,suffix='.webm',quantidade=6):
 """Extrai poucos frames representativos sem guardar cópia permanente do vídeo."""
 frames=[]
 with tempfile.TemporaryDirectory(prefix='ff_vistoria_') as tmp:
  origem=Path(tmp)/('video'+suffix); origem.write_bytes(video_bytes)
  padrao=str(Path(tmp)/'frame-%02d.jpg')
  cmd=['ffmpeg','-hide_banner','-loglevel','error','-i',str(origem),'-vf',f'fps=1/5,scale=960:-2','-frames:v',str(quantidade),'-q:v','4',padrao]
  subprocess.run(cmd,check=True,timeout=25,stdout=subprocess.DEVNULL,stderr=subprocess.PIPE)
  for arq in sorted(Path(tmp).glob('frame-*.jpg'))[:quantidade]: frames.append(arq.read_bytes())
 return frames

def _imagem_vistoria_para_data_url(key,mime=None):
 """Baixa uma evidência da vistoria e devolve data URL para análise visual."""
 conteudo=storage.download(key)
 tipo=(mime or '').strip().lower()
 if tipo not in ('image/jpeg','image/png','image/webp'):
  ext=Path(key or '').suffix.lower()
  tipo='image/png' if ext=='.png' else ('image/webp' if ext=='.webp' else 'image/jpeg')
 return 'data:'+tipo+';base64,'+base64.b64encode(conteudo).decode('ascii')

def _vistoria_fotografica_anterior(item):
 """Retorna a vistoria fotográfica anterior do mesmo veículo, quando houver evidências completas."""
 return (Inspection.query
  .filter(Inspection.tenant_id==item.tenant_id,
          Inspection.vehicle_id==item.vehicle_id,
          Inspection.id < item.id,
          Inspection.tipo_vistoria=='fotos',
          Inspection.front_photo_key.isnot(None),
          Inspection.right_photo_key.isnot(None),
          Inspection.rear_photo_key.isnot(None),
          Inspection.left_photo_key.isnot(None))
  .order_by(Inspection.id.desc())
  .first())

def _resultado_json_responses(data):
 texto=''
 for out in data.get('output',[]):
  for c in out.get('content',[]):
   if c.get('type')=='output_text': texto+=c.get('text','')
 texto=texto.strip()
 if texto.startswith('```json'): texto=texto[7:]
 elif texto.startswith('```'): texto=texto[3:]
 if texto.endswith('```'): texto=texto[:-3]
 return json.loads(texto.strip())

def analisar_avarias_vistoria(item):
 """Triagem assistida. Não aprova/reprova a vistoria e nunca substitui a revisão da locadora."""
 integration=Integration.query.filter_by(tenant_id=item.tenant_id,tipo='whatsapp').first()
 cfg=CommunicationService.parse_config(integration)
 if not cfg.get('inspection_damage_detection_enabled',False):
  item.damage_analysis_status='DESATIVADA'; item.damage_analysis_summary='Análise assistida de avarias desativada nas configurações.'; return False
 api_key=(os.getenv('OPENAI_API_KEY') or '').strip()
 if not api_key:
  item.damage_analysis_status='AGUARDANDO_CONFIGURACAO'; item.damage_analysis_summary='Detecção de avarias ativada, mas o provedor de visão ainda não possui credencial configurada.'; return False
 try:
  modo=(item.tipo_vistoria or '').strip().lower()
  if modo=='fotos':
   atuais=[
    ('FRENTE',item.front_photo_key,item.front_photo_mime),
    ('LATERAL DIREITA',item.right_photo_key,item.right_photo_mime),
    ('TRASEIRA',item.rear_photo_key,item.rear_photo_mime),
    ('LATERAL ESQUERDA',item.left_photo_key,item.left_photo_mime),
   ]
   if not all(k for _,k,_ in atuais): raise RuntimeError('As quatro fotos externas não estão completas.')
   anterior=_vistoria_fotografica_anterior(item)
   content=[{'type':'input_text','text':(
    'Você faz TRIAGEM VISUAL ASSISTIDA de vistoria veicular. A decisão final é humana. '
    'Analise separadamente FRENTE, LATERAL DIREITA, TRASEIRA e LATERAL ESQUERDA. '
    'Procure somente indícios visíveis de avaria externa como amassado, risco relevante, trinca, peça quebrada, desalinhamento, dano em roda, para-choque, farol ou lanterna. '
    'Não invente danos e não trate reflexo, sujeira, sombra ou baixa qualidade como avaria confirmada. '
    'Quando houver fotos ANTERIORES da mesma posição, compare anterior x atual e dê prioridade a POSSÍVEIS AVARIAS NOVAS. '
    'Responda SOMENTE JSON válido no formato: '
    '{"level":"SEM_INDICIOS|POSSIVEL_AVARIA|REVISAO_RECOMENDADA","summary":"resumo curto em português",'
    '"sides":{"frente":{"level":"...","detail":"..."},"direita":{"level":"...","detail":"..."},"traseira":{"level":"...","detail":"..."},"esquerda":{"level":"...","detail":"..."}}}. '
    'Use REVISAO_RECOMENDADA quando a imagem não permitir avaliação confiável.')}]
   if anterior:
    anteriores=[
     ('FRENTE',anterior.front_photo_key,anterior.front_photo_mime),
     ('LATERAL DIREITA',anterior.right_photo_key,anterior.right_photo_mime),
     ('TRASEIRA',anterior.rear_photo_key,anterior.rear_photo_mime),
     ('LATERAL ESQUERDA',anterior.left_photo_key,anterior.left_photo_mime),
    ]
    content.append({'type':'input_text','text':f'VISTORIA ANTERIOR de referência (ID {anterior.id}). As quatro imagens abaixo são ANTERIORES, na ordem indicada.'})
    for rotulo,key,mime in anteriores:
     content.append({'type':'input_text','text':'ANTERIOR — '+rotulo})
     content.append({'type':'input_image','image_url':_imagem_vistoria_para_data_url(key,mime)})
   else:
    content.append({'type':'input_text','text':'Não existe vistoria fotográfica anterior completa. Faça análise isolada, sem afirmar que um dano é novo.'})
   content.append({'type':'input_text','text':'VISTORIA ATUAL. As quatro imagens abaixo são ATUAIS, na ordem indicada.'})
   for rotulo,key,mime in atuais:
    content.append({'type':'input_text','text':'ATUAL — '+rotulo})
    content.append({'type':'input_image','image_url':_imagem_vistoria_para_data_url(key,mime)})
   payload={'model':(os.getenv('FROTA_FACIL_VISION_MODEL') or 'gpt-5.6-luna').strip(),'input':[{'role':'user','content':content}], 'max_output_tokens':700}
   resp=requests.post('https://api.openai.com/v1/responses',headers={'Authorization':f'Bearer {api_key}','Content-Type':'application/json'},json=payload,timeout=55)
   resp.raise_for_status(); result=_resultado_json_responses(resp.json())
   level=str(result.get('level') or 'REVISAO_RECOMENDADA').upper()
   if level not in {'SEM_INDICIOS','POSSIVEL_AVARIA','REVISAO_RECOMENDADA'}: level='REVISAO_RECOMENDADA'
   sides=result.get('sides') if isinstance(result.get('sides'),dict) else {}
   nomes=[('frente','Frente'),('direita','Direita'),('traseira','Traseira'),('esquerda','Esquerda')]
   linhas=[]
   resumo=str(result.get('summary') or '').strip()
   if resumo: linhas.append(resumo)
   for chave,nome in nomes:
    dados=sides.get(chave) if isinstance(sides.get(chave),dict) else {}
    detalhe=str(dados.get('detail') or '').strip()
    nivel=str(dados.get('level') or '').upper().strip()
    if detalhe: linhas.append(f'{nome}: {detalhe}' + (f' [{nivel}]' if nivel else ''))
   linhas.append('Comparação com vistoria anterior: '+('sim (ID '+str(anterior.id)+').' if anterior else 'não havia referência fotográfica anterior completa.'))
   item.damage_analysis_status='CONCLUIDA'; item.damage_analysis_level=level; item.damage_analysis_summary='\n'.join(linhas)[:4000]; item.damage_analysis_at=datetime.utcnow()
   return True

  video_bytes=storage.download(item.video_key)
  suffix=Path(item.video_key or '').suffix.lower() or '.webm'
  frames=_extrair_frames_vistoria(video_bytes,suffix=suffix,quantidade=6)
  if not frames: raise RuntimeError('Nenhum frame pôde ser extraído do vídeo.')
  content=[{'type':'input_text','text':('Você faz triagem visual de vistorias de veículos. Analise os frames do MESMO vídeo e procure apenas indícios visíveis de avaria externa: amassado, risco relevante, peça quebrada, trinca, desalinhamento ou dano em roda/para-choque/farol. Não invente dano quando a imagem não permitir concluir. Responda SOMENTE JSON válido com level igual a SEM_INDICIOS, POSSIVEL_AVARIA ou REVISAO_RECOMENDADA e summary em português, curto, citando a região suspeita. A decisão final é humana.')}]
  for frame in frames:
   content.append({'type':'input_image','image_url':'data:image/jpeg;base64,'+base64.b64encode(frame).decode('ascii')})
  payload={'model':(os.getenv('FROTA_FACIL_VISION_MODEL') or 'gpt-5.6-luna').strip(),'input':[{'role':'user','content':content}], 'max_output_tokens':350}
  resp=requests.post('https://api.openai.com/v1/responses',headers={'Authorization':f'Bearer {api_key}','Content-Type':'application/json'},json=payload,timeout=35)
  resp.raise_for_status(); result=_resultado_json_responses(resp.json())
  level=str(result.get('level') or 'REVISAO_RECOMENDADA').upper()
  if level not in {'SEM_INDICIOS','POSSIVEL_AVARIA','REVISAO_RECOMENDADA'}: level='REVISAO_RECOMENDADA'
  item.damage_analysis_status='CONCLUIDA'; item.damage_analysis_level=level; item.damage_analysis_summary=str(result.get('summary') or '')[:1500]; item.damage_analysis_at=datetime.utcnow()
  return True
 except Exception:
  app.logger.exception('Falha na triagem assistida de avarias da vistoria %s',item.id)
  item.damage_analysis_status='FALHA'; item.damage_analysis_level='REVISAO_RECOMENDADA'; item.damage_analysis_summary='A análise automática não pôde ser concluída. Faça a revisão visual das evidências da vistoria.'; item.damage_analysis_at=datetime.utcnow()
  return False

@app.route('/vistorias',methods=['GET','POST'])
@login_required
def vistorias():
 if request.method=='POST':
  v=Vehicle.query.filter_by(id=request.form.get('vehicle_id'),tenant_id=tid()).first_or_404()
  d=motorista_atual_veiculo(v)
  c=v.current_contract if v.current_contract and v.current_contract.tenant_id==tid() else None
  if not d:
   flash('O veículo precisa ter motorista vinculado a um contrato vigente para gerar a vistoria.','warning')
   return redirect(url_for('vistorias'))
  token=uuid.uuid4().hex+uuid.uuid4().hex[:8]
  expira_horas=int(request.form.get('expira_horas') or 48)
  tipo_vistoria=(request.form.get('tipo_vistoria') or 'fotos').strip().lower()
  if tipo_vistoria not in ('fotos','guiada','simples'):
   tipo_vistoria='fotos'
  item=Inspection(tenant_id=tid(),vehicle_id=v.id,driver_id=d.id,contract_id=(c.id if c else None),token=token,status='Pendente',tipo_vistoria=tipo_vistoria,expires_at=datetime.utcnow()+timedelta(hours=max(1,min(expira_horas,168))))
  db.session.add(item); db.session.commit()
  ok_envio,msg_envio=enviar_vistoria_whatsapp_automatico(item)
  flash(msg_envio,'success' if ok_envio else 'warning')
  return redirect(url_for('vistorias'))
 items=Inspection.query.filter_by(tenant_id=tid()).order_by(Inspection.id.desc()).all()
 # Corrige também pendências que já existiam antes desta versão.
 reconciliados=0
 for _concluida in items:
  if _concluida.status in ('Concluída','Concluida','Aprovada') and (_concluida.submitted_at or _concluida.status=='Aprovada'):
   reconciliados+=_encerrar_vistorias_anteriores(_concluida)
 if reconciliados:
  db.session.commit()
 # O template legado usa strftime diretamente. Localizamos apenas a cópia em memória
 # para exibição; o banco continua preservando requested_at em UTC.
 from sqlalchemy.orm.attributes import set_committed_value
 for _item in items:
  if _item.requested_at:
   _local=_as_tenant_time(_item.requested_at,_item.tenant_id)
   if _local:
    set_committed_value(_item,'requested_at',_local.replace(tzinfo=None))
 veiculos=Vehicle.query.filter_by(tenant_id=tid()).order_by(Vehicle.placa).all()
 return render_template('vistorias.html',items=items,veiculos=veiculos)

@app.route('/vistorias/<int:id>/aprovar',methods=['POST'])
@login_required
def aprovar_vistoria(id):
 item=Inspection.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 evidencia_fotos=((item.tipo_vistoria or '')=='fotos' and item.front_photo_key and item.right_photo_key and item.rear_photo_key and item.left_photo_key and item.painel_photo_key)
 if not item.video_key and not evidencia_fotos:
  flash('A vistoria ainda não possui todas as evidências obrigatórias.','warning'); return redirect(url_for('vistorias'))
 status_anterior=item.status
 item.status='Aprovada'
 _encerrar_vistorias_anteriores(item)
 tentativa=InspectionAttempt.query.filter_by(inspection_id=item.id,tenant_id=tid(),decision='Pendente').order_by(InspectionAttempt.id.desc()).first()
 if not tentativa and item.video_key:
  tentativa=InspectionAttempt(inspection_id=item.id,tenant_id=item.tenant_id,video_key=item.video_key,video_mime=item.video_mime,duration_seconds=item.duration_seconds,painel_photo_key=item.painel_photo_key,painel_photo_mime=item.painel_photo_mime,km_informada=item.km_informada,brightness_avg=item.brightness_avg,submitted_at=item.submitted_at or datetime.utcnow(),decision='Pendente')
  db.session.add(tentativa)
 if tentativa:
  tentativa.decision='Aprovada'; tentativa.decided_at=datetime.utcnow()
 duracao_txt=f'{item.duration_seconds}s' if item.duration_seconds is not None else 'não informada'
 km_txt=(f'{item.km_informada:,} km'.replace(',','.') if item.km_informada is not None else 'não informada')
 descricao=(f'Vistoria fotográfica aprovada pela locadora; 4 lados e painel anexados; KM registrada: {km_txt}.' if (item.tipo_vistoria or '')=='fotos' else f'Vistoria em vídeo aprovada pela locadora; duração {duracao_txt}; foto do painel anexada; KM registrada: {km_txt}.')
 # Corrige também o evento de recebimento já gravado, para que o histórico não continue dizendo
 # 'Aguardando aprovação' depois que a vistoria foi decidida.
 evento_recebimento=VehicleEvent.query.filter_by(
  tenant_id=item.tenant_id,vehicle_id=item.vehicle_id,contract_id=item.contract_id,
  driver_id=item.driver_id,evento='Vistoria em vídeo recebida'
 ).order_by(VehicleEvent.id.desc()).first()
 if evento_recebimento and 'Aguardando aprovação' in (evento_recebimento.descricao or ''):
  evento_recebimento.descricao=(evento_recebimento.descricao or '').replace(
   'Aguardando aprovação.','Posteriormente aprovada pela locadora.'
  ).replace('Aguardando aprovação','Posteriormente aprovada pela locadora')
 db.session.add(VehicleEvent(
  tenant_id=item.tenant_id,
  vehicle_id=item.vehicle_id,
  contract_id=item.contract_id,
  driver_id=item.driver_id,
  user_id=current_user.id,
  evento=('Vistoria fotográfica aprovada' if (item.tipo_vistoria or '')=='fotos' else 'Vistoria em vídeo aprovada'),
  descricao=descricao,
  status_anterior=status_anterior,
  status_novo='Aprovada'
 ))
 db.session.commit()
 flash('Vistoria aprovada.','success'); return redirect(url_for('vistorias'))

@app.route('/vistorias/<int:id>/rejeitar',methods=['POST'])
@login_required
def rejeitar_vistoria(id):
 item=Inspection.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 item.status='Pendente'; item.notes=(request.form.get('motivo') or 'Nova vistoria solicitada pela locadora.').strip()
 tentativa=InspectionAttempt.query.filter_by(inspection_id=item.id,tenant_id=tid(),decision='Pendente').order_by(InspectionAttempt.id.desc()).first()
 if not tentativa and item.video_key:
  tentativa=InspectionAttempt(inspection_id=item.id,tenant_id=item.tenant_id,video_key=item.video_key,video_mime=item.video_mime,duration_seconds=item.duration_seconds,painel_photo_key=item.painel_photo_key,painel_photo_mime=item.painel_photo_mime,km_informada=item.km_informada,brightness_avg=item.brightness_avg,submitted_at=item.submitted_at or datetime.utcnow(),decision='Pendente')
  db.session.add(tentativa)
 if tentativa:
  tentativa.decision='Regravar'; tentativa.decision_notes=item.notes; tentativa.decided_at=datetime.utcnow()
 item.token=uuid.uuid4().hex+uuid.uuid4().hex[:8]
 item.requested_at=datetime.utcnow()
 item.expires_at=datetime.utcnow()+timedelta(hours=48)
 db.session.commit()
 ok_envio,msg_envio=enviar_vistoria_whatsapp_automatico(item)
 if ok_envio:
  flash('Nova vistoria solicitada. O link foi enviado automaticamente pelo WhatsApp.','success')
 else:
  flash('Nova vistoria solicitada e novo link gerado. '+msg_envio,'warning')
 return redirect(url_for('vistorias'))

@app.route('/vistorias/<int:id>/video')
@login_required
def vistoria_video(id):
 item=Inspection.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 if not item.video_key: abort(404)
 try: conteudo=storage.download(item.video_key)
 except StorageNotFoundError: abort(404)
 return send_file(BytesIO(conteudo),mimetype=item.video_mime or 'video/webm',download_name=f'vistoria-{item.id}.webm',conditional=True)

@app.route('/vistorias/<int:id>/painel')
@login_required
def vistoria_painel(id):
 item=Inspection.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 if not item.painel_photo_key: abort(404)
 try: conteudo=storage.download(item.painel_photo_key)
 except StorageNotFoundError: abort(404)
 return send_file(BytesIO(conteudo),mimetype=item.painel_photo_mime or 'image/jpeg',download_name=f'vistoria-{item.id}-painel.jpg',conditional=True)

@app.route('/vistorias/<int:id>/foto/<posicao>')
@login_required
def vistoria_foto(id,posicao):
 item=Inspection.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 mapa={
  'frente':('front_photo_key','front_photo_mime','frente'),
  'direita':('right_photo_key','right_photo_mime','lateral-direita'),
  'traseira':('rear_photo_key','rear_photo_mime','traseira'),
  'esquerda':('left_photo_key','left_photo_mime','lateral-esquerda'),
  'painel':('painel_photo_key','painel_photo_mime','painel'),
 }
 if posicao not in mapa: abort(404)
 key_attr,mime_attr,nome=mapa[posicao]; chave=getattr(item,key_attr,None)
 if not chave: abort(404)
 conteudo=storage.download(chave)
 return send_file(BytesIO(conteudo),mimetype=getattr(item,mime_attr,None) or 'image/jpeg',download_name=f'vistoria-{item.id}-{nome}.jpg',conditional=True)

@app.route('/vistorias/<int:id>/tentativas/<int:attempt_id>/video')
@login_required
def vistoria_tentativa_video(id,attempt_id):
 item=Inspection.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 tentativa=InspectionAttempt.query.filter_by(id=attempt_id,inspection_id=item.id,tenant_id=tid()).first_or_404()
 try: conteudo=storage.download(tentativa.video_key)
 except StorageNotFoundError: abort(404)
 return send_file(BytesIO(conteudo),mimetype=tentativa.video_mime or 'video/webm',download_name=f'vistoria-{item.id}-tentativa-{tentativa.id}.webm',conditional=True)

@app.route('/vistoria/<token>')
def vistoria_publica(token):
 item=Inspection.query.filter_by(token=token).first_or_404()
 evidencia_ok=(item.video_key or ((item.tipo_vistoria or '')=='fotos' and item.front_photo_key and item.right_photo_key and item.rear_photo_key and item.left_photo_key and item.painel_photo_key))
 if item.status in ('Aprovada','Recebida','Concluída','Concluida') and evidencia_ok:
  return render_template('vistoria_publica.html',item=item,done=True)
 if item.expires_at and item.expires_at < datetime.utcnow():
  return render_template('vistoria_publica.html',item=item,expired=True),410
 return render_template('vistoria_publica.html',item=item)

@app.route('/vistoria/<token>/upload',methods=['POST'])
def vistoria_upload(token):
 item=Inspection.query.filter_by(token=token).first_or_404()
 if item.expires_at and item.expires_at < datetime.utcnow():
  return {'ok':False,'error':'Link expirado.'},410
 modo=(item.tipo_vistoria or 'fotos').strip().lower()
 veiculo=Vehicle.query.filter_by(id=item.vehicle_id,tenant_id=item.tenant_id).first()
 if not veiculo:
  return {'ok':False,'error':'Veículo da vistoria não encontrado.'},404
 try: km=int(str(request.form.get('km') or '').replace('.','').replace(',','').strip())
 except Exception:
  return {'ok':False,'error':'Informe a quilometragem atual mostrada no painel.'},400
 if km < 0:
  return {'ok':False,'error':'Quilometragem inválida.'},400
 km_anterior=int(veiculo.km_atual or 0)
 if km < km_anterior:
  return {'ok':False,'error':f'A quilometragem informada ({km:,} km) é menor que a última registrada ({km_anterior:,} km). Confira o painel.'.replace(',','.')},400

 if modo=='fotos':
  campos=[
   ('front_photo','frente','front_photo_key','front_photo_mime'),
   ('right_photo','lateral direita','right_photo_key','right_photo_mime'),
   ('rear_photo','traseira','rear_photo_key','rear_photo_mime'),
   ('left_photo','lateral esquerda','left_photo_key','left_photo_mime'),
   ('painel_photo','painel','painel_photo_key','painel_photo_mime'),
  ]
  arquivos=[]
  for form_name,rotulo,key_attr,mime_attr in campos:
   arq=request.files.get(form_name)
   if not arq:
    return {'ok':False,'error':f'Tire a foto da {rotulo} antes de finalizar a vistoria.'},400
   mime=(arq.mimetype or '').lower()
   if mime not in ('image/jpeg','image/png','image/webp'):
    return {'ok':False,'error':f'Formato da foto da {rotulo} não suportado. Use a câmera da página.'},400
   arquivos.append((arq,rotulo,key_attr,mime_attr,mime))
  pasta=f"{item.tenant_id}/vistorias/{item.vehicle_id}/{datetime.utcnow().strftime('%Y/%m')}"
  uploaded=[]
  try:
   for arq,rotulo,key_attr,mime_attr,mime in arquivos:
    ext='.png' if mime=='image/png' else ('.webp' if mime=='image/webp' else '.jpg')
    slug={'frente':'frente','lateral direita':'direita','traseira':'traseira','lateral esquerda':'esquerda','painel':'painel'}[rotulo]
    chave=f"{pasta}/{uuid.uuid4().hex}-{slug}{ext}"
    storage.upload(arq.stream,chave,mime)
    uploaded.append(chave)
    setattr(item,key_attr,chave); setattr(item,mime_attr,mime)
  except Exception:
   app.logger.exception('Falha ao armazenar fotos da vistoria %s',item.id)
   for chave in uploaded:
    try: storage.delete(chave)
    except Exception: pass
   return {'ok':False,'error':'Não foi possível armazenar todas as fotos. Tente novamente.'},503
  item.video_key=None; item.video_mime=None; item.duration_seconds=None
  origem_odo='Vistoria fotográfica'
  evento_recebido='Vistoria fotográfica recebida'
  descricao=f'Vistoria fotográfica concluída com 4 lados do veículo e foto do painel; KM informada {km:,} km.'.replace(',','.')
 else:
  video=request.files.get('video'); painel=request.files.get('painel_photo')
  if not video: return {'ok':False,'error':'Vídeo não recebido.'},400
  if not painel: return {'ok':False,'error':'Tire uma foto do painel antes de finalizar a vistoria.'},400
  mime=(video.mimetype or '').lower()
  if not (mime.startswith('video/webm') or mime.startswith('video/mp4') or mime.startswith('video/quicktime')):
   return {'ok':False,'error':'Formato de vídeo não suportado.'},400
  painel_mime=(painel.mimetype or '').lower()
  if painel_mime not in ('image/jpeg','image/png','image/webp'):
   return {'ok':False,'error':'Formato da foto do painel não suportado. Use JPG, PNG ou WEBP.'},400
  try: duracao=max(0,int(float(request.form.get('duration_seconds') or 0)))
  except Exception: duracao=0
  if duracao < 15:
   msg_curta='A vistoria ficou muito curta. Grave pelo menos 15 segundos mostrando o veículo.' if modo=='simples' else 'A vistoria ficou muito curta. Grave o veículo seguindo todas as etapas.'
   return {'ok':False,'error':msg_curta},400
  ext='.mp4' if ('mp4' in mime or 'quicktime' in mime) else '.webm'
  painel_ext='.png' if painel_mime=='image/png' else ('.webp' if painel_mime=='image/webp' else '.jpg')
  pasta=f"{item.tenant_id}/vistorias/{item.vehicle_id}/{datetime.utcnow().strftime('%Y/%m')}"
  chave=f"{pasta}/{uuid.uuid4().hex}{ext}"; painel_chave=f"{pasta}/{uuid.uuid4().hex}-painel{painel_ext}"
  try:
   storage.upload(video.stream,chave,mime); storage.upload(painel.stream,painel_chave,painel_mime)
  except Exception:
   app.logger.exception('Falha ao armazenar vídeo/foto da vistoria %s',item.id)
   try: storage.delete(chave)
   except Exception: pass
   try: storage.delete(painel_chave)
   except Exception: pass
   return {'ok':False,'error':'Não foi possível armazenar o vídeo e a foto do painel. Tente novamente.'},503
  item.video_key=chave; item.video_mime=mime; item.duration_seconds=duracao
  item.painel_photo_key=painel_chave; item.painel_photo_mime=painel_mime
  try: brilho=float(request.form.get('brightness_avg') or 0)
  except Exception: brilho=0
  item.brightness_avg=Decimal(str(round(brilho,2)))
  tentativa=InspectionAttempt(inspection_id=item.id,tenant_id=item.tenant_id,video_key=chave,video_mime=mime,duration_seconds=duracao,painel_photo_key=painel_chave,painel_photo_mime=painel_mime,km_informada=km,brightness_avg=item.brightness_avg,submitted_at=datetime.utcnow(),decision='Concluída',decided_at=datetime.utcnow())
  db.session.add(tentativa)
  origem_odo='Vistoria em vídeo'; evento_recebido='Vistoria em vídeo recebida'
  descricao=f'Vistoria em vídeo concluída; duração {duracao}s; foto do painel anexada; KM informada {km:,} km.'.replace(',','.')

 km_anterior_vistoria=item.km_informada
 item.km_informada=km; item.brightness_status='Não avaliada'; item.submitted_at=datetime.utcnow(); item.status='Concluída'; item.notes=None
 # Uma vistoria completa satisfaz a obrigação do contrato. Encerra pedidos
 # anteriores ainda abertos, inclusive registros legados sem contract_id.
 _encerrar_vistorias_anteriores(item)
 if km_anterior_vistoria != km:
  db.session.add(Odometer(tenant_id=item.tenant_id,vehicle_id=item.vehicle_id,km=km,origem=origem_odo))
 if km >= km_anterior: veiculo.km_atual=km
 tenant_km=Tenant.query.get(item.tenant_id)
 km_pendentes=MileageRequest.query.filter_by(tenant_id=item.tenant_id,vehicle_id=item.vehicle_id,driver_id=item.driver_id,status='Pendente').all()
 for _req in km_pendentes:
  _req.km=km; _req.photo=item.painel_photo_key; _req.submitted_at=item.submitted_at
  _req.notes='Atendida automaticamente pela vistoria fotográfica.' if modo=='fotos' else 'Atendida automaticamente pela vistoria em vídeo.'
  _req.status='Aguardando conferência' if (tenant_km and tenant_km.conferir_km_motorista) else 'Concluído'
 db.session.add(VehicleEvent(tenant_id=item.tenant_id,vehicle_id=item.vehicle_id,contract_id=item.contract_id,driver_id=item.driver_id,evento=evento_recebido,descricao=descricao))
 db.session.commit()
 try:
  analisar_avarias_vistoria(item); db.session.commit()
 except Exception:
  db.session.rollback(); app.logger.exception('Falha ao persistir análise de avarias da vistoria %s',item.id)
 try: recalcular_alertas(item.tenant_id)
 except Exception: app.logger.exception('Falha ao recalcular alertas após KM da vistoria %s',item.id)
 if modo=='fotos':
  return {'ok':True,'message':'Vistoria enviada com sucesso: 4 lados, painel e quilometragem recebidos.'}
 return {'ok':True,'message':'Vistoria, foto do painel e quilometragem enviadas com sucesso.'}

@app.route('/alertas')
@login_required
def alertas():
 recalcular_alertas(tid())
 nivel=(request.args.get('nivel') or '').strip()
 q=Alert.query.filter(Alert.tenant_id==tid(),Alert.resolvido_em.is_(None))
 if nivel in ('danger','warning','info'):
  q=q.filter(Alert.nivel==nivel)
 items=q.order_by(Alert.criado_em.desc()).all()
 return render_template('alertas.html',items=items,nivel=nivel)

@app.route('/alertas/<int:id>/abrir')
@login_required
def alerta_abrir(id):
 a=Alert.query.filter_by(id=id,tenant_id=tid()).first_or_404()

 # O destino é resolvido pela entidade do alerta, evitando links genéricos.
 if a.entidade == 'Manutenção' and a.entidade_id:
  m=Maintenance.query.filter_by(id=a.entidade_id,tenant_id=tid()).first()
  if m:
   return redirect(url_for('manutencoes') + f'#manutencao-{m.id}')

 if a.entidade == 'Veículo' and a.entidade_id:
  v=Vehicle.query.filter_by(id=a.entidade_id,tenant_id=tid()).first()
  if v:
   return redirect(url_for('editar_veiculo',id=v.id) + '#troca-oleo')

 if a.entidade == 'Contrato' and a.entidade_id:
  c=Contract.query.filter_by(id=a.entidade_id,tenant_id=tid()).first()
  if c:
   return redirect(url_for('contrato_detalhe',id=c.id))

 # Fallback para alertas antigos ou integrações futuras.
 if a.action_url:
  return redirect(a.action_url)
 return redirect(url_for('alertas'))

@app.route('/alertas/<int:id>/whatsapp-motorista',methods=['POST'])
@login_required
def alerta_whatsapp_motorista(id):
 a=Alert.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 v=None; m=None
 if a.entidade=='Manutenção':
  m=Maintenance.query.filter_by(id=a.entidade_id,tenant_id=tid()).first()
  v=m.vehicle if m else None
 elif a.entidade=='Veículo':
  v=Vehicle.query.filter_by(id=a.entidade_id,tenant_id=tid()).first()
 if not v:
  flash('Este alerta não possui veículo associado.','warning'); return redirect(url_for('alertas'))
 d=motorista_atual_veiculo(v)
 if not d:
  flash('O veículo não possui motorista vinculado a contrato vigente.','warning'); return redirect(url_for('alertas'))
 if m:
  body=maintenance_message(driver_name=d.nome,vehicle=v,maintenance=m,reminder=False)
  msg_type='alerta_manutencao'
 else:
  body=f'Olá, {d.nome}! A locadora identificou um alerta no veículo {v.marca_modelo or "Veículo"} — {v.placa}: {a.titulo}. {a.mensagem}'
  msg_type='alerta_veiculo'
 integration=Integration.query.filter_by(tenant_id=tid(),tipo='whatsapp').first()
 cfg=CommunicationService.parse_config(integration)
 template_name=(cfg.get('maintenance_template_name') or '').strip() or None
 params=[d.nome,v.marca_modelo or 'Veículo',v.placa,m.tipo or 'Manutenção',data_br(m.proxima_data) if m and m.proxima_data else 'a definir',m.proxima_hora or 'a definir'] if m else [d.nome,v.marca_modelo or 'Veículo',v.placa,a.titulo or 'Alerta',a.mensagem or '','a definir']
 fila,redirect_url,err=criar_mensagem_whatsapp(tenant_id=tid(),driver=d,body=body,message_type=msg_type,related_entity=a.entidade or 'Alerta',related_entity_id=a.entidade_id or a.id,template_name=template_name,template_parameters=params)
 db.session.commit()
 if err: flash('Não foi possível enviar o alerta: '+err,'danger')
 elif redirect_url: return redirect(redirect_url)
 else: flash('Alerta enviado ao motorista pelo WhatsApp Business.','success')
 return redirect(url_for('alertas'))

@app.route('/alertas/<int:id>/lido',methods=['POST'])
@login_required
def alerta_lido(id):
 a=Alert.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 a.lido=True; db.session.commit()
 return redirect(request.referrer or url_for('alertas'))

@app.route('/alertas/atualizar',methods=['POST'])
@login_required
def atualizar_alertas():
 recalcular_alertas(tid())
 flash('Alertas atualizados.','success')
 return redirect(request.referrer or url_for('alertas'))

@app.route('/administracao/armazenamento')
@login_required
def administracao_armazenamento():
 documentos=Document.query.filter_by(tenant_id=tid()).all()
 fotos=MileageRequest.query.options(joinedload(MileageRequest.vehicle)).filter_by(tenant_id=tid(),status='Concluído').filter(MileageRequest.photo.isnot(None)).all()
 referencias=[]
 for documento in documentos:
  referencias.append({'tipo':'Documento','nome':documento.nome_original or documento.arquivo,'chave':documento.arquivo})
 for foto in fotos:
  referencias.append({'tipo':'Foto do painel','nome':f'{foto.vehicle.placa if foto.vehicle else "Veículo"} — leitura #{foto.id}','chave':foto.photo})

 ausentes=[]
 erro_verificacao=None
 try:
  for item in referencias:
   if '/' not in (item['chave'] or ''):
    existe=(UPLOAD/str(tid())/'odometros'/item['chave']).exists()
   else:
    existe=storage.exists(item['chave'])
   if not existe:
    ausentes.append(item)
 except Exception as exc:
  app.logger.exception('Falha na verificação de integridade')
  erro_verificacao=str(exc)

 try:
  conectado=storage.check_connection() if storage.using_r2 else False
  uso=storage.tenant_usage(tid())
 except Exception as exc:
  app.logger.exception('Falha ao consultar armazenamento')
  conectado=False
  uso={'objects':0,'bytes':0}
  if not erro_verificacao:
   erro_verificacao=str(exc)

 backups=[]
 try:
  if storage.using_r2:
   prefix=f'{tid()}/backups/'
   pagina=storage._client.list_objects_v2(Bucket=storage.bucket_name,Prefix=prefix)
   backups=sorted(
    [{'chave':o['Key'],'tamanho':o['Size'],'data':o['LastModified']} for o in pagina.get('Contents',[])],
    key=lambda x:x['data'],reverse=True
   )[:10]
 except Exception:
  app.logger.exception('Falha ao listar backups')

 return render_template('armazenamento.html',storage_backend=storage.backend_name,conectado=conectado,uso=uso,referencias_total=len(referencias),ausentes=ausentes,erro_verificacao=erro_verificacao,backups=backups)

def _normalizar_dia_semana(valor):
 texto=unicodedata.normalize('NFKD',str(valor or '')).encode('ascii','ignore').decode('ascii').lower().strip()
 texto=texto.replace('_',' ').replace('-feira','').replace('feira','').replace('-',' ').strip()
 mapa={
  'segunda':0,'segunda feira':0,'2':0,
  'terca':1,'terca feira':1,'3':1,
  'quarta':2,'quarta feira':2,'4':2,
  'quinta':3,'quinta feira':3,'5':3,
  'sexta':4,'sexta feira':4,'6':4,
  'sabado':5,'7':5,
  'domingo':6,'1':6,
 }
 return mapa.get(texto)

def _ultimas_leituras_km(vehicle_id, limit=2, tenant_id=None):
 tenant_id=tenant_id if tenant_id is not None else tid()
 return Odometer.query.filter_by(tenant_id=tenant_id,vehicle_id=vehicle_id).order_by(Odometer.data.desc(),Odometer.id.desc()).limit(limit).all()

def calcular_cobranca_semanal(contract):
 valor_base=Decimal(str(contract.valor_locacao or 0))
 total=valor_base
 info={
  'valor_base':valor_base,'total':total,'km_periodo':None,'limite_km':contract.limite_km,
  'km_excedente':0,'valor_excesso':Decimal('0'),'valor_excesso_teorico':Decimal('0'),
  'tem_historico_km':False,'usa_excesso':False,'tem_excesso':False,
 }
 if not contract.vehicle_id or not contract.limite_km or not contract.valor_km_excedente:
  return info
 leituras=_ultimas_leituras_km(contract.vehicle_id,2,contract.tenant_id)
 if len(leituras)<2:
  return info
 atual,anterior=leituras[0],leituras[1]
 agora=agora_sao_paulo_naive()
 data_atual=_as_sao_paulo(atual.data).replace(tzinfo=None) if atual.data else None
 if data_atual and (agora-data_atual)>timedelta(days=10):
  return info
 km_periodo=max(0,int(atual.km or 0)-int(anterior.km or 0))
 info['tem_historico_km']=True
 info['km_periodo']=km_periodo
 excedente=max(0,km_periodo-int(contract.limite_km or 0))
 info['km_excedente']=excedente
 if excedente>0:
  valor_excesso_teorico=Decimal(excedente)*Decimal(str(contract.valor_km_excedente or 0))
  info['valor_excesso_teorico']=valor_excesso_teorico
  info['tem_excesso']=True
  tenant=Tenant.query.get(contract.tenant_id)
  cobrar=bool(tenant and tenant.cobrar_km_excedente)
  if cobrar:
   info['valor_excesso']=valor_excesso_teorico
   info['total']=valor_base+valor_excesso_teorico
   info['usa_excesso']=True
 return info

def cobranca_vence_hoje(contract):
 dia=_normalizar_dia_semana(contract.dia_vencimento)
 return dia is not None and dia==datetime.now(SAO_PAULO).weekday()

def mensagem_cobranca_semanal(contract, info, comprovante_url=None):
 motorista=contract.driver.nome if contract.driver else 'Motorista'
 veiculo=contract.vehicle.placa if contract.vehicle else 'veículo contratado'
 linhas=[
  f'Olá, {motorista}!',
  '',
  f'Lembrete de pagamento semanal do veículo {veiculo}.',
  f'Vencimento: {"hoje" if cobranca_vence_hoje(contract) else "toda "+str(contract.dia_vencimento or "semana")}.',
  f'Valor semanal: R$ {moeda_br(info["valor_base"])}.',
 ]
 if info.get('usa_excesso'):
  linhas += [
   f'Quilometragem entre as duas últimas leituras: {info["km_periodo"]} km.',
   f'Limite semanal: {info["limite_km"]} km.',
   f'Excesso: {info["km_excedente"]} km (R$ {moeda_br(info["valor_excesso"])}).',
   f'Total desta cobrança: R$ {moeda_br(info["total"])}.',
  ]
 if comprovante_url:
  linhas += ['', 'Após o pagamento, envie o comprovante pelo link:', comprovante_url]
 linhas += ['', 'Obrigado.']
 return '\n'.join(linhas)

def _cobranca_template_params(contract,info,comprovante_url=None,include_link=True):
 d=contract.driver; v=contract.vehicle
 vencimento='hoje' if cobranca_vence_hoje(contract) else str(contract.dia_vencimento or 'semanal')
 if info.get('usa_excesso'):
  # cobranca_semanal_excesso_com_comprovante (pt_BR)
  # 1 motorista, 2 veículo, 3 placa, 4 aluguel, 5 km período, 6 limite,
  # 7 km excedente, 8 valor excesso, 9 total, 10 vencimento, 11 link comprovante.
  params=[
   d.nome if d else 'Motorista',
   v.marca_modelo if v and v.marca_modelo else 'Veículo',
   v.placa if v else '-',
   moeda_br(info['valor_base']),
   str(info.get('km_periodo') or 0),
   str(info.get('limite_km') or 0),
   str(info.get('km_excedente') or 0),
   moeda_br(info['valor_excesso']),
   moeda_br(info['total']),
   vencimento,
  ]
 else:
  # cobranca_semanal_com_comprovante (pt_BR)
  # 1 motorista, 2 veículo, 3 placa, 4 aluguel, 5 vencimento, 6 link comprovante.
  params=[
   d.nome if d else 'Motorista',
   v.marca_modelo if v and v.marca_modelo else 'Veículo',
   v.placa if v else '-',
   moeda_br(info['valor_base']),
   vencimento,
  ]
 if include_link and comprovante_url:
  params.append(comprovante_url)
 return params

def _automation_cfg(tenant_id):
 integration=Integration.query.filter_by(tenant_id=tenant_id,tipo='whatsapp').first()
 cfg=CommunicationService.parse_config(integration)
 return integration,cfg

def _automation_window_open(cfg, kind=None, agora=None):
 agora=agora or datetime.now(SAO_PAULO)
 if not cfg.get('automation_enabled',False): return False
 try: old_weekday=int(cfg.get('automation_weekday',0))
 except Exception: old_weekday=0
 if kind=='km':
  weekdays=cfg.get('km_automation_weekdays') or [old_weekday]
  start_raw=cfg.get('km_start_hour',cfg.get('automation_start_hour',7))
  end_raw=cfg.get('km_end_hour',cfg.get('automation_end_hour',20))
 elif kind=='billing':
  weekdays=cfg.get('billing_automation_weekdays') or [old_weekday]
  start_raw=cfg.get('billing_start_hour',cfg.get('automation_start_hour',7))
  end_raw=cfg.get('billing_end_hour',cfg.get('automation_end_hour',20))
 elif kind=='inspection':
  weekdays=cfg.get('inspection_automation_weekdays') or [old_weekday]
  start_raw=cfg.get('inspection_start_hour',10)
  end_raw=cfg.get('inspection_end_hour',20)
 else:
  weekdays=[old_weekday]
  start_raw=cfg.get('automation_start_hour',7); end_raw=cfg.get('automation_end_hour',20)
 try: weekdays=[int(x) for x in weekdays]
 except Exception: weekdays=[old_weekday]
 try: start_hour=int(start_raw)
 except Exception: start_hour=7
 try: end_hour=int(end_raw)
 except Exception: end_hour=20
 return agora.weekday() in weekdays and start_hour<=agora.hour<=end_hour

def _reminder_interval(cfg, kind=None):
 key='reminder_interval_hours'
 if kind=='km': key='km_reminder_interval_hours'
 elif kind=='billing': key='billing_reminder_interval_hours'
 elif kind=='inspection': key='inspection_reminder_interval_hours'
 try: return max(1,int(cfg.get(key,cfg.get('reminder_interval_hours',1))))
 except Exception: return 1

def _audit_info(audit):
 return {
  'valor_base':Decimal(str(audit.base_amount or 0)),'total':Decimal(str(audit.total_amount or 0)),
  'km_periodo':audit.km_period,'limite_km':audit.km_limit,'km_excedente':audit.km_excess or 0,
  'valor_excesso':Decimal(str(audit.excess_amount or 0)),'tem_historico_km':audit.km_period is not None,
  'usa_excesso':Decimal(str(audit.excess_amount or 0))>0,
 }

def _enviar_cobranca_audit(contract,audit,cfg):
 if not audit or (audit.payment_status or 'PENDENTE').upper()=='PAGO':
  return None,None,'Esta cobrança já está baixada como paga.'
 info=_audit_info(audit)
 comprovante_url=url_comprovante_cobranca(audit)
 body=audit.body or mensagem_cobranca_semanal(contract,info,comprovante_url)
 template_name=((cfg.get('payment_excess_template_name') if info.get('usa_excesso') else cfg.get('payment_template_name')) or '').strip() or audit.template_name
 params=_cobranca_template_params(contract,info,comprovante_url,include_link=True)
 fila,redirect_url,err=criar_mensagem_whatsapp(tenant_id=contract.tenant_id,driver=contract.driver,body=body,message_type='lembrete_pagamento_semanal',related_entity='Cobranca',related_entity_id=audit.id,template_name=template_name,template_parameters=params)
 now=agora_sao_paulo_naive()
 if fila:
  audit.message_id=fila.id; audit.provider=fila.provider; audit.status=fila.status; audit.external_id=fila.external_id; audit.error_message=fila.error_message
  audit.reminder_count=(audit.reminder_count or 0)+1; audit.last_reminder_at=now
 return fila,redirect_url,err

def gerar_e_enviar_cobranca(contract,automatico=False):
 info=calcular_cobranca_semanal(contract)
 integration,cfg=_automation_cfg(contract.tenant_id)
 # Segunda barreira de segurança para chamadas automáticas diretas. O envio
 # manual pelo botão da tela continua permitido mesmo com a automação desligada.
 if automatico and not cfg.get('billing_automation_enabled',cfg.get('automatic_billing_enabled',False)):
  return None,None,None,'Automação de cobranças desabilitada.'
 provider=(cfg.get('provider') or 'web').lower()
 template_name=((cfg.get('payment_excess_template_name') if info.get('usa_excesso') else cfg.get('payment_template_name')) or '').strip() or None
 hoje=datetime.now(_tenant_zone(contract.tenant_id)).date()
 inicio_semana,fim_semana=_periodo_semana_tenant(contract.tenant_id,hoje)
 pago_semana=BillingAudit.query.filter(
  BillingAudit.tenant_id==contract.tenant_id,
  BillingAudit.contract_id==contract.id,
  BillingAudit.billing_date>=inicio_semana,
  BillingAudit.billing_date<=fim_semana,
  BillingAudit.payment_status=='PAGO'
 ).order_by(BillingAudit.id.desc()).first()
 if pago_semana:
  return None,pago_semana,None,'A cobrança desta semana já está baixada como paga.'

 # Reenvio da cobrança do mesmo contrato no mesmo dia deve reutilizar a mesma
 # auditoria. Antes, cada clique manual criava outra BillingAudit e, se ambas
 # fossem baixadas como pagas, o portal do proprietário somava em duplicidade.
 audit=BillingAudit.query.filter(
  BillingAudit.tenant_id==contract.tenant_id,
  BillingAudit.contract_id==contract.id,
  BillingAudit.payment_status!='PAGO',
 ).order_by(BillingAudit.billing_date.desc(),BillingAudit.id.desc()).first()
 if not audit:
  audit=BillingAudit.query.filter_by(tenant_id=contract.tenant_id,contract_id=contract.id,billing_date=hoje).order_by(BillingAudit.id.desc()).first()
 if audit:
  # Mantém o registro já pago intacto. Para pendentes, atualiza a fotografia da
  # cobrança antes do reenvio, sem criar uma segunda competência financeira.
  if (audit.payment_status or 'PENDENTE').upper()!='PAGO':
   audit.driver_name=contract.driver.nome if contract.driver else None
   audit.vehicle_label=contract.vehicle.marca_modelo if contract.vehicle else None
   audit.plate=contract.vehicle.placa if contract.vehicle else None
   audit.base_amount=info['valor_base']; audit.km_period=info.get('km_periodo'); audit.km_limit=info.get('limite_km')
   audit.km_excess=info.get('km_excedente') or 0; audit.excess_rate=contract.valor_km_excedente or 0
   audit.excess_amount=info.get('valor_excesso') or 0; audit.total_amount=info['total']; audit.template_name=template_name
   audit.provider='whatsapp_business' if provider=='business' else 'whatsapp_web'
   comprovante_url=url_comprovante_cobranca(audit)
   audit.body=mensagem_cobranca_semanal(contract,info,comprovante_url)
  else:
   return None,audit,None,'Esta cobrança já está baixada como paga.'
 else:
  audit=BillingAudit(tenant_id=contract.tenant_id,contract_id=contract.id,driver_name=contract.driver.nome if contract.driver else None,vehicle_label=contract.vehicle.marca_modelo if contract.vehicle else None,plate=contract.vehicle.placa if contract.vehicle else None,billing_date=hoje,base_amount=info['valor_base'],km_period=info.get('km_periodo'),km_limit=info.get('limite_km'),km_excess=info.get('km_excedente') or 0,excess_rate=contract.valor_km_excedente or 0,excess_amount=info.get('valor_excesso') or 0,total_amount=info['total'],body='',template_name=template_name,provider='whatsapp_business' if provider=='business' else 'whatsapp_web',status='GERADA',payment_status='PENDENTE',created_at=agora_sao_paulo_naive())
  db.session.add(audit); db.session.flush()
  comprovante_url=url_comprovante_cobranca(audit)
  audit.body=mensagem_cobranca_semanal(contract,info,comprovante_url)
 fila,redirect_url,err=_enviar_cobranca_audit(contract,audit,cfg)
 return fila,audit,redirect_url,err

def processar_cobrancas_automaticas(tenant_id=None):
 q=Contract.query.options(joinedload(Contract.driver),joinedload(Contract.vehicle)).filter(Contract.status.in_(['Assinado','Ativo']))
 if tenant_id is not None: q=q.filter(Contract.tenant_id==tenant_id)
 enviados=0
 for c in q.all():
  agora=datetime.now(_tenant_zone(c.tenant_id)); hoje=agora.date()
  periodicidade=unicodedata.normalize('NFKD',str(c.periodicidade or '')).encode('ascii','ignore').decode('ascii').lower()
  if periodicidade and 'seman' not in periodicidade: continue
  integration,cfg=_automation_cfg(c.tenant_id)
  if not cfg.get('billing_automation_enabled',cfg.get('automatic_billing_enabled',False)) or not _automation_window_open(cfg,'billing',agora): continue
  inicio_semana,fim_semana=_periodo_semana_tenant(c.tenant_id,hoje)
  pago_semana=BillingAudit.query.filter(
   BillingAudit.tenant_id==c.tenant_id,
   BillingAudit.contract_id==c.id,
   BillingAudit.billing_date>=inicio_semana,
   BillingAudit.billing_date<=fim_semana,
   BillingAudit.payment_status=='PAGO'
  ).first()
  if pago_semana: continue
  # Envio realmente automático só é possível pela Business Platform. No Web, o botão manual continua disponível.
  if (cfg.get('provider') or 'web').lower()!='business': continue
  audit=BillingAudit.query.filter(
   BillingAudit.tenant_id==c.tenant_id,
   BillingAudit.contract_id==c.id,
   BillingAudit.payment_status!='PAGO',
  ).order_by(BillingAudit.billing_date.desc(),BillingAudit.id.desc()).first()
  if not audit:
   audit=BillingAudit.query.filter_by(tenant_id=c.tenant_id,contract_id=c.id,billing_date=hoje).order_by(BillingAudit.id.desc()).first()
  if audit and (audit.payment_status or 'PENDENTE')=='PAGO': continue
  if not audit:
   gerar_e_enviar_cobranca(c,automatico=True); enviados+=1; continue
  intervalo=_reminder_interval(cfg,'billing')
  if audit.last_reminder_at and (agora_sao_paulo_naive()-audit.last_reminder_at)<timedelta(hours=intervalo): continue
  _enviar_cobranca_audit(c,audit,cfg); enviados+=1
 db.session.commit(); return enviados

def processar_km_automatico(tenant_id=None):
 q=Contract.query.options(joinedload(Contract.driver),joinedload(Contract.vehicle)).filter(Contract.status.in_(['Assinado','Ativo']))
 if tenant_id is not None: q=q.filter(Contract.tenant_id==tenant_id)
 enviados=0
 for c in q.all():
  agora=datetime.now(_tenant_zone(c.tenant_id))
  inicio_local=datetime.combine(agora.date(),datetime.min.time(),tzinfo=_tenant_zone(c.tenant_id))
  inicio=inicio_local.astimezone(timezone.utc).replace(tzinfo=None)
  if not c.driver or not c.vehicle: continue
  integration,cfg=_automation_cfg(c.tenant_id)

  # A vistoria já solicita vídeo + foto do painel + KM digitada.
  # Quando a automação de vistoria está ativa ela é a fonte principal e
  # bloqueia o envio automático redundante do link exclusivo de KM.
  if cfg.get('inspection_automation_enabled',False):
   continue

  if not cfg.get('km_automation_enabled',cfg.get('automatic_km_enabled',False)) or not _automation_window_open(cfg,'km',agora): continue
  if (cfg.get('provider') or 'web').lower()!='business': continue
  # Se a foto/KM já foi recebida hoje, encerra a automação desse veículo no dia.
  # submitted_at é salvo em UTC sem timezone; convertemos para São Paulo antes de comparar a data.
  ultima_respondida=MileageRequest.query.filter_by(tenant_id=c.tenant_id,vehicle_id=c.vehicle.id,driver_id=c.driver.id).filter(
   MileageRequest.status.in_(['Concluído','Aguardando conferência']),
   MileageRequest.submitted_at.isnot(None),
  ).order_by(MileageRequest.submitted_at.desc()).first()
  if ultima_respondida and ultima_respondida.submitted_at:
   submitted_local=ultima_respondida.submitted_at.replace(tzinfo=timezone.utc).astimezone(_tenant_zone(c.tenant_id))
   if submitted_local.date()==agora.date():
    continue

  req=MileageRequest.query.filter_by(tenant_id=c.tenant_id,vehicle_id=c.vehicle.id,driver_id=c.driver.id,status='Pendente').filter(MileageRequest.expires_at>datetime.utcnow()).order_by(MileageRequest.id.desc()).first()
  if not req:
   req=MileageRequest(tenant_id=c.tenant_id,vehicle_id=c.vehicle.id,driver_id=c.driver.id,token=uuid.uuid4().hex+uuid.uuid4().hex,expires_at=datetime.utcnow()+timedelta(days=7),previous_km=c.vehicle.km_atual); db.session.add(req); db.session.flush()
  ultimo=MessageQueue.query.filter_by(tenant_id=c.tenant_id,message_type='solicitacao_km',related_entity='Veiculo',related_entity_id=c.vehicle.id).filter(MessageQueue.created_at>=inicio).order_by(MessageQueue.id.desc()).first()
  intervalo=_reminder_interval(cfg,'km')
  if ultimo and ultimo.created_at and (agora_sao_paulo_naive()-ultimo.created_at)<timedelta(hours=intervalo): continue
  link=url_for('registrar_quilometragem_publica',token=req.token,_external=True)
  body=f'Olá, {c.driver.nome}! Precisamos da quilometragem atual do veículo {c.vehicle.placa}. Abra o link, tire uma foto do painel e informe o km: {link}'
  template=(cfg.get('mileage_template_name') or '').strip() or None
  params=[c.driver.nome,c.vehicle.marca_modelo or 'Veículo',c.vehicle.placa,link]
  criar_mensagem_whatsapp(tenant_id=c.tenant_id,driver=c.driver,body=body,message_type='solicitacao_km',related_entity='Veiculo',related_entity_id=c.vehicle.id,template_name=template,template_parameters=params); enviados+=1
 db.session.commit(); return enviados

def processar_vistorias_automaticas(tenant_id=None):
 """Cria uma vistoria semanal por contrato/veículo e reenvia o mesmo link somente enquanto estiver pendente.

 O envio completo conclui a obrigação sem depender de aprovação. Usa o fuso do
 tenant e impede criação duplicada do mesmo pedido.
 """
 q=Contract.query.options(joinedload(Contract.driver),joinedload(Contract.vehicle)).filter(Contract.status.in_(['Assinado','Ativo']))
 if tenant_id is not None: q=q.filter(Contract.tenant_id==tenant_id)
 enviados=0

 for contrato_ref in q.all():
  if not contrato_ref.driver or not contrato_ref.vehicle: continue

  # Serializa a decisão por contrato para evitar duas execuções simultâneas do cron.
  try:
   db.session.execute(text('SELECT pg_advisory_xact_lock(:chave)'),{'chave':int(contrato_ref.id)})
  except Exception:
   pass

  c=Contract.query.options(joinedload(Contract.driver),joinedload(Contract.vehicle)).filter_by(
   id=contrato_ref.id,tenant_id=contrato_ref.tenant_id
  ).first()
  if not c or not c.driver or not c.vehicle: continue

  agora_local=datetime.now(_tenant_zone(c.tenant_id))
  integration,cfg=_automation_cfg(c.tenant_id)
  if not cfg.get('inspection_automation_enabled',False) or not _automation_window_open(cfg,'inspection',agora_local): continue
  if (cfg.get('provider') or 'web').lower()!='business': continue

  # Se existe pendente deste contrato/veículo, reutiliza sempre o mesmo token.
  item=Inspection.query.filter_by(
   tenant_id=c.tenant_id,vehicle_id=c.vehicle.id,contract_id=c.id,status='Pendente'
  ).order_by(Inspection.requested_at.desc(),Inspection.id.desc()).first()

  if not item:
   # Depois que a vistoria foi enviada, não cria outra dentro da mesma semana.
   # Consideramos submitted_at para que um pedido antigo, concluído hoje, não
   # provoque uma nova solicitação imediatamente.
   limite_recente=datetime.utcnow()-timedelta(days=7)
   recente=Inspection.query.filter(
    Inspection.tenant_id==c.tenant_id,
    Inspection.vehicle_id==c.vehicle.id,
    Inspection.contract_id==c.id,
    or_(Inspection.submitted_at>=limite_recente,Inspection.requested_at>=limite_recente),
   ).order_by(Inspection.requested_at.desc(),Inspection.id.desc()).first()
   if recente:
    continue

   try: validade=max(24,min(168,int(cfg.get('inspection_expiry_hours',168) or 168)))
   except Exception: validade=168
   item=Inspection(
    tenant_id=c.tenant_id,vehicle_id=c.vehicle.id,driver_id=c.driver.id,contract_id=c.id,
    token=uuid.uuid4().hex+uuid.uuid4().hex[:8],status='Pendente',
    tipo_vistoria=((cfg.get('inspection_automation_type') or 'fotos') if (cfg.get('inspection_automation_type') or 'fotos') in ('fotos','simples','guiada') else 'fotos'),
    requested_at=datetime.utcnow(),
    expires_at=datetime.utcnow()+timedelta(hours=validade),
   )
   db.session.add(item); db.session.flush()
  else:
   item.driver_id=c.driver.id
   try: validade=max(24,min(168,int(cfg.get('inspection_expiry_hours',168) or 168)))
   except Exception: validade=168
   if not item.expires_at or item.expires_at<=datetime.utcnow()+timedelta(hours=1):
    item.expires_at=datetime.utcnow()+timedelta(hours=validade)

  # Se já existe uma vistoria pendente para este contrato/veículo,
  # o link separado de KM deixa de ser necessário. Encerramos pendências
  # antigas para que não apareçam simultaneamente no Portal do Motorista.
  km_pendentes=MileageRequest.query.filter_by(
   tenant_id=c.tenant_id,
   vehicle_id=c.vehicle.id,
   driver_id=c.driver.id,
   status='Pendente',
  ).all()
  for _req in km_pendentes:
   _req.status='Cancelado'
   _req.notes='Substituída pela vistoria completa (fotos externas + painel + KM).'

  intervalo=_reminder_interval(cfg,'inspection')
  ultimo=MessageQueue.query.filter_by(
   tenant_id=c.tenant_id,message_type='vistoria',related_entity='Vistoria',related_entity_id=item.id
  ).order_by(MessageQueue.created_at.desc(),MessageQueue.id.desc()).first()
  if ultimo and ultimo.created_at and (agora_sao_paulo_naive()-ultimo.created_at)<timedelta(hours=intervalo):
   continue

  ok,_=enviar_vistoria_whatsapp_automatico(item)
  if ok: enviados+=1

 db.session.commit()
 return enviados

def processar_alertas_automaticos(tenant_id=None):
 q=Alert.query.filter(Alert.resolvido_em.is_(None))
 if tenant_id is not None: q=q.filter(Alert.tenant_id==tenant_id)
 enviados=0; hoje=datetime.now(SAO_PAULO).date()
 for a in q.order_by(Alert.criado_em.asc()).limit(200).all():
  integration=Integration.query.filter_by(tenant_id=a.tenant_id,tipo='whatsapp').first(); cfg=CommunicationService.parse_config(integration)
  if not cfg.get('automatic_alerts_enabled',False): continue
  v=None; m=None
  if a.entidade=='Manutenção':
   m=Maintenance.query.filter_by(id=a.entidade_id,tenant_id=a.tenant_id).first()
   if m and (m.status or '').strip()=='Cancelada': continue
   v=m.vehicle if m else None
  elif a.entidade=='Veículo': v=Vehicle.query.filter_by(id=a.entidade_id,tenant_id=a.tenant_id).first()
  if not v: continue
  d=motorista_atual_veiculo(v)
  if not d: continue
  existing=MessageQueue.query.filter_by(tenant_id=a.tenant_id,message_type='alerta_automatico',related_entity='Alerta',related_entity_id=a.id).filter(MessageQueue.created_at>=datetime.combine(hoje,datetime.min.time())).first()
  if existing: continue
  body=maintenance_message(driver_name=d.nome,vehicle=v,maintenance=m,reminder=False) if m else f'Olá, {d.nome}! A locadora identificou um alerta no veículo {v.marca_modelo or "Veículo"} — {v.placa}: {a.titulo}. {a.mensagem}'
  template_name=(cfg.get('maintenance_template_name') or '').strip() or None
  params=[d.nome,v.marca_modelo or 'Veículo',v.placa,m.tipo or 'Manutenção',data_br(m.proxima_data) if m and m.proxima_data else 'a definir',m.proxima_hora or 'a definir'] if m else [d.nome,v.marca_modelo or 'Veículo',v.placa,a.titulo or 'Alerta',a.mensagem or '','a definir']
  criar_mensagem_whatsapp(tenant_id=a.tenant_id,driver=d,body=body,message_type='alerta_automatico',related_entity='Alerta',related_entity_id=a.id,template_name=template_name,template_parameters=params); enviados+=1
 db.session.commit(); return enviados

def _periodo_semana_tenant(tenant_id, referencia=None):
 zone=_tenant_zone(tenant_id)
 hoje=referencia or datetime.now(zone).date()
 inicio=hoje-timedelta(days=hoje.weekday())
 return inicio,inicio+timedelta(days=6)

def _auditorias_semana_contrato(tenant_id, contract_id, referencia=None):
 inicio,fim=_periodo_semana_tenant(tenant_id,referencia)
 return BillingAudit.query.filter(
  BillingAudit.tenant_id==tenant_id,
  BillingAudit.contract_id==contract_id,
  BillingAudit.billing_date>=inicio,
  BillingAudit.billing_date<=fim
 ).order_by(BillingAudit.id.desc()).all()

def _cancelar_lembretes_cobranca_ids(tenant_id, audit_ids, motivo):
 ids=[int(x) for x in audit_ids if x]
 if not ids:
  return 0
 filas=MessageQueue.query.filter(
  MessageQueue.tenant_id==tenant_id,
  MessageQueue.related_entity=='Cobranca',
  MessageQueue.related_entity_id.in_(ids),
  MessageQueue.status.in_(['AGENDADA','PENDENTE','AGUARDANDO_MANUAL'])
 ).all()
 agora=agora_sao_paulo_naive()
 for fila in filas:
  fila.status='CANCELADA'
  fila.updated_at=agora
  db.session.add(MessageEvent(
   tenant_id=tenant_id,message_id=fila.id,event='CANCELADA',
   description=motivo,created_at=agora
  ))
 return len(filas)

@app.route('/cobrancas/auditoria/<int:id>/pago',methods=['POST'])
@login_required
def marcar_cobranca_paga(id):
 audit=BillingAudit.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 contrato=Contract.query.filter_by(id=audit.contract_id,tenant_id=tid()).first() if audit.contract_id else None
 if not contrato or contrato.status not in ('Assinado','Ativo'):
  flash('Esta cobrança pertence a um contrato que não está vigente e não pode receber baixa manual por esta tela.','warning')
  return redirect(url_for('cobrancas'))
 if (audit.payment_status or 'PENDENTE').upper()=='PAGO':
  flash('Esta cobrança já estava marcada como paga.','info'); return redirect(url_for('cobrancas'))
 audit.payment_status='PAGO'; audit.paid_at=agora_sao_paulo_naive(); audit.paid_by_id=current_user.id
 audit.payment_method=(request.form.get('payment_method') or '').strip() or None
 audit.payment_notes=(request.form.get('payment_notes') or '').strip() or None
 audit.closed_at=audit.paid_at
 # Ao dar baixa, qualquer lembrete ainda pendente da mesma competência semanal
 # deve morrer imediatamente. Isso cobre inclusive BillingAudit duplicada legada.
 auditorias_semana=_auditorias_semana_contrato(audit.tenant_id,audit.contract_id,audit.billing_date)
 ids_semana=[x.id for x in auditorias_semana]
 canceladas=_cancelar_lembretes_cobranca_ids(
  audit.tenant_id,ids_semana,
  'Lembrete cancelado porque a cobrança da competência foi marcada como paga.'
 )
 db.session.commit()
 msg='Pagamento baixado. Os lembretes desta cobrança foram interrompidos.'
 if canceladas:
  msg+=' {} mensagem(ns) pendente(s) cancelada(s).'.format(canceladas)
 flash(msg,'success')
 return redirect(url_for('cobrancas'))

@app.route('/cobrancas/<int:id>/baixa-manual-criar',methods=['POST'])
@login_required
def criar_baixa_manual_cobranca(id):
 c=Contract.query.options(joinedload(Contract.driver),joinedload(Contract.vehicle)).filter_by(id=id,tenant_id=tid()).first_or_404()
 if c.status not in ('Assinado','Ativo'):
  flash('Somente contratos vigentes podem receber baixa manual.','warning')
  return redirect(url_for('cobrancas'))
 periodicidade=unicodedata.normalize('NFKD',str(c.periodicidade or '')).encode('ascii','ignore').decode('ascii').lower()
 if periodicidade and 'seman' not in periodicidade:
  flash('A baixa manual rápida desta tela está disponível para contratos semanais.','warning')
  return redirect(url_for('cobrancas'))

 agora_local=datetime.now(_tenant_zone(c.tenant_id))
 hoje=agora_local.date()
 inicio_semana=hoje-timedelta(days=hoje.weekday())
 fim_semana=inicio_semana+timedelta(days=6)

 # Trabalha somente com a competência da semana atual. Cobranças antigas não são
 # baixadas por engano e também não impedem a criação da cobrança desta semana.
 audit_semana=BillingAudit.query.filter(
  BillingAudit.tenant_id==c.tenant_id,
  BillingAudit.contract_id==c.id,
  BillingAudit.billing_date>=inicio_semana,
  BillingAudit.billing_date<=fim_semana
 ).order_by(BillingAudit.id.desc()).first()
 if audit_semana:
  if (audit_semana.payment_status or 'PENDENTE')=='PAGO':
   flash('Este contrato já possui cobrança paga nesta semana.','info')
  else:
   audit_semana.payment_status='PAGO'
   audit_semana.paid_at=agora_sao_paulo_naive()
   audit_semana.paid_by_id=current_user.id
   audit_semana.payment_method='Baixa manual'
   audit_semana.payment_notes='Pagamento confirmado manualmente pela locadora, sem comprovante no Frota Fácil.'
   audit_semana.closed_at=audit_semana.paid_at
   auditorias_semana=_auditorias_semana_contrato(c.tenant_id,c.id,hoje)
   _cancelar_lembretes_cobranca_ids(c.tenant_id,[x.id for x in auditorias_semana],'Lembrete cancelado porque a cobrança da competência foi marcada como paga.')
   db.session.commit()
   flash('Cobrança desta semana marcada como paga. Lembretes pendentes foram cancelados e nenhuma nova mensagem foi enviada.','success')
  return redirect(url_for('cobrancas'))

 info=calcular_cobranca_semanal(c)
 pago_em=agora_sao_paulo_naive()
 audit=BillingAudit(
  tenant_id=c.tenant_id,
  contract_id=c.id,
  driver_name=c.driver.nome if c.driver else None,
  vehicle_label=c.vehicle.marca_modelo if c.vehicle else None,
  plate=c.vehicle.placa if c.vehicle else None,
  billing_date=hoje,
  base_amount=info['valor_base'],
  km_period=info.get('km_periodo'),
  km_limit=info.get('limite_km'),
  km_excess=info.get('km_excedente') or 0,
  excess_rate=c.valor_km_excedente or 0,
  excess_amount=info.get('valor_excesso') or 0,
  total_amount=info['total'],
  body='Baixa manual criada pela locadora sem envio de mensagem.',
  template_name=None,
  provider='manual',
  status='BAIXA_MANUAL',
  payment_status='PAGO',
  paid_at=pago_em,
  paid_by_id=current_user.id,
  payment_method='Baixa manual',
  payment_notes='Pagamento confirmado manualmente pela locadora, sem comprovante no Frota Fácil.',
  closed_at=pago_em,
  created_at=pago_em
 )
 db.session.add(audit)
 db.session.flush()
 auditorias_semana=_auditorias_semana_contrato(c.tenant_id,c.id,hoje)
 _cancelar_lembretes_cobranca_ids(c.tenant_id,[x.id for x in auditorias_semana],'Lembrete cancelado porque a cobrança da competência foi marcada como paga.')
 db.session.commit()
 flash('Cobrança criada e marcada como paga. Lembretes pendentes foram cancelados e nenhuma nova mensagem foi enviada.','success')
 return redirect(url_for('cobrancas'))

@app.route('/cobrancas/auditoria/<int:id>/reabrir',methods=['POST'])
@login_required
def reabrir_cobranca(id):
 audit=BillingAudit.query.filter_by(id=id,tenant_id=tid()).first_or_404()
 audit.payment_status='COMPROVANTE_RECEBIDO' if audit.receipt_key else 'PENDENTE'; audit.paid_at=None; audit.paid_by_id=None; audit.payment_method=None; audit.payment_notes=None; audit.closed_at=None
 db.session.commit(); flash('Cobrança reaberta. Ela volta a participar dos lembretes automáticos.','warning')
 return redirect(url_for('cobrancas'))

@app.route('/administracao/backup/criar',methods=['POST'])
@login_required
def criar_backup_tenant():
 payload=tenant_backup_payload(tid())
 conteudo=json.dumps(payload,ensure_ascii=False,indent=2).encode('utf-8')
 chave=f'{tid()}/backups/backup_{datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")}.json'
 try:
  storage.upload(BytesIO(conteudo),chave,'application/json')
  flash('Backup lógico criado no Cloudflare R2.','success')
 except Exception:
  app.logger.exception('Falha ao criar backup')
  flash('Não foi possível criar o backup.','danger')
 return redirect(url_for('administracao_armazenamento'))

@app.route('/administracao/backup/baixar')
@login_required
def baixar_backup_tenant():
 payload=tenant_backup_payload(tid())
 conteudo=json.dumps(payload,ensure_ascii=False,indent=2).encode('utf-8')
 nome=f'backup_frota_facil_{datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")}.json'
 return send_file(BytesIO(conteudo),as_attachment=True,download_name=nome,mimetype='application/json')

def _integration(tipo):
 return Integration.query.filter_by(tenant_id=tid(),tipo=tipo).first()

def _integration_config(item):
 if not item or not item.configuracao:
  return {}
 try:
  value=json.loads(item.configuracao)
  return value if isinstance(value,dict) else {}
 except Exception:
  return {}

def whatsapp_status_label(status):
 labels={
  'PENDENTE':'Pendente',
  'AGENDADA':'Agendada',
  'PREPARADA':'Preparada no WhatsApp Web',
  'ACEITA_META':'Aceita pela Meta — aguardando entrega',
  'ENVIADA':'Enviada pela Meta — aguardando entrega',
  'SENT':'Enviada — aguardando entrega',
  'ENTREGUE':'Entregue',
  'DELIVERED':'Entregue',
  'LIDA':'Lida',
  'READ':'Lida',
  'FALHA':'Falhou',
  'FAILED':'Falhou',
  'AGUARDANDO_MANUAL':'Aguardando envio manual',
 }
 return labels.get((status or '').upper(),status or '-')

def _whatsapp_status_rank(status):
 return {
  'PENDENTE':0,'AGENDADA':0,'ACEITA_META':1,'ENVIADA':2,'SENT':2,
  'ENTREGUE':3,'DELIVERED':3,'LIDA':4,'READ':4,'FALHA':99,'FAILED':99,
 }.get((status or '').upper(),0)

def _meta_error_text(status_payload):
 errors=status_payload.get('errors') or []
 if not errors:
  return None
 parts=[]
 for err in errors:
  if not isinstance(err,dict):
   continue
  code=err.get('code')
  title=err.get('title') or err.get('message')
  details=((err.get('error_data') or {}).get('details') if isinstance(err.get('error_data'),dict) else None)
  chunk=' — '.join(str(x) for x in [title,details] if x)
  if code is not None:
   chunk=(f'#{code}: '+chunk) if chunk else f'#{code}'
  if chunk: parts.append(chunk)
 return ' | '.join(parts) or None

@app.route('/webhooks/whatsapp',methods=['GET','POST'])
def whatsapp_webhook():
 # Verificação do callback feita pela Meta. O token é o configurado no tenant.
 if request.method=='GET':
  mode=(request.args.get('hub.mode') or '').strip()
  verify=(request.args.get('hub.verify_token') or '').strip()
  challenge=request.args.get('hub.challenge') or ''
  if mode!='subscribe' or not verify:
   abort(403)
  for item in Integration.query.filter_by(tipo='whatsapp').all():
   cfg=_integration_config(item)
   expected=(cfg.get('verify_token') or '').strip()
   if expected and verify==expected:
    return challenge,200,{'Content-Type':'text/plain; charset=utf-8'}
  abort(403)

 payload=request.get_json(silent=True) or {}
 try:
  entries=payload.get('entry') or []
  for entry in entries:
   for change in (entry.get('changes') or []):
    value=change.get('value') or {}
    metadata=value.get('metadata') or {}
    phone_number_id=str(metadata.get('phone_number_id') or '')
    for st in (value.get('statuses') or []):
     external_id=str(st.get('id') or '').strip()
     raw_status=str(st.get('status') or '').lower().strip()
     if not external_id or not raw_status:
      continue
     fila=MessageQueue.query.filter_by(external_id=external_id).order_by(MessageQueue.id.desc()).first()
     # Fallback defensivo: se o message id não foi salvo, tenta limitar pelo Phone Number ID e destinatário.
     if not fila and phone_number_id:
      recipient=str(st.get('recipient_id') or '').strip()
      candidates=MessageQueue.query.filter_by(provider='whatsapp_business',recipient=recipient).order_by(MessageQueue.id.desc()).limit(20).all()
      for candidate in candidates:
       integ=Integration.query.filter_by(tenant_id=candidate.tenant_id,tipo='whatsapp').first()
       cfg=_integration_config(integ)
       if str(cfg.get('phone_number_id') or '')==phone_number_id:
        fila=candidate; break
     if not fila:
      app.logger.warning('Webhook WhatsApp sem mensagem correspondente: %s',external_id)
      continue
     mapping={'sent':'ENVIADA','delivered':'ENTREGUE','read':'LIDA','failed':'FALHA'}
     new_status=mapping.get(raw_status,raw_status.upper())
     old_status=fila.status
     # Falha sempre deve prevalecer. Nos demais callbacks, evita regressão (ex.: read seguido de delivered atrasado).
     if new_status=='FALHA' or _whatsapp_status_rank(new_status)>=_whatsapp_status_rank(old_status):
      fila.status=new_status
     fila.updated_at=agora_sao_paulo_naive()
     if new_status in ('ENVIADA','ENTREGUE','LIDA') and not fila.sent_at:
      fila.sent_at=agora_sao_paulo_naive()
     error_text=_meta_error_text(st) if new_status=='FALHA' else None
     if error_text:
      fila.error_message=error_text
     event_description=f'Status WhatsApp atualizado pela Meta: {raw_status}.'
     if error_text:
      event_description+=' '+error_text
     db.session.add(MessageEvent(tenant_id=fila.tenant_id,message_id=fila.id,event=new_status,description=event_description,created_at=agora_sao_paulo_naive()))
     if fila.related_entity=='Cobranca' and fila.related_entity_id:
      audit=BillingAudit.query.filter_by(id=fila.related_entity_id,tenant_id=fila.tenant_id).first()
      if audit:
       audit.status=fila.status; audit.external_id=fila.external_id; audit.error_message=fila.error_message
  db.session.commit()
 except Exception:
  db.session.rollback()
  app.logger.exception('Falha ao processar webhook de status do WhatsApp')
  # Retorna 200 para evitar tempestade de retries enquanto o evento fica registrado em log.
 return {'ok':True},200

@app.route('/integracoes',methods=['GET','POST'])
@login_required
def integracoes():
 if request.method=='POST':
  section=request.form.get('section')
  if section=='whatsapp':
   item=_integration('whatsapp') or Integration(tenant_id=tid(),tipo='whatsapp')
   provider=request.form.get('provider','web')
   cfg=dict(_integration_config(item))
   cfg.update({
    'provider':provider,
    'phone_number_id':request.form.get('phone_number_id','').strip() or cfg.get('phone_number_id',''),
    'business_account_id':request.form.get('business_account_id','').strip() or cfg.get('business_account_id',''),
    'graph_version':request.form.get('graph_version','v23.0').strip() or 'v23.0',
    'contract_template_name':request.form.get('contract_template_name','').strip(),
    'inspection_template_name':request.form.get('inspection_template_name','').strip(),
    'owner_portal_template_name':request.form.get('owner_portal_template_name',cfg.get('owner_portal_template_name','acesso_portal_proprietario')).strip() or 'acesso_portal_proprietario',
    'mileage_template_name':request.form.get('mileage_template_name','').strip(),
    'maintenance_template_name':request.form.get('maintenance_template_name','').strip(),
    'maintenance_reminder_template_name':request.form.get('maintenance_reminder_template_name','').strip(),
    'oil_change_template_name':request.form.get('oil_change_template_name','').strip(),
    'payment_template_name':request.form.get('payment_template_name','').strip(),
    'payment_excess_template_name':request.form.get('payment_excess_template_name','').strip(),
    'template_language':request.form.get('template_language','pt_BR').strip() or 'pt_BR',
    'automation_enabled':bool(request.form.get('automation_enabled')),
    'automation_weekday':int(request.form.get('automation_weekday') or 0),
    'automation_start_hour':int(request.form.get('automation_start_hour') or 7),
    'automation_end_hour':int(request.form.get('automation_end_hour') or 20),
    'reminder_interval_hours':max(1,int(request.form.get('reminder_interval_hours') or 1)),
    'automatic_billing_enabled':bool(request.form.get('automatic_billing_enabled')),
    'automatic_km_enabled':bool(request.form.get('automatic_km_enabled')),
    'automatic_alerts_enabled':bool(request.form.get('automatic_alerts_enabled')),
   })
   novo_access_token=request.form.get('access_token','').strip()
   novo_verify_token=request.form.get('verify_token','').strip()
   if novo_access_token: cfg['access_token']=novo_access_token
   if novo_verify_token: cfg['verify_token']=novo_verify_token
   item.ativo=(provider=='business')
   item.configuracao=json.dumps(cfg,ensure_ascii=False)
   db.session.add(item); db.session.commit(); flash('Configuração do WhatsApp salva.','success')
  elif section=='signature':
   item=_integration('signature') or Integration(tenant_id=tid(),tipo='signature')
   provider=request.form.get('signature_provider','local')
   cfg={'provider':provider}
   if provider=='clicksign':
    cfg.update({'api_token':request.form.get('clicksign_api_token','').strip(),'workspace_key':request.form.get('clicksign_workspace_key','').strip(),'environment':request.form.get('clicksign_environment','sandbox')})
   elif provider=='docusign':
    cfg.update({'account_id':request.form.get('docusign_account_id','').strip(),'integration_key':request.form.get('docusign_integration_key','').strip(),'environment':request.form.get('docusign_environment','demo')})
   item.ativo=(provider!='local')
   item.configuracao=json.dumps(cfg,ensure_ascii=False)
   db.session.add(item); db.session.commit(); flash('Provedor de assinatura salvo.','success')
  return redirect(url_for('integracoes'))
 whatsapp_item=_integration('whatsapp'); signature_item=_integration('signature')
 whatsapp_cfg=_integration_config(whatsapp_item); signature_cfg=_integration_config(signature_item)
 signature_ready,signature_message=SignatureProviderService.readiness(SignatureProviderService.from_integration(signature_item))
 recentes=MessageQueue.query.filter_by(tenant_id=tid()).order_by(MessageQueue.id.desc()).limit(20).all()
 for fila_recente in recentes:
  for campo_data in ('created_at','updated_at','sent_at','scheduled_at'):
   valor=getattr(fila_recente,campo_data,None)
   if valor:
    set_committed_value(fila_recente,campo_data,_message_db_time_as_utc_naive(valor))
 return render_template('integracoes.html',whatsapp=whatsapp_item,whatsapp_cfg=whatsapp_cfg,signature=signature_item,signature_cfg=signature_cfg,signature_ready=signature_ready,signature_message=signature_message,recentes=recentes,status_label=whatsapp_status_label,webhook_url=url_for('whatsapp_webhook',_external=True),meta_embedded_ready=bool(META_APP_ID and META_APP_SECRET and META_WHATSAPP_CONFIG_ID),meta_app_id=META_APP_ID,meta_config_id=META_WHATSAPP_CONFIG_ID,meta_graph_version=META_GRAPH_VERSION)

@app.route('/integracoes/whatsapp/embedded-signup/concluir',methods=['POST'])
@login_required
def concluir_whatsapp_embedded_signup():
 """Conclui o Embedded Signup sem remover o modo manual existente.

 O navegador entrega o authorization code e os IDs da sessão. O App Secret
 permanece somente no backend. A credencial resultante fica isolada no
 Integration do tenant atual.
 """
 if not (META_APP_ID and META_APP_SECRET and META_WHATSAPP_CONFIG_ID):
  return {'ok':False,'error':'Embedded Signup ainda não foi habilitado no ambiente do Frota Fácil.'},400
 data=request.get_json(silent=True) or {}
 code=str(data.get('code') or '').strip()
 waba_id=str(data.get('waba_id') or '').strip()
 phone_number_id=str(data.get('phone_number_id') or '').strip()
 redirect_uri=str(data.get('redirect_uri') or '').strip()
 if not code:
  return {'ok':False,'error':'A Meta não retornou o código de autorização.'},400
 if not redirect_uri:
  return {'ok':False,'error':'Não foi possível identificar a URL que iniciou a autorização da Meta.'},400
 # O redirect_uri usado para trocar o code precisa ser exatamente o mesmo da página
 # que iniciou o FB.login(). Aceita apenas URL do próprio Frota Fácil.
 current_origin=request.host_url.rstrip('/')
 if not redirect_uri.startswith(current_origin + '/'):
  return {'ok':False,'error':'URL de retorno da Meta inválida para este ambiente.'},400
 try:
  resp=requests.get(
   f'https://graph.facebook.com/{META_GRAPH_VERSION}/oauth/access_token',
   params={
    'client_id':META_APP_ID,
    'client_secret':META_APP_SECRET,
    'code':code,
    'redirect_uri':redirect_uri,
   },
   timeout=20,
  )
  payload=resp.json()
 except Exception as exc:
  app.logger.exception('Falha no Embedded Signup Meta')
  return {'ok':False,'error':f'Falha ao concluir autorização Meta: {exc}'},502
 if not resp.ok or not payload.get('access_token'):
  detail=(payload.get('error') or {}).get('message') if isinstance(payload,dict) else None
  return {'ok':False,'error':detail or 'A Meta não retornou um Access Token válido.'},400
 token=payload['access_token']
 # Se os IDs não vieram do evento SESSION_INFO, tenta descobrir WABA via token.
 if not waba_id:
  try:
   dbg=requests.get(f'https://graph.facebook.com/{META_GRAPH_VERSION}/debug_token',params={'input_token':token,'access_token':f'{META_APP_ID}|{META_APP_SECRET}'},timeout=20).json()
   scopes=(dbg.get('data') or {}).get('granular_scopes') or []
   for scope in scopes:
    if scope.get('scope')=='whatsapp_business_management' and scope.get('target_ids'):
     waba_id=str(scope['target_ids'][0]); break
  except Exception:
   app.logger.exception('Não foi possível descobrir WABA pelo debug_token')
 if waba_id and not phone_number_id:
  try:
   nums=requests.get(f'https://graph.facebook.com/{META_GRAPH_VERSION}/{waba_id}/phone_numbers',headers={'Authorization':f'Bearer {token}'},params={'fields':'id,display_phone_number,verified_name'},timeout=20).json()
   if nums.get('data'):
    phone_number_id=str(nums['data'][0].get('id') or '')
  except Exception:
   app.logger.exception('Não foi possível descobrir Phone Number ID')
 if not waba_id or not phone_number_id:
  return {'ok':False,'error':'Autorização concluída, mas não foi possível identificar WABA e número. Use a configuração avançada enquanto revisamos a conta Meta.'},400
 item=_integration('whatsapp') or Integration(tenant_id=tid(),tipo='whatsapp')
 oldcfg=_integration_config(item)
 oldcfg.update({
  'provider':'business','onboarding_mode':'embedded_signup','business_account_id':waba_id,
  'phone_number_id':phone_number_id,'access_token':token,'graph_version':META_GRAPH_VERSION,
  'embedded_connected_at':agora_sao_paulo_naive().isoformat(),
 })
 item.ativo=True; item.configuracao=json.dumps(oldcfg,ensure_ascii=False)
 db.session.add(item); db.session.commit()
 return {'ok':True,'waba_id':waba_id,'phone_number_id':phone_number_id}

@app.route('/integracoes/whatsapp/embedded-signup/desconectar',methods=['POST'])
@login_required
def desconectar_whatsapp_embedded_signup():
 item=_integration('whatsapp')
 if item:
  cfg=_integration_config(item)
  # Não revoga ativos na Meta automaticamente: apenas retira a credencial do tenant.
  cfg.update({'provider':'web','onboarding_mode':'manual','access_token':'','phone_number_id':'','business_account_id':''})
  item.ativo=False; item.configuracao=json.dumps(cfg,ensure_ascii=False); db.session.commit()
 flash('Conexão simplificada removida do Frota Fácil. Os ativos da Meta não foram excluídos.','success')
 return redirect(url_for('integracoes'))

@app.route('/integracoes/whatsapp/templates-meta',methods=['GET'])
@login_required
def templates_meta_whatsapp():
 item=_integration('whatsapp')
 cfg=_integration_config(item)
 waba_id=(cfg.get('business_account_id') or '').strip()
 token=(cfg.get('access_token') or '').strip()
 graph_version=(cfg.get('graph_version') or 'v23.0').strip() or 'v23.0'
 templates=[]
 erro=None
 if not waba_id or not token:
  erro='Business Account ID (WABA) e Access Token precisam estar configurados para consultar os templates.'
 else:
  try:
   url=f'https://graph.facebook.com/{graph_version}/{waba_id}/message_templates'
   params={'fields':'id,name,language,status,category','limit':100}
   headers={'Authorization':f'Bearer {token}'}
   while url:
    resp=requests.get(url,headers=headers,params=params if '?' not in url else None,timeout=20)
    try:
     payload=resp.json()
    except ValueError:
     payload={'raw':resp.text[:2000]}
    if not resp.ok:
     detail=(payload.get('error') or {}).get('message') if isinstance(payload,dict) else None
     raise RuntimeError(detail or f'Falha ao consultar templates na Meta (HTTP {resp.status_code}).')
    templates.extend(payload.get('data') or [])
    url=((payload.get('paging') or {}).get('next'))
    params=None
  except Exception as exc:
   app.logger.exception('Falha ao consultar templates WhatsApp da Meta')
   erro=str(exc)
 templates=sorted(templates,key=lambda x:((x.get('name') or '').lower(),x.get('language') or ''))
 return render_template('meta_templates.html',templates=templates,erro=erro,waba_id=waba_id,graph_version=graph_version,config_template=cfg.get('mileage_template_name') or '',config_language=cfg.get('template_language') or '')

@app.route('/automacoes/processar-mensagens',methods=['POST'])
@login_required
def processar_mensagens_manual():
 km=processar_km_automatico(tid()); cobrancas=processar_cobrancas_automaticas(tid()); vistorias_auto=processar_vistorias_automaticas(tid()); alertas=processar_alertas_automaticos(tid()); quantidade=processar_mensagens_agendadas(tid(),limit=200)
 flash(f'Automação processada: {km} KM, {cobrancas} cobrança(s), {vistorias_auto} vistoria(s), {alertas} alerta(s) e {quantidade} mensagem(ns) agendada(s).','success')
 return redirect(url_for('configuracoes_automacoes'))

@app.route('/jobs/processar-mensagens',methods=['GET','POST'])
def processar_mensagens_job():
 token=(request.args.get('token') or request.headers.get('X-Automation-Token') or '').strip()
 expected=(os.getenv('AUTOMATION_JOB_TOKEN') or '').strip()
 if not expected or token != expected:
  abort(403)
 km=processar_km_automatico(None)
 cobrancas=processar_cobrancas_automaticas(None)
 vistorias_auto=processar_vistorias_automaticas(None)
 alertas=processar_alertas_automaticos(None)
 quantidade=processar_mensagens_agendadas(None,limit=500)
 return {'ok':True,'solicitacoes_km':km,'cobrancas_geradas':cobrancas,'vistorias_enviadas':vistorias_auto,'alertas_enviados':alertas,'mensagens_agendadas_processadas':quantidade,'executado_em':agora_sao_paulo_naive().isoformat()}

@app.route('/integracoes/whatsapp/testar',methods=['POST'])
@login_required
def testar_whatsapp_business():
 item=_integration('whatsapp')
 cfg=_integration_config(item)
 telefone=normalize_phone(request.form.get('telefone'))
 if not telefone:
  flash('Informe um telefone para o teste.','danger'); return redirect(url_for('integracoes'))
 mensagem='Teste de integração enviado pelo Frota Fácil.'
 provider_cfg=(cfg.get('provider') or 'web').lower()
 fila=MessageQueue(tenant_id=tid(),channel='whatsapp',provider='whatsapp_business' if provider_cfg=='business' else 'whatsapp_web',recipient=telefone,recipient_name='Teste',message_type='teste',body=mensagem,status='PENDENTE',created_at=agora_sao_paulo_naive(),updated_at=agora_sao_paulo_naive())
 db.session.add(fila); db.session.flush()
 try:
  result=CommunicationService().send_whatsapp(phone=telefone,message=mensagem,integration=item,template_name=cfg.get('contract_template_name') or None,template_language=cfg.get('template_language') or 'pt_BR')
  fila.provider=result.provider; fila.status=result.status; fila.external_id=result.external_id; fila.attempts=1; fila.sent_at=agora_sao_paulo_naive() if result.status=='ENVIADA' else None
  db.session.add(MessageEvent(tenant_id=tid(),message_id=fila.id,event=result.status,description='Teste manual de integração.',created_at=agora_sao_paulo_naive()))
  db.session.commit(); flash('Teste processado com sucesso.','success')
 except CommunicationError as exc:
  fila.status='FALHA'; fila.error_message=str(exc); fila.attempts=1
  db.session.add(MessageEvent(tenant_id=tid(),message_id=fila.id,event='FALHA',description=str(exc),created_at=agora_sao_paulo_naive()))
  db.session.commit(); flash(str(exc),'danger')
 return redirect(url_for('integracoes'))

with app.app_context(): seed()
if __name__=='__main__': app.run(host='0.0.0.0',port=int(os.getenv('PORT',5000)),debug=True)
